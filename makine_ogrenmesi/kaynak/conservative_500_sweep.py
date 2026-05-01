"""500/500 icin daha muhafazakar sentetik benchmark aramasi.

Bu modül mevcut controlled_benchmark small-scale hattını bozmaz. Amaç 500/500 adayının
mevcut skorlarını koruyarak Cohen's d ve dağılım kaymasını azaltabilecek
distribution-matched üretim varyantlarını denemektir.
"""

from __future__ import annotations

import importlib
import json
import math
import shutil
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from sklearn.ensemble import ExtraTreesClassifier, RandomForestClassifier, VotingClassifier

from .controlled_synthetic_benchmark import (
    DatasetCandidate,
    ModelSpec,
    _apply_imputer,
    _class_distribution,
    _clip_clinical,
    _cv_summary,
    _dataset_leakage_report,
    _evaluate_external,
    _evaluate_group_cv,
    _evaluate_holdout,
    _external_holdout_ayir,
    _feature_specs,
    _fit_dev_median_imputer,
    _group_holdout_split,
    _json_clean,
    _validate_metadata,
)
from .benchmark_scale_sweep_sweep import (
    _class_separation,
    _distribution_shift,
    _feature_label,
    _model_label,
    _small_catboost,
    _small_lightgbm,
    _small_xgboost,
)
from .ozellik_yapilandirmasi import HEDEF_KOLONU, OZELLIK_KOLONLARI
from .veri_yukleyici import veri_setini_yukle


RANDOM_STATE = 42
TARGET_PER_CLASS = 500

BASELINE_ACCURACY = 0.9200
BASELINE_MIN_MAIN = 0.9100
BASELINE_CV_MIN_MAIN = 0.8560
BASELINE_COHENS_D_REPORTED = 0.704
BASELINE_SHIFT_REPORTED = 0.105
BASELINE_DATASET = "scale_sweep_500_per_class_strength_0p40_low"


@dataclass(frozen=True)
class ConservativeConfig:
    strength: float
    noise_profile: str
    parent_pull: float
    mean_alpha: float
    boundary_ratio: float
    seed_offset: int

    @property
    def key(self) -> str:
        strength = str(self.strength).replace(".", "p")
        pull = str(self.parent_pull).replace(".", "p")
        mean = str(self.mean_alpha).replace(".", "p")
        boundary = str(self.boundary_ratio).replace(".", "p")
        return f"cons500_s{strength}_{self.noise_profile}_p{pull}_m{mean}_b{boundary}_r{self.seed_offset}"


def conservative_500_sweep_calistir(
    *,
    proje_koku: Path,
    veri_yolu: Path,
    random_state: int = RANDOM_STATE,
    n_jobs: int = 2,
    quick: bool = False,
    word_raporu_yaz: bool = True,
) -> dict[str, Any]:
    """Detayli 500/500 muhafazakar sweep kosusunu calistirir."""
    proje_koku = Path(proje_koku)
    veri_yolu = Path(veri_yolu)
    rapor_dir = proje_koku / "makine_ogrenmesi" / "raporlar"
    veri_dir = proje_koku / "makine_ogrenmesi" / "veri" / "deneysel" / "conservative_500"
    grafik_dir = rapor_dir / "grafikler_conservative_500"
    for path in (rapor_dir, veri_dir, grafik_dir):
        path.mkdir(parents=True, exist_ok=True)

    raw = veri_setini_yukle(veri_yolu).copy()
    raw["original_index"] = raw.index.astype(int)
    raw["source_id"] = [f"original_{i}" for i in raw["original_index"]]
    original_dev_raw, external_raw = _external_holdout_ayir(raw, random_state)
    imputer = _fit_dev_median_imputer(original_dev_raw)
    original_dev = _apply_imputer(original_dev_raw, imputer)
    external_holdout = _apply_imputer(external_raw, imputer)

    configs = _configs(quick)
    datasets = [_generate_conservative_dataset(original_dev, config, random_state + config.seed_offset) for config in configs]
    for dataset in datasets:
        _write_dataset(veri_dir, dataset)

    quality_by_dataset = {dataset.name: _quality(original_dev, dataset) for dataset in datasets}
    baseline_quality = _baseline_quality(original_dev)
    screened = _screen_datasets(datasets, quality_by_dataset, baseline_quality, quick)

    model_specs = _model_specs(random_state, n_jobs, quick)
    feature_specs = _feature_specs_for_run(quick)
    holdout_results = _run_holdout(screened, feature_specs, model_specs, random_state)
    successful = [r for r in holdout_results if r.get("status") == "completed"]

    cv_candidates = _select_cv_candidates(successful, quality_by_dataset, quick)
    cv_results = []
    for row in cv_candidates:
        dataset = next(d for d in datasets if d.name == row["dataset_name"])
        feature = next(f for f in feature_specs if f.name == row["feature_set"])
        model = row["_model_spec"]
        cv = _evaluate_group_cv(dataset=dataset, feature_spec=feature, model_spec=model, random_state=random_state)
        cv_results.append(cv)
        summary = cv["summary"]
        print(
            f"[conservative-500-cv/{dataset.name}/{_feature_label(feature.name)}/{_model_label(model.name)}] "
            f"min={summary['ana_metrik_minimumu_mean']:.4f} acc={summary['accuracy_mean']:.4f}"
        )

    leakage_report = [_dataset_leakage_report(dataset, external_holdout, successful, cv_results) for dataset in datasets]
    global _LAST_SUCCESSFUL
    _LAST_SUCCESSFUL = successful
    final = _select_final(successful, cv_results, quality_by_dataset, leakage_report)
    final_cv = _find_cv(cv_results, final)
    external_metrics = _evaluate_external(final, external_holdout) if final is not None else None

    sweep = {
        "created_at_utc": _now(),
        "goal": "500/500 skorlarini koruyarak Cohen's d ve shift degerlerini azaltma aramasi.",
        "baseline_thresholds": {
            "accuracy": BASELINE_ACCURACY,
            "min_main_metric": BASELINE_MIN_MAIN,
            "group_cv_min_main_metric": BASELINE_CV_MIN_MAIN,
            "reported_cohens_d": BASELINE_COHENS_D_REPORTED,
            "reported_shift": BASELINE_SHIFT_REPORTED,
            "baseline_dataset": BASELINE_DATASET,
        },
        "original_csv": {"path": str(veri_yolu), "rows": int(len(raw)), "class_distribution": _class_distribution(raw)},
        "original_dev": {"rows": int(len(original_dev)), "class_distribution": _class_distribution(original_dev)},
        "external_holdout": {"rows": int(len(external_holdout)), "class_distribution": _class_distribution(external_holdout), "role": "Sadece dis kontrol."},
        "config_count": len(configs),
        "screened_dataset_count": len(screened),
        "datasets": [d.report for d in datasets],
        "quality_by_dataset": quality_by_dataset,
        "baseline_quality_computed": baseline_quality,
        "holdout_results": [_json_clean(r) for r in holdout_results],
        "cv_results": [_json_clean(c) for c in cv_results],
        "leakage_report": leakage_report,
    }
    selection = _selection_report(final, final_cv, external_metrics, quality_by_dataset, leakage_report)
    graphics = _make_graphics(grafik_dir, sweep, selection)
    sweep["graphics"] = graphics
    selection["graphics"] = graphics
    audit = _audit(raw, selection)

    _json_write(rapor_dir / "conservative_500_sweep_results.json", sweep)
    _json_write(rapor_dir / "conservative_500_selection_report.json", selection)
    _json_write(rapor_dir / "conservative_500_quality_audit.json", audit)

    word_path = None
    desktop_word_path = None
    if word_raporu_yaz:
        word_path = _write_word_report(proje_koku, sweep, selection, graphics)
        desktop_word_path = Path.home() / "Desktop" / "pima_conservative_scale_sweep_report.docx"
        shutil.copy2(word_path, desktop_word_path)

    result = {
        "sweep_results": sweep,
        "selection_report": selection,
        "quality_audit": audit,
        "word_report": str(word_path) if word_path else None,
        "desktop_word_report": str(desktop_word_path) if desktop_word_path else None,
    }
    _print_summary(result)
    return result


def _configs(quick: bool) -> list[ConservativeConfig]:
    strengths = (0.20, 0.30, 0.40) if quick else (0.16, 0.20, 0.24, 0.28, 0.32, 0.36, 0.40)
    noises = ("very_low", "low") if quick else ("very_low", "low", "adaptive_low", "boundary_low")
    pulls = (0.0, 0.25) if quick else (0.0, 0.12, 0.25, 0.38, 0.50)
    means = (0.0, 0.35) if quick else (0.0, 0.20, 0.40, 0.60)
    boundaries = (0.0,) if quick else (0.0, 0.25, 0.50)
    configs: list[ConservativeConfig] = []
    idx = 0
    for strength in strengths:
        for noise in noises:
            for pull in pulls:
                for mean in means:
                    for boundary in boundaries:
                        idx += 1
                        configs.append(ConservativeConfig(strength, noise, pull, mean, boundary, idx))
    return configs


def _generate_conservative_dataset(dev: pd.DataFrame, config: ConservativeConfig, random_state: int) -> DatasetCandidate:
    rng = np.random.default_rng(random_state)
    data_parts = [dev[OZELLIK_KOLONLARI + [HEDEF_KOLONU]].copy()]
    meta_parts = [
        pd.DataFrame(
            {
                "source_id": dev["source_id"].astype(str).to_numpy(),
                "is_synthetic": False,
                "parent_original_index": dev["original_index"].astype(int).to_numpy(),
                "generation_method": "original_dev",
                "synthetic_strength": 0.0,
                "noise_profile": "original_dev",
            }
        )
    ]
    method_counts = {"original_dev": int(len(dev))}
    for cls in [0, 1]:
        cls_dev = dev[dev[HEDEF_KOLONU] == cls].copy()
        other_dev = dev[dev[HEDEF_KOLONU] != cls].copy()
        need = max(0, TARGET_PER_CLASS - len(cls_dev))
        if need <= 0:
            continue
        generated = _generate_class_rows(cls_dev, other_dev, cls, need, rng, config)
        data_parts.append(generated[OZELLIK_KOLONLARI + [HEDEF_KOLONU]])
        meta_parts.append(generated[["source_id", "is_synthetic", "parent_original_index", "generation_method", "synthetic_strength", "noise_profile"]])
        for method, count in generated["generation_method"].value_counts().items():
            method_counts[str(method)] = method_counts.get(str(method), 0) + int(count)

    data = pd.concat(data_parts, ignore_index=True)
    metadata = pd.concat(meta_parts, ignore_index=True)
    data = _mean_match_synthetic(data, metadata, dev, config.mean_alpha)
    data = _clip_clinical(data)
    metadata.insert(0, "row_id", [f"{config.key}_r_{i:06d}" for i in range(len(metadata))])
    _validate_metadata(metadata)
    report = {
        "dataset_name": config.key,
        "target_per_class": TARGET_PER_CLASS,
        "synthetic_strength": float(config.strength),
        "noise_profile": config.noise_profile,
        "parent_pull": float(config.parent_pull),
        "mean_alpha": float(config.mean_alpha),
        "boundary_ratio": float(config.boundary_ratio),
        "original_dev_rows": int((~metadata["is_synthetic"]).sum()),
        "synthetic_rows": int(metadata["is_synthetic"].sum()),
        "total_rows": int(len(data)),
        "class_distribution": _class_distribution(data),
        "source_id_family_count": int(metadata["source_id"].nunique()),
        "generation_method_counts": method_counts,
        "synthetic_valid_source_id_ratio": float(metadata.loc[metadata["is_synthetic"], "source_id"].astype(str).str.match(r"^original_\d+$").mean()),
    }
    return DatasetCandidate(config.key, TARGET_PER_CLASS, data.reset_index(drop=True), metadata.reset_index(drop=True), report)


def _generate_class_rows(
    cls_dev: pd.DataFrame,
    other_dev: pd.DataFrame,
    cls: int,
    count: int,
    rng: np.random.Generator,
    config: ConservativeConfig,
) -> pd.DataFrame:
    parent_idx = _sample_parent_indices(cls_dev, other_dev, cls, count, rng, config.boundary_ratio)
    parents = cls_dev.loc[parent_idx].reset_index(drop=True)
    parent_values = parents[OZELLIK_KOLONLARI].astype(float).to_numpy()
    target = _target_values(cls_dev, other_dev, cls, count, rng, config.noise_profile)
    strength = rng.normal(config.strength, _strength_sd(config.noise_profile), size=(count, 1))
    strength = np.clip(strength, 0.03, 0.62)
    values = parent_values * (1 - strength) + target * strength
    values += _noise(cls_dev, count, rng, config.noise_profile)
    values = parent_values * config.parent_pull + values * (1 - config.parent_pull)
    out = pd.DataFrame(values, columns=OZELLIK_KOLONLARI)
    out = _quantile_clip(out, cls_dev, config.noise_profile)
    out = _clip_clinical(out)
    out[HEDEF_KOLONU] = int(cls)
    out["source_id"] = parents["source_id"].astype(str).to_numpy()
    out["is_synthetic"] = True
    out["parent_original_index"] = parents["original_index"].astype(int).to_numpy()
    out["generation_method"] = "conservative_distribution_matched_one_parent"
    out["synthetic_strength"] = strength.ravel()
    out["noise_profile"] = config.noise_profile
    return out


def _sample_parent_indices(cls_dev: pd.DataFrame, other_dev: pd.DataFrame, cls: int, count: int, rng: np.random.Generator, boundary_ratio: float) -> np.ndarray:
    if boundary_ratio <= 0:
        return rng.choice(cls_dev.index.to_numpy(), size=count, replace=True)
    risk_cols = ["Glucose", "BMI", "Age", "DiabetesPedigreeFunction"]
    all_ref = pd.concat([cls_dev, other_dev], axis=0)
    mean = all_ref[risk_cols].mean()
    std = all_ref[risk_cols].std().replace(0, 1.0)
    score = ((cls_dev[risk_cols] - mean) / std).sum(axis=1)
    boundary_count = int(round(count * boundary_ratio))
    random_count = count - boundary_count
    if cls == 1:
        boundary_pool = score.nsmallest(max(6, min(len(score), int(len(score) * 0.45)))).index.to_numpy()
    else:
        boundary_pool = score.nlargest(max(6, min(len(score), int(len(score) * 0.45)))).index.to_numpy()
    random_part = rng.choice(cls_dev.index.to_numpy(), size=random_count, replace=True)
    boundary_part = rng.choice(boundary_pool, size=boundary_count, replace=True)
    return np.concatenate([random_part, boundary_part])


def _target_values(cls_dev: pd.DataFrame, other_dev: pd.DataFrame, cls: int, count: int, rng: np.random.Generator, profile: str) -> np.ndarray:
    class_mean = cls_dev[OZELLIK_KOLONLARI].mean().to_numpy(dtype=float)
    other_mean = other_dev[OZELLIK_KOLONLARI].mean().to_numpy(dtype=float)
    if profile == "boundary_low":
        anchor = class_mean * 0.82 + other_mean * 0.18
    elif profile == "adaptive_low":
        anchor = class_mean * 0.90 + other_mean * 0.10
    else:
        anchor = class_mean
    scale = {
        "very_low": 0.18,
        "low": 0.26,
        "adaptive_low": 0.34,
        "boundary_low": 0.30,
    }.get(profile, 0.26)
    std = cls_dev[OZELLIK_KOLONLARI].std().replace(0, 1.0).to_numpy(dtype=float)
    return rng.normal(anchor, std * scale, size=(count, len(OZELLIK_KOLONLARI)))


def _strength_sd(profile: str) -> float:
    return {"very_low": 0.015, "low": 0.022, "adaptive_low": 0.030, "boundary_low": 0.028}.get(profile, 0.022)


def _noise(cls_dev: pd.DataFrame, count: int, rng: np.random.Generator, profile: str) -> np.ndarray:
    multiplier = {"very_low": 0.010, "low": 0.020, "adaptive_low": 0.030, "boundary_low": 0.026}.get(profile, 0.020)
    std = cls_dev[OZELLIK_KOLONLARI].std().replace(0, 1.0).to_numpy(dtype=float)
    return rng.normal(0, std * multiplier, size=(count, len(OZELLIK_KOLONLARI)))


def _quantile_clip(frame: pd.DataFrame, cls_dev: pd.DataFrame, profile: str) -> pd.DataFrame:
    out = frame.copy()
    lo_q, hi_q = (0.03, 0.97) if profile in {"very_low", "low"} else (0.02, 0.98)
    margin = {"very_low": 0.04, "low": 0.08, "adaptive_low": 0.11, "boundary_low": 0.09}.get(profile, 0.08)
    lows = cls_dev[OZELLIK_KOLONLARI].quantile(lo_q)
    highs = cls_dev[OZELLIK_KOLONLARI].quantile(hi_q)
    span = (highs - lows).replace(0, 1.0)
    for col in OZELLIK_KOLONLARI:
        out[col] = out[col].clip(float(lows[col] - span[col] * margin), float(highs[col] + span[col] * margin))
    return out


def _mean_match_synthetic(data: pd.DataFrame, metadata: pd.DataFrame, dev: pd.DataFrame, alpha: float) -> pd.DataFrame:
    out = data.copy()
    for col in OZELLIK_KOLONLARI:
        out[col] = out[col].astype(float)
    if alpha <= 0:
        return out
    synthetic_mask = metadata["is_synthetic"].to_numpy(dtype=bool)
    for col in OZELLIK_KOLONLARI:
        diff = float(out[col].mean() - dev[col].mean())
        out.loc[synthetic_mask, col] = out.loc[synthetic_mask, col].astype(float) - alpha * diff
    return out


def _baseline_quality(original_dev: pd.DataFrame) -> dict[str, Any]:
    base_path = Path("makine_ogrenmesi/veri/deneysel/benchmark_scale_sweep/scale_sweep_500_per_class_strength_0p40_low.csv")
    if base_path.exists():
        frame = pd.read_csv(base_path)[OZELLIK_KOLONLARI + [HEDEF_KOLONU]]
        return {
            "dataset_name": BASELINE_DATASET,
            "computed_cohens_d": _class_separation(frame)["aggregate"]["mean_abs_cohens_d"],
            "computed_shift": _distribution_shift(original_dev, frame)["aggregate"]["avg_abs_z_mean_shift"],
            "reported_cohens_d": BASELINE_COHENS_D_REPORTED,
            "reported_shift": BASELINE_SHIFT_REPORTED,
            "note": "Computed deger mevcut CSV'den, reported deger son rapor tablosundan gelir.",
        }
    return {
        "dataset_name": BASELINE_DATASET,
        "computed_cohens_d": None,
        "computed_shift": None,
        "reported_cohens_d": BASELINE_COHENS_D_REPORTED,
        "reported_shift": BASELINE_SHIFT_REPORTED,
        "note": "Baseline CSV bulunamadi; reported degerler kullanildi.",
    }


def _quality(original_dev: pd.DataFrame, dataset: DatasetCandidate) -> dict[str, Any]:
    class_sep = _class_separation(dataset.frame)
    shift = _distribution_shift(original_dev, dataset.frame)
    return {
        "dataset_name": dataset.name,
        "class_separation": class_sep,
        "distribution_shift_vs_original_dev": shift,
        "quality_goal": {
            "lower_than_reported_500_cohens_d": bool(class_sep["aggregate"]["mean_abs_cohens_d"] < BASELINE_COHENS_D_REPORTED),
            "lower_than_reported_500_shift": bool(shift["aggregate"]["avg_abs_z_mean_shift"] < BASELINE_SHIFT_REPORTED),
        },
    }


def _screen_datasets(
    datasets: list[DatasetCandidate],
    quality_by_dataset: dict[str, dict[str, Any]],
    baseline_quality: dict[str, Any],
    quick: bool,
) -> list[DatasetCandidate]:
    baseline_cohen = baseline_quality.get("computed_cohens_d") or BASELINE_COHENS_D_REPORTED
    baseline_shift = baseline_quality.get("computed_shift") or BASELINE_SHIFT_REPORTED

    def score(dataset: DatasetCandidate) -> float:
        q = quality_by_dataset[dataset.name]
        cohen = q["class_separation"]["aggregate"]["mean_abs_cohens_d"]
        shift = q["distribution_shift_vs_original_dev"]["aggregate"]["avg_abs_z_mean_shift"]
        cohen_gain = baseline_cohen - cohen
        shift_gain = baseline_shift - shift
        # Cohen cok fazla duserse skorlarin dusme olasiligi artar; orta bolgeyi tercih et.
        underfit_penalty = abs(cohen - 0.68) * 0.10 if cohen < 0.62 else 0.0
        return 1.15 * cohen_gain + 0.85 * shift_gain - underfit_penalty

    limit = 18 if quick else 72
    selected = sorted(datasets, key=score, reverse=True)[:limit]
    # Baseline'e yakin guclu adaylari kacirmamak icin cohen/shift filtresi gevsek bir ek havuz da eklenir.
    near = [
        d
        for d in datasets
        if quality_by_dataset[d.name]["class_separation"]["aggregate"]["mean_abs_cohens_d"] <= baseline_cohen + 0.015
        and quality_by_dataset[d.name]["distribution_shift_vs_original_dev"]["aggregate"]["avg_abs_z_mean_shift"] <= baseline_shift + 0.010
    ]
    selected = list({d.name: d for d in [*selected, *near[: 8 if quick else 28]]}.values())
    print(f"[conservative-500] Üretilen veri adayı: {len(datasets)}, modele sokulan aday: {len(selected)}")
    return selected


def _feature_specs_for_run(quick: bool) -> list[Any]:
    names = {"no_skinthickness", "all_features"} if quick else {"no_skinthickness", "all_features", "high_signal_features", "compact_best_features"}
    return [f for f in _feature_specs() if f.name in names]


def _model_specs(random_state: int, n_jobs: int, quick: bool) -> list[ModelSpec]:
    estimators = 120 if quick else 260
    specs: list[ModelSpec] = [
        ModelSpec(
            "extra_trees",
            ExtraTreesClassifier(n_estimators=estimators, max_features="sqrt", min_samples_leaf=1, class_weight="balanced", random_state=random_state, n_jobs=n_jobs),
            "tree_models_no_scaler",
            "class_weight=balanced",
            "ExtraTrees.",
        ),
        ModelSpec(
            "random_forest",
            RandomForestClassifier(n_estimators=estimators, max_features="sqrt", min_samples_leaf=1, class_weight="balanced", random_state=random_state + 1, n_jobs=n_jobs),
            "tree_models_no_scaler",
            "class_weight=balanced",
            "RandomForest.",
        ),
    ]
    xgb = _small_xgboost(random_state + 2, n_jobs, quick)
    lgbm = _small_lightgbm(random_state + 3, n_jobs, quick)
    cat = _small_catboost(random_state + 4, quick)
    if xgb is not None:
        specs.append(ModelSpec("xgboost", xgb, "tree_models_no_scaler", "none", "XGBoost."))
    if lgbm is not None:
        specs.append(ModelSpec("lightgbm", lgbm, "tree_models_no_scaler", "none", "LightGBM."))
    if cat is not None and not quick:
        specs.append(ModelSpec("catboost", cat, "tree_models_no_scaler", "none", "CatBoost."))
    if xgb is not None and lgbm is not None and not quick:
        et = ExtraTreesClassifier(n_estimators=220, max_features="sqrt", class_weight="balanced", random_state=random_state + 5, n_jobs=1)
        specs.append(
            ModelSpec(
                "soft_voting_xgb_lgbm_et",
                VotingClassifier(
                    [("xgb", _small_xgboost(random_state + 6, 1, quick)), ("lgbm", _small_lightgbm(random_state + 7, 1, quick)), ("et", et)],
                    voting="soft",
                    n_jobs=1,
                ),
                "tree_models_no_scaler",
                "none",
                "XGBoost + LightGBM + ExtraTrees soft voting.",
            )
        )
    return specs


def _run_holdout(
    datasets: list[DatasetCandidate],
    feature_specs: list[Any],
    model_specs: list[ModelSpec],
    random_state: int,
) -> list[dict[str, Any]]:
    results: list[dict[str, Any]] = []
    for dataset in datasets:
        split = _group_holdout_split(dataset, random_state)
        for feature in feature_specs:
            for model in model_specs:
                try:
                    result = _evaluate_holdout(dataset=dataset, split=split, feature_spec=feature, model_spec=model, random_state=random_state)
                    results.append(result)
                    m = result["tuned_threshold_metrics"]
                    print(f"[cons500/{dataset.name}/{_feature_label(feature.name)}/{_model_label(model.name)}] min={m['ana_metrik_minimumu']:.4f} acc={m['accuracy']:.4f} auc={m['roc_auc']:.4f}")
                except Exception as exc:
                    results.append(
                        {
                            "status": "error",
                            "dataset_name": dataset.name,
                            "target_per_class": TARGET_PER_CLASS,
                            "model": model.name,
                            "feature_set": feature.name,
                            "error": f"{type(exc).__name__}: {exc}",
                        }
                    )
                    print(f"[cons500/{dataset.name}/{feature.name}/{model.name}] hata: {exc}")
    return results


def _select_cv_candidates(successful: list[dict[str, Any]], quality_by_dataset: dict[str, dict[str, Any]], quick: bool) -> list[dict[str, Any]]:
    def rank(row: dict[str, Any]) -> tuple[float, float, float, float]:
        m = row["tuned_threshold_metrics"]
        q = quality_by_dataset[row["dataset_name"]]
        cohen = q["class_separation"]["aggregate"]["mean_abs_cohens_d"]
        shift = q["distribution_shift_vs_original_dev"]["aggregate"]["avg_abs_z_mean_shift"]
        pass_bonus = 1.0 if m["accuracy"] >= BASELINE_ACCURACY and m["ana_metrik_minimumu"] >= BASELINE_MIN_MAIN else 0.0
        return (pass_bonus, m["ana_metrik_minimumu"], BASELINE_COHENS_D_REPORTED - cohen, BASELINE_SHIFT_REPORTED - shift)

    good = [
        r
        for r in successful
        if r["tuned_threshold_metrics"]["accuracy"] >= 0.895
        and r["tuned_threshold_metrics"]["ana_metrik_minimumu"] >= 0.875
    ]
    pool = good if good else successful
    limit = 6 if quick else 20
    selected: dict[tuple[str, str, str], dict[str, Any]] = {}
    for row in sorted(pool, key=rank, reverse=True)[:limit]:
        selected[(row["dataset_name"], row["feature_set"], row["model"])] = row
    return list(selected.values())


def _select_final(
    successful: list[dict[str, Any]],
    cv_results: list[dict[str, Any]],
    quality_by_dataset: dict[str, dict[str, Any]],
    leakage_report: list[dict[str, Any]],
) -> dict[str, Any] | None:
    successful_map = {(r["dataset_name"], r["feature_set"], r["model"]): r for r in successful}
    cv_map = {(r["dataset_name"], r["feature_set"], r["model"]): r for r in cv_results}
    leak_map = {r["dataset_name"]: r for r in leakage_report}
    candidates = []
    for cv in cv_results:
        key = (cv["dataset_name"], cv["feature_set"], cv["model"])
        leak = leak_map.get(cv["dataset_name"])
        if not leak or leak["leakage_status"] != "clean":
            continue
        q = quality_by_dataset[cv["dataset_name"]]
        # Holdout sonucu cv adaylarindan bulunur.
        h = successful_map.get(key)
        if h is None:
            continue
        m = h["tuned_threshold_metrics"]
        summary = cv_map[key]["summary"]
        cohen = q["class_separation"]["aggregate"]["mean_abs_cohens_d"]
        shift = q["distribution_shift_vs_original_dev"]["aggregate"]["avg_abs_z_mean_shift"]
        score = (
            0.36 * (m["ana_metrik_minimumu"] - BASELINE_MIN_MAIN)
            + 0.24 * (m["accuracy"] - BASELINE_ACCURACY)
            + 0.18 * (summary["ana_metrik_minimumu_mean"] - BASELINE_CV_MIN_MAIN)
            + 0.12 * (BASELINE_COHENS_D_REPORTED - cohen)
            + 0.10 * (BASELINE_SHIFT_REPORTED - shift)
        )
        row = dict(h)
        row["selection_score"] = float(score)
        row["selection_flags"] = {
            "accuracy_not_below_500": bool(m["accuracy"] >= BASELINE_ACCURACY),
            "min_main_not_below_500": bool(m["ana_metrik_minimumu"] >= BASELINE_MIN_MAIN),
            "group_cv_not_below_500": bool(summary["ana_metrik_minimumu_mean"] >= BASELINE_CV_MIN_MAIN),
            "computed_cohens_below_reported_500": bool(cohen < BASELINE_COHENS_D_REPORTED),
            "shift_below_reported_500": bool(shift < BASELINE_SHIFT_REPORTED),
        }
        candidates.append(row)
    if not candidates:
        return None
    full_pass = [
        c
        for c in candidates
        if c["selection_flags"]["accuracy_not_below_500"]
        and c["selection_flags"]["min_main_not_below_500"]
        and c["selection_flags"]["group_cv_not_below_500"]
    ]
    if full_pass:
        return sorted(full_pass, key=lambda c: c["selection_score"], reverse=True)[0]
    return sorted(candidates, key=lambda c: c["selection_score"], reverse=True)[0]


_LAST_SUCCESSFUL: list[dict[str, Any]] = []


def _find_cv(cv_results: list[dict[str, Any]], final: dict[str, Any] | None) -> dict[str, Any] | None:
    if final is None:
        return None
    for cv in cv_results:
        if cv["dataset_name"] == final["dataset_name"] and cv["feature_set"] == final["feature_set"] and cv["model"] == final["model"]:
            return cv
    return None


def _selection_report(
    final: dict[str, Any] | None,
    final_cv: dict[str, Any] | None,
    external: dict[str, Any] | None,
    quality_by_dataset: dict[str, dict[str, Any]],
    leakage: list[dict[str, Any]],
) -> dict[str, Any]:
    if final is None or final_cv is None:
        return {
            "created_at_utc": _now(),
            "final": None,
            "target_hit": False,
            "message": "CV adayi uretilemedi.",
        }
    q = quality_by_dataset[final["dataset_name"]]
    leak = next(x for x in leakage if x["dataset_name"] == final["dataset_name"])
    m = final["tuned_threshold_metrics"]
    cv_summary = final_cv["summary"]
    return {
        "created_at_utc": _now(),
        "final": _json_clean(final),
        "final_cv": _json_clean(final_cv),
        "external_holdout": external,
        "quality_summary": q,
        "leakage_summary": leak,
        "baseline_thresholds": {
            "accuracy": BASELINE_ACCURACY,
            "min_main_metric": BASELINE_MIN_MAIN,
            "group_cv_min_main_metric": BASELINE_CV_MIN_MAIN,
            "reported_cohens_d": BASELINE_COHENS_D_REPORTED,
            "reported_shift": BASELINE_SHIFT_REPORTED,
        },
        "target_hit": {
            "accuracy_not_below_500": bool(m["accuracy"] >= BASELINE_ACCURACY),
            "min_main_not_below_500": bool(m["ana_metrik_minimumu"] >= BASELINE_MIN_MAIN),
            "group_cv_not_below_500": bool(cv_summary["ana_metrik_minimumu_mean"] >= BASELINE_CV_MIN_MAIN),
            "computed_cohens_below_reported_500": bool(q["class_separation"]["aggregate"]["mean_abs_cohens_d"] < BASELINE_COHENS_D_REPORTED),
            "shift_below_reported_500": bool(q["distribution_shift_vs_original_dev"]["aggregate"]["avg_abs_z_mean_shift"] < BASELINE_SHIFT_REPORTED),
        },
        "interpretation": "Eger skor kosullari korunurken cohen/shift duserse 500/500 icin daha muhafazakar aday bulunmus olur; aksi durumda mevcut 500/500 denge noktasi korunur.",
    }


def _make_graphics(grafik_dir: Path, sweep: dict[str, Any], selection: dict[str, Any]) -> dict[str, str]:
    grafik_dir.mkdir(parents=True, exist_ok=True)
    paths = {
        "quality_scatter": str(_plot_quality_scatter(grafik_dir, sweep, selection)),
        "score_vs_quality": str(_plot_score_vs_quality(grafik_dir, sweep, selection)),
        "final_confusion": str(_plot_final_confusion(grafik_dir, selection)),
    }
    return paths


def _plot_quality_scatter(grafik_dir: Path, sweep: dict[str, Any], selection: dict[str, Any]) -> Path:
    path = grafik_dir / "conservative_500_quality_scatter.png"
    rows = []
    for dataset, q in sweep["quality_by_dataset"].items():
        rows.append(
            (
                q["class_separation"]["aggregate"]["mean_abs_cohens_d"],
                q["distribution_shift_vs_original_dev"]["aggregate"]["avg_abs_z_mean_shift"],
                dataset,
            )
        )
    fig, ax = plt.subplots(figsize=(7.4, 4.2))
    ax.scatter([r[0] for r in rows], [r[1] for r in rows], color="#9BAAAA", alpha=0.55, s=22, label="Üretilen aday")
    ax.axvline(BASELINE_COHENS_D_REPORTED, color="#C0392B", linestyle="--", linewidth=1.2, label="500/500 rapor Cohen's d")
    ax.axhline(BASELINE_SHIFT_REPORTED, color="#D35400", linestyle="--", linewidth=1.2, label="500/500 rapor shift")
    final = selection.get("final")
    if final:
        q = selection["quality_summary"]
        ax.scatter(
            [q["class_separation"]["aggregate"]["mean_abs_cohens_d"]],
            [q["distribution_shift_vs_original_dev"]["aggregate"]["avg_abs_z_mean_shift"]],
            color="#2D6A4F",
            s=80,
            label="Seçilen aday",
        )
    ax.set_xlabel("Cohen's d")
    ax.set_ylabel("Dağılım kayması")
    ax.set_title("500/500 Muhafazakar Adaylarda Cohen's d ve Shift")
    ax.legend(fontsize=8)
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _plot_score_vs_quality(grafik_dir: Path, sweep: dict[str, Any], selection: dict[str, Any]) -> Path:
    path = grafik_dir / "conservative_500_score_quality.png"
    best_by_dataset: dict[str, dict[str, Any]] = {}
    for row in sweep["holdout_results"]:
        if row.get("status") != "completed":
            continue
        current = best_by_dataset.get(row["dataset_name"])
        if current is None or row["tuned_threshold_metrics"]["ana_metrik_minimumu"] > current["tuned_threshold_metrics"]["ana_metrik_minimumu"]:
            best_by_dataset[row["dataset_name"]] = row
    xs, ys, colors = [], [], []
    for dataset, row in best_by_dataset.items():
        q = sweep["quality_by_dataset"][dataset]
        xs.append(q["class_separation"]["aggregate"]["mean_abs_cohens_d"])
        ys.append(row["tuned_threshold_metrics"]["ana_metrik_minimumu"])
        colors.append("#2D6A4F" if selection.get("final", {}).get("dataset_name") == dataset else "#9BAAAA")
    fig, ax = plt.subplots(figsize=(7.4, 4.0))
    ax.scatter(xs, ys, color=colors, s=34, alpha=0.8)
    ax.axhline(BASELINE_MIN_MAIN, color="#C0392B", linestyle="--", linewidth=1.2, label="500/500 min ana")
    ax.axvline(BASELINE_COHENS_D_REPORTED, color="#D35400", linestyle="--", linewidth=1.2, label="500/500 Cohen's d")
    ax.set_xlabel("Cohen's d")
    ax.set_ylabel("Holdout min ana metrik")
    ax.set_title("Skor Koruma ile Sınıf Ayrımı Arasındaki Takas")
    ax.legend(fontsize=8)
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _plot_final_confusion(grafik_dir: Path, selection: dict[str, Any]) -> Path:
    path = grafik_dir / "conservative_500_selected_confusion.png"
    final = selection.get("final")
    if not final:
        fig, ax = plt.subplots(figsize=(4, 3))
        ax.text(0.5, 0.5, "Final aday yok", ha="center", va="center")
        ax.axis("off")
        fig.savefig(path, dpi=180)
        plt.close(fig)
        return path
    cm = final["tuned_threshold_metrics"]["confusion_matrix"]
    matrix = np.array([[cm["tn"], cm["fp"]], [cm["fn"], cm["tp"]]])
    fig, ax = plt.subplots(figsize=(4.8, 3.8))
    ax.imshow(matrix, cmap="YlGnBu")
    ax.set_title("Seçilen Conservative 500 Hata Matrisi")
    ax.set_xticks([0, 1], ["Tahmin 0", "Tahmin 1"])
    ax.set_yticks([0, 1], ["Gerçek 0", "Gerçek 1"])
    for i in range(2):
        for j in range(2):
            ax.text(j, i, str(matrix[i, j]), ha="center", va="center", fontweight="bold", fontsize=12)
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _write_word_report(proje_koku: Path, sweep: dict[str, Any], selection: dict[str, Any], graphics: dict[str, str]) -> Path:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.shared import Inches, Pt

    path = proje_koku / "makine_ogrenmesi" / "raporlar" / "conservative_500_sweep_raporu.docx"
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(9)
    doc.add_heading("500/500 Muhafazakar Sentetik Benchmark Deneyi", level=1)
    doc.add_paragraph(
        "Bu raporda 500/500 adayının mevcut skorlarını koruyarak Cohen's d ve dağılım kaymasını azaltabilecek "
        "distribution-matched üretim varyantları denenmiştir. Amaç 1000/1000 skorlarını kovalamak değil, 500/500 "
        "seviyesinde daha doğal ve savunulabilir bir aday bulmaktır."
    )
    _doc_table(
        doc,
        ["Ölçüt", "Mevcut 500/500 eşiği"],
        [
            ["Accuracy", f"%{BASELINE_ACCURACY*100:.2f}"],
            ["Min ana metrik", f"%{BASELINE_MIN_MAIN*100:.2f}"],
            ["Group CV min ana", f"%{BASELINE_CV_MIN_MAIN*100:.2f}"],
            ["Rapor Cohen's d", f"{BASELINE_COHENS_D_REPORTED:.3f}"],
            ["Rapor shift", f"{BASELINE_SHIFT_REPORTED:.3f}"],
        ],
    )
    final = selection.get("final")
    if final:
        q = selection["quality_summary"]
        cv = selection["final_cv"]["summary"]
        m = final["tuned_threshold_metrics"]
        doc.add_heading("Seçilen En İyi Aday", level=2)
        _doc_table(
            doc,
            ["Alan", "Değer"],
            [
                ["Dataset", final["dataset_name"]],
                ["Model", _model_label(final["model"])],
                ["Feature set", _feature_label(final["feature_set"])],
                ["Accuracy", f"%{m['accuracy']*100:.2f}"],
                ["Min ana", f"%{m['ana_metrik_minimumu']*100:.2f}"],
                ["Group CV min ana", f"%{cv['ana_metrik_minimumu_mean']*100:.2f} ± {cv['ana_metrik_minimumu_std']:.3f}"],
                ["Cohen's d", f"{q['class_separation']['aggregate']['mean_abs_cohens_d']:.3f}"],
                ["Shift", f"{q['distribution_shift_vs_original_dev']['aggregate']['avg_abs_z_mean_shift']:.3f}"],
            ],
        )
        flags = selection["target_hit"]
        doc.add_paragraph(
            "Sonuç yorumu: "
            + (
                "Skor eşikleri korunurken kalite ölçütlerinde iyileşme bulundu."
                if all(flags.values())
                else "Tüm hedefler aynı anda sağlanamadı; bu durum 500/500 seviyesinde skor ile doğallık arasında belirgin bir takas olduğunu gösterir."
            )
        )
    for key in ["quality_scatter", "score_vs_quality", "final_confusion"]:
        doc.add_picture(graphics[key], width=Inches(6.4))
    doc.add_heading("Akademik Yorum", level=2)
    doc.add_paragraph(
        "Daha düşük Cohen's d ve shift hedefi, sentetik verinin orijinal PIMA dağılımına daha yakın kalmasını sağlar; "
        "ancak sınıf ayrımı azaldığında model performansının düşmesi beklenir. Bu deney, 500/500 adayının mevcut "
        "skorlarının ne kadar doğal dağılım korunarak sürdürülebildiğini ölçmek için yapılmıştır."
    )
    doc.save(path)
    return path


def _doc_table(doc: Any, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    doc.add_paragraph()


def _audit(raw: pd.DataFrame, selection: dict[str, Any]) -> dict[str, Any]:
    leak = selection.get("leakage_summary") or {}
    return {
        "created_at_utc": _now(),
        "diabetes_csv_rows_distribution_unchanged": bool(int(len(raw)) == 768 and _class_distribution(raw) == {"0": 500, "1": 268}),
        "final_dataset_name": selection.get("final", {}).get("dataset_name"),
        "leakage_clean": bool(leak.get("leakage_status") == "clean"),
        "target_hit": selection.get("target_hit"),
        "final_quality_status": "pass" if selection.get("final") else "warning",
        "notes": [
            "Bu audit conservative_500 deney hattına aittir.",
            "External holdout seçim skoruna dahil edilmedi.",
        ],
    }


def _write_dataset(veri_dir: Path, dataset: DatasetCandidate) -> None:
    combined = pd.concat([dataset.frame.reset_index(drop=True), dataset.metadata.reset_index(drop=True)], axis=1)
    combined.to_csv(veri_dir / f"{dataset.name}.csv", index=False)
    dataset.metadata.to_csv(veri_dir / f"{dataset.name}_metadata.csv", index=False)


def _json_write(path: Path, value: dict[str, Any]) -> None:
    path.write_text(json.dumps(_json_clean(value), ensure_ascii=False, indent=2), encoding="utf-8")


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _print_summary(result: dict[str, Any]) -> None:
    selection = result["selection_report"]
    final = selection.get("final")
    print("\n=== Conservative 500 Sweep Özeti ===")
    if not final:
        print("- Final aday: yok")
        return
    q = selection["quality_summary"]
    cv = selection["final_cv"]["summary"]
    m = final["tuned_threshold_metrics"]
    print(f"- Final aday: {final['dataset_name']}")
    print(f"- Model: {_model_label(final['model'])}")
    print(f"- Feature set: {_feature_label(final['feature_set'])}")
    print(f"- Accuracy: {m['accuracy']:.4f}")
    print(f"- Min ana: {m['ana_metrik_minimumu']:.4f}")
    print(f"- Group CV min ana: {cv['ana_metrik_minimumu_mean']:.4f} ± {cv['ana_metrik_minimumu_std']:.4f}")
    print(f"- Cohen's d: {q['class_separation']['aggregate']['mean_abs_cohens_d']:.4f}")
    print(f"- Shift: {q['distribution_shift_vs_original_dev']['aggregate']['avg_abs_z_mean_shift']:.4f}")
    print(f"- Hedefler: {selection['target_hit']}")
    print(f"- Word raporu: {result['word_report']}")
    print(f"- Masaüstü raporu: {result['desktop_word_report']}")
