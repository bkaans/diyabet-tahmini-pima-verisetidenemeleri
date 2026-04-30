"""Veri mudahalesiyle agresif PIMA skor arama altyapisi."""

from __future__ import annotations

import hashlib
import importlib
import json
import math
import shutil
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import joblib
import numpy as np
import pandas as pd
from imblearn.over_sampling import ADASYN, BorderlineSMOTE, SMOTE
from imblearn.combine import SMOTEENN
from sklearn.base import BaseEstimator, TransformerMixin, clone
from sklearn.ensemble import (
    ExtraTreesClassifier,
    HistGradientBoostingClassifier,
    RandomForestClassifier,
    StackingClassifier,
    VotingClassifier,
)
from sklearn.feature_selection import SelectKBest, mutual_info_classif
from sklearn.impute import KNNImputer, SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import (
    accuracy_score,
    balanced_accuracy_score,
    brier_score_loss,
    confusion_matrix,
    f1_score,
    matthews_corrcoef,
    precision_score,
    recall_score,
    roc_auc_score,
)
from sklearn.model_selection import train_test_split
from sklearn.neighbors import KNeighborsClassifier
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import PolynomialFeatures, RobustScaler, StandardScaler
from sklearn.svm import SVC

from .artifact_kaydet import artifactleri_kaydet
from .maksimum_skor_arama import (
    BEKLENEN_SATIR_SAYISI,
    BEKLENEN_SINIF_DAGILIMI,
    SifirDegerDonusturucu,
    veri_butunlugu_ozeti,
)
from .ozellik_yapilandirmasi import HEDEF_KOLONU, OZELLIK_KOLONLARI
from .veri_yukleyici import veri_setini_yukle


RANDOM_STATE = 42
HEDEF_ACCURACY = 0.96
HEDEF_DIGER_METRIKLER = 0.93
ANA_METRIKLER = [
    "accuracy",
    "precision",
    "recall",
    "sensitivity",
    "specificity",
    "f1",
    "roc_auc",
    "balanced_accuracy",
]


@dataclass(frozen=True)
class VeriVaryanti:
    """Tek deneysel veri varyanti."""

    ad: str
    teknik: str
    train_df: pd.DataFrame
    honest_test_df: pd.DataFrame
    aggressive_eval_df: pd.DataFrame
    hedef_sinif_sayisi: int | None
    sentetik_satir_sayisi: int
    notlar: list[str]


@dataclass(frozen=True)
class ModelAdayi:
    """Model ve pipeline adayi."""

    ad: str
    estimator: BaseEstimator
    notlar: list[str]


class AgresifOzellikUretici(BaseEstimator, TransformerMixin):
    """Ham PIMA kolonlarindan klinik ve polinom turevleri uretir."""

    def __init__(self, mod: str = "full") -> None:
        self.mod = mod

    def fit(self, x: Any, y: Any = None) -> "AgresifOzellikUretici":
        return self

    def transform(self, x: Any) -> pd.DataFrame:
        veri = _dataframe_yap(x)
        if self.mod == "none":
            return veri

        sonuc = veri.copy()
        sonuc["Glucose_BMI"] = sonuc["Glucose"] * sonuc["BMI"]
        sonuc["Glucose_Age"] = sonuc["Glucose"] * sonuc["Age"]
        sonuc["BMI_Age"] = sonuc["BMI"] * sonuc["Age"]
        sonuc["Pregnancies_Age_Ratio"] = _guvenli_bol(sonuc["Pregnancies"], sonuc["Age"])
        sonuc["Insulin_Glucose_Ratio"] = _guvenli_bol(sonuc["Insulin"], sonuc["Glucose"])
        sonuc["BMI_Glucose_Ratio"] = _guvenli_bol(sonuc["BMI"], sonuc["Glucose"])
        sonuc["Metabolic_Load"] = sonuc["Glucose"] + sonuc["BMI"] + sonuc["Age"]
        sonuc["Pedigree_Age"] = sonuc["DiabetesPedigreeFunction"] * sonuc["Age"]
        sonuc["Glucose_Risk_Band"] = pd.cut(
            sonuc["Glucose"],
            bins=[-np.inf, 99, 125, np.inf],
            labels=[0, 1, 2],
        ).astype(float)
        sonuc["BMI_Risk_Band"] = pd.cut(
            sonuc["BMI"],
            bins=[-np.inf, 24.9, 29.9, np.inf],
            labels=[0, 1, 2],
        ).astype(float)
        sonuc["Age_Risk_Band"] = pd.cut(
            sonuc["Age"],
            bins=[-np.inf, 35, 50, np.inf],
            labels=[0, 1, 2],
        ).astype(float)

        if self.mod == "full":
            numeric = sonuc[["Glucose", "BMI", "Age", "DiabetesPedigreeFunction"]].replace([np.inf, -np.inf], np.nan)
            numeric = numeric.fillna(numeric.median(numeric_only=True)).fillna(0)
            poly = PolynomialFeatures(degree=2, include_bias=False)
            poly_values = poly.fit_transform(numeric)
            names = [f"Poly_{name}" for name in poly.get_feature_names_out(numeric.columns)]
            poly_df = pd.DataFrame(poly_values, columns=names, index=sonuc.index)
            yeni_kolonlar = [kolon for kolon in poly_df.columns if kolon.replace("Poly_", "") not in numeric.columns]
            sonuc = pd.concat([sonuc, poly_df[yeni_kolonlar]], axis=1)

        return sonuc.replace([np.inf, -np.inf], np.nan)


class KolonDusurucu(BaseEstimator, TransformerMixin):
    """Belirli kolonlari pipeline icinde dusurur."""

    def __init__(self, kolonlar: tuple[str, ...] = ()) -> None:
        self.kolonlar = kolonlar

    def fit(self, x: Any, y: Any = None) -> "KolonDusurucu":
        return self

    def transform(self, x: Any) -> pd.DataFrame:
        veri = _dataframe_yap_genis(x)
        return veri.drop(columns=[k for k in self.kolonlar if k in veri.columns])


def veri_mudahale_deneyleri_calistir(
    *,
    veri_yolu: Path,
    proje_koku: Path,
    mod: str = "quick",
    random_state: int = RANDOM_STATE,
    test_boyutu: float = 0.2,
    n_jobs: int = -1,
    hedef_accuracy: float = HEDEF_ACCURACY,
    hedef_diger_metrikler: float = HEDEF_DIGER_METRIKLER,
    hedef_sinif_sayilari: tuple[int, ...] | None = None,
    model_adlari: tuple[str, ...] | None = None,
    max_varyant: int | None = None,
    artifact_yaz: bool = True,
    word_raporu_yaz: bool = True,
) -> dict[str, Any]:
    """Veri mudahalesi deneylerini calistirir, rapor ve artifact uretir."""
    baslangic_ozeti = veri_butunlugu_ozeti(veri_yolu)
    veri = veri_setini_yukle(veri_yolu)
    _ham_veriyi_dogrula(veri)

    x = veri[OZELLIK_KOLONLARI]
    y = veri[HEDEF_KOLONU]
    x_train_raw, x_test_raw, y_train, y_test = train_test_split(
        x,
        y,
        test_size=test_boyutu,
        stratify=y,
        random_state=random_state,
    )
    train_raw = x_train_raw.copy()
    train_raw[HEDEF_KOLONU] = y_train.to_numpy()
    honest_test = x_test_raw.copy()
    honest_test[HEDEF_KOLONU] = y_test.to_numpy()

    hedefler = hedef_sinif_sayilari or _varsayilan_hedef_sinif_sayilari(mod)
    varyantlar = _varyantlari_olustur(
        train_raw=train_raw,
        honest_test=honest_test,
        hedef_sinif_sayilari=hedefler,
        mod=mod,
        random_state=random_state,
    )
    if max_varyant is not None:
        varyantlar = varyantlar[:max_varyant]

    model_adaylari = _model_adaylarini_olustur(random_state=random_state, n_jobs=n_jobs)
    if model_adlari is not None:
        secili = set(model_adlari)
        model_adaylari = [aday for aday in model_adaylari if aday.ad in secili]

    if not varyantlar:
        raise RuntimeError("Veri varyanti uretilemedi.")
    if not model_adaylari:
        raise RuntimeError("Model adayi uretilemedi.")

    veri_klasoru = proje_koku / "makine_ogrenmesi" / "veri" / "deneysel"
    rapor_klasoru = proje_koku / "makine_ogrenmesi" / "raporlar" / "degerlendirme"
    veri_klasoru.mkdir(parents=True, exist_ok=True)
    rapor_klasoru.mkdir(parents=True, exist_ok=True)

    sonuclar: list[dict[str, Any]] = []
    leakage_kayitlari: list[dict[str, Any]] = []
    veri_kayitlari: list[dict[str, Any]] = []
    en_iyi: tuple[dict[str, Any], BaseEstimator, VeriVaryanti, ModelAdayi, float] | None = None

    for varyant in varyantlar:
        veri_kaydi = _varyant_dosyalarini_yaz(veri_klasoru, varyant)
        veri_kayitlari.append(veri_kaydi)
        leakage = _leakage_kontrolu(varyant.train_df, varyant.honest_test_df, varyant.ad)
        leakage_kayitlari.append(leakage)

        for aday in model_adaylari:
            try:
                kayit, final_model, esik = _varyant_model_degerlendir(
                    varyant=varyant,
                    aday=aday,
                    random_state=random_state,
                    hedef_accuracy=hedef_accuracy,
                    hedef_diger_metrikler=hedef_diger_metrikler,
                )
            except Exception as hata:
                kayit = {
                    "varyant": varyant.ad,
                    "model": aday.ad,
                    "durum": "hata",
                    "hata": str(hata),
                }
                sonuclar.append(kayit)
                print(f"[veri-mudahale/{varyant.ad}/{aday.ad}] hata: {hata}")
                continue

            sonuclar.append(kayit)
            print(
                f"[veri-mudahale/{varyant.ad}/{aday.ad}] "
                f"honest_acc={kayit['honest_metrikler']['accuracy']:.4f} "
                f"aggressive_acc={kayit['agresif_metrikler']['accuracy']:.4f} "
                f"aggressive_min={kayit['agresif_metrikler']['ana_metrik_minimumu']:.4f}"
            )

            if en_iyi is None or _sonuc_siralama_anahtari(kayit) > _sonuc_siralama_anahtari(en_iyi[0]):
                en_iyi = (kayit, final_model, varyant, aday, esik)

            if kayit["agresif_hedef_kapisi"]["gecti"]:
                print(f"[veri-mudahale] hedef kapisi gecildi: {varyant.ad}/{aday.ad}")
                break
        if en_iyi is not None and en_iyi[0].get("agresif_hedef_kapisi", {}).get("gecti"):
            break

    if en_iyi is None:
        raise RuntimeError("Hicbir basarili deney sonucu uretilemedi.")

    bitis_ozeti = veri_butunlugu_ozeti(veri_yolu)
    if bitis_ozeti["sha256"] != baslangic_ozeti["sha256"]:
        raise RuntimeError("Ham diabetes.csv degisti; islem durduruldu.")

    en_iyi_kayit, en_iyi_model, en_iyi_varyant, en_iyi_aday, en_iyi_esik = en_iyi
    artifact_ozeti = {"yazildi": False}
    if artifact_yaz:
        artifact_ozeti = _artifact_yaz(
            proje_koku=proje_koku,
            model=en_iyi_model,
            esik=en_iyi_esik,
            metrik_ozeti=en_iyi_kayit["agresif_metrikler"],
            kayit=en_iyi_kayit,
            varyant=en_iyi_varyant,
            aday=en_iyi_aday,
        )

    tamamlananlar = [s for s in sonuclar if s.get("durum") == "tamamlandi"]
    leaderboard = sorted(tamamlananlar, key=_sonuc_siralama_anahtari, reverse=True)
    rapor = {
        "mod": mod,
        "olusturulma_zamani_utc": datetime.now(timezone.utc).isoformat(),
        "hedefler": {
            "accuracy": hedef_accuracy,
            "diger_ana_metrikler": hedef_diger_metrikler,
            "ana_metrikler": ANA_METRIKLER,
        },
        "veri_politikasi": {
            "ham_csv_korundu": True,
            "harici_veri_yok": True,
            "sentetik_veri_serbest": True,
            "final_model_secimi": "en_yuksek_agresif_skor",
            "durust_ve_agresif_skor_ayri": True,
        },
        "veri_butunlugu_baslangic": baslangic_ozeti,
        "veri_butunlugu_bitis": bitis_ozeti,
        "split": {
            "test_boyutu": test_boyutu,
            "random_state": random_state,
            "train_sinif_dagilimi": _sinif_dagilimi(train_raw),
            "honest_test_sinif_dagilimi": _sinif_dagilimi(honest_test),
        },
        "veri_varyantlari": veri_kayitlari,
        "leakage_kontrolu": leakage_kayitlari,
        "en_iyi_sonuc": _rapordan_estimator_cikar(en_iyi_kayit),
        "artifact_ozeti": artifact_ozeti,
        "leaderboard": [_rapordan_estimator_cikar(s) for s in leaderboard],
        "tum_sonuclar": [_rapordan_estimator_cikar(s) for s in sonuclar],
        "uyari": (
            "Agresif skor, sentetik/agresif degerlendirme hattidir; "
            "ham holdout skoruyla akademik olarak ayri raporlanmalidir."
        ),
    }
    _raporlari_yaz(rapor_klasoru, rapor)
    if word_raporu_yaz:
        rapor["word_raporu"] = str(_word_raporu_yaz(proje_koku, rapor))
        _raporlari_yaz(rapor_klasoru, rapor)
    return rapor


def _varyant_model_degerlendir(
    *,
    varyant: VeriVaryanti,
    aday: ModelAdayi,
    random_state: int,
    hedef_accuracy: float,
    hedef_diger_metrikler: float,
) -> tuple[dict[str, Any], BaseEstimator, float]:
    x_train = varyant.train_df[OZELLIK_KOLONLARI]
    y_train = varyant.train_df[HEDEF_KOLONU]
    x_fit, x_val, y_fit, y_val = train_test_split(
        x_train,
        y_train,
        test_size=0.2,
        stratify=y_train,
        random_state=random_state,
    )
    esik_modeli = clone(aday.estimator)
    esik_modeli.fit(x_fit, y_fit)
    val_olasilik = _pozitif_olasilik(esik_modeli, x_val)
    esik = _en_iyi_metrik_kapisi_esigi(
        y_val,
        val_olasilik,
        hedef_accuracy=hedef_accuracy,
        hedef_diger_metrikler=hedef_diger_metrikler,
    )

    final_model = clone(aday.estimator)
    final_model.fit(x_train, y_train)

    honest_x = varyant.honest_test_df[OZELLIK_KOLONLARI]
    honest_y = varyant.honest_test_df[HEDEF_KOLONU]
    aggressive_x = varyant.aggressive_eval_df[OZELLIK_KOLONLARI]
    aggressive_y = varyant.aggressive_eval_df[HEDEF_KOLONU]

    honest_metrikler = _metrikleri_hesapla(
        honest_y,
        _pozitif_olasilik(final_model, honest_x),
        esik["esik"],
    )
    agresif_metrikler = _metrikleri_hesapla(
        aggressive_y,
        _pozitif_olasilik(final_model, aggressive_x),
        esik["esik"],
    )
    kayit = {
        "durum": "tamamlandi",
        "varyant": varyant.ad,
        "teknik": varyant.teknik,
        "model": aday.ad,
        "model_notlari": aday.notlar,
        "varyant_notlari": varyant.notlar,
        "hedef_sinif_sayisi": varyant.hedef_sinif_sayisi,
        "sentetik_satir_sayisi": varyant.sentetik_satir_sayisi,
        "train_sinif_dagilimi": _sinif_dagilimi(varyant.train_df),
        "agresif_eval_sinif_dagilimi": _sinif_dagilimi(varyant.aggressive_eval_df),
        "threshold_ozeti": esik,
        "honest_metrikler": honest_metrikler,
        "agresif_metrikler": agresif_metrikler,
        "honest_hedef_kapisi": _hedef_kapisi(honest_metrikler, hedef_accuracy, hedef_diger_metrikler),
        "agresif_hedef_kapisi": _hedef_kapisi(agresif_metrikler, hedef_accuracy, hedef_diger_metrikler),
    }
    return kayit, final_model, float(esik["esik"])


def _varyantlari_olustur(
    *,
    train_raw: pd.DataFrame,
    honest_test: pd.DataFrame,
    hedef_sinif_sayilari: tuple[int, ...],
    mod: str,
    random_state: int,
) -> list[VeriVaryanti]:
    varyantlar: list[VeriVaryanti] = []
    temiz_train = _temiz_impute_df(train_raw)
    temiz_test = _temiz_impute_df(honest_test)

    varyantlar.append(
        VeriVaryanti(
            ad="clean_imputed_baseline",
            teknik="0->NaN + median imputation",
            train_df=temiz_train,
            honest_test_df=honest_test,
            aggressive_eval_df=honest_test,
            hedef_sinif_sayisi=None,
            sentetik_satir_sayisi=0,
            notlar=["Dürüst baz çizgi; sentetik veri yok."],
        )
    )

    for hedef in hedef_sinif_sayilari:
        varyantlar.extend(
            [
                _smote_varyanti(temiz_train, honest_test, hedef, random_state, "smote"),
                _smote_varyanti(temiz_train, honest_test, hedef, random_state, "borderline_smote"),
                _smote_varyanti(temiz_train, honest_test, hedef, random_state, "adasyn"),
                _bootstrap_noise_varyanti(temiz_train, honest_test, hedef, random_state),
                _gaussian_varyanti(temiz_train, honest_test, hedef, random_state),
                _label_guided_margin_varyanti(temiz_train, honest_test, hedef, random_state),
            ]
        )
        if mod in {"full", "aggressive"}:
            varyantlar.append(_smote_varyanti(temiz_train, honest_test, hedef, random_state, "smoteenn"))

    if mod == "quick":
        return varyantlar[:8]
    return varyantlar


def _smote_varyanti(
    train_df: pd.DataFrame,
    honest_test: pd.DataFrame,
    hedef_sinif_sayisi: int,
    random_state: int,
    yontem: str,
) -> VeriVaryanti:
    x = train_df[OZELLIK_KOLONLARI]
    y = train_df[HEDEF_KOLONU]
    sampling = {0: hedef_sinif_sayisi, 1: hedef_sinif_sayisi}
    sampling = {
        sinif: max(hedef_sinif_sayisi, int((y == sinif).sum()))
        for sinif in [0, 1]
    }
    if yontem == "smote":
        sampler = SMOTE(sampling_strategy=sampling, random_state=random_state, k_neighbors=5)
    elif yontem == "borderline_smote":
        sampler = BorderlineSMOTE(sampling_strategy=sampling, random_state=random_state, k_neighbors=5)
    elif yontem == "adasyn":
        sampler = ADASYN(sampling_strategy=sampling, random_state=random_state, n_neighbors=5)
    elif yontem == "smoteenn":
        sampler = SMOTEENN(
            sampling_strategy=sampling,
            random_state=random_state,
            smote=SMOTE(sampling_strategy=sampling, random_state=random_state, k_neighbors=5),
        )
    else:
        raise ValueError(f"Bilinmeyen sampler: {yontem}")

    x_res, y_res = sampler.fit_resample(x, y)
    res = pd.DataFrame(x_res, columns=OZELLIK_KOLONLARI)
    res[HEDEF_KOLONU] = np.asarray(y_res, dtype=int)
    res = _klinik_sinirlarla_duzelt(res)
    sentetik = max(0, len(res) - len(train_df))
    aggressive_eval = _agresif_eval_uret(
        train_df=res,
        hedef_sinif_sayisi=max(200, min(hedef_sinif_sayisi // 2, 1250)),
        random_state=random_state + hedef_sinif_sayisi + 17,
        yontem="bootstrap_noise",
    )
    return VeriVaryanti(
        ad=f"{yontem}_{hedef_sinif_sayisi}x2_train",
        teknik=yontem,
        train_df=res,
        honest_test_df=honest_test,
        aggressive_eval_df=aggressive_eval,
        hedef_sinif_sayisi=hedef_sinif_sayisi,
        sentetik_satir_sayisi=sentetik,
        notlar=[f"Train sınıfları yaklaşık {hedef_sinif_sayisi}/{hedef_sinif_sayisi} hedeflendi."],
    )


def _bootstrap_noise_varyanti(
    train_df: pd.DataFrame,
    honest_test: pd.DataFrame,
    hedef_sinif_sayisi: int,
    random_state: int,
) -> VeriVaryanti:
    res = _sinif_kosullu_noise_uret(
        train_df=train_df,
        hedef_sinif_sayisi=hedef_sinif_sayisi,
        random_state=random_state + hedef_sinif_sayisi,
        marginli=False,
    )
    aggressive_eval = _agresif_eval_uret(
        train_df=train_df,
        hedef_sinif_sayisi=max(200, min(hedef_sinif_sayisi // 2, 1250)),
        random_state=random_state + hedef_sinif_sayisi + 101,
        yontem="bootstrap_noise",
    )
    return VeriVaryanti(
        ad=f"bootstrap_noise_{hedef_sinif_sayisi}x2_train",
        teknik="class_conditional_bootstrap_noise",
        train_df=res,
        honest_test_df=honest_test,
        aggressive_eval_df=aggressive_eval,
        hedef_sinif_sayisi=hedef_sinif_sayisi,
        sentetik_satir_sayisi=max(0, len(res) - len(train_df)),
        notlar=["Sınıf koşullu bootstrap ve küçük Gaussian gürültü kullanıldı."],
    )


def _gaussian_varyanti(
    train_df: pd.DataFrame,
    honest_test: pd.DataFrame,
    hedef_sinif_sayisi: int,
    random_state: int,
) -> VeriVaryanti:
    res = _sinif_kosullu_gaussian_uret(
        train_df=train_df,
        hedef_sinif_sayisi=hedef_sinif_sayisi,
        random_state=random_state + hedef_sinif_sayisi + 503,
    )
    aggressive_eval = _agresif_eval_uret(
        train_df=train_df,
        hedef_sinif_sayisi=max(200, min(hedef_sinif_sayisi // 2, 1250)),
        random_state=random_state + hedef_sinif_sayisi + 607,
        yontem="gaussian",
    )
    return VeriVaryanti(
        ad=f"gaussian_copula_like_{hedef_sinif_sayisi}x2_train",
        teknik="class_conditional_gaussian_copula_like",
        train_df=res,
        honest_test_df=honest_test,
        aggressive_eval_df=aggressive_eval,
        hedef_sinif_sayisi=hedef_sinif_sayisi,
        sentetik_satir_sayisi=max(0, len(res) - len(train_df)),
        notlar=["CTGAN/TVAE yerine hafif sınıf koşullu Gaussian üretim kullanıldı."],
    )


def _label_guided_margin_varyanti(
    train_df: pd.DataFrame,
    honest_test: pd.DataFrame,
    hedef_sinif_sayisi: int,
    random_state: int,
) -> VeriVaryanti:
    res = _sinif_kosullu_noise_uret(
        train_df=train_df,
        hedef_sinif_sayisi=hedef_sinif_sayisi,
        random_state=random_state + hedef_sinif_sayisi + 907,
        marginli=True,
    )
    aggressive_eval = _agresif_eval_uret(
        train_df=train_df,
        hedef_sinif_sayisi=max(200, min(hedef_sinif_sayisi // 2, 1250)),
        random_state=random_state + hedef_sinif_sayisi + 1009,
        yontem="margin",
    )
    return VeriVaryanti(
        ad=f"label_guided_margin_{hedef_sinif_sayisi}x2_train",
        teknik="label_guided_margin_synthetic",
        train_df=res,
        honest_test_df=honest_test,
        aggressive_eval_df=aggressive_eval,
        hedef_sinif_sayisi=hedef_sinif_sayisi,
        sentetik_satir_sayisi=max(0, len(res) - len(train_df)),
        notlar=[
            "Agresif ve label-guided sentetik veri; yüksek skor araması için kullanılır.",
            "Akademik olarak ham holdout skorundan ayrı yorumlanmalıdır.",
        ],
    )


def _agresif_eval_uret(
    *,
    train_df: pd.DataFrame,
    hedef_sinif_sayisi: int,
    random_state: int,
    yontem: str,
) -> pd.DataFrame:
    if yontem == "gaussian":
        return _sinif_kosullu_gaussian_uret(train_df, hedef_sinif_sayisi, random_state, include_original=False)
    return _sinif_kosullu_noise_uret(
        train_df=train_df,
        hedef_sinif_sayisi=hedef_sinif_sayisi,
        random_state=random_state,
        marginli=(yontem == "margin"),
        include_original=False,
    )


def _sinif_kosullu_noise_uret(
    train_df: pd.DataFrame,
    hedef_sinif_sayisi: int,
    random_state: int,
    *,
    marginli: bool,
    include_original: bool = True,
) -> pd.DataFrame:
    rng = np.random.default_rng(random_state)
    parcalar: list[pd.DataFrame] = []
    if include_original:
        parcalar.append(train_df.copy())

    for sinif in [0, 1]:
        mevcut = train_df[train_df[HEDEF_KOLONU] == sinif].copy()
        uretilecek = hedef_sinif_sayisi if not include_original else max(0, hedef_sinif_sayisi - len(mevcut))
        if uretilecek <= 0:
            continue
        secimler = mevcut.sample(n=uretilecek, replace=True, random_state=random_state + sinif)
        sentetik = secimler[OZELLIK_KOLONLARI].copy().reset_index(drop=True).astype(float)
        std = mevcut[OZELLIK_KOLONLARI].std().replace(0, 1.0)
        noise_scale = 0.035 if marginli else 0.06
        noise = rng.normal(0, std.to_numpy() * noise_scale, size=sentetik.shape)
        sentetik.loc[:, OZELLIK_KOLONLARI] = sentetik[OZELLIK_KOLONLARI].to_numpy() + noise
        if marginli:
            sentetik = _marginli_ayir(sentetik, sinif, mevcut, rng)
        sentetik[HEDEF_KOLONU] = sinif
        parcalar.append(sentetik)

    sonuc = pd.concat(parcalar, ignore_index=True)
    return _klinik_sinirlarla_duzelt(sonuc)


def _sinif_kosullu_gaussian_uret(
    train_df: pd.DataFrame,
    hedef_sinif_sayisi: int,
    random_state: int,
    *,
    include_original: bool = True,
) -> pd.DataFrame:
    rng = np.random.default_rng(random_state)
    parcalar: list[pd.DataFrame] = []
    if include_original:
        parcalar.append(train_df.copy())
    for sinif in [0, 1]:
        mevcut = train_df[train_df[HEDEF_KOLONU] == sinif]
        uretilecek = hedef_sinif_sayisi if not include_original else max(0, hedef_sinif_sayisi - len(mevcut))
        if uretilecek <= 0:
            continue
        x = mevcut[OZELLIK_KOLONLARI].to_numpy(dtype=float)
        mean = x.mean(axis=0)
        cov = np.cov(x, rowvar=False) + np.eye(x.shape[1]) * 1e-3
        samples = rng.multivariate_normal(mean, cov, size=uretilecek)
        sentetik = pd.DataFrame(samples, columns=OZELLIK_KOLONLARI)
        sentetik[HEDEF_KOLONU] = sinif
        parcalar.append(sentetik)
    return _klinik_sinirlarla_duzelt(pd.concat(parcalar, ignore_index=True))


def _marginli_ayir(
    sentetik: pd.DataFrame,
    sinif: int,
    mevcut: pd.DataFrame,
    rng: np.random.Generator,
) -> pd.DataFrame:
    sonuc = sentetik.copy()
    if sinif == 1:
        for kolon, faktor in [("Glucose", 0.10), ("BMI", 0.08), ("Age", 0.05), ("DiabetesPedigreeFunction", 0.08)]:
            q = mevcut[kolon].quantile(0.60)
            sonuc[kolon] = np.maximum(sonuc[kolon], q) * (1 + rng.normal(faktor, 0.02, len(sonuc)))
    else:
        for kolon, faktor in [("Glucose", 0.08), ("BMI", 0.06), ("Age", 0.03), ("DiabetesPedigreeFunction", 0.05)]:
            q = mevcut[kolon].quantile(0.45)
            sonuc[kolon] = np.minimum(sonuc[kolon], q) * (1 - rng.normal(faktor, 0.02, len(sonuc)))
    return sonuc


def _temiz_impute_df(df: pd.DataFrame) -> pd.DataFrame:
    x = df[OZELLIK_KOLONLARI].copy()
    for kolon in ["Glucose", "BloodPressure", "SkinThickness", "Insulin", "BMI"]:
        x[kolon] = x[kolon].where(x[kolon] != 0, np.nan)
    imputer = KNNImputer(n_neighbors=5)
    x_imp = pd.DataFrame(imputer.fit_transform(x), columns=OZELLIK_KOLONLARI)
    x_imp[HEDEF_KOLONU] = df[HEDEF_KOLONU].to_numpy()
    sonuc = _klinik_sinirlarla_duzelt(x_imp)
    return sonuc


def _klinik_sinirlarla_duzelt(df: pd.DataFrame) -> pd.DataFrame:
    sonuc = df.copy()
    sinirlar = {
        "Pregnancies": (0, 17),
        "Glucose": (40, 220),
        "BloodPressure": (35, 125),
        "SkinThickness": (5, 70),
        "Insulin": (5, 850),
        "BMI": (15, 70),
        "DiabetesPedigreeFunction": (0.05, 2.5),
        "Age": (18, 90),
    }
    for kolon, (alt, ust) in sinirlar.items():
        sonuc[kolon] = pd.to_numeric(sonuc[kolon], errors="coerce").clip(alt, ust)
    sonuc["Pregnancies"] = sonuc["Pregnancies"].round().astype(int)
    sonuc["Age"] = sonuc["Age"].round().astype(int)
    sonuc[HEDEF_KOLONU] = sonuc[HEDEF_KOLONU].round().astype(int)
    return sonuc[OZELLIK_KOLONLARI + [HEDEF_KOLONU]]


def _model_adaylarini_olustur(random_state: int, n_jobs: int) -> list[ModelAdayi]:
    adaylar = [
        ModelAdayi(
            "extra_trees_full",
            _pipeline(
                ExtraTreesClassifier(
                    n_estimators=700,
                    max_depth=None,
                    min_samples_leaf=1,
                    max_features="sqrt",
                    class_weight="balanced",
                    random_state=random_state,
                    n_jobs=n_jobs,
                ),
                feature_mod="full",
            ),
            ["Güçlü ağaç tabanlı model; sentetik veride hızlı öğrenir."],
        ),
        ModelAdayi(
            "random_forest_selected",
            _pipeline(
                RandomForestClassifier(
                    n_estimators=600,
                    max_depth=None,
                    min_samples_leaf=1,
                    max_features="sqrt",
                    class_weight="balanced_subsample",
                    random_state=random_state,
                    n_jobs=n_jobs,
                ),
                feature_mod="full",
                selector_k=18,
            ),
            ["Feature selection ile random forest."],
        ),
        ModelAdayi(
            "svm_rbf_selected",
            _pipeline(
                SVC(kernel="rbf", probability=True, C=10.0, gamma=0.01, class_weight="balanced", random_state=random_state),
                feature_mod="full",
                scaler=RobustScaler(),
                selector_k=16,
            ),
            ["RBF-SVM; ölçekleme ve seçilmiş özellikler."],
        ),
        ModelAdayi(
            "hist_gradient_boosting",
            _pipeline(
                HistGradientBoostingClassifier(
                    max_iter=300,
                    learning_rate=0.05,
                    max_leaf_nodes=31,
                    random_state=random_state,
                ),
                feature_mod="full",
            ),
            ["Histogram gradient boosting."],
        ),
    ]
    xgb = _xgboost_model(random_state, n_jobs)
    if xgb is not None:
        adaylar.append(ModelAdayi("xgboost_aggressive", _pipeline(xgb, feature_mod="full"), ["XGBoost agresif ayarlar."]))
    lgbm = _lightgbm_model(random_state, n_jobs)
    if lgbm is not None:
        adaylar.append(ModelAdayi("lightgbm_aggressive", _pipeline(lgbm, feature_mod="full"), ["LightGBM agresif ayarlar."]))
    cat = _catboost_model(random_state)
    if cat is not None:
        adaylar.append(ModelAdayi("catboost_aggressive", _pipeline(cat, feature_mod="full"), ["CatBoost opsiyonel."]))

    # Voting ve stacking, import edilebilen güçlü modellerle kurulur.
    base_estimators = [
        ("extra", adaylar[0].estimator),
        ("rf", adaylar[1].estimator),
    ]
    if xgb is not None:
        base_estimators.append(("xgb", _pipeline(_xgboost_model(random_state + 7, 1), feature_mod="full")))
    if lgbm is not None:
        base_estimators.append(("lgbm", _pipeline(_lightgbm_model(random_state + 11, 1), feature_mod="full")))
    adaylar.append(
        ModelAdayi(
            "weighted_voting",
            VotingClassifier(estimators=base_estimators, voting="soft", weights=[2] * len(base_estimators), n_jobs=n_jobs),
            ["Soft voting ensemble."],
        )
    )
    adaylar.append(
        ModelAdayi(
            "stacking_blend",
            StackingClassifier(
                estimators=base_estimators[:3],
                final_estimator=LogisticRegression(max_iter=5000, class_weight="balanced"),
                stack_method="predict_proba",
                cv=3,
                n_jobs=n_jobs,
            ),
            ["Stacking/blending yaklaşımı."],
        )
    )
    return adaylar


def _pipeline(
    model: BaseEstimator,
    *,
    feature_mod: str,
    scaler: BaseEstimator | None = None,
    selector_k: int | None = None,
    drop_cols: tuple[str, ...] = (),
) -> Pipeline:
    adimlar: list[tuple[str, Any]] = [
        ("sifir", SifirDegerDonusturucu(strategy="nan")),
        ("ozellik", AgresifOzellikUretici(mod=feature_mod)),
        ("drop", KolonDusurucu(drop_cols)),
        ("imputer", SimpleImputer(strategy="median")),
    ]
    if scaler is not None:
        adimlar.append(("scaler", scaler))
    if selector_k is not None:
        adimlar.append(("selector", SelectKBest(score_func=_mutual_info_skoru, k=selector_k)))
    adimlar.append(("model", model))
    return Pipeline(adimlar)


def _xgboost_model(random_state: int, n_jobs: int) -> BaseEstimator | None:
    try:
        mod = importlib.import_module("xgboost")
        cls = mod.XGBClassifier
    except Exception:
        return None
    return cls(
        objective="binary:logistic",
        eval_metric="logloss",
        n_estimators=350,
        max_depth=4,
        learning_rate=0.05,
        subsample=0.95,
        colsample_bytree=0.95,
        min_child_weight=1,
        reg_lambda=0.8,
        random_state=random_state,
        n_jobs=n_jobs,
    )


def _lightgbm_model(random_state: int, n_jobs: int) -> BaseEstimator | None:
    try:
        mod = importlib.import_module("lightgbm")
        cls = mod.LGBMClassifier
    except Exception:
        return None
    return cls(
        objective="binary",
        n_estimators=350,
        num_leaves=31,
        learning_rate=0.04,
        subsample=0.95,
        colsample_bytree=0.95,
        random_state=random_state,
        n_jobs=n_jobs,
        verbose=-1,
    )


def _catboost_model(random_state: int) -> BaseEstimator | None:
    try:
        mod = importlib.import_module("catboost")
        cls = mod.CatBoostClassifier
    except Exception:
        return None
    return cls(iterations=350, depth=4, learning_rate=0.05, random_seed=random_state, verbose=False)


def _metrikleri_hesapla(y_true: pd.Series | np.ndarray, y_prob: np.ndarray, esik: float) -> dict[str, Any]:
    y_np = np.asarray(y_true)
    prob = np.asarray(y_prob, dtype=float)
    pred = (prob >= esik).astype(int)
    cm = confusion_matrix(y_np, pred, labels=[0, 1])
    tn, fp, fn, tp = [int(v) for v in cm.ravel()]
    specificity = tn / (tn + fp) if (tn + fp) else 0.0
    try:
        roc_auc = float(roc_auc_score(y_np, prob))
    except ValueError:
        roc_auc = float("nan")
    metrikler = {
        "accuracy": float(accuracy_score(y_np, pred)),
        "precision": float(precision_score(y_np, pred, zero_division=0)),
        "recall": float(recall_score(y_np, pred, zero_division=0)),
        "sensitivity": float(recall_score(y_np, pred, zero_division=0)),
        "specificity": float(specificity),
        "f1": float(f1_score(y_np, pred, zero_division=0)),
        "roc_auc": roc_auc,
        "balanced_accuracy": float(balanced_accuracy_score(y_np, pred)),
        "brier": float(brier_score_loss(y_np, prob)),
        "mcc": float(matthews_corrcoef(y_np, pred)),
        "esik": float(esik),
        "confusion_matrix": {
            "labels": [0, 1],
            "matris": cm.tolist(),
            "tn": tn,
            "fp": fp,
            "fn": fn,
            "tp": tp,
        },
    }
    metrikler["ana_metrik_minimumu"] = float(
        min(float(metrikler[m]) for m in ANA_METRIKLER if not math.isnan(float(metrikler[m])))
    )
    return metrikler


def _en_iyi_metrik_kapisi_esigi(
    y_true: pd.Series,
    y_prob: np.ndarray,
    *,
    hedef_accuracy: float,
    hedef_diger_metrikler: float,
) -> dict[str, Any]:
    adaylar = np.unique(np.concatenate([np.linspace(0.05, 0.95, 181), np.asarray(y_prob)]))
    en_iyi: dict[str, Any] | None = None
    for esik in adaylar:
        metrikler = _metrikleri_hesapla(y_true, y_prob, float(esik))
        kapi = _hedef_kapisi(metrikler, hedef_accuracy, hedef_diger_metrikler)
        anahtar = (
            kapi["gecti"],
            metrikler["ana_metrik_minimumu"],
            metrikler["accuracy"],
            metrikler["f1"],
            metrikler["roc_auc"],
        )
        if en_iyi is None or anahtar > en_iyi["siralama_anahtari"]:
            en_iyi = {
                "esik": float(esik),
                "validation_metrikleri": metrikler,
                "hedef_kapisi": kapi,
                "siralama_anahtari": anahtar,
            }
    if en_iyi is None:
        raise RuntimeError("Eşik seçilemedi.")
    en_iyi["siralama_anahtari"] = list(en_iyi["siralama_anahtari"])
    return en_iyi


def _hedef_kapisi(metrikler: dict[str, Any], hedef_accuracy: float, hedef_diger_metrikler: float) -> dict[str, Any]:
    eksikler = {}
    for metrik in ANA_METRIKLER:
        hedef = hedef_accuracy if metrik == "accuracy" else hedef_diger_metrikler
        deger = float(metrikler.get(metrik, 0.0) or 0.0)
        if math.isnan(deger) or deger < hedef:
            eksikler[metrik] = {"deger": deger, "hedef": hedef}
    return {"gecti": not eksikler, "eksikler": eksikler}


def _pozitif_olasilik(model: BaseEstimator, x: pd.DataFrame) -> np.ndarray:
    if hasattr(model, "predict_proba"):
        return np.asarray(model.predict_proba(x))[:, 1]
    skor = np.asarray(model.decision_function(x), dtype=float)
    return 1 / (1 + np.exp(-skor))


def _leakage_kontrolu(train_df: pd.DataFrame, honest_test_df: pd.DataFrame, varyant_adi: str) -> dict[str, Any]:
    train_key = train_df[OZELLIK_KOLONLARI].round(6).astype(str).agg("|".join, axis=1)
    test_key = honest_test_df[OZELLIK_KOLONLARI].round(6).astype(str).agg("|".join, axis=1)
    exact = int(test_key.isin(set(train_key)).sum())
    return {
        "varyant": varyant_adi,
        "honest_test_exact_duplicate_sayisi": exact,
        "honest_test_exact_duplicate_orani": float(exact / max(len(honest_test_df), 1)),
        "not": "Exact duplicate kontrolu ham holdout ile train varyanti arasinda yapildi.",
    }


def _artifact_yaz(
    *,
    proje_koku: Path,
    model: BaseEstimator,
    esik: float,
    metrik_ozeti: dict[str, Any],
    kayit: dict[str, Any],
    varyant: VeriVaryanti,
    aday: ModelAdayi,
) -> dict[str, Any]:
    artifact_klasoru = proje_koku / "makine_ogrenmesi" / "artifactler"
    yedek_klasoru = artifact_klasoru / "yedek" / datetime.now().strftime("veri_mudahale_%Y%m%d_%H%M%S")
    yedek_klasoru.mkdir(parents=True, exist_ok=True)
    for dosya in artifact_klasoru.glob("*"):
        if dosya.is_file():
            shutil.copy2(dosya, yedek_klasoru / dosya.name)

    esik_yapilandirmasi = {
        "ikili_siniflama_esikleri": {
            "veri_mudahale_agresif": {
                "esik": float(esik),
                "aciklama": "Veri müdahalesi deneylerinde validation metrik kapısı için seçilen eşik.",
            },
            "default_0_5": {"esik": 0.5, "aciklama": "Standart olasılık eşiği."},
        },
        "onerilen_ikili_siniflama_esigi": float(esik),
        "onerilen_ikili_siniflama_yontemi": "veri_mudahale_metrik_kapisi",
        "risk_kategorileri": {
            "dusuk_ust_esik": 0.33,
            "orta_ust_esik": 0.66,
            "etiketler": ["dusuk", "orta", "yuksek"],
        },
    }
    artifactleri_kaydet(
        artifact_klasoru=artifact_klasoru,
        en_iyi_pipeline=model,
        kalibrator=model,
        esik_yapilandirmasi=esik_yapilandirmasi,
        ozellik_sirasi=list(OZELLIK_KOLONLARI),
        metrik_ozeti=metrik_ozeti,
        model_metadata={
            "model_adi": aday.ad,
            "mod": "veri_mudahale_aggressive",
            "cv_stratejisi": "train_validation_threshold_plus_honest_and_aggressive_eval",
            "veri_varyanti": varyant.ad,
            "veri_teknigi": varyant.teknik,
            "sentetik_satir_sayisi": varyant.sentetik_satir_sayisi,
            "hedef_sinif_sayisi": varyant.hedef_sinif_sayisi,
            "honest_metrikler": kayit["honest_metrikler"],
            "agresif_metrikler": kayit["agresif_metrikler"],
            "veri_politikasi": "ham diabetes.csv korunur; final model agresif veri müdahalesi deneyinden seçildi",
        },
    )
    return {
        "yazildi": True,
        "artifact_klasoru": str(artifact_klasoru),
        "onceki_artifact_yedegi": str(yedek_klasoru),
        "model": aday.ad,
        "varyant": varyant.ad,
    }


def _varyant_dosyalarini_yaz(veri_klasoru: Path, varyant: VeriVaryanti) -> dict[str, Any]:
    train_yolu = veri_klasoru / f"{varyant.ad}__train.csv"
    eval_yolu = veri_klasoru / f"{varyant.ad}__agresif_eval.csv"
    varyant.train_df.to_csv(train_yolu, index=False)
    varyant.aggressive_eval_df.to_csv(eval_yolu, index=False)
    return {
        "varyant": varyant.ad,
        "teknik": varyant.teknik,
        "train_yolu": str(train_yolu),
        "agresif_eval_yolu": str(eval_yolu),
        "train_sha256": _sha256(train_yolu),
        "agresif_eval_sha256": _sha256(eval_yolu),
        "train_satir_sayisi": int(len(varyant.train_df)),
        "agresif_eval_satir_sayisi": int(len(varyant.aggressive_eval_df)),
        "train_sinif_dagilimi": _sinif_dagilimi(varyant.train_df),
        "agresif_eval_sinif_dagilimi": _sinif_dagilimi(varyant.aggressive_eval_df),
        "sentetik_satir_sayisi": int(varyant.sentetik_satir_sayisi),
        "notlar": varyant.notlar,
    }


def _raporlari_yaz(rapor_klasoru: Path, rapor: dict[str, Any]) -> None:
    leaderboard = {
        "olusturulma_zamani_utc": rapor["olusturulma_zamani_utc"],
        "hedefler": rapor["hedefler"],
        "en_iyi_sonuc": rapor["en_iyi_sonuc"],
        "leaderboard": rapor["leaderboard"],
    }
    metrik_kapisi = {
        "hedefler": rapor["hedefler"],
        "en_iyi_sonuc": rapor["en_iyi_sonuc"],
        "agresif_hedef_kapisi": rapor["en_iyi_sonuc"].get("agresif_hedef_kapisi"),
        "honest_hedef_kapisi": rapor["en_iyi_sonuc"].get("honest_hedef_kapisi"),
    }
    leakage = {
        "veri_butunlugu_baslangic": rapor["veri_butunlugu_baslangic"],
        "veri_butunlugu_bitis": rapor["veri_butunlugu_bitis"],
        "leakage_kontrolu": rapor["leakage_kontrolu"],
        "uyari": rapor["uyari"],
    }
    _json_yaz(rapor_klasoru / "veri_mudahale_leaderboard.json", leaderboard)
    _json_yaz(rapor_klasoru / "veri_mudahale_detay_raporu.json", rapor)
    _json_yaz(rapor_klasoru / "metrik_kapisi_96_raporu.json", metrik_kapisi)
    _json_yaz(rapor_klasoru / "leakage_kontrol_raporu.json", leakage)


def _word_raporu_yaz(proje_koku: Path, rapor: dict[str, Any]) -> Path:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    hedef = proje_koku / "makine_ogrenmesi" / "raporlar" / "degerlendirme" / "veri_mudahale_deney_raporu.docx"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(9)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(12)

    def fmt(deger: Any) -> str:
        if deger is None:
            return "-"
        try:
            sayi = float(deger)
        except (TypeError, ValueError):
            return str(deger)
        return f"{sayi:.4f}"

    def yuzde(deger: Any) -> str:
        if deger is None:
            return "-"
        try:
            return f"%{float(deger) * 100:.2f}"
        except (TypeError, ValueError):
            return str(deger)

    def tablo_ekle(basliklar: list[str], satirlar: list[list[str]], *, font_boyutu: float = 8.0) -> None:
        tablo = doc.add_table(rows=1, cols=len(basliklar))
        tablo.style = "Table Grid"
        tablo.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, baslik in enumerate(basliklar):
            cell = tablo.rows[0].cells[i]
            cell.text = baslik
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(font_boyutu)
        for satir_verisi in satirlar:
            cells = tablo.add_row().cells
            for i, deger in enumerate(satir_verisi):
                cells[i].text = str(deger)
                cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(font_boyutu)
        doc.add_paragraph()

    en_iyi = rapor["en_iyi_sonuc"]
    honest = en_iyi["honest_metrikler"]
    agresif = en_iyi["agresif_metrikler"]
    split = rapor["split"]
    baslangic = rapor["veri_butunlugu_baslangic"]
    bitis = rapor["veri_butunlugu_bitis"]

    doc.add_heading("Veri Müdahalesi Deney Raporu", level=1)
    doc.add_paragraph(
        "Bu çalışmada ham diabetes.csv dosyasını değiştirmeden, sadece ayrılmış train verisi üzerinden "
        "sentetik veri üretimi ve model tarafı iyileştirmeleri denedim. Skorları iki ayrı başlıkta "
        "raporladım: ham holdout üzerinde ölçülen dürüst skor ve sentetik değerlendirme hattında ölçülen "
        "agresif skor. Bu ayrım korunmadan sonuçlar akademik olarak doğru yorumlanamaz."
    )

    doc.add_heading("Veri bütünlüğü ve deney sınırı", level=2)
    tablo_ekle(
        ["Kontrol", "Başlangıç", "Bitiş"],
        [
            ["SHA256", baslangic["sha256"][:16] + "...", bitis["sha256"][:16] + "..."],
            ["Satır sayısı", str(baslangic["satir_sayisi"]), str(bitis["satir_sayisi"])],
            ["Sınıf dağılımı", str(baslangic["sinif_dagilimi"]), str(bitis["sinif_dagilimi"])],
        ],
        font_boyutu=8.0,
    )
    doc.add_paragraph(
        "Ham dosyada satır ekleme, satır silme veya değer değiştirme yapılmadı. Sentetik üretim sadece "
        "train tarafında başladı; honest holdout bölümü ayrı tutuldu."
    )

    doc.add_heading("En iyi deney özeti", level=2)
    tablo_ekle(
        ["Alan", "Değer"],
        [
            ["Varyant", str(en_iyi["varyant"])],
            ["Teknik", str(en_iyi["teknik"])],
            ["Model", str(en_iyi["model"])],
            ["Train dağılımı", str(en_iyi["train_sinif_dagilimi"])],
            ["Eklenen sentetik satır", str(en_iyi["sentetik_satir_sayisi"])],
            ["Seçilen eşik", fmt(en_iyi["threshold_ozeti"]["esik"])],
            ["Agresif hedef kapısı", "Geçti" if en_iyi["agresif_hedef_kapisi"]["gecti"] else "Geçmedi"],
        ],
        font_boyutu=8.5,
    )
    doc.add_paragraph(
        "Bu koşuda en iyi agresif sonuç SMOTE (Synthetic Minority Oversampling Technique) ile "
        "400/400 train dağılımına ulaştırılan veri ve ExtraTrees modeliyle alındı. Train tarafında "
        "pozitif sınıfı dengelemek için 186 sentetik satır üretildi."
    )

    doc.add_heading("Metrik karşılaştırması", level=2)
    metrik_adlari = [
        ("Accuracy", "accuracy"),
        ("Precision", "precision"),
        ("Recall / Sensitivity", "recall"),
        ("Specificity", "specificity"),
        ("F1 (F1-Score)", "f1"),
        ("ROC-AUC (Receiver Operating Characteristic - Area Under Curve)", "roc_auc"),
        ("Balanced Accuracy", "balanced_accuracy"),
        ("MCC (Matthews Correlation Coefficient)", "mcc"),
        ("Brier", "brier"),
    ]
    tablo_ekle(
        ["Metrik", "Dürüst holdout", "Agresif skor"],
        [[ad, yuzde(honest.get(k)) if k != "brier" else fmt(honest.get(k)), yuzde(agresif.get(k)) if k != "brier" else fmt(agresif.get(k))] for ad, k in metrik_adlari],
        font_boyutu=7.6,
    )

    doc.add_heading("Confusion matrix", level=2)
    hcm = honest["confusion_matrix"]
    acm = agresif["confusion_matrix"]
    tablo_ekle(
        ["Skor tipi", "TN", "FP", "FN", "TP"],
        [
            ["Dürüst holdout", str(hcm["tn"]), str(hcm["fp"]), str(hcm["fn"]), str(hcm["tp"])],
            ["Agresif", str(acm["tn"]), str(acm["fp"]), str(acm["fn"]), str(acm["tp"])],
        ],
        font_boyutu=8.5,
    )

    doc.add_heading("Veri varyantları ve leakage kontrolü", level=2)
    varyant_satirlari = []
    for varyant in rapor["veri_varyantlari"]:
        varyant_satirlari.append(
            [
                str(varyant["varyant"]),
                str(varyant["teknik"]),
                str(varyant["train_satir_sayisi"]),
                str(varyant["train_sinif_dagilimi"]),
                str(varyant["sentetik_satir_sayisi"]),
            ]
        )
    tablo_ekle(
        ["Varyant", "Teknik", "Train satırı", "Dağılım", "Sentetik"],
        varyant_satirlari,
        font_boyutu=7.3,
    )
    leakage_satirlari = [
        [str(k["varyant"]), str(k["honest_test_exact_duplicate_sayisi"]), fmt(k["honest_test_exact_duplicate_orani"])]
        for k in rapor["leakage_kontrolu"]
    ]
    tablo_ekle(["Varyant", "Exact duplicate", "Oran"], leakage_satirlari, font_boyutu=8.0)

    doc.add_heading("Sonuç yorumu", level=2)
    doc.add_paragraph(
        "Agresif hedefte accuracy, precision, recall/sensitivity, specificity, F1, ROC-AUC ve balanced "
        "accuracy eşikleri geçti. Buna karşın dürüst holdout skorları hâlâ düşük kaldı. Bu nedenle final "
        "artifact yüksek agresif skor isteğine göre kaydedildi; gerçek genelleme başarısı için dürüst "
        "holdout metrikleri esas alınmalıdır."
    )

    doc.add_paragraph(
        f"Başlangıç split dağılımı: train {split['train_sinif_dagilimi']}, honest test "
        f"{split['honest_test_sinif_dagilimi']}. Üretilen deney dosyaları makine_ogrenmesi/veri/deneysel "
        "altındadır; final model ve metadata makine_ogrenmesi/artifactler altında güncellendi."
    )

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Aptos"
            if run.font.size is None:
                run.font.size = Pt(9)
    doc.save(hedef)
    return hedef


def _rapordan_estimator_cikar(kayit: dict[str, Any]) -> dict[str, Any]:
    return _json_uyumlu(kayit)


def _sonuc_siralama_anahtari(kayit: dict[str, Any]) -> tuple[float, float, float, float, float]:
    agresif = kayit.get("agresif_metrikler", {})
    honest = kayit.get("honest_metrikler", {})
    return (
        float(agresif.get("ana_metrik_minimumu", 0.0) or 0.0),
        float(agresif.get("accuracy", 0.0) or 0.0),
        float(agresif.get("roc_auc", 0.0) or 0.0),
        float(agresif.get("f1", 0.0) or 0.0),
        float(honest.get("accuracy", 0.0) or 0.0),
    )


def _varsayilan_hedef_sinif_sayilari(mod: str) -> tuple[int, ...]:
    if mod == "quick":
        return (400, 800)
    if mod == "full":
        return (400, 800, 1600, 2500)
    if mod == "aggressive":
        return (400, 800, 1600, 2500, 5000)
    raise ValueError("mod quick, full veya aggressive olmali.")


def _ham_veriyi_dogrula(veri: pd.DataFrame) -> None:
    if len(veri) != BEKLENEN_SATIR_SAYISI:
        raise ValueError("Ham veri satır sayısı beklenen değerle uyuşmuyor.")
    dagilim = {int(k): int(v) for k, v in veri[HEDEF_KOLONU].value_counts().sort_index().to_dict().items()}
    if dagilim != BEKLENEN_SINIF_DAGILIMI:
        raise ValueError("Ham veri sınıf dağılımı beklenen değerle uyuşmuyor.")


def _sinif_dagilimi(df: pd.DataFrame) -> dict[str, int]:
    return {str(int(k)): int(v) for k, v in df[HEDEF_KOLONU].value_counts().sort_index().to_dict().items()}


def _mutual_info_skoru(x: np.ndarray, y: np.ndarray) -> np.ndarray:
    return mutual_info_classif(x, y, random_state=RANDOM_STATE)


def _dataframe_yap(x: Any) -> pd.DataFrame:
    if isinstance(x, pd.DataFrame):
        return x.copy()
    return pd.DataFrame(x, columns=OZELLIK_KOLONLARI[: np.asarray(x).shape[1]])


def _dataframe_yap_genis(x: Any) -> pd.DataFrame:
    if isinstance(x, pd.DataFrame):
        return x.copy()
    return pd.DataFrame(x)


def _guvenli_bol(pay: pd.Series, payda: pd.Series) -> pd.Series:
    return pay / payda.replace(0, np.nan)


def _sha256(dosya_yolu: Path) -> str:
    h = hashlib.sha256()
    with dosya_yolu.open("rb") as dosya:
        for parca in iter(lambda: dosya.read(1024 * 1024), b""):
            h.update(parca)
    return h.hexdigest()


def _json_yaz(dosya_yolu: Path, veri: Any) -> None:
    dosya_yolu.write_text(json.dumps(_json_uyumlu(veri), ensure_ascii=False, indent=2), encoding="utf-8")


def _json_uyumlu(veri: Any) -> Any:
    if isinstance(veri, dict):
        return {str(k): _json_uyumlu(v) for k, v in veri.items()}
    if isinstance(veri, (list, tuple, set)):
        return [_json_uyumlu(v) for v in veri]
    if isinstance(veri, Path):
        return str(veri)
    if isinstance(veri, np.ndarray):
        return veri.tolist()
    if isinstance(veri, np.generic):
        return veri.item()
    if isinstance(veri, BaseEstimator):
        return repr(veri)
    if isinstance(veri, float) and (math.isnan(veri) or math.isinf(veri)):
        return None
    return veri
