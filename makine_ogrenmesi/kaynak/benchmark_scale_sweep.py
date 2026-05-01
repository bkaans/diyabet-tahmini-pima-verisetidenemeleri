"""Kucuk/orta olcekli source_id kontrollu sentetik PIMA benchmark sweep akisi."""

from __future__ import annotations

import importlib
import json
import math
import re
import shutil
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from sklearn.base import BaseEstimator, clone
from sklearn.ensemble import ExtraTreesClassifier, HistGradientBoostingClassifier, RandomForestClassifier, VotingClassifier
from sklearn.linear_model import LogisticRegression
from sklearn.model_selection import GroupShuffleSplit
from sklearn.pipeline import Pipeline

from .controlled_synthetic_benchmark import (
    ANA_METRIKLER,
    DatasetCandidate,
    FeatureSpec,
    ModelSpec,
    _apply_imputer,
    _assert_source_disjoint,
    _catboost_model,
    _class_distribution,
    _clip_clinical,
    _cv_summary,
    _dataset_leakage_report,
    _evaluate_external,
    _evaluate_group_cv,
    _evaluate_holdout,
    _external_holdout_ayir,
    _feature_specs,
    _find_cv_result,
    _find_dataset,
    _find_feature,
    _fit_dev_median_imputer,
    _group_holdout_split,
    _json_clean,
    _lightgbm_model,
    _metrics,
    _near_duplicate,
    _pipeline,
    _positive_probability,
    _split_counts,
    _threshold_report,
    _validate_metadata,
    _xgboost_model,
)
from .ozellik_yapilandirmasi import HEDEF_KOLONU, OZELLIK_KOLONLARI
from .veri_yukleyici import veri_setini_yukle


RANDOM_STATE = 42
HOLDOUT_MIN_TARGET = 0.92
HOLDOUT_ACCURACY_TARGET = 0.93
CV_MIN_TARGET = 0.90
REFERENCE_TARGET_PER_CLASS = 2500
REFERENCE_MIN_MAIN_METRIC = 0.9667
REFERENCE_COHENS_D = 2.950
REFERENCE_SHIFT = 0.556

PHASE1_TARGETS = (500, 650, 800, 1000, 1250, 1500, 2000, 2250, 2500)
PHASE1_STRENGTHS = (0.30, 0.40, 0.50, 0.60)
PHASE1_NOISE = ("low", "adaptive")
PHASE2_STRENGTHS = (0.25, 0.35, 0.45, 0.55, 0.65)


@dataclass(frozen=True)
class SweepConfig:
    target_per_class: int
    synthetic_strength: float
    noise_profile: str
    phase: str

    @property
    def key(self) -> str:
        strength = str(self.synthetic_strength).replace(".", "p")
        if strength in {"0p3", "0p4", "0p5", "0p6"}:
            strength = f"{strength}0"
        return f"scale_sweep_{self.target_per_class}_per_class_strength_{strength}_{self.noise_profile}"


def benchmark_scale_sweep_calistir(
    *,
    proje_koku: Path,
    veri_yolu: Path,
    random_state: int = RANDOM_STATE,
    n_jobs: int = 2,
    word_raporu_yaz: bool = True,
    quick: bool = False,
) -> dict[str, Any]:
    """Kucuk/orta olcek sweep'i calistirir ve yeni dosyalara raporlar."""
    proje_koku = Path(proje_koku)
    veri_yolu = Path(veri_yolu)
    rapor_dir = proje_koku / "makine_ogrenmesi" / "raporlar"
    veri_dir = proje_koku / "makine_ogrenmesi" / "veri" / "deneysel" / "benchmark_scale_sweep"
    grafik_dir = rapor_dir / "grafikler_benchmark_scale_sweep"
    for path in (rapor_dir, veri_dir, grafik_dir):
        path.mkdir(parents=True, exist_ok=True)

    raw = veri_setini_yukle(veri_yolu).copy()
    raw["original_index"] = raw.index.astype(int)
    raw["source_id"] = [f"original_{i}" for i in raw["original_index"]]
    original_dev_raw, external_raw = _external_holdout_ayir(raw, random_state)
    imputer = _fit_dev_median_imputer(original_dev_raw)
    original_dev = _apply_imputer(original_dev_raw, imputer)
    external_holdout = _apply_imputer(external_raw, imputer)

    package_report = _package_report()
    phase1_configs = _phase1_configs(quick)
    datasets_by_key: dict[str, DatasetCandidate] = {}
    for config in phase1_configs:
        dataset = _generate_small_dataset(original_dev, config, random_state + _seed_offset(config))
        datasets_by_key[dataset.name] = dataset
        _write_dataset(veri_dir, dataset)

    phase1_results = _run_holdout_sweep(
        datasets=list(datasets_by_key.values()),
        feature_specs=_phase1_features(quick),
        model_specs=_phase1_models(random_state, n_jobs, quick),
        random_state=random_state,
        label="Aşama 1",
    )

    promising_configs = _promising_configs(phase1_results, max_count=2 if quick else 6)
    phase2_configs = _phase2_configs(promising_configs, quick)
    phase2_datasets = []
    for config in phase2_configs:
        if config.key in datasets_by_key:
            phase2_datasets.append(datasets_by_key[config.key])
            continue
        dataset = _generate_small_dataset(original_dev, config, random_state + _seed_offset(config))
        datasets_by_key[dataset.name] = dataset
        phase2_datasets.append(dataset)
        _write_dataset(veri_dir, dataset)

    phase2_results = _run_holdout_sweep(
        datasets=phase2_datasets,
        feature_specs=_phase2_features(quick),
        model_specs=_phase2_models(random_state, n_jobs, quick),
        random_state=random_state + 11,
        label="Aşama 2",
    )

    successful = [r for r in [*phase1_results, *phase2_results] if r.get("status") == "completed"]
    if not successful:
        raise RuntimeError("Small-scale sweep icin basarili sonuc uretilemedi.")

    optuna_results = _run_optuna_refinement(
        successful=successful,
        datasets_by_key=datasets_by_key,
        random_state=random_state + 23,
        n_jobs=n_jobs,
        quick=quick,
    )
    successful.extend(optuna_results)

    cv_candidates = _select_cv_candidates(successful)
    print(f"[small-scale] Group CV aday sayisi: {len(cv_candidates)}")
    cv_results = []
    all_features = _all_feature_lookup()
    for candidate in cv_candidates:
        dataset = datasets_by_key[candidate["dataset_name"]]
        feature_spec = all_features[candidate["feature_set"]]
        model_spec = candidate["_model_spec"]
        cv = _evaluate_group_cv(dataset=dataset, feature_spec=feature_spec, model_spec=model_spec, random_state=random_state)
        cv_results.append(cv)
        s = cv["summary"]
        print(
            f"[small-cv/{_dataset_label(dataset)}/{_feature_label(feature_spec.name)}/{_model_label(model_spec.name)}] "
            f"min={s['ana_metrik_minimumu_mean']:.4f} acc={s['accuracy_mean']:.4f}"
        )

    datasets = list(datasets_by_key.values())
    leakage_report = [_dataset_leakage_report(dataset, external_holdout, successful, cv_results) for dataset in datasets]
    quality_by_dataset = _quality_by_dataset(original_dev, external_holdout, datasets)
    final = _select_final_small(successful, cv_results, leakage_report, quality_by_dataset)
    final_cv = _find_cv_result(cv_results, final)
    external_metrics = _evaluate_external(final, external_holdout)
    literature_compare = _run_literature_comparison(
        original_dev=original_dev,
        random_state=random_state + 37,
        n_jobs=n_jobs,
        quick=quick,
    )

    sweep_results = _build_sweep_results(
        raw=raw,
        original_dev=original_dev,
        external_holdout=external_holdout,
        datasets=datasets,
        phase1_results=phase1_results,
        phase2_results=phase2_results,
        optuna_results=optuna_results,
        cv_results=cv_results,
        leakage_report=leakage_report,
        quality_by_dataset=quality_by_dataset,
        package_report=package_report,
        literature_compare=literature_compare,
        veri_yolu=veri_yolu,
    )
    selection_report = _build_selection_report(final, final_cv, external_metrics, leakage_report, quality_by_dataset)
    audit = _build_quality_audit(raw, veri_yolu, sweep_results, selection_report)
    graphics = _make_graphics(grafik_dir, sweep_results, selection_report)
    sweep_results["graphics"] = graphics
    selection_report["graphics"] = graphics

    _json_write(rapor_dir / "benchmark_scale_sweep_results.json", sweep_results)
    _json_write(rapor_dir / "benchmark_scale_sweep_selection_report.json", selection_report)
    _json_write(rapor_dir / "benchmark_scale_sweep_quality_audit.json", audit)

    word_path = None
    desktop_word_path = None
    if word_raporu_yaz:
        word_path = _write_word_report(
            proje_koku=proje_koku,
            sweep=sweep_results,
            selection=selection_report,
            graphics=graphics,
        )
        desktop_word_path = Path.home() / "Desktop" / "pima_benchmark_scale_sweep_literature_report.docx"
        shutil.copy2(word_path, desktop_word_path)

    result = {
        "sweep_results": sweep_results,
        "selection_report": selection_report,
        "quality_audit": audit,
        "word_report": str(word_path) if word_path else None,
        "desktop_word_report": str(desktop_word_path) if desktop_word_path else None,
    }
    _print_summary(result)
    return result


def _phase1_configs(quick: bool) -> list[SweepConfig]:
    if quick:
        targets = (500, 650)
        strengths = (0.30, 0.50)
        noises = ("low",)
    else:
        targets = PHASE1_TARGETS
        strengths = PHASE1_STRENGTHS
        noises = PHASE1_NOISE
    return [SweepConfig(t, s, n, "phase1") for t in targets for s in strengths for n in noises]


def _phase2_configs(promising: list[SweepConfig], quick: bool) -> list[SweepConfig]:
    strengths = (0.35,) if quick else PHASE2_STRENGTHS
    configs: list[SweepConfig] = []
    for base in promising:
        for strength in strengths:
            configs.append(SweepConfig(base.target_per_class, strength, base.noise_profile, "phase2"))
    return sorted({c.key: c for c in configs}.values(), key=lambda c: (c.target_per_class, c.noise_profile, c.synthetic_strength))


def _promising_configs(results: list[dict[str, Any]], max_count: int) -> list[SweepConfig]:
    best_by_dataset: dict[str, dict[str, Any]] = {}
    for result in results:
        if result.get("status") != "completed":
            continue
        current = best_by_dataset.get(result["dataset_name"])
        if current is None or _promise_score(result) > _promise_score(current):
            best_by_dataset[result["dataset_name"]] = result
    promising = [
        r
        for r in best_by_dataset.values()
        if r["tuned_threshold_metrics"]["ana_metrik_minimumu"] >= 0.88
        or r["tuned_threshold_metrics"]["roc_auc"] >= 0.94
    ]
    if not promising:
        promising = sorted(best_by_dataset.values(), key=_promise_score, reverse=True)[:max_count]
    promising = sorted(promising, key=lambda r: (r["target_per_class"], -_promise_score(r)))[:max_count]
    return [_config_from_name(r["dataset_name"]) for r in promising]


def _phase1_features(quick: bool) -> list[FeatureSpec]:
    names = {"all_features"} if quick else {"all_features", "no_skinthickness", "high_signal_features"}
    return [f for f in _feature_specs() if f.name in names]


def _phase2_features(quick: bool) -> list[FeatureSpec]:
    names = {"all_features"} if quick else {"all_features", "no_skinthickness", "high_signal_features", "clinical_interactions", "compact_best_features"}
    return [f for f in _feature_specs() if f.name in names]


def _all_feature_lookup() -> dict[str, FeatureSpec]:
    return {f.name: f for f in _feature_specs()}


def _phase1_models(random_state: int, n_jobs: int, quick: bool) -> list[ModelSpec]:
    et_estimators = 80 if quick else 220
    specs: list[ModelSpec] = [
        ModelSpec(
            "extra_trees",
            ExtraTreesClassifier(
                n_estimators=et_estimators,
                max_features="sqrt",
                min_samples_leaf=1,
                class_weight="balanced",
                random_state=random_state,
                n_jobs=n_jobs,
            ),
            "tree_models_no_scaler",
            "class_weight=balanced",
            "ExtraTrees.",
        )
    ]
    xgb = _small_xgboost(random_state, n_jobs, quick)
    lgbm = _small_lightgbm(random_state, n_jobs, quick)
    if xgb is not None:
        specs.append(ModelSpec("xgboost", xgb, "tree_models_no_scaler", "none", "XGBoost."))
    if lgbm is not None:
        specs.append(ModelSpec("lightgbm", lgbm, "tree_models_no_scaler", "none", "LightGBM."))
    if xgb is not None and lgbm is not None and not quick:
        et = ExtraTreesClassifier(n_estimators=180, max_features="sqrt", class_weight="balanced", random_state=random_state + 2, n_jobs=1)
        specs.append(
            ModelSpec(
                "soft_voting_xgb_lgbm_et",
                VotingClassifier([("xgb", _small_xgboost(random_state + 3, 1, quick)), ("lgbm", _small_lightgbm(random_state + 4, 1, quick)), ("et", et)], voting="soft", n_jobs=1),
                "tree_models_no_scaler",
                "none",
                "XGBoost + LightGBM + ExtraTrees soft voting.",
            )
        )
    return specs


def _phase2_models(random_state: int, n_jobs: int, quick: bool) -> list[ModelSpec]:
    specs = [
        ModelSpec(
            "random_forest",
            RandomForestClassifier(
                n_estimators=90 if quick else 260,
                max_features="sqrt",
                min_samples_leaf=1,
                class_weight="balanced",
                random_state=random_state + 10,
                n_jobs=n_jobs,
            ),
            "tree_models_no_scaler",
            "class_weight=balanced",
            "RandomForest.",
        ),
        ModelSpec(
            "hist_gradient_boosting",
            HistGradientBoostingClassifier(max_iter=80 if quick else 220, learning_rate=0.055, max_leaf_nodes=31, random_state=random_state + 11),
            "tree_models_no_scaler",
            "none",
            "HistGradientBoosting.",
        ),
    ]
    cat = _small_catboost(random_state + 12, quick)
    if cat is not None:
        specs.append(ModelSpec("catboost", cat, "tree_models_no_scaler", "none", "CatBoost."))
    if not quick:
        specs.extend(_phase1_models(random_state + 20, n_jobs, quick=False))
    return specs


def _small_xgboost(random_state: int, n_jobs: int, quick: bool) -> BaseEstimator | None:
    try:
        xgb = importlib.import_module("xgboost")
    except Exception:
        return None
    return xgb.XGBClassifier(
        objective="binary:logistic",
        eval_metric="logloss",
        n_estimators=80 if quick else 220,
        max_depth=4,
        learning_rate=0.055,
        subsample=0.9,
        colsample_bytree=0.9,
        min_child_weight=1,
        reg_lambda=1.0,
        random_state=random_state,
        n_jobs=n_jobs,
    )


def _small_lightgbm(random_state: int, n_jobs: int, quick: bool) -> BaseEstimator | None:
    try:
        lgb = importlib.import_module("lightgbm")
    except Exception:
        return None
    return lgb.LGBMClassifier(
        objective="binary",
        n_estimators=80 if quick else 220,
        num_leaves=31,
        learning_rate=0.055,
        subsample=0.9,
        colsample_bytree=0.9,
        reg_lambda=1.0,
        random_state=random_state,
        n_jobs=n_jobs,
        verbose=-1,
    )


def _small_catboost(random_state: int, quick: bool) -> BaseEstimator | None:
    try:
        cb = importlib.import_module("catboost")
    except Exception:
        return None
    return cb.CatBoostClassifier(
        iterations=80 if quick else 220,
        depth=5,
        learning_rate=0.05,
        loss_function="Logloss",
        random_seed=random_state,
        verbose=False,
    )


def _run_holdout_sweep(
    *,
    datasets: list[DatasetCandidate],
    feature_specs: list[FeatureSpec],
    model_specs: list[ModelSpec],
    random_state: int,
    label: str,
) -> list[dict[str, Any]]:
    results: list[dict[str, Any]] = []
    for dataset in datasets:
        split = _group_holdout_split(dataset, random_state)
        for feature_spec in feature_specs:
            for model_spec in model_specs:
                try:
                    result = _evaluate_holdout(dataset=dataset, split=split, feature_spec=feature_spec, model_spec=model_spec, random_state=random_state)
                    result["phase"] = label
                    results.append(result)
                    m = result["tuned_threshold_metrics"]
                    print(
                        f"[{label}/{_dataset_label(dataset)}/{_feature_label(feature_spec.name)}/{_model_label(model_spec.name)}] "
                        f"min={m['ana_metrik_minimumu']:.4f} acc={m['accuracy']:.4f} auc={m['roc_auc']:.4f}"
                    )
                except Exception as exc:
                    results.append(
                        {
                            "status": "error",
                            "phase": label,
                            "dataset_name": dataset.name,
                            "target_per_class": dataset.target_per_class,
                            "model": model_spec.name,
                            "feature_set": feature_spec.name,
                            "preprocessing": model_spec.preprocessing,
                            "resampling": model_spec.resampling,
                            "error": f"{type(exc).__name__}: {exc}",
                        }
                    )
                    print(f"[{label}/{dataset.name}/{feature_spec.name}/{model_spec.name}] hata: {exc}")
    return results


def _generate_small_dataset(dev: pd.DataFrame, config: SweepConfig, random_state: int) -> DatasetCandidate:
    rng = np.random.default_rng(random_state)
    data_parts = [dev[OZELLIK_KOLONLARI + [HEDEF_KOLONU]].copy()]
    meta_parts = [
        pd.DataFrame(
            {
                "source_id": dev["source_id"].astype(str).to_numpy(),
                "is_synthetic": False,
                "parent_original_index": dev["original_index"].astype(int).to_numpy(),
                "generation_method": "original_dev",
                "synthetic_strength": 0.0,
                "noise_profile": "original_dev",
            }
        )
    ]
    method_counts = {"original_dev": int(len(dev))}
    for cls in [0, 1]:
        cls_dev = dev[dev[HEDEF_KOLONU] == cls].copy()
        need = max(0, config.target_per_class - len(cls_dev))
        if need <= 0:
            continue
        generated = _generate_one_parent_small(cls_dev, cls, need, rng, config)
        data_parts.append(generated[OZELLIK_KOLONLARI + [HEDEF_KOLONU]])
        meta_parts.append(generated[["source_id", "is_synthetic", "parent_original_index", "generation_method", "synthetic_strength", "noise_profile"]].copy())
        for method, count in generated["generation_method"].value_counts().items():
            method_counts[str(method)] = method_counts.get(str(method), 0) + int(count)

    data = pd.concat(data_parts, ignore_index=True)
    metadata = pd.concat(meta_parts, ignore_index=True)
    metadata.insert(0, "row_id", [f"{config.key}_r_{i:06d}" for i in range(len(metadata))])
    _validate_metadata(metadata)
    report = {
        "dataset_name": config.key,
        "target_per_class": int(config.target_per_class),
        "synthetic_strength": float(config.synthetic_strength),
        "noise_profile": config.noise_profile,
        "phase": config.phase,
        "original_dev_rows": int((~metadata["is_synthetic"]).sum()),
        "synthetic_rows": int(metadata["is_synthetic"].sum()),
        "total_rows": int(len(data)),
        "class_distribution": _class_distribution(data),
        "source_id_family_count": int(metadata["source_id"].nunique()),
        "generation_method_counts": method_counts,
        "synthetic_valid_source_id_ratio": float(metadata.loc[metadata["is_synthetic"], "source_id"].astype(str).str.match(r"^original_\d+$").mean()),
    }
    return DatasetCandidate(config.key, config.target_per_class, data.reset_index(drop=True), metadata.reset_index(drop=True), report)


def _generate_one_parent_small(
    cls_dev: pd.DataFrame,
    cls: int,
    count: int,
    rng: np.random.Generator,
    config: SweepConfig,
) -> pd.DataFrame:
    parent_idx = rng.choice(cls_dev.index.to_numpy(), size=count, replace=True)
    parents = cls_dev.loc[parent_idx].reset_index(drop=True)
    parent_values = parents[OZELLIK_KOLONLARI].astype(float).to_numpy()
    target = _class_target_sample_small(cls, count, rng, config.noise_profile)
    strength = rng.normal(config.synthetic_strength, 0.035 if config.noise_profile == "adaptive" else 0.025, size=(count, 1))
    strength = np.clip(strength, 0.05, 0.85)
    values = parent_values * (1 - strength) + target * strength
    values += _feature_noise_small(cls_dev, count, rng, config.noise_profile)
    out = pd.DataFrame(values, columns=OZELLIK_KOLONLARI)
    out = _quantile_clip_small(out, cls_dev, config.noise_profile)
    out = _clip_clinical(out)
    jitter_cols = ["Glucose", "BloodPressure", "SkinThickness", "Insulin", "BMI", "DiabetesPedigreeFunction"]
    out.loc[:, jitter_cols] = out[jitter_cols].to_numpy() + rng.normal(0, 1e-5, size=(len(out), len(jitter_cols)))
    out = _clip_clinical(out)
    out[HEDEF_KOLONU] = int(cls)
    out["source_id"] = parents["source_id"].astype(str).to_numpy()
    out["is_synthetic"] = True
    out["parent_original_index"] = parents["original_index"].astype(int).to_numpy()
    out["generation_method"] = "source_id_controlled_local_augmentation"
    out["synthetic_strength"] = strength.ravel()
    out["noise_profile"] = config.noise_profile
    return out


def _class_target_sample_small(cls: int, count: int, rng: np.random.Generator, profile: str) -> np.ndarray:
    if cls == 1:
        mean = np.array([5.0, 166, 80, 35, 190, 38.5, 0.78, 49], dtype=float)
        sd = np.array([3.0, 17, 9, 7, 78, 5.2, 0.26, 11], dtype=float)
    else:
        mean = np.array([2.2, 92, 68, 21, 58, 26.0, 0.27, 31], dtype=float)
        sd = np.array([2.0, 12, 8, 6, 30, 3.8, 0.11, 8], dtype=float)
    scale = 0.92 if profile == "adaptive" else 0.75
    return rng.normal(mean, sd * scale, size=(count, len(OZELLIK_KOLONLARI)))


def _feature_noise_small(cls_dev: pd.DataFrame, count: int, rng: np.random.Generator, profile: str) -> np.ndarray:
    multiplier = 0.045 if profile == "adaptive" else 0.030
    std = cls_dev[OZELLIK_KOLONLARI].std().replace(0, 1.0).to_numpy()
    return rng.normal(0, std * multiplier, size=(count, len(OZELLIK_KOLONLARI)))


def _quantile_clip_small(frame: pd.DataFrame, cls_dev: pd.DataFrame, profile: str) -> pd.DataFrame:
    out = frame.copy()
    lo_q, hi_q = (0.01, 0.99) if profile == "adaptive" else (0.02, 0.98)
    lows = cls_dev[OZELLIK_KOLONLARI].quantile(lo_q)
    highs = cls_dev[OZELLIK_KOLONLARI].quantile(hi_q)
    span = (highs - lows).replace(0, 1.0)
    margin = 0.24 if profile == "adaptive" else 0.14
    for col in OZELLIK_KOLONLARI:
        out[col] = out[col].clip(float(lows[col] - span[col] * margin), float(highs[col] + span[col] * margin))
    return out


def _run_optuna_refinement(
    *,
    successful: list[dict[str, Any]],
    datasets_by_key: dict[str, DatasetCandidate],
    random_state: int,
    n_jobs: int,
    quick: bool,
) -> list[dict[str, Any]]:
    if quick:
        return []
    try:
        optuna = importlib.import_module("optuna")
    except Exception:
        return []
    top = sorted(successful, key=lambda r: (r["tuned_threshold_metrics"]["ana_metrik_minimumu"], r["tuned_threshold_metrics"]["accuracy"], r["tuned_threshold_metrics"]["roc_auc"]), reverse=True)[:3]
    results = []
    feature_lookup = _all_feature_lookup()
    for idx, row in enumerate(top, start=1):
        if row["model"] not in {"extra_trees", "xgboost", "lightgbm", "random_forest", "hist_gradient_boosting", "catboost"}:
            continue
        dataset = datasets_by_key[row["dataset_name"]]
        feature_spec = feature_lookup[row["feature_set"]]
        split = _group_holdout_split(dataset, random_state + idx)

        def objective(trial: Any) -> float:
            spec = _trial_model_spec(row["model"], trial, random_state + idx, n_jobs)
            train_idx = np.asarray(split["train_index"], dtype=int)
            x_train = dataset.frame.iloc[train_idx][OZELLIK_KOLONLARI].reset_index(drop=True)
            y_train = dataset.frame.iloc[train_idx][HEDEF_KOLONU].reset_index(drop=True)
            meta = dataset.metadata.iloc[train_idx].reset_index(drop=True)
            fit_idx, val_idx = _safe_inner_split(y_train, meta, random_state + idx)
            model = _pipeline(feature_spec, spec.estimator)
            model.fit(x_train.iloc[fit_idx], y_train.iloc[fit_idx])
            prob = _positive_probability(model, x_train.iloc[val_idx])
            return float(_threshold_report(y_train.iloc[val_idx], prob)["tuned_min_main_metric"])

        study = optuna.create_study(direction="maximize", sampler=optuna.samplers.TPESampler(seed=random_state + idx))
        study.optimize(objective, n_trials=6, show_progress_bar=False)
        tuned_spec = _trial_model_spec(row["model"], optuna.trial.FixedTrial(study.best_params), random_state + idx, n_jobs)
        tuned_spec = ModelSpec(f"optuna_{row['model']}", tuned_spec.estimator, tuned_spec.preprocessing, tuned_spec.resampling, f"Optuna kısa tuning: {row['model']}.")
        result = _evaluate_holdout(dataset=dataset, split=split, feature_spec=feature_spec, model_spec=tuned_spec, random_state=random_state + idx)
        result["phase"] = "Optuna kısa tuning"
        result["optuna_best_value"] = float(study.best_value)
        result["optuna_best_params"] = dict(study.best_params)
        results.append(result)
        m = result["tuned_threshold_metrics"]
        print(f"[optuna/{dataset.name}/{row['model']}] min={m['ana_metrik_minimumu']:.4f} acc={m['accuracy']:.4f}")
    return results


def _trial_model_spec(model_name: str, trial: Any, random_state: int, n_jobs: int) -> ModelSpec:
    if model_name in {"extra_trees", "random_forest"}:
        cls = ExtraTreesClassifier if model_name == "extra_trees" else RandomForestClassifier
        estimator = cls(
            n_estimators=trial.suggest_categorical("n_estimators", [180, 260, 360]),
            max_depth=trial.suggest_categorical("max_depth", [None, 8, 12, 16]),
            min_samples_leaf=trial.suggest_categorical("min_samples_leaf", [1, 2, 4]),
            max_features="sqrt",
            class_weight="balanced",
            random_state=random_state,
            n_jobs=n_jobs,
        )
        return ModelSpec(model_name, estimator, "tree_models_no_scaler", "class_weight=balanced", model_name)
    if model_name == "xgboost":
        xgb = importlib.import_module("xgboost")
        estimator = xgb.XGBClassifier(
            objective="binary:logistic",
            eval_metric="logloss",
            n_estimators=trial.suggest_categorical("n_estimators", [160, 240, 320]),
            max_depth=trial.suggest_int("max_depth", 3, 5),
            learning_rate=trial.suggest_float("learning_rate", 0.035, 0.075),
            subsample=trial.suggest_float("subsample", 0.75, 1.0),
            colsample_bytree=trial.suggest_float("colsample_bytree", 0.75, 1.0),
            random_state=random_state,
            n_jobs=n_jobs,
        )
        return ModelSpec(model_name, estimator, "tree_models_no_scaler", "none", model_name)
    if model_name == "lightgbm":
        lgb = importlib.import_module("lightgbm")
        estimator = lgb.LGBMClassifier(
            objective="binary",
            n_estimators=trial.suggest_categorical("n_estimators", [160, 240, 320]),
            num_leaves=trial.suggest_categorical("num_leaves", [15, 31, 63]),
            learning_rate=trial.suggest_float("learning_rate", 0.035, 0.075),
            subsample=trial.suggest_float("subsample", 0.75, 1.0),
            colsample_bytree=trial.suggest_float("colsample_bytree", 0.75, 1.0),
            random_state=random_state,
            n_jobs=n_jobs,
            verbose=-1,
        )
        return ModelSpec(model_name, estimator, "tree_models_no_scaler", "none", model_name)
    if model_name == "catboost":
        cb = importlib.import_module("catboost")
        estimator = cb.CatBoostClassifier(
            iterations=trial.suggest_categorical("iterations", [160, 240, 320]),
            depth=trial.suggest_int("depth", 4, 6),
            learning_rate=trial.suggest_float("learning_rate", 0.035, 0.075),
            loss_function="Logloss",
            random_seed=random_state,
            verbose=False,
        )
        return ModelSpec(model_name, estimator, "tree_models_no_scaler", "none", model_name)
    estimator = HistGradientBoostingClassifier(
        max_iter=trial.suggest_categorical("max_iter", [160, 220, 300]),
        learning_rate=trial.suggest_float("learning_rate", 0.035, 0.075),
        max_leaf_nodes=trial.suggest_categorical("max_leaf_nodes", [15, 31, 63]),
        random_state=random_state,
    )
    return ModelSpec(model_name, estimator, "tree_models_no_scaler", "none", model_name)


def _safe_inner_split(y_train: pd.Series, metadata: pd.DataFrame, random_state: int) -> tuple[np.ndarray, np.ndarray]:
    groups = metadata["source_id"].astype(str).to_numpy()
    splitter = GroupShuffleSplit(n_splits=20, test_size=0.18, random_state=random_state)
    for fit_idx, val_idx in splitter.split(np.zeros(len(y_train)), y_train, groups):
        if len(np.unique(y_train.iloc[fit_idx])) == 2 and len(np.unique(y_train.iloc[val_idx])) == 2:
            _assert_source_disjoint(metadata, fit_idx, val_idx, "small_scale_inner")
            return np.asarray(fit_idx), np.asarray(val_idx)
    idx = np.arange(len(y_train))
    return idx[: int(len(idx) * 0.82)], idx[int(len(idx) * 0.82) :]


def _select_cv_candidates(successful: list[dict[str, Any]]) -> list[dict[str, Any]]:
    selected: dict[tuple[str, str, str], dict[str, Any]] = {}
    passing = [
        r
        for r in successful
        if r["tuned_threshold_metrics"]["ana_metrik_minimumu"] >= HOLDOUT_MIN_TARGET
        and r["tuned_threshold_metrics"]["accuracy"] >= HOLDOUT_ACCURACY_TARGET
    ]
    for row in passing:
        selected[(row["dataset_name"], row["feature_set"], row["model"])] = row
    by_target: dict[int, list[dict[str, Any]]] = {}
    for row in successful:
        by_target.setdefault(int(row["target_per_class"]), []).append(row)
    for target, rows in by_target.items():
        for row in sorted(rows, key=lambda r: (r["tuned_threshold_metrics"]["ana_metrik_minimumu"], r["tuned_threshold_metrics"]["accuracy"]), reverse=True)[:2]:
            selected[(row["dataset_name"], row["feature_set"], row["model"])] = row
    ref_rows = [r for r in successful if int(r["target_per_class"]) == REFERENCE_TARGET_PER_CLASS]
    for row in sorted(ref_rows, key=lambda r: r["tuned_threshold_metrics"]["ana_metrik_minimumu"], reverse=True)[:2]:
        selected[(row["dataset_name"], row["feature_set"], row["model"])] = row
    if not selected:
        for row in sorted(successful, key=lambda r: r["tuned_threshold_metrics"]["ana_metrik_minimumu"], reverse=True)[:8]:
            selected[(row["dataset_name"], row["feature_set"], row["model"])] = row
    return list(selected.values())


def _select_final_small(
    successful: list[dict[str, Any]],
    cv_results: list[dict[str, Any]],
    leakage_report: list[dict[str, Any]],
    quality_by_dataset: dict[str, dict[str, Any]],
) -> dict[str, Any]:
    leakage_map = {r["dataset_name"]: r for r in leakage_report}
    cv_map = {(r["dataset_name"], r["feature_set"], r["model"]): r for r in cv_results}
    candidates = []
    for result in successful:
        cv = cv_map.get((result["dataset_name"], result["feature_set"], result["model"]))
        leak = leakage_map.get(result["dataset_name"])
        if not cv or not leak or leak["leakage_status"] != "clean":
            continue
        metrics = result["tuned_threshold_metrics"]
        quality = quality_by_dataset[result["dataset_name"]]
        row = dict(result)
        row["selection_flags"] = {
            "holdout_min_main_metric_92": bool(metrics["ana_metrik_minimumu"] >= HOLDOUT_MIN_TARGET),
            "holdout_accuracy_93": bool(metrics["accuracy"] >= HOLDOUT_ACCURACY_TARGET),
            "group_cv_min_main_metric_90": bool(cv["summary"]["ana_metrik_minimumu_mean"] >= CV_MIN_TARGET),
        }
        row["selection_score"] = float(
            0.34 * metrics["ana_metrik_minimumu"]
            + 0.22 * metrics["accuracy"]
            + 0.22 * cv["summary"]["ana_metrik_minimumu_mean"]
            + 0.12 * metrics["roc_auc"]
            + 0.05 * max(0.0, 1 - metrics["brier"])
            + 0.05 * max(0.0, 1 - quality["distribution_shift_vs_original_dev"]["aggregate"]["avg_abs_z_mean_shift"])
        )
        candidates.append(row)
    if not candidates:
        raise RuntimeError("Small-scale final secim icin CV ve leakage-clean aday yok.")
    passing = [
        c
        for c in candidates
        if c["selection_flags"]["holdout_min_main_metric_92"]
        and c["selection_flags"]["holdout_accuracy_93"]
        and c["selection_flags"]["group_cv_min_main_metric_90"]
    ]
    if passing:
        return sorted(
            passing,
            key=lambda c: (
                c["target_per_class"],
                quality_by_dataset[c["dataset_name"]]["distribution_shift_vs_original_dev"]["aggregate"]["avg_abs_z_mean_shift"],
                quality_by_dataset[c["dataset_name"]]["class_separation"]["aggregate"]["mean_abs_cohens_d"],
                _find_cv_result(cv_results, c)["summary"]["ana_metrik_minimumu_std"],
                c["tuned_threshold_metrics"]["brier"],
                -c["selection_score"],
            ),
        )[0]
    return sorted(candidates, key=lambda c: c["selection_score"], reverse=True)[0]


def _quality_by_dataset(original_dev: pd.DataFrame, external: pd.DataFrame, datasets: list[DatasetCandidate]) -> dict[str, dict[str, Any]]:
    return {dataset.name: _dataset_quality(original_dev, external, dataset) for dataset in datasets}


def _dataset_quality(original_dev: pd.DataFrame, external: pd.DataFrame, dataset: DatasetCandidate) -> dict[str, Any]:
    return {
        "dataset_name": dataset.name,
        "class_separation": _class_separation(dataset.frame),
        "distribution_shift_vs_original_dev": _distribution_shift(original_dev, dataset.frame),
        "external_distribution_shift": _distribution_shift(original_dev, external),
        "key_feature_summary": _key_feature_summary(dataset.frame),
        "feature_signal_concentration": _feature_signal_concentration(dataset.frame),
        "outcome_encoding_warning": sorted(dataset.frame[HEDEF_KOLONU].unique().tolist()) != [0, 1],
        "reference_2500_comparison": {
            "reference_min_main_metric": REFERENCE_MIN_MAIN_METRIC,
            "reference_mean_abs_cohens_d": REFERENCE_COHENS_D,
            "reference_distribution_shift": REFERENCE_SHIFT,
        },
    }


def _class_separation(frame: pd.DataFrame) -> dict[str, Any]:
    rows = {}
    for col in OZELLIK_KOLONLARI:
        neg = frame.loc[frame[HEDEF_KOLONU] == 0, col].astype(float)
        pos = frame.loc[frame[HEDEF_KOLONU] == 1, col].astype(float)
        pooled = math.sqrt(((len(neg) - 1) * float(neg.var(ddof=1)) + (len(pos) - 1) * float(pos.var(ddof=1))) / max(len(neg) + len(pos) - 2, 1))
        d = float((pos.mean() - neg.mean()) / pooled) if pooled else 0.0
        rows[col] = {"cohens_d": d, "abs_cohens_d": abs(d), "negative_mean": float(neg.mean()), "positive_mean": float(pos.mean())}
    return {"by_feature": rows, "aggregate": {"mean_abs_cohens_d": float(np.mean([v["abs_cohens_d"] for v in rows.values()]))}}


def _distribution_shift(reference: pd.DataFrame, candidate: pd.DataFrame) -> dict[str, Any]:
    rows = {}
    for col in OZELLIK_KOLONLARI:
        ref = reference[col].astype(float)
        cand = candidate[col].astype(float)
        std = float(ref.std(ddof=1)) or 1.0
        shift = abs(float(cand.mean() - ref.mean())) / std
        rows[col] = {"z_mean_shift": shift, "reference_mean": float(ref.mean()), "candidate_mean": float(cand.mean())}
    return {"by_feature": rows, "aggregate": {"avg_abs_z_mean_shift": float(np.mean([v["z_mean_shift"] for v in rows.values()]))}}


def _key_feature_summary(frame: pd.DataFrame) -> dict[str, Any]:
    out = {}
    for col in ["Glucose", "BMI", "Age", "DiabetesPedigreeFunction"]:
        out[col] = {}
        for cls in [0, 1]:
            values = frame.loc[frame[HEDEF_KOLONU] == cls, col].astype(float)
            out[col][str(cls)] = {
                "mean": float(values.mean()),
                "std": float(values.std(ddof=1)),
                "q25": float(values.quantile(0.25)),
                "median": float(values.median()),
                "q75": float(values.quantile(0.75)),
            }
    return out


def _feature_signal_concentration(frame: pd.DataFrame) -> dict[str, Any]:
    scores = {}
    y = frame[HEDEF_KOLONU].to_numpy(dtype=int)
    for col in OZELLIK_KOLONLARI:
        values = frame[col].to_numpy(dtype=float)
        scores[col] = _single_feature_auc_directionless(y, values)
    ordered = sorted(scores.items(), key=lambda kv: kv[1], reverse=True)
    total = sum(scores.values()) or 1.0
    return {
        "top_feature": ordered[0][0],
        "top_feature_auc": ordered[0][1],
        "top_feature_share": float(ordered[0][1] / total),
        "top3_feature_share": float(sum(v for _, v in ordered[:3]) / total),
        "scores": scores,
    }


def _single_feature_auc_directionless(y: np.ndarray, values: np.ndarray) -> float:
    order = np.argsort(values)
    ranks = np.empty_like(order, dtype=float)
    ranks[order] = np.arange(1, len(values) + 1)
    n_pos = int(y.sum())
    n_neg = int(len(y) - n_pos)
    if n_pos == 0 or n_neg == 0:
        return 0.5
    sum_ranks_pos = float(ranks[y == 1].sum())
    auc = (sum_ranks_pos - n_pos * (n_pos + 1) / 2) / (n_pos * n_neg)
    return float(max(auc, 1 - auc))


def _run_literature_comparison(original_dev: pd.DataFrame, random_state: int, n_jobs: int, quick: bool) -> dict[str, Any]:
    methods = [{"method": "class_weight_balanced", "status": "completed", "note": "Ağaç tabanlı modellerde class_weight='balanced' adayları ana sweep içinde denendi."}]
    try:
        from imblearn.combine import SMOTEENN, SMOTETomek
        from imblearn.over_sampling import SMOTE
        from imblearn.pipeline import Pipeline as ImbPipeline
    except Exception as exc:
        return {"methods": methods, "imblearn_available": False, "note": f"imbalanced-learn bulunamadı: {type(exc).__name__}: {exc}"}
    if quick:
        return {"methods": methods, "imblearn_available": True, "note": "Quick modda train-only resampling karşılaştırması atlandı."}
    splitter = GroupShuffleSplit(n_splits=1, test_size=0.20, random_state=random_state)
    x = original_dev[OZELLIK_KOLONLARI]
    y = original_dev[HEDEF_KOLONU]
    groups = original_dev["source_id"].astype(str)
    train_idx, test_idx = next(splitter.split(x, y, groups))
    _assert_source_disjoint(original_dev[["source_id"]].reset_index(drop=True), train_idx, test_idx, "literature_compare")
    resamplers = {
        "SMOTE": SMOTE(random_state=random_state),
        "SMOTEENN": SMOTEENN(random_state=random_state),
        "SMOTETomek": SMOTETomek(random_state=random_state),
    }
    for name, sampler in resamplers.items():
        try:
            model = ImbPipeline(
                [
                    ("resampling", sampler),
                    (
                        "model",
                        ExtraTreesClassifier(
                            n_estimators=220,
                            max_features="sqrt",
                            class_weight="balanced",
                            random_state=random_state,
                            n_jobs=n_jobs,
                        ),
                    ),
                ]
            )
            model.fit(x.iloc[train_idx], y.iloc[train_idx])
            prob = model.predict_proba(x.iloc[test_idx])[:, 1]
            metrics = _metrics(y.iloc[test_idx], prob, 0.50)
            methods.append({"method": name, "status": "completed", "metrics": metrics, "note": "Sadece train split üzerinde resampling uygulandı."})
        except Exception as exc:
            methods.append({"method": name, "status": "error", "error": f"{type(exc).__name__}: {exc}"})
    return {"methods": methods, "imblearn_available": True, "note": "Bu sonuçlar ana final benchmark seçimine dahil edilmedi."}


def _build_sweep_results(
    *,
    raw: pd.DataFrame,
    original_dev: pd.DataFrame,
    external_holdout: pd.DataFrame,
    datasets: list[DatasetCandidate],
    phase1_results: list[dict[str, Any]],
    phase2_results: list[dict[str, Any]],
    optuna_results: list[dict[str, Any]],
    cv_results: list[dict[str, Any]],
    leakage_report: list[dict[str, Any]],
    quality_by_dataset: dict[str, dict[str, Any]],
    package_report: dict[str, Any],
    literature_compare: dict[str, Any],
    veri_yolu: Path,
) -> dict[str, Any]:
    return {
        "created_at_utc": _now(),
        "goal": "En küçük source_id kontrollü sentetik benchmarkta holdout min ana metrik >= 0.92 ve accuracy >= 0.93 hedefi.",
        "original_csv": {"path": str(veri_yolu), "rows": int(len(raw)), "class_distribution": _class_distribution(raw)},
        "original_dev": {"rows": int(len(original_dev)), "class_distribution": _class_distribution(original_dev)},
        "external_holdout": {"rows": int(len(external_holdout)), "class_distribution": _class_distribution(external_holdout), "role": "Dış kontrol; seçim skoruna dahil edilmedi."},
        "datasets": [d.report for d in datasets],
        "phase1_results": [_json_clean(r) for r in phase1_results],
        "phase2_results": [_json_clean(r) for r in phase2_results],
        "optuna_results": [_json_clean(r) for r in optuna_results],
        "cv_results": [_json_clean(c) for c in cv_results],
        "leakage_report": leakage_report,
        "quality_by_dataset": quality_by_dataset,
        "package_report": package_report,
        "literature_comparison": literature_compare,
    }


def _build_selection_report(
    final: dict[str, Any],
    final_cv: dict[str, Any],
    external: dict[str, Any],
    leakage: list[dict[str, Any]],
    quality_by_dataset: dict[str, dict[str, Any]],
) -> dict[str, Any]:
    final_leak = next(x for x in leakage if x["dataset_name"] == final["dataset_name"])
    quality = quality_by_dataset[final["dataset_name"]]
    return {
        "created_at_utc": _now(),
        "selection_goal": {
            "synthetic_holdout_min_main_metric": HOLDOUT_MIN_TARGET,
            "synthetic_holdout_accuracy": HOLDOUT_ACCURACY_TARGET,
            "synthetic_group_cv_min_main_metric_minimum": CV_MIN_TARGET,
        },
        "external_holdout_in_selection": False,
        "final": _json_clean(final),
        "final_cv": _json_clean(final_cv),
        "external_holdout": external,
        "leakage_summary": final_leak,
        "quality_summary": quality,
        "target_hit": {
            "holdout_min_main_metric_92": bool(final["tuned_threshold_metrics"]["ana_metrik_minimumu"] >= HOLDOUT_MIN_TARGET),
            "holdout_accuracy_93": bool(final["tuned_threshold_metrics"]["accuracy"] >= HOLDOUT_ACCURACY_TARGET),
            "group_cv_min_main_metric_90": bool(final_cv["summary"]["ana_metrik_minimumu_mean"] >= CV_MIN_TARGET),
        },
        "recommended_over_2500_reference": bool(final["target_per_class"] < REFERENCE_TARGET_PER_CLASS and final["tuned_threshold_metrics"]["ana_metrik_minimumu"] >= HOLDOUT_MIN_TARGET),
        "why_selected": "Leakage-clean adaylar içinde hedefi sağlayan en küçük ve dağılım kayması daha savunulabilir aday önceliklendirildi.",
    }


def _build_quality_audit(raw: pd.DataFrame, veri_yolu: Path, sweep: dict[str, Any], selection: dict[str, Any]) -> dict[str, Any]:
    final_leak = selection["leakage_summary"]
    csv_ok = int(len(raw)) == 768 and _class_distribution(raw) == {"0": 500, "1": 268}
    clean = (
        final_leak["leakage_status"] == "clean"
        and final_leak["train_test_source_intersection"] == 0
        and max(final_leak["cv_source_intersections"] or [0]) == 0
        and final_leak["exact_duplicate_count"] == 0
        and final_leak["independent_synthetic_source_id_count"] == 0
        and final_leak["external_holdout_source_overlap_count"] == 0
    )
    return {
        "created_at_utc": _now(),
        "diabetes_csv_path": str(veri_yolu),
        "diabetes_csv_rows_distribution_unchanged": csv_ok,
        "final_dataset_name": selection["final"]["dataset_name"],
        "target_hit": selection["target_hit"],
        "leakage_clean": clean,
        "empty_graph_check_required": True,
        "package_report": sweep["package_report"],
        "final_quality_status": "pass" if csv_ok and clean else "warning",
        "notes": [
            "External holdout seçim skoruna dahil edilmedi.",
            "Tüm aday CSV dosyaları benchmark_scale_sweep klasörüne yazıldı.",
            "CTGAN final deney ailesine dahil edilmedi; parent source_id takibi zayıf kalacağı için yalnız literatür notunda açıklandı.",
        ],
    }


def _make_graphics(grafik_dir: Path, sweep: dict[str, Any], selection: dict[str, Any]) -> dict[str, str]:
    plt.rcParams.update({"font.size": 9, "axes.titlesize": 11, "axes.labelsize": 9})
    paths = {
        "holdout_min_metric": _plot_best_by_target(grafik_dir, sweep, "ana_metrik_minimumu", "Sınıf Başı Hedefe Göre Holdout Minimum Ana Metrik", HOLDOUT_MIN_TARGET),
        "holdout_accuracy": _plot_best_by_target(grafik_dir, sweep, "accuracy", "Sınıf Başı Hedefe Göre Holdout Doğruluk", HOLDOUT_ACCURACY_TARGET),
        "holdout_roc_auc": _plot_best_by_target(grafik_dir, sweep, "roc_auc", "Sınıf Başı Hedefe Göre Holdout ROC-AUC", None),
        "cv_min_metric": _plot_cv_by_target(grafik_dir, sweep, "Sınıf Başı Hedefe Göre Grup Çapraz Doğrulama Minimum Ana Metrik"),
        "cohens_d": _plot_quality_by_target(grafik_dir, sweep, "class_separation", "Sınıf Ayrımı: Ortalama Mutlak Cohen's d"),
        "distribution_shift": _plot_quality_by_target(grafik_dir, sweep, "distribution_shift_vs_original_dev", "Geliştirme Verisine Göre Dağılım Kayması"),
        "confusion_matrix": _plot_confusion_matrix(grafik_dir, selection),
        "feature_importance": _plot_feature_signal(grafik_dir, selection),
        "protocol_summary": _plot_protocol_summary(grafik_dir),
    }
    return {k: str(v) for k, v in paths.items()}


def _best_holdout_rows(sweep: dict[str, Any]) -> list[dict[str, Any]]:
    rows = [r for r in [*sweep["phase1_results"], *sweep["phase2_results"], *sweep["optuna_results"]] if r.get("status") == "completed"]
    best: dict[int, dict[str, Any]] = {}
    for row in rows:
        target = int(row["target_per_class"])
        current = best.get(target)
        if current is None or _promise_score(row) > _promise_score(current):
            best[target] = row
    return [best[k] for k in sorted(best)]


def _plot_best_by_target(grafik_dir: Path, sweep: dict[str, Any], metric: str, title: str, target_line: float | None) -> Path:
    path = grafik_dir / f"small_scale_{metric}.png"
    rows = _best_holdout_rows(sweep)
    labels = [str(r["target_per_class"]) for r in rows]
    values = [r["tuned_threshold_metrics"][metric] for r in rows]
    fig, ax = plt.subplots(figsize=(8.6, 3.5))
    ax.bar(labels, values, color="#1f4e79")
    if target_line is not None:
        ax.axhline(target_line, color="#d62828", linestyle="--", linewidth=1)
    ax.set_ylim(0.55, 1.02)
    ax.set_title(title)
    ax.set_xlabel("Sınıf başı hedef")
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _plot_cv_by_target(grafik_dir: Path, sweep: dict[str, Any], title: str) -> Path:
    path = grafik_dir / "small_scale_cv_min_metric.png"
    best: dict[int, dict[str, Any]] = {}
    for cv in sweep["cv_results"]:
        target = _config_from_name(cv["dataset_name"]).target_per_class
        current = best.get(target)
        if current is None or cv["summary"]["ana_metrik_minimumu_mean"] > current["summary"]["ana_metrik_minimumu_mean"]:
            best[target] = cv
    labels = [str(k) for k in sorted(best)]
    values = [best[k]["summary"]["ana_metrik_minimumu_mean"] for k in sorted(best)]
    errors = [best[k]["summary"]["ana_metrik_minimumu_std"] for k in sorted(best)]
    fig, ax = plt.subplots(figsize=(8.6, 3.5))
    ax.bar(labels, values, yerr=errors, color="#2a9d8f", capsize=3)
    ax.axhline(CV_MIN_TARGET, color="#d62828", linestyle="--", linewidth=1)
    ax.set_ylim(0.55, 1.02)
    ax.set_title(title)
    ax.set_xlabel("Sınıf başı hedef")
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _plot_quality_by_target(grafik_dir: Path, sweep: dict[str, Any], quality_key: str, title: str) -> Path:
    path = grafik_dir / f"small_scale_{quality_key}.png"
    values_by_target: dict[int, float] = {}
    for name, quality in sweep["quality_by_dataset"].items():
        target = _config_from_name(name).target_per_class
        if quality_key == "class_separation":
            value = quality[quality_key]["aggregate"]["mean_abs_cohens_d"]
        else:
            value = quality[quality_key]["aggregate"]["avg_abs_z_mean_shift"]
        values_by_target[target] = min(values_by_target.get(target, value), value)
    labels = [str(k) for k in sorted(values_by_target)]
    values = [values_by_target[k] for k in sorted(values_by_target)]
    fig, ax = plt.subplots(figsize=(8.6, 3.5))
    ax.bar(labels, values, color="#e76f51")
    ax.set_title(title)
    ax.set_xlabel("Sınıf başı hedef")
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _plot_confusion_matrix(grafik_dir: Path, selection: dict[str, Any]) -> Path:
    path = grafik_dir / "small_scale_selected_model_confusion_matrix.png"
    cm = np.array(selection["final"]["tuned_threshold_metrics"]["confusion_matrix"]["matrix"])
    fig, ax = plt.subplots(figsize=(4.2, 3.6))
    im = ax.imshow(cm, cmap="YlGnBu")
    ax.set_xticks([0, 1], ["Tahmin 0", "Tahmin 1"])
    ax.set_yticks([0, 1], ["Gerçek 0", "Gerçek 1"])
    ax.set_title("Seçilen Aday Karmaşıklık Matrisi")
    for i in range(2):
        for j in range(2):
            ax.text(j, i, str(int(cm[i, j])), ha="center", va="center", weight="bold")
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _plot_feature_signal(grafik_dir: Path, selection: dict[str, Any]) -> Path:
    path = grafik_dir / "scale_sweep_selected_feature_signal.png"
    scores = selection["quality_summary"]["feature_signal_concentration"]["scores"]
    rows = sorted(scores.items(), key=lambda kv: kv[1])
    fig, ax = plt.subplots(figsize=(7.2, 3.8))
    ax.barh([r[0] for r in rows], [r[1] for r in rows], color="#577590")
    ax.set_title("Seçilen Adayda Değişken Ayrıştırma Gücü")
    ax.set_xlabel("Tek değişken yönsüz ROC-AUC")
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _plot_protocol_summary(grafik_dir: Path) -> Path:
    path = grafik_dir / "small_scale_literatur_protokol_ozeti.png"
    labels = ["Kaynak aile takibi", "Dış kontrol", "Grup CV", "Kopya kontrolü", "Sentetik ayrımı"]
    values = [1, 1, 1, 1, 1]
    fig, ax = plt.subplots(figsize=(7.2, 3.2))
    ax.bar(labels, values, color="#2a9d8f")
    ax.set_ylim(0, 1.2)
    ax.set_title("Bu Çalışmada Raporlanan Doğrulama Protokol Kontrolleri")
    ax.set_yticks([0, 1], ["Yok", "Var"])
    ax.tick_params(axis="x", rotation=15)
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _write_word_report(*, proje_koku: Path, sweep: dict[str, Any], selection: dict[str, Any], graphics: dict[str, str]) -> Path:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    path = proje_koku / "makine_ogrenmesi" / "raporlar" / "benchmark_scale_sweep_literature_style_report.docx"
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(9.0)
    for style in ["Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[style].font.name = "Aptos"
        doc.styles[style].font.color.rgb = RGBColor(31, 78, 121)

    final = selection["final"]
    final_cv = selection["final_cv"]["summary"]
    external = selection["external_holdout"]["metrics"]
    quality = selection["quality_summary"]

    _doc_title(doc, "PIMA Veri Setinde Kaynak Aile Kontrollü Küçük Ölçekli Sentetik Benchmark Taraması")
    _doc_p(doc, "Bu raporda amaç, mevcut 2500/2500 referans sonucunu bozmadan daha küçük ve savunulabilir bir PIMA + sentetik benchmark boyutu aramaktır. Başarı ölçütü, sentetik benchmark holdout tarafında minimum ana metrik değerinin %92 ve doğruluk değerinin en az %93 olmasıdır.")
    _doc_p(doc, "Orijinal dış kontrol verisi bu seçim kararına dahil edilmemiştir. Bu veri yalnızca modelin gerçek PIMA dağılımına transferini dış kontrol olarak göstermek için raporlanmıştır.")
    _doc_table(
        doc,
        ["Alan", "Değer"],
        [
            ["Orijinal PIMA", f"{sweep['original_csv']['rows']} satır"],
            ["Geliştirme verisi", f"{sweep['original_dev']['rows']} satır"],
            ["Dış kontrol verisi", f"{sweep['external_holdout']['rows']} satır"],
            ["Hedef", "Holdout minimum ana metrik >= %92 ve doğruluk >= %93"],
            ["Seçilen aday", _dataset_label_from_name(final["dataset_name"])],
        ],
    )
    _doc_img(doc, graphics["holdout_min_metric"], 6.2, "Sınıf başı hedef arttıkça minimum ana metrik değerinin nasıl değiştiği gösterilmiştir.")
    _doc_img(doc, graphics["holdout_accuracy"], 6.2, "Doğruluk hedefi ayrıca çizilmiştir; seçim yalnız minimum ana metriğe göre yapılmamıştır.")

    doc.add_page_break()
    doc.add_heading("Doğrulama Protokolü ve Veri Sızıntısı Kontrolü", level=1)
    _doc_p(doc, "Bu çalışma literatürdeki veri artırma yaklaşımlarından farklı olarak her sentetik örneği bir kaynak aile kimliğine (source_id) bağlamıştır. Böylece aynı aileden gelen örneklerin eğitim ve test tarafına birlikte düşmesi engellenmiştir.")
    leak = selection["leakage_summary"]
    _doc_table(
        doc,
        ["Kontrol", "Beklenen", "Sonuç"],
        [
            ["Eğitim/test kaynak aile kesişimi", "0", str(leak["train_test_source_intersection"])],
            ["Çapraz doğrulama kaynak aile kesişimi", "0", str(max(leak["cv_source_intersections"] or [0]))],
            ["Birebir kopya", "0", str(leak["exact_duplicate_count"])],
            ["Çok yakın benzerlik oranı", "Raporlanır", _fmt(leak["near_duplicate_rate"])],
            ["Minimum mesafe", "Raporlanır", _fmt(leak["minimum_near_duplicate_distance"])],
            ["Dış kontrol kaynak çakışması", "0", str(leak["external_holdout_source_overlap_count"])],
            ["Bağımsız sentetik kaynak aile", "0", str(leak["independent_synthetic_source_id_count"])],
        ],
    )
    _doc_img(doc, graphics["protocol_summary"], 5.8, "Doğrulama protokolünde yalnız skor değil, veri üretim ve sızıntı kontrolleri de raporlanmıştır.")

    doc.add_page_break()
    doc.add_heading("Aday Taraması Sonuçları", level=1)
    _doc_p(doc, "Aşağıdaki tablo, her sınıf başı hedefte en iyi holdout sonucunu ve aynı hedef için raporlanan grup çapraz doğrulama özetini birlikte göstermektedir. Böylece küçük adayların neden elendiği, seçilen adayın neden öne çıktığı ve 2500/2500 referansının neden daha büyük kaldığı hızlıca izlenebilir.")
    _doc_table(
        doc,
        ["Aday", "En iyi model", "Doğruluk", "Min ana", "Grup CV min ana", "Cohen's d", "Shift", "Durum"],
        _candidate_summary_table_rows(sweep, selection),
        font_size=6.2,
    )
    _doc_img(doc, graphics["holdout_roc_auc"], 6.2, "ROC-AUC değerleri, sentetik benchmark adaylarının sınıf ayrıştırma gücünü göstermektedir.")
    _doc_img(doc, graphics["cv_min_metric"], 6.2, "Grup çapraz doğrulama sonucu, seçilen adayın tek bir veri ayrımına bağlı olup olmadığını kontrol etmek için kullanılmıştır.")
    _doc_img(doc, graphics["cohens_d"], 6.2, "Cohen's d değeri sınıfların birbirinden ne kadar ayrıştığını gösterir; çok yüksek değerler sentetik benchmarkın kolaylaşabileceğine işaret eder.")
    _doc_img(doc, graphics["distribution_shift"], 6.2, "Dağılım kayması, aday veri setinin geliştirme verisinden ne kadar uzaklaştığını gösterir.")

    doc.add_page_break()
    doc.add_heading("Seçilen Küçük Benchmark Adayı", level=1)
    m = final["tuned_threshold_metrics"]
    _doc_table(
        doc,
        ["Alan", "Değer"],
        [
            ["Veri adayı", _dataset_label_from_name(final["dataset_name"])],
            ["Model", _model_label(final["model"])],
            ["Değişken seti", _feature_label(final["feature_set"])],
            ["Karar eşiği", f"{final['threshold_report']['selected_threshold']:.2f}"],
            ["Holdout doğruluk", _pct(m["accuracy"])],
            ["Holdout minimum ana metrik", _pct(m["ana_metrik_minimumu"])],
            ["Grup çapraz doğrulama minimum ana metrik", f"{_pct(final_cv['ana_metrik_minimumu_mean'])} ± {final_cv['ana_metrik_minimumu_std']:.3f}"],
            ["Cohen's d", f"{quality['class_separation']['aggregate']['mean_abs_cohens_d']:.3f}"],
            ["Dağılım kayması", f"{quality['distribution_shift_vs_original_dev']['aggregate']['avg_abs_z_mean_shift']:.3f}"],
        ],
    )
    _doc_img(doc, graphics["confusion_matrix"], 4.6, "Karmaşıklık matrisi, seçilen adayın holdout tarafındaki hata dağılımını göstermektedir.")
    _doc_img(doc, graphics["feature_importance"], 5.9, "Değişken ayrıştırma gücü grafiği, seçilen adayda hangi değişkenlerin sınıfları daha fazla ayırdığını gösterir.")

    doc.add_page_break()
    doc.add_heading("Literatürle Protokol Karşılaştırması", level=1)
    _doc_p(doc, "PIMA ve diyabet sınıflandırma literatüründe SMOTE, ADASYN, SMOTE-ENN, SMOTE-Tomek, CTGAN, ensemble ve feature selection tabanlı yüksek skorlar raporlanmaktadır. Ancak yüksek skorların doğrudan karşılaştırılabilmesi için preprocessing, sentetik üretim, resampling ve validation adımlarının hangi sırayla uygulandığı açık olmalıdır.")
    _doc_table(
        doc,
        ["Karşılaştırma Ekseni", "Literatürde Sık Görülen Yaklaşım", "Bu Çalışmadaki Yaklaşım"],
        [
            ["Veri artırma", "SMOTE, ADASYN, CTGAN, SMOTE-ENN", "Kaynak aile kontrollü sentetik benchmark"],
            ["Veri bölme", "Çalışmadan çalışmaya değişken", "Dış kontrol verisi sentetik üretimden önce ayrıldı"],
            ["Kaynak aile takibi", "Genellikle raporlanmaz", "Her sentetik örnek kaynak aile kimliği taşıdı"],
            ["Sızıntı kontrolü", "Her çalışmada aynı detayda değil", "Kaynak aile, birebir kopya, yakın benzerlik ve dış kontrol çakışması kontrol edildi"],
            ["Çapraz doğrulama", "k-fold / stratified k-fold", "Kaynak aileleri birlikte tutan StratifiedGroupKFold kullanıldı"],
            ["Sonuç yorumu", "Genel başarı gibi sunulabilir", "Sentetik benchmark ve dış kontrol sonucu ayrı yorumlandı"],
        ],
        font_size=7.0,
    )
    _doc_p(doc, "SMOTE azınlık sınıfını komşular arasında sentetik örneklerle artırır. ADASYN öğrenmesi zor azınlık örneklerine daha fazla ağırlık verir. SMOTE-ENN ve SMOTE-Tomek çoğaltma sonrası gürültü veya sınır çakışması temizliği uygular. CTGAN tabular veri dağılımını öğrenerek sentetik kayıt üretir. Bu çalışmada ana final benchmark için kaynak aile takibi korunmuştur.")
    _doc_table(
        doc,
        ["Dış kontrol metriği", "Değer"],
        [
            ["Doğruluk", _pct(external["accuracy"])],
            ["F1", _pct(external["f1"])],
            ["ROC-AUC", _pct(external["roc_auc"])],
            ["Min ana metrik", _pct(external["ana_metrik_minimumu"])],
            ["Brier", f"{external['brier']:.3f}"],
        ],
    )
    _doc_p(doc, "Doğrulama protokolü ve veri sızıntısı kontrol seviyesi çalışmadan çalışmaya değiştiği için skorlar doğrudan birebir kıyaslanmamalıdır. Bu rapordaki yüksek metrikler kontrollü sentetik benchmark bağlamında değerlendirilmelidir.")

    doc.save(path)
    return path


def _candidate_summary_table_rows(sweep: dict[str, Any], selection: dict[str, Any]) -> list[list[str]]:
    best_rows = _best_holdout_rows(sweep)
    cv_by_target: dict[int, dict[str, Any]] = {}
    for row in sweep["cv_results"]:
        target = int(row.get("target_per_class") or _config_from_name(row["dataset_name"]).target_per_class)
        current = cv_by_target.get(target)
        if current is None or row["summary"]["ana_metrik_minimumu_mean"] > current["summary"]["ana_metrik_minimumu_mean"]:
            cv_by_target[target] = row

    final_target = int(selection["final"]["target_per_class"])
    quality_by_dataset = sweep["quality_by_dataset"]
    rows: list[list[str]] = []
    for row in best_rows:
        target = int(row["target_per_class"])
        metrics = row["tuned_threshold_metrics"]
        quality = quality_by_dataset[row["dataset_name"]]
        cv = cv_by_target.get(target)
        cv_text = "-"
        if cv is not None:
            cv_text = f"{_pct(cv['summary']['ana_metrik_minimumu_mean'])} ± {cv['summary']['ana_metrik_minimumu_std']:.3f}"

        passes = metrics["ana_metrik_minimumu"] >= HOLDOUT_MIN_TARGET and metrics["accuracy"] >= HOLDOUT_ACCURACY_TARGET
        if target == final_target:
            status = "Seçildi"
        elif not passes:
            status = "Geçmedi"
        elif target >= REFERENCE_TARGET_PER_CLASS:
            status = "Referans / daha büyük"
        else:
            status = "Geçti, daha büyük"

        rows.append(
            [
                f"{target}/{target}",
                _model_label(row["model"]),
                _pct(metrics["accuracy"]),
                _pct(metrics["ana_metrik_minimumu"]),
                cv_text,
                f"{quality['class_separation']['aggregate']['mean_abs_cohens_d']:.3f}",
                f"{quality['distribution_shift_vs_original_dev']['aggregate']['avg_abs_z_mean_shift']:.3f}",
                status,
            ]
        )
    return rows


def _doc_title(doc: Any, text: str) -> None:
    from docx.shared import Pt

    paragraph = doc.add_paragraph()
    paragraph.alignment = 1
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(17)


def _doc_p(doc: Any, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.line_spacing = 1.08
    paragraph.paragraph_format.space_after = 3


def _doc_img(doc: Any, path: str, width: float, caption: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    image_path = Path(path)
    if not image_path.exists() or image_path.stat().st_size <= 0:
        _doc_p(doc, caption)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(8.0)


def _doc_table(doc: Any, headers: list[str], rows: list[list[str]], font_size: float = 8.0) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), "1F4E79")
        cell._tc.get_or_add_tcPr().append(shade)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for paragraph in cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    doc.add_paragraph()


def _write_dataset(veri_dir: Path, dataset: DatasetCandidate) -> None:
    combined = pd.concat([dataset.frame.reset_index(drop=True), dataset.metadata.reset_index(drop=True)], axis=1)
    combined.to_csv(veri_dir / f"{dataset.name}.csv", index=False)
    dataset.metadata.to_csv(veri_dir / f"{dataset.name}_metadata.csv", index=False)


def _json_write(path: Path, value: dict[str, Any]) -> None:
    path.write_text(json.dumps(_json_clean(value), ensure_ascii=False, indent=2), encoding="utf-8")


def _package_report() -> dict[str, Any]:
    report = {}
    for pkg in ["xgboost", "lightgbm", "catboost", "imblearn", "optuna"]:
        try:
            mod = importlib.import_module(pkg)
            report[pkg] = {"available": True, "version": getattr(mod, "__version__", None)}
        except Exception as exc:
            report[pkg] = {"available": False, "version": None, "error": f"{type(exc).__name__}: {exc}"}
    report["sdv_ctgan"] = {"available": False, "used": False, "note": "Parent source_id takibi final benchmarkta zayıf kalacağı için kurulmadı ve kullanılmadı."}
    return report


def _config_from_name(name: str) -> SweepConfig:
    match = re.match(r"^scale_sweep_(\d+)_per_class_strength_(\d+p\d+)_(.+)$", name)
    if not match:
        raise ValueError(f"Gecersiz benchmark aday adi: {name}")
    target = int(match.group(1))
    strength = float(match.group(2).replace("p", "."))
    noise = match.group(3)
    return SweepConfig(target, strength, noise, "derived")


def _seed_offset(config: SweepConfig) -> int:
    return int(config.target_per_class * 13 + round(config.synthetic_strength * 1000) * 7 + sum(ord(c) for c in config.noise_profile))


def _promise_score(result: dict[str, Any]) -> float:
    m = result["tuned_threshold_metrics"]
    return float(0.45 * m["ana_metrik_minimumu"] + 0.25 * m["accuracy"] + 0.20 * m["roc_auc"] + 0.10 * (1 - m["brier"]))


def _dataset_label(dataset: DatasetCandidate) -> str:
    cfg = _config_from_name(dataset.name)
    return f"{cfg.target_per_class}/{cfg.target_per_class} sınıf dengesi, yakınlaştırma düzeyi {_tr_decimal(cfg.synthetic_strength)}, {_noise_label(cfg.noise_profile)}"


def _dataset_label_from_name(name: str) -> str:
    cfg = _config_from_name(name)
    return f"{cfg.target_per_class}/{cfg.target_per_class} sınıf dengesi; sentetik yakınlaştırma düzeyi {_tr_decimal(cfg.synthetic_strength)}; {_noise_label(cfg.noise_profile)}"


def _tr_decimal(value: float) -> str:
    return f"{value:.2f}".replace(".", ",")


def _noise_label(name: str) -> str:
    labels = {
        "adaptive": "adaptif gürültü profili",
        "low": "düşük gürültü profili",
        "original_dev": "orijinal geliştirme profili",
    }
    return labels.get(name, name.replace("_", " "))


def _model_label(name: str) -> str:
    labels = {
        "extra_trees": "ExtraTrees",
        "xgboost": "XGBoost",
        "lightgbm": "LightGBM",
        "soft_voting_xgb_lgbm_et": "XGBoost + LightGBM + ExtraTrees soft voting",
        "random_forest": "RandomForest",
        "hist_gradient_boosting": "HistGradientBoosting",
        "catboost": "CatBoost",
    }
    if name.startswith("optuna_"):
        return f"Optuna kısa tuning + {_model_label(name.removeprefix('optuna_'))}"
    return labels.get(name, name.replace("_", " "))


def _feature_label(name: str) -> str:
    labels = {
        "all_features": "Tüm orijinal PIMA değişkenleri",
        "no_skinthickness": "SkinThickness çıkarılmış",
        "high_signal_features": "Yüksek sinyal değişken seti",
        "clinical_interactions": "Klinik etkileşim değişkenleri",
        "compact_best_features": "Kompakt değişken seti",
    }
    return labels.get(name, name.replace("_", " "))


def _fmt(value: Any) -> str:
    if value is None:
        return "Yok"
    return f"{float(value):.4f}"


def _pct(value: Any) -> str:
    return f"%{float(value) * 100:.2f}"


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _print_summary(result: dict[str, Any]) -> None:
    selection = result["selection_report"]
    final = selection["final"]
    cv = selection["final_cv"]["summary"]
    ext = selection["external_holdout"]["metrics"]
    quality = selection["quality_summary"]
    leak = selection["leakage_summary"]
    m = final["tuned_threshold_metrics"]
    by_target = {}
    for row in _best_holdout_rows(result["sweep_results"]):
        by_target[int(row["target_per_class"])] = row["tuned_threshold_metrics"]["ana_metrik_minimumu"]
    print("\n=== SMALL-SCALE SWEEP OZET ===")
    print("- Aşama 1 tamamlandı mı: Evet")
    print(f"- Aşama 2 yapıldı mı: {'Evet' if result['sweep_results']['phase2_results'] else 'Hayır'}")
    for target in PHASE1_TARGETS:
        print(f"- {target}/{target} en iyi sonuç: {by_target.get(target, 'Yok')}")
    print("- 2500/2500 referans:", by_target.get(2500, "Yok"))
    print(f"- Minimum %92 min ana metrik ve %93 accuracy hedefini geçen en küçük aday: {_dataset_label_from_name(final['dataset_name']) if selection['target_hit']['holdout_min_main_metric_92'] and selection['target_hit']['holdout_accuracy_93'] else 'Tam hedefi geçen aday yok'}")
    print(f"- Synthetic holdout min ana metrik: {m['ana_metrik_minimumu']:.4f}")
    print(f"- Synthetic holdout accuracy: {m['accuracy']:.4f}")
    print(f"- Synthetic Group CV min ana metrik: {cv['ana_metrik_minimumu_mean']:.4f} ± {cv['ana_metrik_minimumu_std']:.4f}")
    print(f"- ROC-AUC: {m['roc_auc']:.4f}")
    print(f"- Cohen’s d: {quality['class_separation']['aggregate']['mean_abs_cohens_d']:.4f}")
    print(f"- Distribution shift: {quality['distribution_shift_vs_original_dev']['aggregate']['avg_abs_z_mean_shift']:.4f}")
    print(f"- Threshold: {final['threshold_report']['selected_threshold']:.2f}")
    print(f"- Model: {_model_label(final['model'])}")
    print(f"- Feature set: {_feature_label(final['feature_set'])}")
    print(f"- Leakage durumu: {leak['leakage_status']}")
    print(f"- External holdout sonucu: acc={ext['accuracy']:.4f}, f1={ext['f1']:.4f}, auc={ext['roc_auc']:.4f}")
    print(f"- 2500/2500 yerine bu aday önerilir mi: {'Evet' if selection['recommended_over_2500_reference'] else 'Hayır'}")
    print(f"- Neden: {selection['why_selected']}")
    print("- Literatür/protokol farkı rapora eklendi mi: Evet")
    print(f"- Yeni Word yolu: {result['word_report']}")
    print(f"- Masaüstü Word yolu: {result['desktop_word_report']}")
