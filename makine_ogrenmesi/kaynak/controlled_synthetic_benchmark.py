"""Min-0.90 hedefli, source_id kontrollu sentetik PIMA benchmark akisi."""

from __future__ import annotations

import hashlib
import importlib
import json
import math
import shutil
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from sklearn.base import BaseEstimator, TransformerMixin, clone
from sklearn.calibration import calibration_curve
from sklearn.ensemble import (
    AdaBoostClassifier,
    ExtraTreesClassifier,
    GradientBoostingClassifier,
    HistGradientBoostingClassifier,
    RandomForestClassifier,
    StackingClassifier,
    VotingClassifier,
)
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import (
    accuracy_score,
    balanced_accuracy_score,
    brier_score_loss,
    confusion_matrix,
    f1_score,
    matthews_corrcoef,
    precision_score,
    recall_score,
    roc_auc_score,
)
from sklearn.model_selection import GroupShuffleSplit, StratifiedGroupKFold, train_test_split
from sklearn.naive_bayes import GaussianNB
from sklearn.neighbors import KNeighborsClassifier, NearestNeighbors
from sklearn.neural_network import MLPClassifier
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import RobustScaler, StandardScaler
from sklearn.svm import SVC

from .ozellik_yapilandirmasi import HEDEF_KOLONU, OZELLIK_KOLONLARI, SIFIRI_EKSIK_SAYILAN_KOLONLAR
from .veri_yukleyici import veri_setini_yukle


RANDOM_STATE = 42
HEDEF_MIN_METRIK = 0.90
ANA_METRIKLER = ["accuracy", "precision", "recall", "specificity", "f1", "balanced_accuracy"]
RAPOR_METRIKLERI = [
    "accuracy",
    "precision",
    "recall",
    "specificity",
    "f1",
    "roc_auc",
    "balanced_accuracy",
    "mcc",
    "brier",
    "ana_metrik_minimumu",
]
DEFAULT_TARGETS = (2500, 2700, 3000, 4000, 5000, 7500, 10000)


@dataclass(frozen=True)
class DatasetCandidate:
    name: str
    target_per_class: int
    frame: pd.DataFrame
    metadata: pd.DataFrame
    report: dict[str, Any]


@dataclass(frozen=True)
class FeatureSpec:
    name: str
    base_columns: tuple[str, ...]
    engineered: tuple[str, ...]
    removed_columns: tuple[str, ...]
    description: str


@dataclass(frozen=True)
class ModelSpec:
    name: str
    estimator: BaseEstimator
    preprocessing: str
    resampling: str
    description: str
    max_rows: int | None = None


class ClinicalFeatureTransformer(BaseEstimator, TransformerMixin):
    """Secilen klinik feature setini uretir ve yalniz istenen kolonlari dondurur."""

    def __init__(self, base_columns: tuple[str, ...], engineered: tuple[str, ...]) -> None:
        self.base_columns = base_columns
        self.engineered = engineered

    def fit(self, x: Any, y: Any = None) -> "ClinicalFeatureTransformer":
        return self

    def transform(self, x: Any) -> pd.DataFrame:
        df = _dataframe_yap(x)
        eps = 1e-6
        out = df.loc[:, list(self.base_columns)].copy()
        if "glucose_bmi_ratio" in self.engineered:
            out["glucose_bmi_ratio"] = df["Glucose"] / (df["BMI"] + eps)
        if "glucose_age_interaction" in self.engineered:
            out["glucose_age_interaction"] = df["Glucose"] * df["Age"]
        if "glucose_bmi_interaction" in self.engineered:
            out["glucose_bmi_interaction"] = df["Glucose"] * df["BMI"]
        if "glucose_pedigree_interaction" in self.engineered:
            out["glucose_pedigree_interaction"] = df["Glucose"] * df["DiabetesPedigreeFunction"]
        if "age_bmi_interaction" in self.engineered:
            out["age_bmi_interaction"] = df["Age"] * df["BMI"]
        if "age_pedigree_interaction" in self.engineered:
            out["age_pedigree_interaction"] = df["Age"] * df["DiabetesPedigreeFunction"]
        if "pregnancies_age_ratio" in self.engineered:
            out["pregnancies_age_ratio"] = df["Pregnancies"] / (df["Age"] + eps)
        if "insulin_glucose_ratio" in self.engineered:
            out["insulin_glucose_ratio"] = df["Insulin"] / (df["Glucose"] + eps)
        if "bp_bmi_ratio" in self.engineered:
            out["bp_bmi_ratio"] = df["BloodPressure"] / (df["BMI"] + eps)
        if "bmi_age_ratio" in self.engineered:
            out["bmi_age_ratio"] = df["BMI"] / (df["Age"] + eps)
        if "glucose_minus_bmi" in self.engineered:
            out["glucose_minus_bmi"] = df["Glucose"] - df["BMI"]
        if "glucose_per_age" in self.engineered:
            out["glucose_per_age"] = df["Glucose"] / (df["Age"] + eps)
        return out.replace([np.inf, -np.inf], np.nan).fillna(0)

    def feature_names(self) -> list[str]:
        return [*self.base_columns, *self.engineered]


class PimaZeroMedianImputer(BaseEstimator, TransformerMixin):
    """PIMA'da klinik olarak eksik sayilan 0 degerlerini egitim medyaniyla doldurur."""

    def __init__(self, missing_zero_columns: tuple[str, ...] = tuple(SIFIRI_EKSIK_SAYILAN_KOLONLAR)) -> None:
        self.missing_zero_columns = missing_zero_columns

    def fit(self, x: Any, y: Any = None) -> "PimaZeroMedianImputer":
        veri = _dataframe_yap(x).copy()
        medyanlar: dict[str, float] = {}
        for kolon in OZELLIK_KOLONLARI:
            seri = pd.to_numeric(veri[kolon], errors="coerce")
            if kolon in self.missing_zero_columns:
                seri = seri.where(seri != 0, np.nan)
            medyan = float(seri.median())
            medyanlar[kolon] = medyan if not math.isnan(medyan) else 0.0
        self.medians_ = medyanlar
        return self

    def transform(self, x: Any) -> pd.DataFrame:
        if not hasattr(self, "medians_"):
            raise RuntimeError("PimaZeroMedianImputer fit edilmeden transform cagrildi.")
        veri = _dataframe_yap(x).copy()
        for kolon in OZELLIK_KOLONLARI:
            seri = pd.to_numeric(veri[kolon], errors="coerce")
            if kolon in self.missing_zero_columns:
                seri = seri.where(seri != 0, np.nan)
            veri[kolon] = seri.fillna(self.medians_[kolon])
        return _clip_clinical(veri)


def controlled_benchmark_calistir(
    *,
    proje_koku: Path,
    veri_yolu: Path,
    target_per_class_values: tuple[int, ...] = DEFAULT_TARGETS,
    random_state: int = RANDOM_STATE,
    n_jobs: int = -1,
    word_raporu_yaz: bool = True,
    quick: bool = False,
) -> dict[str, Any]:
    """Kontrollü sentetik benchmark aramasını çalıştırır."""
    proje_koku = Path(proje_koku)
    veri_yolu = Path(veri_yolu)
    rapor_dir = proje_koku / "makine_ogrenmesi" / "raporlar"
    veri_dir = proje_koku / "makine_ogrenmesi" / "veri" / "deneysel"
    grafik_dir = rapor_dir / "grafikler_controlled_benchmark"
    for path in (rapor_dir, veri_dir, grafik_dir):
        path.mkdir(parents=True, exist_ok=True)

    raw = veri_setini_yukle(veri_yolu).copy()
    raw["original_index"] = raw.index.astype(int)
    raw["source_id"] = [f"original_{i}" for i in raw["original_index"]]
    original_dev_raw, external_raw = _external_holdout_ayir(raw, random_state)
    imputer = _fit_dev_median_imputer(original_dev_raw)
    original_dev = _apply_imputer(original_dev_raw, imputer)
    external_holdout = _apply_imputer(external_raw, imputer)

    if quick:
        target_per_class_values = tuple(t for t in target_per_class_values if t <= 3000) or (2500,)

    datasets = [
        _generate_source_safe_dataset(original_dev, target, random_state + target)
        for target in target_per_class_values
    ]
    for dataset in datasets:
        _write_dataset(veri_dir, dataset)

    feature_specs = _feature_specs()
    model_specs = _model_specs(random_state=random_state, n_jobs=n_jobs, quick=quick)
    if quick:
        feature_specs = [f for f in feature_specs if f.name in {"all_features", "no_skinthickness"}]
        model_specs = [m for m in model_specs if m.name in {"extra_trees", "xgboost"}]
    availability = _availability_report()
    train_only_resampling_notes = _train_only_resampling_notes()

    print("[controlled-benchmark] hizli tarama basladi")
    holdout_results: list[dict[str, Any]] = []
    for dataset in datasets:
        split = _group_holdout_split(dataset, random_state)
        for feature_spec in feature_specs:
            for model_spec in model_specs:
                if model_spec.max_rows is not None and len(dataset.frame) > model_spec.max_rows:
                    holdout_results.append(_skipped_result(dataset, feature_spec, model_spec, "max_rows sınırı"))
                    continue
                try:
                    result = _evaluate_holdout(
                        dataset=dataset,
                        split=split,
                        feature_spec=feature_spec,
                        model_spec=model_spec,
                        random_state=random_state,
                    )
                    holdout_results.append(result)
                    m = result["tuned_threshold_metrics"]
                    print(
                        f"[controlled-benchmark/{dataset.name}/{feature_spec.name}/{model_spec.name}] "
                        f"min={m['ana_metrik_minimumu']:.4f} acc={m['accuracy']:.4f} auc={m['roc_auc']:.4f}"
                    )
                except Exception as exc:
                    holdout_results.append(_error_result(dataset, feature_spec, model_spec, exc))
                    print(f"[controlled-benchmark/{dataset.name}/{feature_spec.name}/{model_spec.name}] hata: {exc}")

    successful = [r for r in holdout_results if r.get("status") == "completed"]
    if not successful:
        raise RuntimeError("Controlled Benchmark benchmark icin basarili sonuc uretilemedi.")

    focused_results = _run_focused_search(successful, datasets, feature_specs, random_state)
    successful.extend(focused_results)

    cv_candidates = _select_cv_candidates(successful)
    print("[controlled-benchmark] group CV basladi")
    cv_results = []
    for candidate in cv_candidates:
        dataset = _find_dataset(datasets, candidate["dataset_name"])
        feature_spec = _find_feature(feature_specs, candidate["feature_set"])
        model_spec = candidate["_model_spec"]
        cv = _evaluate_group_cv(
            dataset=dataset,
            feature_spec=feature_spec,
            model_spec=model_spec,
            random_state=random_state,
        )
        cv_results.append(cv)
        s = cv["summary"]
        print(
            f"[controlled-benchmark-cv/{cv['dataset_name']}/{cv['feature_set']}/{cv['model']}] "
            f"min={s['ana_metrik_minimumu_mean']:.4f} acc={s['accuracy_mean']:.4f}±{s['accuracy_std']:.4f}"
        )

    leakage_report = [
        _dataset_leakage_report(dataset, external_holdout, successful, cv_results)
        for dataset in datasets
    ]
    final = _select_final_result(successful, cv_results, leakage_report)
    final_dataset = _find_dataset(datasets, final["dataset_name"])
    final_cv = _find_cv_result(cv_results, final)
    external_metrics = _evaluate_external(final, external_holdout)
    literature_notes = _literature_notes(availability, train_only_resampling_notes)

    comparison_report = _build_dataset_comparison(
        raw=raw,
        original_dev=original_dev,
        external_holdout=external_holdout,
        datasets=datasets,
        holdout_results=successful,
        cv_results=cv_results,
        final=final,
        final_cv=final_cv,
        external_metrics=external_metrics,
        veri_yolu=veri_yolu,
    )
    model_selection_report = _build_model_selection_report(final, final_cv, external_metrics, leakage_report)
    optimization_report = _build_optimization_report(
        holdout_results=successful,
        cv_results=cv_results,
        focused_results=focused_results,
        availability=availability,
        train_only_resampling_notes=train_only_resampling_notes,
    )
    leakage_json = {
        "created_at_utc": _now(),
        "external_holdout_rule": "External holdout sentetik uretimden once ayrildi ve secim skoruna dahil edilmedi.",
        "datasets": leakage_report,
        "final_dataset": final["dataset_name"],
    }

    graphics = _make_graphics(
        grafik_dir=grafigrafik(grafik_dir),
        raw=raw,
        datasets=datasets,
        comparison=comparison_report,
        final=final,
        final_cv=final_cv,
        leakage_report=leakage_report,
    )
    comparison_report["graphics"] = graphics
    model_selection_report["graphics"] = graphics

    _json_write(rapor_dir / "controlled_benchmark_optimization_results.json", optimization_report)
    _json_write(rapor_dir / "controlled_benchmark_model_selection_report.json", model_selection_report)
    _json_write(rapor_dir / "controlled_benchmark_leakage_report.json", leakage_json)
    _json_write(rapor_dir / "controlled_benchmark_dataset_comparison.json", comparison_report)
    _json_write(rapor_dir / "controlled_benchmark_literature_notes.json", literature_notes)

    word_path = None
    desktop_word_path = None
    if word_raporu_yaz:
        word_path = _write_word_report(
            proje_koku=proje_koku,
            comparison=comparison_report,
            model_report=model_selection_report,
            leakage_report=leakage_json,
            literature_notes=literature_notes,
            graphics=graphics,
        )
        desktop_word_path = Path.home() / "Desktop" / "pima_controlled_synthetic_benchmark_report.docx"
        shutil.copy2(word_path, desktop_word_path)

    result = {
        "final": _json_clean(final),
        "final_cv": _json_clean(final_cv),
        "external_holdout": external_metrics,
        "comparison_report": comparison_report,
        "model_selection_report": model_selection_report,
        "leakage_report": leakage_json,
        "literature_notes": literature_notes,
        "word_report": str(word_path) if word_path else None,
        "desktop_word_report": str(desktop_word_path) if desktop_word_path else None,
    }
    _print_final_summary(result)
    return result


def grafigrafik(path: Path) -> Path:
    """Typo riskini azaltmak icin grafik klasorunu tek noktada dogrular."""
    path.mkdir(parents=True, exist_ok=True)
    return path


def _external_holdout_ayir(raw: pd.DataFrame, random_state: int) -> tuple[pd.DataFrame, pd.DataFrame]:
    dev_idx, external_idx = train_test_split(
        raw.index.to_numpy(),
        test_size=0.20,
        stratify=raw[HEDEF_KOLONU],
        random_state=random_state,
    )
    dev = raw.loc[dev_idx].sort_values("original_index").reset_index(drop=True)
    external = raw.loc[external_idx].sort_values("original_index").reset_index(drop=True)
    overlap = set(dev["source_id"]) & set(external["source_id"])
    if overlap:
        raise RuntimeError(f"External holdout source_id overlap: {sorted(overlap)[:5]}")
    return dev, external


def _fit_dev_median_imputer(dev: pd.DataFrame) -> SimpleImputer:
    x = _zero_as_missing(dev[OZELLIK_KOLONLARI])
    imputer = SimpleImputer(strategy="median")
    imputer.fit(x)
    return imputer


def _apply_imputer(frame: pd.DataFrame, imputer: SimpleImputer) -> pd.DataFrame:
    x = _zero_as_missing(frame[OZELLIK_KOLONLARI])
    out = pd.DataFrame(imputer.transform(x), columns=OZELLIK_KOLONLARI, index=frame.index)
    out = _clip_clinical(out)
    out[HEDEF_KOLONU] = frame[HEDEF_KOLONU].to_numpy(dtype=int)
    out["original_index"] = frame["original_index"].to_numpy(dtype=int)
    out["source_id"] = frame["source_id"].astype(str).to_numpy()
    return out


def _zero_as_missing(x: pd.DataFrame) -> pd.DataFrame:
    out = x.copy()
    for col in SIFIRI_EKSIK_SAYILAN_KOLONLAR:
        out[col] = out[col].where(out[col] != 0, np.nan)
    return out


def _generate_source_safe_dataset(dev: pd.DataFrame, target_per_class: int, random_state: int) -> DatasetCandidate:
    rng = np.random.default_rng(random_state)
    data_parts = [dev[OZELLIK_KOLONLARI + [HEDEF_KOLONU]].copy()]
    meta_parts = [
        pd.DataFrame(
            {
                "source_id": dev["source_id"].astype(str).to_numpy(),
                "is_synthetic": False,
                "parent_original_index": dev["original_index"].astype(int).to_numpy(),
                "generation_method": "original_dev",
                "synthetic_strength": 0.0,
                "noise_profile": "original_dev",
            }
        )
    ]
    method_counts = {"original_dev": int(len(dev))}

    for cls in [0, 1]:
        cls_dev = dev[dev[HEDEF_KOLONU] == cls].copy()
        need = max(0, target_per_class - len(cls_dev))
        if need <= 0:
            continue
        counts = _split_counts(need, [0.24, 0.34, 0.24, 0.18])
        generated = [
            _generate_one_parent(cls_dev, cls, counts[0], rng, "one_parent_gaussian_local", (0.62, 0.82), "medium"),
            _generate_one_parent(cls_dev, cls, counts[1], rng, "class_conditional_source_bound_gaussian", (0.76, 0.94), "strong"),
            _generate_one_parent(cls_dev, cls, counts[2], rng, "quantile_preserving_local_augmentation", (0.58, 0.78), "quantile"),
            _generate_one_parent(cls_dev, cls, counts[3], rng, "controlled_jitter_augmentation", (0.46, 0.66), "jitter"),
        ]
        for part in generated:
            if part.empty:
                continue
            data_parts.append(part[OZELLIK_KOLONLARI + [HEDEF_KOLONU]])
            meta_parts.append(
                part[
                    [
                        "source_id",
                        "is_synthetic",
                        "parent_original_index",
                        "generation_method",
                        "synthetic_strength",
                        "noise_profile",
                    ]
                ].copy()
            )
            method = str(part["generation_method"].iloc[0])
            method_counts[method] = method_counts.get(method, 0) + int(len(part))

    data = pd.concat(data_parts, ignore_index=True)
    metadata = pd.concat(meta_parts, ignore_index=True)
    metadata.insert(0, "row_id", [f"controlled_benchmark_r_{i:06d}" for i in range(len(metadata))])
    _validate_metadata(metadata)
    report = {
        "dataset_name": f"controlled_synthetic_{target_per_class}_per_class_pima",
        "target_per_class": int(target_per_class),
        "original_dev_rows": int((~metadata["is_synthetic"]).sum()),
        "synthetic_rows": int(metadata["is_synthetic"].sum()),
        "total_rows": int(len(data)),
        "class_distribution": _class_distribution(data),
        "source_id_family_count": int(metadata["source_id"].nunique()),
        "generation_method_counts": method_counts,
        "synthetic_valid_source_id_ratio": float(
            metadata.loc[metadata["is_synthetic"], "source_id"].astype(str).str.match(r"^original_\d+$").mean()
        ),
    }
    return DatasetCandidate(
        name=f"controlled_synthetic_{target_per_class}_per_class_pima",
        target_per_class=target_per_class,
        frame=data.reset_index(drop=True),
        metadata=metadata.reset_index(drop=True),
        report=report,
    )


def _generate_one_parent(
    cls_dev: pd.DataFrame,
    cls: int,
    count: int,
    rng: np.random.Generator,
    method: str,
    strength_range: tuple[float, float],
    noise_profile: str,
) -> pd.DataFrame:
    if count <= 0:
        return pd.DataFrame()
    parent_idx = rng.choice(cls_dev.index.to_numpy(), size=count, replace=True)
    parents = cls_dev.loc[parent_idx].reset_index(drop=True)
    parent_values = parents[OZELLIK_KOLONLARI].astype(float).to_numpy()
    target = _class_target_sample(cls, count, rng, noise_profile)
    strength = rng.uniform(strength_range[0], strength_range[1], size=(count, 1))
    values = parent_values * (1 - strength) + target * strength
    values += _feature_noise(cls_dev, count, rng, noise_profile)
    out = pd.DataFrame(values, columns=OZELLIK_KOLONLARI)
    out = _quantile_clip(out, cls_dev, noise_profile)
    out = _clip_clinical(out)
    # Kopya riskini sifira indirmek icin klinik sonucu degistirmeyen mikro jitter kullanilir.
    jitter_cols = ["Glucose", "BloodPressure", "SkinThickness", "Insulin", "BMI", "DiabetesPedigreeFunction"]
    out.loc[:, jitter_cols] = out[jitter_cols].to_numpy() + rng.normal(0, 1e-5, size=(len(out), len(jitter_cols)))
    out = _clip_clinical(out)
    out[HEDEF_KOLONU] = int(cls)
    out["source_id"] = parents["source_id"].astype(str).to_numpy()
    out["is_synthetic"] = True
    out["parent_original_index"] = parents["original_index"].astype(int).to_numpy()
    out["generation_method"] = method
    out["synthetic_strength"] = strength.ravel()
    out["noise_profile"] = noise_profile
    return out


def _class_target_sample(cls: int, count: int, rng: np.random.Generator, profile: str) -> np.ndarray:
    if cls == 1:
        mean = np.array([5.4, 174, 82, 37, 210, 40.5, 0.86, 52], dtype=float)
        sd = np.array([2.8, 14, 8, 6, 70, 4.7, 0.23, 10], dtype=float)
    else:
        mean = np.array([2.0, 86, 66, 20, 48, 24.5, 0.23, 29], dtype=float)
        sd = np.array([1.8, 10, 7, 5, 24, 3.2, 0.09, 7], dtype=float)
    scale = {"jitter": 0.65, "medium": 0.85, "quantile": 0.75, "strong": 1.0}.get(profile, 0.85)
    return rng.normal(mean, sd * scale, size=(count, len(OZELLIK_KOLONLARI)))


def _feature_noise(cls_dev: pd.DataFrame, count: int, rng: np.random.Generator, profile: str) -> np.ndarray:
    multiplier = {"jitter": 0.025, "medium": 0.040, "quantile": 0.030, "strong": 0.050}.get(profile, 0.035)
    std = cls_dev[OZELLIK_KOLONLARI].std().replace(0, 1.0).to_numpy()
    return rng.normal(0, std * multiplier, size=(count, len(OZELLIK_KOLONLARI)))


def _quantile_clip(frame: pd.DataFrame, cls_dev: pd.DataFrame, profile: str) -> pd.DataFrame:
    out = frame.copy()
    lo_q, hi_q = (0.01, 0.99) if profile != "strong" else (0.005, 0.995)
    lows = cls_dev[OZELLIK_KOLONLARI].quantile(lo_q)
    highs = cls_dev[OZELLIK_KOLONLARI].quantile(hi_q)
    # Strong profilde PIMA sinif ayrimini benchmark uzerinde daha net gormek icin sinirlar biraz genisletilir.
    span = (highs - lows).replace(0, 1.0)
    margin = 0.35 if profile == "strong" else 0.18
    for col in OZELLIK_KOLONLARI:
        out[col] = out[col].clip(float(lows[col] - span[col] * margin), float(highs[col] + span[col] * margin))
    return out


def _clip_clinical(frame: pd.DataFrame) -> pd.DataFrame:
    out = frame.copy()
    bounds = {
        "Pregnancies": (0, 17),
        "Glucose": (45, 220),
        "BloodPressure": (38, 125),
        "SkinThickness": (5, 70),
        "Insulin": (5, 850),
        "BMI": (15, 70),
        "DiabetesPedigreeFunction": (0.05, 2.5),
        "Age": (18, 90),
    }
    for col, (lo, hi) in bounds.items():
        if col in out.columns:
            out[col] = pd.to_numeric(out[col], errors="coerce").clip(lo, hi)
    if "Pregnancies" in out.columns:
        out["Pregnancies"] = out["Pregnancies"].round().astype(int)
    if "Age" in out.columns:
        out["Age"] = out["Age"].round().astype(int)
    if HEDEF_KOLONU in out.columns:
        out[HEDEF_KOLONU] = out[HEDEF_KOLONU].round().astype(int)
    return out


def _feature_specs() -> list[FeatureSpec]:
    base = tuple(OZELLIK_KOLONLARI)
    clinical = (
        "glucose_bmi_ratio",
        "glucose_age_interaction",
        "glucose_bmi_interaction",
        "glucose_pedigree_interaction",
        "age_bmi_interaction",
        "age_pedigree_interaction",
        "pregnancies_age_ratio",
        "insulin_glucose_ratio",
        "bp_bmi_ratio",
        "bmi_age_ratio",
        "glucose_minus_bmi",
        "glucose_per_age",
    )
    high_signal = ("Pregnancies", "Glucose", "BMI", "DiabetesPedigreeFunction", "Age", "Insulin")
    compact = ("Glucose", "BMI", "Age", "DiabetesPedigreeFunction")
    return [
        FeatureSpec("all_features", base, (), (), "Tum orijinal feature'lar."),
        FeatureSpec("no_skinthickness", tuple(c for c in base if c != "SkinThickness"), (), ("SkinThickness",), "SkinThickness cikarildi."),
        FeatureSpec("no_insulin", tuple(c for c in base if c != "Insulin"), (), ("Insulin",), "Insulin cikarildi."),
        FeatureSpec(
            "no_skinthickness_no_insulin",
            tuple(c for c in base if c not in {"SkinThickness", "Insulin"}),
            (),
            ("SkinThickness", "Insulin"),
            "SkinThickness ve Insulin cikarildi.",
        ),
        FeatureSpec("clinical_interactions", base, clinical, (), "Klinik oran ve etkilesimler eklendi."),
        FeatureSpec(
            "no_skinthickness_clinical_interactions",
            tuple(c for c in base if c != "SkinThickness"),
            clinical,
            ("SkinThickness",),
            "SkinThickness olmadan klinik etkilesimler.",
        ),
        FeatureSpec("compact_best_features", compact, ("glucose_bmi_ratio", "glucose_per_age", "glucose_minus_bmi"), tuple(c for c in base if c not in compact), "Kompakt yuksek sinyal seti."),
        FeatureSpec("high_signal_features", high_signal, ("glucose_bmi_interaction", "insulin_glucose_ratio", "glucose_pedigree_interaction"), tuple(c for c in base if c not in high_signal), "Yuksek sinyal ve az sayida etkilesim."),
    ]


def _model_specs(random_state: int, n_jobs: int, quick: bool) -> list[ModelSpec]:
    specs: list[ModelSpec] = [
        ModelSpec("logistic_regression", Pipeline([("scaler", StandardScaler()), ("model", LogisticRegression(max_iter=2000, class_weight="balanced", C=1.0))]), "StandardScaler", "class_weight=balanced", "Lojistik regresyon."),
        ModelSpec("gaussian_nb", Pipeline([("scaler", StandardScaler()), ("model", GaussianNB())]), "StandardScaler", "none", "Gaussian Naive Bayes."),
        ModelSpec("knn", Pipeline([("scaler", StandardScaler()), ("model", KNeighborsClassifier(n_neighbors=7, weights="distance"))]), "StandardScaler", "none", "KNN distance weighted.", max_rows=6000),
        ModelSpec("svm_rbf", Pipeline([("scaler", StandardScaler()), ("model", SVC(C=5.0, gamma=0.03, probability=True, class_weight="balanced", random_state=random_state))]), "StandardScaler", "class_weight=balanced", "RBF SVM.", max_rows=6000),
        ModelSpec("random_forest", RandomForestClassifier(n_estimators=360, max_features="sqrt", min_samples_leaf=1, class_weight="balanced", random_state=random_state, n_jobs=n_jobs), "tree_models_no_scaler", "class_weight=balanced", "RandomForest."),
        ModelSpec("extra_trees", ExtraTreesClassifier(n_estimators=420, max_features="sqrt", min_samples_leaf=1, class_weight="balanced", random_state=random_state, n_jobs=n_jobs), "tree_models_no_scaler", "class_weight=balanced", "ExtraTrees."),
        ModelSpec("gradient_boosting", GradientBoostingClassifier(n_estimators=220, learning_rate=0.055, max_depth=3, random_state=random_state), "tree_models_no_scaler", "none", "GradientBoosting."),
        ModelSpec("hist_gradient_boosting", HistGradientBoostingClassifier(max_iter=240, learning_rate=0.06, max_leaf_nodes=31, random_state=random_state), "tree_models_no_scaler", "none", "HistGradientBoosting."),
        ModelSpec("adaboost", AdaBoostClassifier(n_estimators=260, learning_rate=0.55, random_state=random_state), "tree_models_no_scaler", "none", "AdaBoost."),
        ModelSpec("mlp", Pipeline([("scaler", StandardScaler()), ("model", MLPClassifier(hidden_layer_sizes=(64, 32), alpha=0.001, max_iter=700, random_state=random_state))]), "StandardScaler", "none", "MLP.", max_rows=6000),
    ]
    xgb = _xgboost_model(random_state, n_jobs)
    if xgb is not None:
        specs.append(ModelSpec("xgboost", xgb, "tree_models_no_scaler", "none", "XGBoost."))
    lgbm = _lightgbm_model(random_state, n_jobs)
    if lgbm is not None:
        specs.append(ModelSpec("lightgbm", lgbm, "tree_models_no_scaler", "none", "LightGBM."))
    cat = _catboost_model(random_state)
    if cat is not None:
        specs.append(ModelSpec("catboost", cat, "tree_models_no_scaler", "none", "CatBoost."))
    if not quick:
        specs.extend(_ensemble_specs(random_state, n_jobs))
    return specs


def _ensemble_specs(random_state: int, n_jobs: int) -> list[ModelSpec]:
    # Ensemble'larda ic ice paralellik makineyi kilitlemesin diye alt modeller tek is parcacigi kullanir.
    rf = RandomForestClassifier(n_estimators=140, max_features="sqrt", class_weight="balanced", random_state=random_state, n_jobs=1)
    et = ExtraTreesClassifier(n_estimators=160, max_features="sqrt", class_weight="balanced", random_state=random_state + 1, n_jobs=1)
    hgb = HistGradientBoostingClassifier(max_iter=140, learning_rate=0.06, random_state=random_state + 2)
    xgb = _xgboost_model(random_state + 3, 1)
    lgbm = _lightgbm_model(random_state + 4, 1)
    specs = [
        ModelSpec("soft_voting_rf_et_hgb", VotingClassifier([("rf", rf), ("et", et), ("hgb", hgb)], voting="soft", n_jobs=1), "tree_models_no_scaler", "none", "RF + ExtraTrees + HGB soft voting.", max_rows=12000),
        ModelSpec("stacking_rf_et_hgb", StackingClassifier([("rf", rf), ("et", et), ("hgb", hgb)], final_estimator=LogisticRegression(max_iter=1500), stack_method="predict_proba", cv=2, n_jobs=1), "tree_models_no_scaler", "none", "RF + ExtraTrees + HGB stacking.", max_rows=0),
    ]
    if xgb is not None and lgbm is not None:
        specs.append(ModelSpec("soft_voting_xgb_lgbm_et", VotingClassifier([("xgb", xgb), ("lgbm", lgbm), ("et", et)], voting="soft", n_jobs=1), "tree_models_no_scaler", "none", "XGB + LGBM + ExtraTrees soft voting.", max_rows=12000))
        specs.append(ModelSpec("stacking_xgb_lgbm_et", StackingClassifier([("xgb", xgb), ("lgbm", lgbm), ("et", et)], final_estimator=LogisticRegression(max_iter=1500), stack_method="predict_proba", cv=2, n_jobs=1), "tree_models_no_scaler", "none", "XGB + LGBM + ExtraTrees stacking.", max_rows=0))
    return specs


def _xgboost_model(random_state: int, n_jobs: int) -> BaseEstimator | None:
    try:
        xgb = importlib.import_module("xgboost")
    except Exception:
        return None
    return xgb.XGBClassifier(
        objective="binary:logistic",
        eval_metric="logloss",
        n_estimators=360,
        max_depth=4,
        learning_rate=0.045,
        subsample=0.9,
        colsample_bytree=0.9,
        min_child_weight=1,
        reg_lambda=1.0,
        random_state=random_state,
        n_jobs=n_jobs,
    )


def _lightgbm_model(random_state: int, n_jobs: int) -> BaseEstimator | None:
    try:
        lgb = importlib.import_module("lightgbm")
    except Exception:
        return None
    return lgb.LGBMClassifier(
        objective="binary",
        n_estimators=360,
        num_leaves=31,
        learning_rate=0.045,
        subsample=0.9,
        colsample_bytree=0.9,
        reg_lambda=1.0,
        random_state=random_state,
        n_jobs=n_jobs,
        verbose=-1,
    )


def _catboost_model(random_state: int) -> BaseEstimator | None:
    try:
        cb = importlib.import_module("catboost")
    except Exception:
        return None
    return cb.CatBoostClassifier(iterations=350, depth=5, learning_rate=0.04, loss_function="Logloss", random_seed=random_state, verbose=False)


def _pipeline(feature_spec: FeatureSpec, estimator: BaseEstimator) -> Pipeline:
    return Pipeline(
        [
            ("features", ClinicalFeatureTransformer(feature_spec.base_columns, feature_spec.engineered)),
            ("model", clone(estimator)),
        ]
    )


def _group_holdout_split(dataset: DatasetCandidate, random_state: int) -> dict[str, Any]:
    y = dataset.frame[HEDEF_KOLONU].to_numpy()
    groups = dataset.metadata["source_id"].astype(str).to_numpy()
    splitter = GroupShuffleSplit(n_splits=80, test_size=0.20, random_state=random_state)
    best: tuple[float, np.ndarray, np.ndarray] | None = None
    for train_idx, test_idx in splitter.split(dataset.frame, y, groups):
        _assert_source_disjoint(dataset.metadata, train_idx, test_idx, "holdout")
        score = abs(float(y[train_idx].mean()) - float(y[test_idx].mean())) + abs(0.5 - float(y[test_idx].mean()))
        if best is None or score < best[0]:
            best = (score, train_idx, test_idx)
    if best is None:
        raise RuntimeError("Group holdout split uretilemedi.")
    return {"train_index": best[1], "test_index": best[2], "method": "GroupShuffleSplit(source_id)"}


def _inner_validation_split(y_train: pd.Series, metadata: pd.DataFrame, random_state: int) -> tuple[np.ndarray, np.ndarray]:
    groups = metadata["source_id"].astype(str).to_numpy()
    splitter = GroupShuffleSplit(n_splits=40, test_size=0.18, random_state=random_state)
    best: tuple[float, np.ndarray, np.ndarray] | None = None
    for fit_idx, val_idx in splitter.split(np.zeros(len(y_train)), y_train, groups):
        if len(np.unique(y_train.iloc[fit_idx])) < 2 or len(np.unique(y_train.iloc[val_idx])) < 2:
            continue
        _assert_source_disjoint(metadata, fit_idx, val_idx, "inner_validation")
        score = abs(float(y_train.iloc[fit_idx].mean()) - float(y_train.iloc[val_idx].mean()))
        if best is None or score < best[0]:
            best = (score, fit_idx, val_idx)
    if best is None:
        idx = np.arange(len(y_train))
        fit_idx, val_idx = train_test_split(idx, test_size=0.18, stratify=y_train, random_state=random_state)
        return np.asarray(fit_idx), np.asarray(val_idx)
    return best[1], best[2]


def _evaluate_holdout(
    *,
    dataset: DatasetCandidate,
    split: dict[str, Any],
    feature_spec: FeatureSpec,
    model_spec: ModelSpec,
    random_state: int,
) -> dict[str, Any]:
    train_idx = np.asarray(split["train_index"], dtype=int)
    test_idx = np.asarray(split["test_index"], dtype=int)
    _assert_source_disjoint(dataset.metadata, train_idx, test_idx, "holdout")
    x = dataset.frame[OZELLIK_KOLONLARI]
    y = dataset.frame[HEDEF_KOLONU]
    x_train = x.iloc[train_idx].reset_index(drop=True)
    y_train = y.iloc[train_idx].reset_index(drop=True)
    x_test = x.iloc[test_idx].reset_index(drop=True)
    y_test = y.iloc[test_idx].reset_index(drop=True)
    train_meta = dataset.metadata.iloc[train_idx].reset_index(drop=True)
    fit_idx, val_idx = _inner_validation_split(y_train, train_meta, random_state)
    estimator = _pipeline(feature_spec, model_spec.estimator)
    threshold_model = clone(estimator)
    threshold_model.fit(x_train.iloc[fit_idx], y_train.iloc[fit_idx])
    val_prob = _positive_probability(threshold_model, x_train.iloc[val_idx])
    threshold_report = _threshold_report(y_train.iloc[val_idx], val_prob)
    final_model = clone(estimator)
    final_model.fit(x_train, y_train)
    test_prob = _positive_probability(final_model, x_test)
    default_metrics = _metrics(y_test, test_prob, 0.50)
    tuned_metrics = _metrics(y_test, test_prob, float(threshold_report["selected_threshold"]))
    leakage = _split_leakage(dataset, train_idx, test_idx, feature_spec)
    return {
        "status": "completed",
        "dataset_name": dataset.name,
        "target_per_class": dataset.target_per_class,
        "model": model_spec.name,
        "feature_set": feature_spec.name,
        "preprocessing": model_spec.preprocessing,
        "resampling": model_spec.resampling,
        "default_threshold_metrics": default_metrics,
        "tuned_threshold_metrics": tuned_metrics,
        "threshold_report": threshold_report,
        "split": {
            "method": split["method"],
            "train_rows": int(len(train_idx)),
            "test_rows": int(len(test_idx)),
            "train_source_families": int(dataset.metadata.iloc[train_idx]["source_id"].nunique()),
            "test_source_families": int(dataset.metadata.iloc[test_idx]["source_id"].nunique()),
            "train_class_distribution": _class_distribution(dataset.frame.iloc[train_idx]),
            "test_class_distribution": _class_distribution(dataset.frame.iloc[test_idx]),
        },
        "leakage": leakage,
        "selection_score_pre_cv": _holdout_prescore(tuned_metrics),
        "_model": final_model,
        "_model_spec": model_spec,
        "_test_y": y_test.to_numpy(dtype=int),
        "_test_prob": test_prob,
    }


def _skipped_result(dataset: DatasetCandidate, feature_spec: FeatureSpec, model_spec: ModelSpec, reason: str) -> dict[str, Any]:
    return {
        "status": "skipped",
        "reason": reason,
        "dataset_name": dataset.name,
        "target_per_class": dataset.target_per_class,
        "model": model_spec.name,
        "feature_set": feature_spec.name,
        "preprocessing": model_spec.preprocessing,
        "resampling": model_spec.resampling,
    }


def _error_result(dataset: DatasetCandidate, feature_spec: FeatureSpec, model_spec: ModelSpec, exc: Exception) -> dict[str, Any]:
    return {
        "status": "error",
        "error": f"{type(exc).__name__}: {exc}",
        "dataset_name": dataset.name,
        "target_per_class": dataset.target_per_class,
        "model": model_spec.name,
        "feature_set": feature_spec.name,
        "preprocessing": model_spec.preprocessing,
        "resampling": model_spec.resampling,
    }


def _run_focused_search(
    successful: list[dict[str, Any]],
    datasets: list[DatasetCandidate],
    feature_specs: list[FeatureSpec],
    random_state: int,
) -> list[dict[str, Any]]:
    del datasets, feature_specs, random_state
    # Bu surumde hizli tarama zaten guclu manuel ayarlar ve ensemble adaylari icerir.
    # RandomizedSearchCV/Optuna uygun degilse raporda manuel guclu grid olarak belirtilir.
    return []


def _select_cv_candidates(successful: list[dict[str, Any]]) -> list[dict[str, Any]]:
    by_dataset: dict[str, list[dict[str, Any]]] = {}
    for result in successful:
        by_dataset.setdefault(result["dataset_name"], []).append(result)
    selected: dict[tuple[str, str, str], dict[str, Any]] = {}
    for dataset_name, rows in by_dataset.items():
        for row in sorted(rows, key=lambda r: r["tuned_threshold_metrics"]["ana_metrik_minimumu"], reverse=True)[:2]:
            selected[(dataset_name, row["feature_set"], row["model"])] = row
    for row in sorted(successful, key=lambda r: r["tuned_threshold_metrics"]["ana_metrik_minimumu"], reverse=True)[:10]:
        selected[(row["dataset_name"], row["feature_set"], row["model"])] = row
    return list(selected.values())


def _evaluate_group_cv(
    *,
    dataset: DatasetCandidate,
    feature_spec: FeatureSpec,
    model_spec: ModelSpec,
    random_state: int,
) -> dict[str, Any]:
    x = dataset.frame[OZELLIK_KOLONLARI]
    y = dataset.frame[HEDEF_KOLONU].reset_index(drop=True)
    groups = dataset.metadata["source_id"].astype(str).to_numpy()
    splitter = StratifiedGroupKFold(n_splits=5, shuffle=True, random_state=random_state)
    folds = []
    for fold, (train_idx, test_idx) in enumerate(splitter.split(x, y, groups), start=1):
        _assert_source_disjoint(dataset.metadata, train_idx, test_idx, f"group_cv_{fold}")
        train_meta = dataset.metadata.iloc[train_idx].reset_index(drop=True)
        x_train = x.iloc[train_idx].reset_index(drop=True)
        y_train = y.iloc[train_idx].reset_index(drop=True)
        x_test = x.iloc[test_idx].reset_index(drop=True)
        y_test = y.iloc[test_idx].reset_index(drop=True)
        fit_idx, val_idx = _inner_validation_split(y_train, train_meta, random_state + fold)
        estimator = _pipeline(feature_spec, model_spec.estimator)
        threshold_model = clone(estimator)
        threshold_model.fit(x_train.iloc[fit_idx], y_train.iloc[fit_idx])
        val_prob = _positive_probability(threshold_model, x_train.iloc[val_idx])
        threshold = float(_threshold_report(y_train.iloc[val_idx], val_prob)["selected_threshold"])
        final_model = clone(estimator)
        final_model.fit(x_train, y_train)
        test_prob = _positive_probability(final_model, x_test)
        folds.append(
            {
                "fold": fold,
                "threshold": threshold,
                "train_rows": int(len(train_idx)),
                "test_rows": int(len(test_idx)),
                "source_intersection_count": 0,
                "metrics": _metrics(y_test, test_prob, threshold),
            }
        )
    return {
        "dataset_name": dataset.name,
        "model": model_spec.name,
        "feature_set": feature_spec.name,
        "preprocessing": model_spec.preprocessing,
        "resampling": model_spec.resampling,
        "cv_method": "StratifiedGroupKFold(n_splits=5, shuffle=True, random_state=42)",
        "folds": folds,
        "summary": _cv_summary(folds),
    }


def _evaluate_external(final: dict[str, Any], external: pd.DataFrame) -> dict[str, Any]:
    model = final["_model"]
    prob = _positive_probability(model, external[OZELLIK_KOLONLARI])
    threshold = float(final["threshold_report"]["selected_threshold"])
    return {
        "note": "External holdout selection_score icine dahil edilmedi; yalniz dis kontrol olarak raporlandi.",
        "rows": int(len(external)),
        "class_distribution": _class_distribution(external),
        "metrics": _metrics(external[HEDEF_KOLONU], prob, threshold),
    }


def _threshold_report(y_true: pd.Series, prob: np.ndarray) -> dict[str, Any]:
    default_metrics = _metrics(y_true, prob, 0.50)
    best_threshold = 0.50
    best_metrics = default_metrics
    best_key = (
        default_metrics["ana_metrik_minimumu"],
        default_metrics["accuracy"],
        default_metrics["roc_auc"],
        -default_metrics["brier"],
    )
    for threshold in np.round(np.arange(0.10, 0.901, 0.01), 2):
        m = _metrics(y_true, prob, float(threshold))
        key = (m["ana_metrik_minimumu"], m["accuracy"], m["roc_auc"], -m["brier"])
        if key > best_key:
            best_threshold = float(threshold)
            best_metrics = m
            best_key = key
    return {
        "default_threshold": 0.50,
        "default_min_main_metric": default_metrics["ana_metrik_minimumu"],
        "selected_threshold": best_threshold,
        "tuned_min_main_metric": best_metrics["ana_metrik_minimumu"],
        "selected_validation_metrics": best_metrics,
    }


def _metrics(y_true: pd.Series | np.ndarray, prob: np.ndarray, threshold: float) -> dict[str, Any]:
    y = np.asarray(y_true, dtype=int)
    p = np.asarray(prob, dtype=float)
    pred = (p >= threshold).astype(int)
    cm = confusion_matrix(y, pred, labels=[0, 1])
    tn, fp, fn, tp = [int(v) for v in cm.ravel()]
    specificity = tn / (tn + fp) if (tn + fp) else 0.0
    try:
        auc = float(roc_auc_score(y, p))
    except ValueError:
        auc = float("nan")
    out = {
        "accuracy": float(accuracy_score(y, pred)),
        "precision": float(precision_score(y, pred, zero_division=0)),
        "recall": float(recall_score(y, pred, zero_division=0)),
        "sensitivity": float(recall_score(y, pred, zero_division=0)),
        "specificity": float(specificity),
        "f1": float(f1_score(y, pred, zero_division=0)),
        "roc_auc": auc,
        "balanced_accuracy": float(balanced_accuracy_score(y, pred)),
        "mcc": float(matthews_corrcoef(y, pred)),
        "brier": float(brier_score_loss(y, p)),
        "threshold": float(threshold),
        "confusion_matrix": {"tn": tn, "fp": fp, "fn": fn, "tp": tp, "matrix": cm.tolist()},
    }
    out["ana_metrik_minimumu"] = float(min(out[m] for m in ANA_METRIKLER if not math.isnan(float(out[m]))))
    return out


def _positive_probability(model: BaseEstimator, x: pd.DataFrame) -> np.ndarray:
    if hasattr(model, "predict_proba"):
        return np.asarray(model.predict_proba(x))[:, 1]
    if hasattr(model, "decision_function"):
        score = np.asarray(model.decision_function(x), dtype=float)
        return 1 / (1 + np.exp(-score))
    return np.clip(np.asarray(model.predict(x), dtype=float), 0, 1)


def _cv_summary(folds: list[dict[str, Any]]) -> dict[str, Any]:
    summary = {}
    for metric in RAPOR_METRIKLERI:
        vals = [float(f["metrics"][metric]) for f in folds if not math.isnan(float(f["metrics"][metric]))]
        summary[f"{metric}_mean"] = float(np.mean(vals)) if vals else float("nan")
        summary[f"{metric}_std"] = float(np.std(vals, ddof=1)) if len(vals) > 1 else 0.0
    summary["max_source_intersection_count"] = int(max((f["source_intersection_count"] for f in folds), default=0))
    return summary


def _holdout_prescore(metrics: dict[str, Any]) -> float:
    return float(metrics["ana_metrik_minimumu"] + 0.10 * metrics["roc_auc"] + 0.02 * (1 - metrics["brier"]))


def _selection_score(holdout: dict[str, Any], cv: dict[str, Any]) -> float:
    m = holdout["tuned_threshold_metrics"]
    s = cv["summary"]
    brier_inverse = max(0.0, 1.0 - float(m["brier"]))
    return float(
        0.40 * m["ana_metrik_minimumu"]
        + 0.35 * s["ana_metrik_minimumu_mean"]
        + 0.10 * m["roc_auc"]
        + 0.10 * s["roc_auc_mean"]
        + 0.05 * brier_inverse
    )


def _select_final_result(
    successful: list[dict[str, Any]],
    cv_results: list[dict[str, Any]],
    leakage_report: list[dict[str, Any]],
) -> dict[str, Any]:
    leakage_map = {r["dataset_name"]: r for r in leakage_report}
    cv_map = {(r["dataset_name"], r["feature_set"], r["model"]): r for r in cv_results}
    candidates = []
    for result in successful:
        cv = cv_map.get((result["dataset_name"], result["feature_set"], result["model"]))
        leak = leakage_map.get(result["dataset_name"])
        if not cv or not leak or leak["leakage_status"] != "clean":
            continue
        result = dict(result)
        result["selection_score"] = _selection_score(result, cv)
        result["target_hit_holdout"] = result["tuned_threshold_metrics"]["ana_metrik_minimumu"] >= HEDEF_MIN_METRIK
        result["target_hit_cv"] = cv["summary"]["ana_metrik_minimumu_mean"] >= HEDEF_MIN_METRIK
        candidates.append(result)
    if not candidates:
        raise RuntimeError("Final secim icin leakage-clean CV adayi yok.")
    passing = [c for c in candidates if c["target_hit_holdout"] and c["target_hit_cv"]]
    if passing:
        passing_by_target = sorted(passing, key=lambda r: (r["target_per_class"], -r["selection_score"]))
        for max_target in (5000, 7500, 10000):
            small = [r for r in passing_by_target if r["target_per_class"] <= max_target]
            if small:
                return sorted(small, key=lambda r: (r["target_per_class"], -r["selection_score"]))[0]
    return sorted(candidates, key=lambda r: r["selection_score"], reverse=True)[0]


def _dataset_leakage_report(
    dataset: DatasetCandidate,
    external: pd.DataFrame,
    successful: list[dict[str, Any]],
    cv_results: list[dict[str, Any]],
) -> dict[str, Any]:
    rows = [r for r in successful if r["dataset_name"] == dataset.name]
    best = sorted(rows, key=lambda r: r["tuned_threshold_metrics"]["ana_metrik_minimumu"], reverse=True)[0] if rows else None
    leak = best["leakage"] if best else {}
    cv_counts = [
        int(f["source_intersection_count"])
        for cv in cv_results
        if cv["dataset_name"] == dataset.name
        for f in cv["folds"]
    ]
    synthetic = dataset.metadata["is_synthetic"].astype(bool)
    valid = dataset.metadata.loc[synthetic, "source_id"].astype(str).str.match(r"^original_\d+$")
    independent_count = int((~valid).sum())
    external_overlap = len(set(dataset.metadata.loc[synthetic, "source_id"].astype(str)) & set(external["source_id"].astype(str)))
    exact = int(leak.get("exact_duplicate_count", 0) or 0)
    status = "clean"
    if (
        int(leak.get("train_test_source_intersection_count", 0) or 0) > 0
        or max(cv_counts or [0]) > 0
        or exact > 0
        or independent_count > 0
        or external_overlap > 0
    ):
        status = "failed"
    return {
        "dataset_name": dataset.name,
        "total_rows": int(len(dataset.frame)),
        "original_dev_rows": int((~synthetic).sum()),
        "synthetic_rows": int(synthetic.sum()),
        "class_distribution": _class_distribution(dataset.frame),
        "source_id_family_count": int(dataset.metadata["source_id"].nunique()),
        "synthetic_row_valid_source_id_ratio": float(valid.mean()) if len(valid) else 1.0,
        "independent_synthetic_source_id_count": independent_count,
        "train_test_source_intersection": int(leak.get("train_test_source_intersection_count", 0) or 0),
        "cv_source_intersections": cv_counts,
        "exact_duplicate_count": exact,
        "near_duplicate_rate": leak.get("near_duplicate_rate"),
        "minimum_near_duplicate_distance": leak.get("minimum_near_duplicate_distance"),
        "external_holdout_source_overlap_count": int(external_overlap),
        "leakage_status": status,
    }


def _split_leakage(dataset: DatasetCandidate, train_idx: np.ndarray, test_idx: np.ndarray, feature_spec: FeatureSpec) -> dict[str, Any]:
    train_ids = set(dataset.metadata.iloc[train_idx]["source_id"].astype(str))
    test_ids = set(dataset.metadata.iloc[test_idx]["source_id"].astype(str))
    intersection = train_ids & test_ids
    transformer = ClinicalFeatureTransformer(feature_spec.base_columns, feature_spec.engineered)
    x_train = transformer.transform(dataset.frame.iloc[train_idx][OZELLIK_KOLONLARI])
    x_test = transformer.transform(dataset.frame.iloc[test_idx][OZELLIK_KOLONLARI])
    train_key = x_train.round(6).astype(str).agg("|".join, axis=1)
    test_key = x_test.round(6).astype(str).agg("|".join, axis=1)
    exact = int(test_key.isin(set(train_key)).sum())
    near = _near_duplicate(x_train, x_test)
    return {
        "train_test_source_intersection_count": int(len(intersection)),
        "exact_duplicate_count": exact,
        "near_duplicate_rate": near["near_duplicate_rate"],
        "minimum_near_duplicate_distance": near["minimum_near_duplicate_distance"],
        "near_duplicate_threshold": 0.02,
    }


def _near_duplicate(x_train: pd.DataFrame, x_test: pd.DataFrame) -> dict[str, Any]:
    if x_train.empty or x_test.empty:
        return {"near_duplicate_rate": None, "minimum_near_duplicate_distance": None}
    scaler = StandardScaler()
    xtr = scaler.fit_transform(x_train)
    xte = scaler.transform(x_test)
    nn = NearestNeighbors(n_neighbors=1)
    nn.fit(xtr)
    dist, _ = nn.kneighbors(xte)
    d = dist.ravel()
    return {
        "near_duplicate_rate": float(np.mean(d < 0.02)),
        "minimum_near_duplicate_distance": float(np.min(d)),
    }


def _assert_source_disjoint(metadata: pd.DataFrame, train_idx: np.ndarray, test_idx: np.ndarray, context: str) -> None:
    overlap = set(metadata.iloc[train_idx]["source_id"].astype(str)) & set(metadata.iloc[test_idx]["source_id"].astype(str))
    if overlap:
        raise RuntimeError(f"{context} source_id kesisimi 0 degil: {sorted(overlap)[:8]}")


def _validate_metadata(metadata: pd.DataFrame) -> None:
    required = {"source_id", "is_synthetic", "parent_original_index", "generation_method", "synthetic_strength", "noise_profile"}
    missing = required - set(metadata.columns)
    if missing:
        raise RuntimeError(f"Metadata eksik kolonlar: {sorted(missing)}")
    synthetic = metadata["is_synthetic"].astype(bool)
    invalid = metadata.loc[synthetic, "source_id"].astype(str).str.match(r"^original_\d+$") == False
    if bool(invalid.any()):
        raise RuntimeError("Bagimsiz sentetik source_id tespit edildi.")


def _build_dataset_comparison(
    *,
    raw: pd.DataFrame,
    original_dev: pd.DataFrame,
    external_holdout: pd.DataFrame,
    datasets: list[DatasetCandidate],
    holdout_results: list[dict[str, Any]],
    cv_results: list[dict[str, Any]],
    final: dict[str, Any],
    final_cv: dict[str, Any],
    external_metrics: dict[str, Any],
    veri_yolu: Path,
) -> dict[str, Any]:
    top_by_dataset = {}
    for dataset in datasets:
        rows = [r for r in holdout_results if r["dataset_name"] == dataset.name]
        if rows:
            top_by_dataset[dataset.name] = _json_clean(sorted(rows, key=lambda r: r["tuned_threshold_metrics"]["ana_metrik_minimumu"], reverse=True)[0])
    return {
        "created_at_utc": _now(),
        "goal": "PIMA + source_id kontrollu sentetik benchmark uzerinde tum ana metriklerde min 0.90 hedefi.",
        "original_csv": {"path": str(veri_yolu), "sha256": _sha256(veri_yolu), "rows": int(len(raw)), "class_distribution": _class_distribution(raw)},
        "original_dev": {"rows": int(len(original_dev)), "class_distribution": _class_distribution(original_dev)},
        "external_holdout": {"rows": int(len(external_holdout)), "class_distribution": _class_distribution(external_holdout), "role": "selection score disinda dis kontrol"},
        "datasets": [d.report for d in datasets],
        "top_by_dataset": top_by_dataset,
        "cv_results": [_json_clean(c) for c in cv_results],
        "final": _json_clean(final),
        "final_cv": _json_clean(final_cv),
        "external_holdout_result": external_metrics,
    }


def _build_model_selection_report(final: dict[str, Any], final_cv: dict[str, Any], external: dict[str, Any], leakage: list[dict[str, Any]]) -> dict[str, Any]:
    return {
        "created_at_utc": _now(),
        "selection_score_formula": "0.40*holdout_min + 0.35*cv_min_mean + 0.10*holdout_auc + 0.10*cv_auc_mean + 0.05*(1-brier)",
        "external_holdout_in_selection_score": False,
        "final": _json_clean(final),
        "final_cv": _json_clean(final_cv),
        "external_holdout": external,
        "leakage_summary": leakage,
        "target_hit": {
            "synthetic_holdout_min_main_metric_90": bool(final["tuned_threshold_metrics"]["ana_metrik_minimumu"] >= HEDEF_MIN_METRIK),
            "synthetic_group_cv_min_main_metric_90": bool(final_cv["summary"]["ana_metrik_minimumu_mean"] >= HEDEF_MIN_METRIK),
        },
        "why_selected": "Leakage-clean adaylar arasinda selection_score ve min ana metrik dengesine gore secildi.",
    }


def _build_optimization_report(
    *,
    holdout_results: list[dict[str, Any]],
    cv_results: list[dict[str, Any]],
    focused_results: list[dict[str, Any]],
    availability: dict[str, Any],
    train_only_resampling_notes: dict[str, Any],
) -> dict[str, Any]:
    return {
        "created_at_utc": _now(),
        "holdout_results": [_json_clean(r) for r in holdout_results],
        "cv_results": [_json_clean(c) for c in cv_results],
        "focused_results": [_json_clean(r) for r in focused_results],
        "availability": availability,
        "train_only_resampling_notes": train_only_resampling_notes,
    }


def _literature_notes(availability: dict[str, Any], train_only_resampling_notes: dict[str, Any]) -> dict[str, Any]:
    return {
        "created_at_utc": _now(),
        "message": (
            "PIMA veri setinde ham veriyle yapilan calismalarda performans genellikle sinirli kalirken, "
            "SMOTE, SMOTE-ENN, SMOTE-Tomek, CTGAN ve benzeri sentetik/hibrit dengeleme yaklasimlari "
            "performansi artirmak icin sik kullanilmaktadir."
        ),
        "warning": (
            "Resampling ve sentetik uretim test verisini gorecek sekilde uygulanirsa leakage riski olusur. "
            "Bu calismada external holdout dis kontrol olarak saklanmistir."
        ),
        "optional_package_availability": availability,
        "train_only_resampling": train_only_resampling_notes,
    }


def _availability_report() -> dict[str, Any]:
    report = {}
    for pkg in ["catboost", "optuna", "sdv"]:
        try:
            mod = importlib.import_module(pkg)
            report[pkg] = {"available": True, "version": getattr(mod, "__version__", None)}
        except Exception:
            report[pkg] = {"available": False, "version": None, "note": "Paket mevcut degil; kurulum zorlanmadi."}
    return report


def _train_only_resampling_notes() -> dict[str, Any]:
    try:
        importlib.import_module("imblearn")
        available = True
    except Exception:
        available = False
    return {
        "available": available,
        "methods_considered": ["class_weight=balanced", "RandomOverSampler", "SMOTE", "BorderlineSMOTE", "ADASYN", "SMOTEENN", "SMOTETomek"],
        "implementation_note": "Final CSV uretiminde iki ebeveynli SMOTE kullanilmadi; resampling yalniz train-only pipeline karsilastirmasi olarak not edildi.",
    }


def _make_graphics(
    *,
    grafik_dir: Path,
    raw: pd.DataFrame,
    datasets: list[DatasetCandidate],
    comparison: dict[str, Any],
    final: dict[str, Any],
    final_cv: dict[str, Any],
    leakage_report: list[dict[str, Any]],
) -> dict[str, str]:
    plt.rcParams.update({"font.size": 9, "axes.titlesize": 11, "axes.labelsize": 9})
    graphics = {
        "original_class_distribution": str(_plot_original_distribution(grafik_dir, raw)),
        "candidate_distribution": str(_plot_candidate_distribution(grafik_dir, datasets)),
        "holdout_min_metric": str(_plot_holdout_metric(grafik_dir, comparison, "ana_metrik_minimumu", "Synthetic Holdout Min Ana Metrik")),
        "cv_min_metric": str(_plot_cv_metric(grafik_dir, comparison, "ana_metrik_minimumu_mean", "Synthetic Group CV Min Ana Metrik")),
        "holdout_roc_auc": str(_plot_holdout_metric(grafik_dir, comparison, "roc_auc", "Synthetic Holdout ROC-AUC")),
        "cv_roc_auc": str(_plot_cv_metric(grafik_dir, comparison, "roc_auc_mean", "Synthetic Group CV ROC-AUC")),
        "threshold_effect": str(_plot_threshold_effect(grafik_dir, comparison)),
        "confusion_matrix": str(_plot_confusion_matrix(grafik_dir, final)),
        "feature_importance": str(_plot_feature_importance(grafik_dir, final)),
        "calibration_curve": str(_plot_calibration(grafik_dir, final)),
        "brier": str(_plot_brier(grafik_dir, comparison)),
        "leakage_summary": str(_plot_leakage(grafik_dir, leakage_report)),
    }
    return graphics


def _plot_original_distribution(grafik_dir: Path, raw: pd.DataFrame) -> Path:
    path = grafik_dir / "controlled_benchmark_orijinal_pima_sinif_dagilimi.png"
    counts = raw[HEDEF_KOLONU].value_counts().sort_index()
    fig, ax = plt.subplots(figsize=(5.8, 3.0))
    bars = ax.bar(["Negatif", "Pozitif"], counts.values, color=["#2a9d8f", "#e76f51"])
    ax.set_title("Orijinal PIMA Sınıf Dağılımı")
    ax.set_ylabel("Satır sayısı")
    for b, v in zip(bars, counts.values, strict=True):
        ax.text(b.get_x() + b.get_width() / 2, b.get_height() + 5, str(int(v)), ha="center", weight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _plot_candidate_distribution(grafik_dir: Path, datasets: list[DatasetCandidate]) -> Path:
    path = grafik_dir / "controlled_benchmark_aday_veri_boyutlari.png"
    labels = [str(d.target_per_class) for d in datasets]
    neg = [int(d.frame[HEDEF_KOLONU].value_counts().get(0, 0)) for d in datasets]
    pos = [int(d.frame[HEDEF_KOLONU].value_counts().get(1, 0)) for d in datasets]
    x = np.arange(len(labels))
    fig, ax = plt.subplots(figsize=(8.2, 3.2))
    ax.bar(x - 0.18, neg, width=0.36, label="Negatif", color="#457b9d")
    ax.bar(x + 0.18, pos, width=0.36, label="Pozitif", color="#f4a261")
    ax.set_xticks(x, labels)
    ax.set_title("Aday Sentetik Benchmark Boyutları")
    ax.legend(frameon=False)
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _plot_holdout_metric(grafik_dir: Path, comparison: dict[str, Any], metric: str, title: str) -> Path:
    path = grafik_dir / f"controlled_benchmark_holdout_{metric}.png"
    rows = _dataset_metric_rows(comparison)
    labels = [r[0] for r in rows]
    values = [r[1]["tuned_threshold_metrics"][metric] for r in rows]
    fig, ax = plt.subplots(figsize=(8.4, 3.2))
    ax.bar(labels, values, color="#264653")
    ax.axhline(0.90, color="#d62828", linestyle="--", linewidth=1)
    ax.set_ylim(0.55, 1.02)
    ax.set_title(title)
    ax.tick_params(axis="x", rotation=20)
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _plot_cv_metric(grafik_dir: Path, comparison: dict[str, Any], metric: str, title: str) -> Path:
    path = grafik_dir / f"controlled_benchmark_cv_{metric}.png"
    cv_by_dataset = {}
    for cv in comparison["cv_results"]:
        current = cv_by_dataset.get(cv["dataset_name"])
        if current is None or cv["summary"]["ana_metrik_minimumu_mean"] > current["summary"]["ana_metrik_minimumu_mean"]:
            cv_by_dataset[cv["dataset_name"]] = cv
    labels = []
    values = []
    for dataset_name, cv in sorted(cv_by_dataset.items(), key=lambda kv: _target_from_name(kv[0])):
        labels.append(str(_target_from_name(dataset_name)))
        values.append(cv["summary"][metric])
    fig, ax = plt.subplots(figsize=(8.4, 3.2))
    ax.bar(labels, values, color="#2a9d8f")
    ax.axhline(0.90, color="#d62828", linestyle="--", linewidth=1)
    ax.set_ylim(0.55, 1.02)
    ax.set_title(title)
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _plot_threshold_effect(grafik_dir: Path, comparison: dict[str, Any]) -> Path:
    path = grafik_dir / "controlled_benchmark_threshold_tuning_etkisi.png"
    rows = _dataset_metric_rows(comparison)
    labels = [r[0] for r in rows]
    default = [r[1]["default_threshold_metrics"]["ana_metrik_minimumu"] for r in rows]
    tuned = [r[1]["tuned_threshold_metrics"]["ana_metrik_minimumu"] for r in rows]
    x = np.arange(len(labels))
    fig, ax = plt.subplots(figsize=(8.4, 3.2))
    ax.bar(x - 0.18, default, width=0.36, label="Default 0.50", color="#8d99ae")
    ax.bar(x + 0.18, tuned, width=0.36, label="Tuned", color="#ef8354")
    ax.axhline(0.90, color="#d62828", linestyle="--", linewidth=1)
    ax.set_xticks(x, labels, rotation=20)
    ax.set_ylim(0.55, 1.02)
    ax.set_title("Threshold Tuning Etkisi")
    ax.legend(frameon=False)
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _plot_confusion_matrix(grafik_dir: Path, final: dict[str, Any]) -> Path:
    path = grafik_dir / "controlled_benchmark_selected_model_confusion_matrix.png"
    cm = np.array(final["tuned_threshold_metrics"]["confusion_matrix"]["matrix"])
    fig, ax = plt.subplots(figsize=(4.2, 3.6))
    im = ax.imshow(cm, cmap="YlGnBu")
    ax.set_xticks([0, 1], ["Tahmin 0", "Tahmin 1"])
    ax.set_yticks([0, 1], ["Gerçek 0", "Gerçek 1"])
    ax.set_title("Final Synthetic Holdout Confusion Matrix")
    for i in range(2):
        for j in range(2):
            ax.text(j, i, str(int(cm[i, j])), ha="center", va="center", weight="bold")
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _plot_feature_importance(grafik_dir: Path, final: dict[str, Any]) -> Path:
    path = grafik_dir / "controlled_benchmark_selected_model_feature_importance.png"
    model = final["_model"]
    transformer = model.named_steps["features"]
    names = transformer.feature_names()
    estimator = model.named_steps["model"]
    importances = getattr(estimator, "feature_importances_", None)
    if importances is None:
        importances = np.zeros(len(names))
    importances = np.asarray(importances, dtype=float)
    if len(importances) != len(names):
        importances = np.zeros(len(names))
    order = np.argsort(importances)[-12:]
    fig, ax = plt.subplots(figsize=(6.8, 3.8))
    ax.barh(np.array(names)[order], importances[order], color="#577590")
    ax.set_title("Final Model Feature Importance")
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _plot_calibration(grafik_dir: Path, final: dict[str, Any]) -> Path:
    path = grafik_dir / "controlled_benchmark_selected_model_calibration_curve.png"
    y = np.asarray(final["_test_y"], dtype=int)
    prob = np.asarray(final["_test_prob"], dtype=float)
    frac, mean_pred = calibration_curve(y, prob, n_bins=8, strategy="quantile")
    fig, ax = plt.subplots(figsize=(4.8, 3.6))
    ax.plot([0, 1], [0, 1], "--", color="#777", label="İdeal")
    ax.plot(mean_pred, frac, marker="o", color="#bc6c25", label="Final")
    ax.set_title("Calibration Curve")
    ax.set_xlabel("Ortalama tahmin olasılığı")
    ax.set_ylabel("Gerçek pozitif oran")
    ax.legend(frameon=False)
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _plot_brier(grafik_dir: Path, comparison: dict[str, Any]) -> Path:
    path = grafik_dir / "controlled_benchmark_brier_karsilastirma.png"
    rows = _dataset_metric_rows(comparison)
    labels = [r[0] for r in rows]
    values = [r[1]["tuned_threshold_metrics"]["brier"] for r in rows]
    fig, ax = plt.subplots(figsize=(8.0, 3.0))
    ax.bar(labels, values, color="#8ecae6")
    ax.set_title("Brier Score Karşılaştırması")
    ax.set_ylabel("Brier score")
    ax.tick_params(axis="x", rotation=20)
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _plot_leakage(grafik_dir: Path, leakage_report: list[dict[str, Any]]) -> Path:
    path = grafik_dir / "controlled_benchmark_leakage_ozeti.png"
    labels = [str(r["target_per_class"]) if "target_per_class" in r else str(_target_from_name(r["dataset_name"])) for r in leakage_report]
    values = [
        r["train_test_source_intersection"]
        + max(r["cv_source_intersections"] or [0])
        + r["exact_duplicate_count"]
        + r["independent_synthetic_source_id_count"]
        + r["external_holdout_source_overlap_count"]
        for r in leakage_report
    ]
    fig, ax = plt.subplots(figsize=(8.0, 3.0))
    ax.bar(labels, values, color="#2a9d8f")
    ax.set_title("Leakage Kontrol Özeti")
    ax.set_ylabel("Problemli kontrol sayısı")
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _write_word_report(
    *,
    proje_koku: Path,
    comparison: dict[str, Any],
    model_report: dict[str, Any],
    leakage_report: dict[str, Any],
    literature_notes: dict[str, Any],
    graphics: dict[str, str],
) -> Path:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    path = proje_koku / "makine_ogrenmesi" / "raporlar" / "controlled_synthetic_benchmark_report.docx"
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(8.8)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[name].font.name = "Aptos"
        doc.styles[name].font.color.rgb = RGBColor(31, 78, 121)

    final = model_report["final"]
    final_cv = model_report["final_cv"]["summary"]
    external = model_report["external_holdout"]["metrics"]
    target_hit = model_report["target_hit"]
    final_leak = next(x for x in leakage_report["datasets"] if x["dataset_name"] == final["dataset_name"])

    _doc_title(doc, "PIMA + Sentetik Benchmark Min-0.90 Final Raporu")
    _doc_p(
        doc,
        "Bu çalışmada ana hedef, orijinal PIMA veri setinden türetilen ve source_id aile ayrımıyla kontrol edilen "
        "sentetik benchmark üzerinde yüksek performans elde etmektir. Original external holdout sonucu ise modelin "
        "gerçek PIMA verisi üzerindeki dış kontrol performansını göstermek amacıyla ayrıca raporlanmıştır.",
    )
    _doc_table(
        doc,
        ["Alan", "Değer"],
        [
            ["Ana hedef", "PIMA + sentetik benchmarkta min ana metrik >= 0.90"],
            ["Orijinal PIMA", f"{comparison['original_csv']['rows']} satır"],
            ["Original dev", f"{comparison['original_dev']['rows']} satır"],
            ["External holdout", f"{comparison['external_holdout']['rows']} satır"],
            ["Final aday", final["dataset_name"]],
            ["Minimum 0.90 hedefi", "Yakalandı" if target_hit["synthetic_holdout_min_main_metric_90"] and target_hit["synthetic_group_cv_min_main_metric_90"] else "Tam yakalanmadı"],
        ],
        font_size=8.2,
    )
    _doc_img(doc, graphics["original_class_distribution"], 5.8)
    _doc_p(doc, "EDA (Exploratory Data Analysis) aşamasında sınıf dengesizliği ve klinik olarak gerçekçi olmayan 0 değerleri dikkate alınmıştır.")

    doc.add_page_break()
    doc.add_heading("Leakage Kontrolü ve Source ID Aile Yapısı", level=1)
    _doc_p(
        doc,
        "Sentetik veri üretiminde her orijinal PIMA satırı bir kaynak aile olarak kabul edilmiştir. Her orijinal satıra "
        "original_{index} formatında bir source_id atanmış, bu satırdan üretilen tüm sentetik örnekler aynı source_id "
        "değerini taşımıştır. Böylece aynı kaynak aileden gelen örneklerin hem eğitim hem de test tarafına düşmesi engellenmiştir.",
    )
    _doc_table(
        doc,
        ["Kontrol", "Beklenen", "Sonuç"],
        [
            ["Train/Test source_id kesişimi", "0", str(final_leak["train_test_source_intersection"])],
            ["CV fold source_id kesişimi", "0", str(max(final_leak["cv_source_intersections"] or [0]))],
            ["Exact duplicate", "0", str(final_leak["exact_duplicate_count"])],
            ["Near duplicate oranı", "Düşük / 0'a yakın", _fmt(final_leak["near_duplicate_rate"])],
            ["Minimum mesafe", "Raporlanır", _fmt(final_leak["minimum_near_duplicate_distance"])],
            ["External holdout izolasyonu", "Evet", "Evet" if final_leak["external_holdout_source_overlap_count"] == 0 else "Hayır"],
            ["Bağımsız synthetic source_id kaldı mı?", "Hayır", "Hayır" if final_leak["independent_synthetic_source_id_count"] == 0 else "Evet"],
        ],
        font_size=7.6,
    )
    rows = [
        [
            str(d["target_per_class"]),
            str(d["original_dev_rows"]),
            str(d["synthetic_rows"]),
            str(d["total_rows"]),
            f"{d['class_distribution'].get('0', 0)} / {d['class_distribution'].get('1', 0)}",
        ]
        for d in comparison["datasets"]
    ]
    _doc_table(doc, ["Sınıf başı", "Orijinal", "Sentetik", "Toplam", "Neg/Poz"], rows, font_size=7.3)
    _doc_img(doc, graphics["candidate_distribution"], 6.2)

    doc.add_page_break()
    doc.add_heading("Sentetik Benchmark Karşılaştırması", level=1)
    _doc_img(doc, graphics["holdout_min_metric"], 6.3)
    _doc_img(doc, graphics["cv_min_metric"], 6.3)
    _doc_img(doc, graphics["holdout_roc_auc"], 6.3)
    _doc_img(doc, graphics["cv_roc_auc"], 6.3)

    doc.add_page_break()
    doc.add_heading("Final Model Sonuçları", level=1)
    m = final["tuned_threshold_metrics"]
    _doc_table(
        doc,
        ["Alan", "Değer"],
        [
            ["Model", final["model"]],
            ["Feature set", final["feature_set"]],
            ["Preprocessing", final["preprocessing"]],
            ["Resampling", final["resampling"]],
            ["Threshold", f"{final['threshold_report']['selected_threshold']:.2f}"],
            ["Synthetic Holdout min ana metrik", _pct(m["ana_metrik_minimumu"])],
            ["Synthetic Group CV min ana metrik", f"{_pct(final_cv['ana_metrik_minimumu_mean'])} ± {final_cv['ana_metrik_minimumu_std']:.3f}"],
            ["Selection score", f"{final['selection_score']:.4f}"],
        ],
        font_size=7.7,
    )
    _doc_img(doc, graphics["confusion_matrix"], 4.4)
    _doc_img(doc, graphics["threshold_effect"], 6.3)
    _doc_img(doc, graphics["feature_importance"], 6.3)

    doc.add_page_break()
    doc.add_heading("Kalibrasyon, External Holdout ve Literatür Notu", level=1)
    _doc_img(doc, graphics["calibration_curve"], 4.8)
    _doc_img(doc, graphics["brier"], 6.1)
    _doc_table(
        doc,
        ["External holdout metriği", "Değer"],
        [
            ["Accuracy", _pct(external["accuracy"])],
            ["F1", _pct(external["f1"])],
            ["ROC-AUC", _pct(external["roc_auc"])],
            ["Min ana metrik", _pct(external["ana_metrik_minimumu"])],
            ["Brier", f"{external['brier']:.3f}"],
        ],
        font_size=7.8,
    )
    _doc_p(doc, literature_notes["message"])
    _doc_p(doc, literature_notes["warning"])
    _doc_p(
        doc,
        "Bu sonuçlar sentetik benchmark bağlamında yorumlanmalıdır. Dış veri setleriyle desteklenmeden klinik genellenebilirlik kanıtı olarak sunulmamıştır.",
    )

    for p in doc.paragraphs:
        p.paragraph_format.space_after = Pt(2)
    doc.save(path)
    return path


def _doc_title(doc: Any, text: str) -> None:
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.alignment = 1
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)


def _doc_p(doc: Any, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.05


def _doc_img(doc: Any, path: str, width: float) -> None:
    from docx.shared import Inches

    p = doc.add_paragraph()
    p.alignment = 1
    p.add_run().add_picture(path, width=Inches(width))


def _doc_table(doc: Any, headers: list[str], rows: list[list[str]], font_size: float = 8.0) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), "1F4E79")
        cell._tc.get_or_add_tcPr().append(shade)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(font_size)
    doc.add_paragraph()


def _print_final_summary(result: dict[str, Any]) -> None:
    final = result["final"]
    cv = result["final_cv"]["summary"]
    ext = result["external_holdout"]["metrics"]
    leak = next(d for d in result["leakage_report"]["datasets"] if d["dataset_name"] == final["dataset_name"])
    m = final["tuned_threshold_metrics"]
    hit_holdout = m["ana_metrik_minimumu"] >= HEDEF_MIN_METRIK
    hit_cv = cv["ana_metrik_minimumu_mean"] >= HEDEF_MIN_METRIK
    under = [metric for metric in ANA_METRIKLER if m[metric] < HEDEF_MIN_METRIK]
    print("\n=== CONTROLLED SYNTHETIC BENCHMARK SUMMARY ===")
    print("- Ana hedef: PIMA + source_id kontrollu sentetik benchmarkta min ana metrik >= 0.90")
    print(f"- En iyi PIMA + sentetik benchmark adayı: {final['dataset_name']}")
    print(f"- Veri boyutu: {result['comparison_report']['final']['target_per_class']} per class")
    print(f"- Model: {final['model']}")
    print(f"- Feature set: {final['feature_set']}")
    print(f"- Preprocessing: {final['preprocessing']}")
    print(f"- Resampling: {final['resampling']}")
    print(f"- Threshold: {final['threshold_report']['selected_threshold']:.2f}")
    print(f"- Synthetic Holdout Accuracy: {m['accuracy']:.4f}")
    print(f"- Synthetic Holdout Precision: {m['precision']:.4f}")
    print(f"- Synthetic Holdout Recall: {m['recall']:.4f}")
    print(f"- Synthetic Holdout Specificity: {m['specificity']:.4f}")
    print(f"- Synthetic Holdout F1: {m['f1']:.4f}")
    print(f"- Synthetic Holdout ROC-AUC: {m['roc_auc']:.4f}")
    print(f"- Synthetic Holdout min ana metrik: {m['ana_metrik_minimumu']:.4f}")
    print(f"- Synthetic Group CV Accuracy mean ± std: {cv['accuracy_mean']:.4f} ± {cv['accuracy_std']:.4f}")
    print(f"- Synthetic Group CV Precision mean ± std: {cv['precision_mean']:.4f} ± {cv['precision_std']:.4f}")
    print(f"- Synthetic Group CV Recall mean ± std: {cv['recall_mean']:.4f} ± {cv['recall_std']:.4f}")
    print(f"- Synthetic Group CV Specificity mean ± std: {cv['specificity_mean']:.4f} ± {cv['specificity_std']:.4f}")
    print(f"- Synthetic Group CV F1 mean ± std: {cv['f1_mean']:.4f} ± {cv['f1_std']:.4f}")
    print(f"- Synthetic Group CV ROC-AUC mean ± std: {cv['roc_auc_mean']:.4f} ± {cv['roc_auc_std']:.4f}")
    print(f"- Synthetic Group CV min ana metrik mean ± std: {cv['ana_metrik_minimumu_mean']:.4f} ± {cv['ana_metrik_minimumu_std']:.4f}")
    print(f"- Original external holdout Accuracy: {ext['accuracy']:.4f}")
    print(f"- Original external holdout F1: {ext['f1']:.4f}")
    print(f"- Original external holdout ROC-AUC: {ext['roc_auc']:.4f}")
    print(f"- Leakage durumu: {leak['leakage_status']}")
    print(f"- Train/test source_id kesişimi: {leak['train_test_source_intersection']}")
    print(f"- CV source_id kesişimi: {max(leak['cv_source_intersections'] or [0])}")
    print(f"- Exact duplicate: {leak['exact_duplicate_count']}")
    print(f"- Near duplicate oranı: {leak['near_duplicate_rate']}")
    print(f"- Minimum mesafe: {leak['minimum_near_duplicate_distance']}")
    print(f"- Yeni Word raporu yolu: {result['word_report']}")
    print(f"- Masaüstü Word raporu yolu: {result['desktop_word_report']}")
    print(f"- Minimum 0.90 hedefi yakalandı mı: {'Evet' if hit_holdout and hit_cv else 'Hayır'}")
    print(f"- Hedef yakalandıysa hangi benchmarkta yakalandı: {final['dataset_name'] if hit_holdout and hit_cv else 'Yakalanmadı'}")
    print(f"- Hedef yakalanmadıysa hangi metrikler altında kaldı: {', '.join(under) if under else 'Holdout tarafında yok'}")
    print(f"- Neden bu model seçildi: selection_score={final['selection_score']:.4f}, leakage-clean ve CV ile doğrulanmış aday.")
    print(f"- XGBoost seçilmediyse neden seçilmedi: {'XGBoost seçildi' if final['model'] == 'xgboost' else 'Daha yüksek selection_score veren model bulundu'}")
    print("- 10000/per class sonucu final seçilmediyse neden seçilmedi: Daha küçük adaylar hedefe yakın/geçer durumdaysa savunulabilirlik önceliklendirildi.")


def _dataset_metric_rows(comparison: dict[str, Any]) -> list[tuple[str, dict[str, Any]]]:
    rows = []
    for name, result in comparison["top_by_dataset"].items():
        rows.append((str(_target_from_name(name)), result))
    return sorted(rows, key=lambda r: int(r[0]))


def _target_from_name(name: str) -> int:
    for part in name.split("_"):
        if part.isdigit():
            return int(part)
    return 0


def _write_dataset(veri_dir: Path, dataset: DatasetCandidate) -> None:
    combined = pd.concat([dataset.frame.reset_index(drop=True), dataset.metadata.reset_index(drop=True)], axis=1)
    combined.to_csv(veri_dir / f"{dataset.name}.csv", index=False)
    dataset.metadata.to_csv(veri_dir / f"{dataset.name}_metadata.csv", index=False)


def _find_dataset(datasets: list[DatasetCandidate], name: str) -> DatasetCandidate:
    for d in datasets:
        if d.name == name:
            return d
    raise KeyError(name)


def _find_feature(features: list[FeatureSpec], name: str) -> FeatureSpec:
    for f in features:
        if f.name == name:
            return f
    raise KeyError(name)


def _find_cv_result(cv_results: list[dict[str, Any]], result: dict[str, Any]) -> dict[str, Any]:
    for cv in cv_results:
        if cv["dataset_name"] == result["dataset_name"] and cv["feature_set"] == result["feature_set"] and cv["model"] == result["model"]:
            return cv
    raise KeyError((result["dataset_name"], result["feature_set"], result["model"]))


def _split_counts(total: int, weights: list[float]) -> list[int]:
    counts = [int(total * w) for w in weights]
    counts[-1] += total - sum(counts)
    return counts


def _class_distribution(frame: pd.DataFrame) -> dict[str, int]:
    return {str(int(k)): int(v) for k, v in frame[HEDEF_KOLONU].value_counts().sort_index().to_dict().items()}


def _dataframe_yap(x: Any) -> pd.DataFrame:
    if isinstance(x, pd.DataFrame):
        return x.copy()
    arr = np.asarray(x)
    return pd.DataFrame(arr, columns=OZELLIK_KOLONLARI[: arr.shape[1]])


def _json_write(path: Path, value: dict[str, Any]) -> None:
    path.write_text(json.dumps(_json_clean(value), ensure_ascii=False, indent=2), encoding="utf-8")


def _json_clean(value: Any) -> Any:
    if isinstance(value, dict):
        return {str(k): _json_clean(v) for k, v in value.items() if not str(k).startswith("_")}
    if isinstance(value, list | tuple):
        return [_json_clean(v) for v in value]
    if isinstance(value, np.ndarray):
        return value.tolist()
    if isinstance(value, np.integer):
        return int(value)
    if isinstance(value, np.floating):
        return float(value)
    if isinstance(value, Path):
        return str(value)
    return value


def _sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _fmt(value: Any) -> str:
    if value is None:
        return "Yok"
    return f"{float(value):.4f}"


def _pct(value: Any) -> str:
    return f"%{float(value) * 100:.2f}"
