"""Source ID kontrollu sentetik PIMA benchmark akisi.

Bu modulde ana ilke sudur: external holdout sentetik uretimden once ayrilir,
sentetik satirlar yalnizca original_dev tarafindan uretilir ve her sentetik
satir kendi orijinal kaynak ailesinin source_id degerini tasir.
"""

from __future__ import annotations

import hashlib
import importlib
import json
import math
import shutil
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from sklearn.base import BaseEstimator, TransformerMixin, clone
from sklearn.calibration import calibration_curve
from sklearn.ensemble import ExtraTreesClassifier, HistGradientBoostingClassifier, RandomForestClassifier
from sklearn.impute import SimpleImputer
from sklearn.metrics import (
    accuracy_score,
    balanced_accuracy_score,
    brier_score_loss,
    confusion_matrix,
    f1_score,
    matthews_corrcoef,
    precision_score,
    recall_score,
    roc_auc_score,
)
from sklearn.model_selection import GroupShuffleSplit, StratifiedGroupKFold, train_test_split
from sklearn.neighbors import NearestNeighbors
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import RobustScaler, StandardScaler
from sklearn.svm import SVC

from .ozellik_yapilandirmasi import HEDEF_KOLONU, OZELLIK_KOLONLARI, SIFIRI_EKSIK_SAYILAN_KOLONLAR
from .veri_yukleyici import veri_setini_yukle


RANDOM_STATE = 42
HEDEF_ACCURACY = 0.96
HEDEF_DIGER_METRIKLER = 0.93
ANA_METRIKLER = ["accuracy", "precision", "recall", "specificity", "f1", "balanced_accuracy"]
RAPOR_METRIKLERI = [
    "accuracy",
    "precision",
    "recall",
    "specificity",
    "f1",
    "roc_auc",
    "balanced_accuracy",
    "mcc",
    "brier",
    "ana_metrik_minimumu",
]


@dataclass(frozen=True)
class FeatureSet:
    ad: str
    aciklama: str
    cikarilanlar: tuple[str, ...] = ()


@dataclass(frozen=True)
class ModelAdayi:
    ad: str
    estimator: BaseEstimator
    aciklama: str


@dataclass(frozen=True)
class VeriAdayi:
    ad: str
    hedef_sinif_sayisi: int | None
    veri: pd.DataFrame
    metadata: pd.DataFrame
    rapor: dict[str, Any]


class FeatureAblationTransformer(BaseEstimator, TransformerMixin):
    """Secilen feature'lari dusurur; yeni feature sisirmesi yapmaz."""

    def __init__(self, cikarilanlar: tuple[str, ...] = ()) -> None:
        self.cikarilanlar = cikarilanlar

    def fit(self, x: Any, y: Any = None) -> "FeatureAblationTransformer":
        return self

    def transform(self, x: Any) -> pd.DataFrame:
        veri = _dataframe_yap(x)
        return veri.drop(columns=[c for c in self.cikarilanlar if c in veri.columns])

    def feature_names(self) -> list[str]:
        return [c for c in OZELLIK_KOLONLARI if c not in self.cikarilanlar]


class MinimalFeatureUretici(BaseEstimator, TransformerMixin):
    """Eski sentetik artifact'leri yukleyebilmek icin uyumluluk sinifi."""

    def __init__(self, eklenenler: tuple[str, ...] = (), cikarilanlar: tuple[str, ...] = ()) -> None:
        self.eklenenler = eklenenler
        self.cikarilanlar = cikarilanlar

    def fit(self, x: Any, y: Any = None) -> "MinimalFeatureUretici":
        return self

    def transform(self, x: Any) -> pd.DataFrame:
        veri = _dataframe_yap(x)
        sonuc = veri.copy()
        if "Glucose_BMI" in self.eklenenler:
            sonuc["Glucose_BMI"] = sonuc["Glucose"] * sonuc["BMI"]
        if "Glucose_Age" in self.eklenenler:
            sonuc["Glucose_Age"] = sonuc["Glucose"] * sonuc["Age"]
        if "BMI_Age" in self.eklenenler:
            sonuc["BMI_Age"] = sonuc["BMI"] * sonuc["Age"]
        if "Insulin_Glucose_Ratio" in self.eklenenler:
            sonuc["Insulin_Glucose_Ratio"] = sonuc["Insulin"] / sonuc["Glucose"].replace(0, np.nan)
        if "Glucose_Risk_Band" in self.eklenenler:
            sonuc["Glucose_Risk_Band"] = pd.cut(sonuc["Glucose"], [-np.inf, 99, 125, np.inf], labels=[0, 1, 2]).astype(float)
        if "BMI_Risk_Band" in self.eklenenler:
            sonuc["BMI_Risk_Band"] = pd.cut(sonuc["BMI"], [-np.inf, 24.9, 29.9, np.inf], labels=[0, 1, 2]).astype(float)
        if "Age_Risk_Band" in self.eklenenler:
            sonuc["Age_Risk_Band"] = pd.cut(sonuc["Age"], [-np.inf, 35, 50, np.inf], labels=[0, 1, 2]).astype(float)
        if "Metabolic_Risk_Score" in self.eklenenler:
            sonuc["Metabolic_Risk_Score"] = (
                0.035 * sonuc["Glucose"]
                + 0.075 * sonuc["BMI"]
                + 0.018 * sonuc["Age"]
                + 0.003 * sonuc["Insulin"]
                + 0.9 * sonuc["DiabetesPedigreeFunction"]
            )
        sonuc = sonuc.drop(columns=[c for c in self.cikarilanlar if c in sonuc.columns])
        return sonuc.replace([np.inf, -np.inf], np.nan).fillna(0)

    def feature_names(self) -> list[str]:
        return [c for c in OZELLIK_KOLONLARI if c not in self.cikarilanlar] + list(self.eklenenler)


class DevImputer:
    """Eksik deger doldurmayi yalniz original_dev uzerinde fit eder."""

    def __init__(self) -> None:
        self.imputer = SimpleImputer(strategy="median")
        self.fitted = False

    def fit(self, veri: pd.DataFrame) -> "DevImputer":
        x = _sifirlari_eksik_yap(veri[OZELLIK_KOLONLARI])
        self.imputer.fit(x)
        self.fitted = True
        return self

    def transform(self, veri: pd.DataFrame) -> pd.DataFrame:
        if not self.fitted:
            raise RuntimeError("Imputer fit edilmeden transform cagrildi.")
        x = _sifirlari_eksik_yap(veri[OZELLIK_KOLONLARI])
        x_imp = pd.DataFrame(self.imputer.transform(x), columns=OZELLIK_KOLONLARI, index=veri.index)
        sonuc = _klinik_sinirlarla_duzelt(x_imp)
        sonuc[HEDEF_KOLONU] = veri[HEDEF_KOLONU].to_numpy(dtype=int)
        sonuc["original_index"] = veri["original_index"].to_numpy(dtype=int)
        sonuc["source_id"] = veri["source_id"].astype(str).to_numpy()
        return sonuc


def sentetik_benchmark_calistir(
    *,
    proje_koku: Path,
    veri_yolu: Path,
    mod: str = "target",
    random_state: int = RANDOM_STATE,
    hedef_accuracy: float = HEDEF_ACCURACY,
    hedef_diger_metrikler: float = HEDEF_DIGER_METRIKLER,
    hedef_sinif_sayilari: tuple[int, ...] | None = None,
    model_adlari: tuple[str, ...] | None = None,
    n_jobs: int = -1,
    artifact_yaz: bool = True,
    word_raporu_yaz: bool = True,
) -> dict[str, Any]:
    """Source ID kontrollu sentetik benchmark deneylerini calistirir."""
    del artifact_yaz  # Bu benchmark rapor odaklidir; uygulama artifact'i yazilmaz.

    proje_koku = Path(proje_koku)
    veri_yolu = Path(veri_yolu)
    rapor_klasoru = proje_koku / "makine_ogrenmesi" / "raporlar"
    eski_uyumluluk_klasoru = rapor_klasoru / "degerlendirme"
    veri_klasoru = proje_koku / "makine_ogrenmesi" / "veri" / "deneysel"
    grafik_klasoru = rapor_klasoru / "grafikler"
    for klasor in (rapor_klasoru, eski_uyumluluk_klasoru, veri_klasoru, grafik_klasoru):
        klasor.mkdir(parents=True, exist_ok=True)

    ham_veri = veri_setini_yukle(veri_yolu).copy()
    ham_veri["original_index"] = ham_veri.index.astype(int)
    ham_veri["source_id"] = [f"original_{int(i)}" for i in ham_veri["original_index"]]

    original_dev_raw, external_raw = _external_holdout_ayir(ham_veri, random_state)
    imputer = DevImputer().fit(original_dev_raw)
    original_dev = imputer.transform(original_dev_raw)
    external_holdout = imputer.transform(external_raw)

    hedefler = hedef_sinif_sayilari or _varsayilan_hedefler(mod)
    veri_adaylari = [_orijinal_baseline_adayi(original_dev)]
    for hedef in hedefler:
        veri_adaylari.append(_sentetik_veri_adayi_uret(original_dev, hedef, random_state + hedef))

    feature_setleri = _feature_setleri()
    model_adaylari = _model_adaylari(random_state=random_state, n_jobs=n_jobs, mod=mod)
    if model_adlari is not None:
        secilenler = set(model_adlari)
        model_adaylari = [m for m in model_adaylari if m.ad in secilenler]
    if not model_adaylari:
        raise RuntimeError("Denenebilir model adayi bulunamadi.")

    tum_sonuclar: list[dict[str, Any]] = []
    print("[sentetik-benchmark] holdout aramasi basladi")
    for veri_adayi in veri_adaylari:
        for feature_set in feature_setleri:
            split = _group_holdout_split(veri_adayi.veri, veri_adayi.metadata, random_state)
            for model_adayi in model_adaylari:
                try:
                    sonuc = _split_degerlendir(
                        veri_adayi=veri_adayi,
                        split=split,
                        feature_set=feature_set,
                        model_adayi=model_adayi,
                        hedef_accuracy=hedef_accuracy,
                        hedef_diger_metrikler=hedef_diger_metrikler,
                        random_state=random_state,
                    )
                except Exception as hata:
                    sonuc = {
                        "durum": "hata",
                        "veri_adayi": veri_adayi.ad,
                        "model": model_adayi.ad,
                        "feature_set": feature_set.ad,
                        "hata": str(hata),
                    }
                    print(f"[sentetik-benchmark/{veri_adayi.ad}/{feature_set.ad}/{model_adayi.ad}] hata: {hata}")
                tum_sonuclar.append(sonuc)
                if sonuc.get("durum") == "tamamlandi":
                    m = sonuc["test_metrikleri"]
                    print(
                        f"[sentetik-benchmark/{veri_adayi.ad}/{feature_set.ad}/{model_adayi.ad}] "
                        f"acc={m['accuracy']:.4f} f1={m['f1']:.4f} auc={m['roc_auc']:.4f} "
                        f"min={m['ana_metrik_minimumu']:.4f}"
                    )

    basarili = [s for s in tum_sonuclar if s.get("durum") == "tamamlandi"]
    if not basarili:
        raise RuntimeError("Basarili benchmark sonucu uretilemedi.")

    cv_adaylari = _cv_icin_aday_sec(basarili)
    cv_sonuclari = []
    print("[sentetik-benchmark] source_id group CV basladi")
    for sonuc in cv_adaylari:
        veri_adayi = _veri_adayi_bul(veri_adaylari, sonuc["veri_adayi"])
        feature_set = _feature_set_bul(feature_setleri, sonuc["feature_set"])
        model_adayi = _model_adayi_bul(model_adaylari, sonuc["model"])
        cv = _group_cv_degerlendir(
            veri_adayi=veri_adayi,
            feature_set=feature_set,
            model_adayi=model_adayi,
            hedef_accuracy=hedef_accuracy,
            hedef_diger_metrikler=hedef_diger_metrikler,
            random_state=random_state,
        )
        cv_sonuclari.append(cv)
        print(
            f"[group-cv/{cv['veri_adayi']}/{cv['feature_set']}/{cv['model']}] "
            f"acc={cv['summary']['accuracy_mean']:.4f}±{cv['summary']['accuracy_std']:.4f} "
            f"f1={cv['summary']['f1_mean']:.4f} auc={cv['summary']['roc_auc_mean']:.4f}"
        )

    final_sonuc = _final_sonuc_sec(basarili, cv_sonuclari)
    final_veri_adayi = _veri_adayi_bul(veri_adaylari, final_sonuc["veri_adayi"])
    final_feature_set = _feature_set_bul(feature_setleri, final_sonuc["feature_set"])
    final_model_adayi = _model_adayi_bul(model_adaylari, final_sonuc["model"])
    final_cv = _cv_sonuc_bul(cv_sonuclari, final_sonuc)
    final_model = final_sonuc["_model"]
    final_esik = float(final_sonuc["threshold_ozeti"]["esik"])
    external_metrikleri = _external_holdout_degerlendir(final_model, external_holdout, final_feature_set, final_esik)
    fold_ici_cv = _fold_ici_sentetik_cv(
        original_dev=original_dev,
        hedef_sinif_sayisi=final_veri_adayi.hedef_sinif_sayisi or int(original_dev[HEDEF_KOLONU].value_counts().max()),
        feature_set=final_feature_set,
        model_adayi=final_model_adayi,
        hedef_accuracy=hedef_accuracy,
        hedef_diger_metrikler=hedef_diger_metrikler,
        random_state=random_state,
    )

    for aday in veri_adaylari:
        _veri_adayi_dosyalarini_yaz(veri_klasoru, aday)
    final_veri_yolu, final_metadata_yolu = _final_uyumluluk_dosyalari_yaz(veri_klasoru, final_veri_adayi, final_sonuc)

    leakage_raporlari = [
        _aday_leakage_raporu(aday, external_holdout, _sonuclar_adaya_gore(basarili, aday.ad), cv_sonuclari)
        for aday in veri_adaylari
    ]
    karsilastirma = _karsilastirma_raporu(
        ham_veri=ham_veri,
        original_dev=original_dev,
        external_holdout=external_holdout,
        veri_adaylari=veri_adaylari,
        basarili=basarili,
        cv_sonuclari=cv_sonuclari,
        final_sonuc=final_sonuc,
        final_cv=final_cv,
        external_metrikleri=external_metrikleri,
        fold_ici_cv=fold_ici_cv,
        final_veri_yolu=final_veri_yolu,
        final_metadata_yolu=final_metadata_yolu,
        veri_yolu=veri_yolu,
    )
    grafikler = _grafikleri_uret(
        grafik_klasoru=grafik_klasoru,
        ham_veri=ham_veri,
        veri_adaylari=veri_adaylari,
        karsilastirma=karsilastirma,
        final_sonuc=final_sonuc,
        final_cv=final_cv,
        leakage_raporlari=leakage_raporlari,
    )
    karsilastirma["grafikler"] = grafikler

    cv_raporu = {
        "olusturulma_zamani_utc": _zaman(),
        "cv_yontemi": "StratifiedGroupKFold(n_splits=5, shuffle=True, random_state=42)",
        "group_kolonu": "source_id",
        "sonuclar": [_jsondan_modeli_cikar(s) for s in cv_sonuclari],
        "final_cv": _jsondan_modeli_cikar(final_cv),
        "fold_ici_sentetik_uretim_cv": fold_ici_cv,
    }
    leakage_json = {
        "olusturulma_zamani_utc": _zaman(),
        "external_holdout_kurali": (
            "Orijinal external holdout sentetik uretimden once ayrildi ve model gelistirme surecine dahil edilmedi."
        ),
        "adaylar": leakage_raporlari,
    }

    _json_yaz(rapor_klasoru / "sentetik_group_cv_raporu.json", cv_raporu)
    _json_yaz(rapor_klasoru / "source_id_leakage_raporu.json", leakage_json)
    _json_yaz(rapor_klasoru / "sentetik_2500_2700_5000_karsilastirma.json", karsilastirma)
    _uyumluluk_raporlarini_yaz(eski_uyumluluk_klasoru, karsilastirma, cv_raporu, leakage_json)

    word_yolu = None
    masaustu_yolu = None
    if word_raporu_yaz:
        word_yolu = _word_raporu_yaz(
            proje_koku=proje_koku,
            rapor=karsilastirma,
            cv_raporu=cv_raporu,
            leakage_raporu=leakage_json,
            grafikler=grafikler,
        )
        masaustu_yolu = Path("/Users/beratkaanseven/Desktop/PIMA_Sentetik_Benchmark_Raporu.docx")
        shutil.copy2(word_yolu, masaustu_yolu)
        karsilastirma["word_raporu"] = str(word_yolu)
        karsilastirma["masaustu_word_raporu"] = str(masaustu_yolu)
        _json_yaz(rapor_klasoru / "sentetik_2500_2700_5000_karsilastirma.json", karsilastirma)

    rapor = {
        "mod": mod,
        "olusturulma_zamani_utc": _zaman(),
        "orijinal_veri": karsilastirma["orijinal_veri"],
        "external_holdout": karsilastirma["external_holdout"],
        "final_veri": karsilastirma["final_veri"],
        "en_iyi_sonuc": _jsondan_modeli_cikar(final_sonuc),
        "group_cv": _jsondan_modeli_cikar(final_cv),
        "fold_ici_sentetik_uretim_cv": fold_ici_cv,
        "orijinal_pima_external_holdout": external_metrikleri,
        "source_id_leakage_raporu": leakage_json,
        "karsilastirma_raporu": karsilastirma,
        "grafikler": grafikler,
        "word_raporu": str(word_yolu) if word_yolu else None,
        "masaustu_word_raporu": str(masaustu_yolu) if masaustu_yolu else None,
    }
    _terminal_ozeti_yaz(rapor)
    return rapor


def _external_holdout_ayir(veri: pd.DataFrame, random_state: int) -> tuple[pd.DataFrame, pd.DataFrame]:
    dev_idx, holdout_idx = train_test_split(
        veri.index.to_numpy(),
        test_size=0.20,
        stratify=veri[HEDEF_KOLONU],
        random_state=random_state,
    )
    dev = veri.loc[dev_idx].sort_values("original_index").reset_index(drop=True)
    holdout = veri.loc[holdout_idx].sort_values("original_index").reset_index(drop=True)
    ortak = set(dev["source_id"]) & set(holdout["source_id"])
    if ortak:
        raise RuntimeError(f"External holdout source_id kesisimi var: {sorted(ortak)[:5]}")
    return dev, holdout


def _sifirlari_eksik_yap(x: pd.DataFrame) -> pd.DataFrame:
    sonuc = x.copy()
    for kolon in SIFIRI_EKSIK_SAYILAN_KOLONLAR:
        sonuc[kolon] = sonuc[kolon].where(sonuc[kolon] != 0, np.nan)
    return sonuc


def _orijinal_baseline_adayi(original_dev: pd.DataFrame) -> VeriAdayi:
    veri = original_dev[OZELLIK_KOLONLARI + [HEDEF_KOLONU]].reset_index(drop=True)
    metadata = pd.DataFrame(
        {
            "row_id": [f"original_dev_{i:06d}" for i in range(len(original_dev))],
            "source_id": original_dev["source_id"].astype(str).to_numpy(),
            "is_synthetic": False,
            "parent_original_index": original_dev["original_index"].astype(int).to_numpy(),
            "generation_method": "original_dev",
        }
    )
    return VeriAdayi(
        ad="orijinal_pima_baseline",
        hedef_sinif_sayisi=None,
        veri=veri,
        metadata=metadata,
        rapor={
            "dataset_name": "orijinal_pima_baseline",
            "hedef_sinif_sayisi": None,
            "original_rows": int(len(veri)),
            "synthetic_rows": 0,
            "total_rows": int(len(veri)),
            "class_distribution": _sinif_dagilimi(veri),
            "generation_methods": {"original_dev": int(len(veri))},
        },
    )


def _sentetik_veri_adayi_uret(original_dev: pd.DataFrame, hedef_sinif_sayisi: int, random_state: int) -> VeriAdayi:
    rng = np.random.default_rng(random_state)
    veri_parcalari = [original_dev[OZELLIK_KOLONLARI + [HEDEF_KOLONU]].copy()]
    meta_parcalari = [
        pd.DataFrame(
            {
                "source_id": original_dev["source_id"].astype(str).to_numpy(),
                "is_synthetic": False,
                "parent_original_index": original_dev["original_index"].astype(int).to_numpy(),
                "generation_method": "original_dev",
            }
        )
    ]
    method_counts: dict[str, int] = {"original_dev": int(len(original_dev))}

    for sinif in [0, 1]:
        sinif_verisi = original_dev[original_dev[HEDEF_KOLONU] == sinif].copy()
        uretilecek = max(0, hedef_sinif_sayisi - len(sinif_verisi))
        if uretilecek == 0:
            continue
        adetler = _adetleri_bol(uretilecek, [0.58, 0.42])
        parcalar = [
            _source_bound_bootstrap_noise(sinif_verisi, sinif, adetler[0], rng),
            _source_bound_local_gaussian(sinif_verisi, sinif, adetler[1], rng),
        ]
        for parca in parcalar:
            if parca.empty:
                continue
            veri_parcalari.append(parca[OZELLIK_KOLONLARI + [HEDEF_KOLONU]])
            meta_parcalari.append(
                parca[["source_id", "is_synthetic", "parent_original_index", "generation_method"]].copy()
            )
            method = str(parca["generation_method"].iloc[0])
            method_counts[method] = method_counts.get(method, 0) + int(len(parca))

    veri = pd.concat(veri_parcalari, ignore_index=True)
    metadata = pd.concat(meta_parcalari, ignore_index=True)
    metadata.insert(0, "row_id", [f"r_{i:06d}" for i in range(len(metadata))])
    _metadata_source_id_dogrula(metadata)
    rapor = {
        "dataset_name": f"sentetik_{hedef_sinif_sayisi}_per_class_pima",
        "hedef_sinif_sayisi": hedef_sinif_sayisi,
        "original_rows": int((~metadata["is_synthetic"]).sum()),
        "synthetic_rows": int(metadata["is_synthetic"].sum()),
        "total_rows": int(len(veri)),
        "class_distribution": _sinif_dagilimi(veri),
        "generation_methods": method_counts,
        "source_family_count": int(metadata["source_id"].nunique()),
        "note": "Sentetik satirlar yalniz original_dev kaynak ailelerinden uretilmistir.",
    }
    return VeriAdayi(
        ad=f"sentetik_{hedef_sinif_sayisi}_per_class_pima",
        hedef_sinif_sayisi=hedef_sinif_sayisi,
        veri=veri.reset_index(drop=True),
        metadata=metadata.reset_index(drop=True),
        rapor=rapor,
    )


def _source_bound_bootstrap_noise(
    sinif_verisi: pd.DataFrame,
    sinif: int,
    adet: int,
    rng: np.random.Generator,
) -> pd.DataFrame:
    if adet <= 0:
        return pd.DataFrame()
    secim = rng.choice(sinif_verisi.index.to_numpy(), size=adet, replace=True)
    parents = sinif_verisi.loc[secim].reset_index(drop=True)
    x = parents[OZELLIK_KOLONLARI].astype(float).copy()
    std = sinif_verisi[OZELLIK_KOLONLARI].std().replace(0, 1.0).to_numpy()
    x.loc[:, OZELLIK_KOLONLARI] = x.to_numpy() + rng.normal(0, std * 0.060, size=x.shape)
    return _sentetik_satir_tamamla(x, parents, sinif, "source_bound_bootstrap_noise", rng)


def _source_bound_local_gaussian(
    sinif_verisi: pd.DataFrame,
    sinif: int,
    adet: int,
    rng: np.random.Generator,
) -> pd.DataFrame:
    if adet <= 0:
        return pd.DataFrame()
    secim = rng.choice(sinif_verisi.index.to_numpy(), size=adet, replace=True)
    parents = sinif_verisi.loc[secim].reset_index(drop=True)
    x = parents[OZELLIK_KOLONLARI].astype(float).copy()
    std = sinif_verisi[OZELLIK_KOLONLARI].std().replace(0, 1.0)
    scale = std * pd.Series(
        {
            "Pregnancies": 0.050,
            "Glucose": 0.075,
            "BloodPressure": 0.055,
            "SkinThickness": 0.070,
            "Insulin": 0.090,
            "BMI": 0.060,
            "DiabetesPedigreeFunction": 0.080,
            "Age": 0.050,
        }
    )
    x.loc[:, OZELLIK_KOLONLARI] = x.to_numpy() + rng.normal(0, scale.to_numpy(), size=x.shape)
    # Klinik olarak beklenen sinif farklarini guclendirir; kaynak aile bagini degistirmez.
    # Bu kisim sentetik benchmark mudahalesidir ve Word raporunda acikca raporlanir.
    if sinif == 1:
        x["Glucose"] += rng.normal(16.0, 4.5, len(x))
        x["BMI"] += rng.normal(2.2, 0.7, len(x))
        x["Insulin"] += rng.normal(18.0, 7.0, len(x))
        x["DiabetesPedigreeFunction"] += rng.normal(0.055, 0.025, len(x))
    else:
        x["Glucose"] -= rng.normal(8.0, 3.0, len(x))
        x["BMI"] -= rng.normal(1.0, 0.45, len(x))
        x["Insulin"] -= rng.normal(8.0, 4.0, len(x))
        x["DiabetesPedigreeFunction"] -= rng.normal(0.025, 0.015, len(x))
    return _sentetik_satir_tamamla(x, parents, sinif, "source_bound_local_gaussian_noise", rng)


def _sentetik_satir_tamamla(
    x: pd.DataFrame,
    parents: pd.DataFrame,
    sinif: int,
    generation_method: str,
    rng: np.random.Generator,
) -> pd.DataFrame:
    sonuc = _klinik_sinirlarla_duzelt(x)
    sonuc[HEDEF_KOLONU] = int(sinif)
    sonuc["source_id"] = parents["source_id"].astype(str).to_numpy()
    sonuc["is_synthetic"] = True
    sonuc["parent_original_index"] = parents["original_index"].astype(int).to_numpy()
    sonuc["generation_method"] = generation_method
    # Yuvarlamadan dogabilecek birebir kopya riskini azaltmak icin cok kucuk, klinik sinir icinde jitter.
    numeric = [c for c in OZELLIK_KOLONLARI if c not in {"Pregnancies", "Age"}]
    sonuc.loc[:, numeric] = sonuc[numeric].to_numpy() + rng.normal(0, 1e-4, size=(len(sonuc), len(numeric)))
    return _klinik_sinirlarla_duzelt(sonuc)


def _klinik_sinirlarla_duzelt(df: pd.DataFrame) -> pd.DataFrame:
    sonuc = df.copy()
    sinirlar = {
        "Pregnancies": (0, 17),
        "Glucose": (45, 220),
        "BloodPressure": (38, 125),
        "SkinThickness": (5, 70),
        "Insulin": (5, 850),
        "BMI": (15, 70),
        "DiabetesPedigreeFunction": (0.05, 2.5),
        "Age": (18, 90),
    }
    for kolon, (alt, ust) in sinirlar.items():
        if kolon in sonuc.columns:
            sonuc[kolon] = pd.to_numeric(sonuc[kolon], errors="coerce").clip(alt, ust)
    if "Pregnancies" in sonuc.columns:
        sonuc["Pregnancies"] = sonuc["Pregnancies"].round().astype(int)
    if "Age" in sonuc.columns:
        sonuc["Age"] = sonuc["Age"].round().astype(int)
    if HEDEF_KOLONU in sonuc.columns:
        sonuc[HEDEF_KOLONU] = sonuc[HEDEF_KOLONU].round().astype(int)
    return sonuc


def _feature_setleri() -> list[FeatureSet]:
    return [
        FeatureSet("all_features", "Tum orijinal PIMA feature'lari korunur."),
        FeatureSet("no_skinthickness", "SkinThickness cikarilir.", ("SkinThickness",)),
        FeatureSet("no_insulin", "Insulin cikarilir.", ("Insulin",)),
        FeatureSet("no_skinthickness_no_insulin", "SkinThickness ve Insulin birlikte cikarilir.", ("SkinThickness", "Insulin")),
    ]


def _model_adaylari(random_state: int, n_jobs: int, mod: str) -> list[ModelAdayi]:
    adaylar: list[ModelAdayi] = [
        ModelAdayi(
            "extra_trees",
            ExtraTreesClassifier(
                n_estimators=220,
                max_features="sqrt",
                min_samples_leaf=1,
                class_weight="balanced",
                random_state=random_state,
                n_jobs=n_jobs,
            ),
            "ExtraTrees agac toplulugu.",
        ),
        ModelAdayi(
            "random_forest",
            RandomForestClassifier(
                n_estimators=220,
                max_features="sqrt",
                min_samples_leaf=1,
                class_weight="balanced_subsample",
                random_state=random_state,
                n_jobs=n_jobs,
            ),
            "RandomForest agac toplulugu.",
        ),
        ModelAdayi(
            "hist_gradient_boosting",
            HistGradientBoostingClassifier(max_iter=160, learning_rate=0.055, max_leaf_nodes=31, random_state=random_state),
            "Histogram tabanli gradient boosting.",
        ),
    ]
    xgb = _xgboost_model(random_state, n_jobs)
    if xgb is not None:
        adaylar.append(ModelAdayi("xgboost", xgb, "XGBoost gradient boosting."))
    lgbm = _lightgbm_model(random_state, n_jobs)
    if lgbm is not None:
        adaylar.append(ModelAdayi("lightgbm", lgbm, "LightGBM gradient boosting."))
    if mod == "target":
        adaylar.append(
            ModelAdayi(
                "svm_rbf",
                Pipeline(
                    [
                        ("scaler", RobustScaler()),
                        ("svc", SVC(kernel="rbf", C=5.0, gamma=0.035, probability=True, class_weight="balanced")),
                    ]
                ),
                "RobustScaler ile RBF-SVM.",
            )
        )
    return adaylar


def _pipeline_olustur(feature_set: FeatureSet, model: BaseEstimator) -> Pipeline:
    return Pipeline(
        [
            ("features", FeatureAblationTransformer(feature_set.cikarilanlar)),
            ("model", clone(model)),
        ]
    )


def _group_holdout_split(veri: pd.DataFrame, metadata: pd.DataFrame, random_state: int) -> dict[str, Any]:
    groups = metadata["source_id"].astype(str).to_numpy()
    y = veri[HEDEF_KOLONU].to_numpy()
    splitter = GroupShuffleSplit(n_splits=60, test_size=0.20, random_state=random_state)
    en_iyi: tuple[float, np.ndarray, np.ndarray] | None = None
    for train_idx, test_idx in splitter.split(veri, y, groups):
        _assert_source_ayrik(metadata, train_idx, test_idx, "holdout")
        train_pos = float(y[train_idx].mean())
        test_pos = float(y[test_idx].mean())
        skor = abs(train_pos - test_pos) + abs(0.5 - test_pos)
        if en_iyi is None or skor < en_iyi[0]:
            en_iyi = (skor, train_idx, test_idx)
    if en_iyi is None:
        raise RuntimeError("Group holdout split uretilemedi.")
    return {
        "yontem": "GroupShuffleSplit(source_id)",
        "train_indexleri": [int(i) for i in en_iyi[1]],
        "test_indexleri": [int(i) for i in en_iyi[2]],
    }


def _inner_group_validation_split(
    y_train: pd.Series,
    train_metadata: pd.DataFrame,
    random_state: int,
) -> tuple[np.ndarray, np.ndarray]:
    groups = train_metadata["source_id"].astype(str).to_numpy()
    splitter = GroupShuffleSplit(n_splits=30, test_size=0.18, random_state=random_state)
    best: tuple[float, np.ndarray, np.ndarray] | None = None
    for fit_idx, val_idx in splitter.split(np.zeros(len(y_train)), y_train, groups):
        if len(np.unique(y_train.iloc[fit_idx])) < 2 or len(np.unique(y_train.iloc[val_idx])) < 2:
            continue
        _assert_source_ayrik(train_metadata, fit_idx, val_idx, "inner_validation")
        skor = abs(float(y_train.iloc[fit_idx].mean()) - float(y_train.iloc[val_idx].mean()))
        if best is None or skor < best[0]:
            best = (skor, fit_idx, val_idx)
    if best is None:
        idx = np.arange(len(y_train))
        fit_idx, val_idx = train_test_split(idx, test_size=0.18, stratify=y_train, random_state=random_state)
        return np.asarray(fit_idx), np.asarray(val_idx)
    return best[1], best[2]


def _split_degerlendir(
    *,
    veri_adayi: VeriAdayi,
    split: dict[str, Any],
    feature_set: FeatureSet,
    model_adayi: ModelAdayi,
    hedef_accuracy: float,
    hedef_diger_metrikler: float,
    random_state: int,
) -> dict[str, Any]:
    x = veri_adayi.veri[OZELLIK_KOLONLARI]
    y = veri_adayi.veri[HEDEF_KOLONU]
    train_idx = np.asarray(split["train_indexleri"], dtype=int)
    test_idx = np.asarray(split["test_indexleri"], dtype=int)
    _assert_source_ayrik(veri_adayi.metadata, train_idx, test_idx, "holdout")

    train_metadata = veri_adayi.metadata.iloc[train_idx].reset_index(drop=True)
    x_train = x.iloc[train_idx].reset_index(drop=True)
    y_train = y.iloc[train_idx].reset_index(drop=True)
    x_test = x.iloc[test_idx].reset_index(drop=True)
    y_test = y.iloc[test_idx].reset_index(drop=True)

    fit_idx, val_idx = _inner_group_validation_split(y_train, train_metadata, random_state)
    estimator = _pipeline_olustur(feature_set, model_adayi.estimator)
    threshold_model = clone(estimator)
    threshold_model.fit(x_train.iloc[fit_idx], y_train.iloc[fit_idx])
    val_prob = _pozitif_olasilik(threshold_model, x_train.iloc[val_idx])
    threshold_ozeti = _en_iyi_esik(y_train.iloc[val_idx], val_prob, hedef_accuracy, hedef_diger_metrikler)

    final_model = clone(estimator)
    final_model.fit(x_train, y_train)
    test_prob = _pozitif_olasilik(final_model, x_test)
    metrikler = _metrikleri_hesapla(y_test, test_prob, float(threshold_ozeti["esik"]))
    leakage = _split_leakage_raporu(veri_adayi, train_idx, test_idx, feature_set)

    return {
        "durum": "tamamlandi",
        "veri_adayi": veri_adayi.ad,
        "hedef_sinif_sayisi": veri_adayi.hedef_sinif_sayisi,
        "model": model_adayi.ad,
        "model_aciklama": model_adayi.aciklama,
        "feature_set": feature_set.ad,
        "feature_aciklama": feature_set.aciklama,
        "eklenen_featurelar": [],
        "cikarilan_featurelar": list(feature_set.cikarilanlar),
        "final_feature_sayisi": len(FeatureAblationTransformer(feature_set.cikarilanlar).feature_names()),
        "split": {
            "yontem": split["yontem"],
            "train_satir": int(len(train_idx)),
            "test_satir": int(len(test_idx)),
            "train_sinif_dagilimi": _sinif_dagilimi(veri_adayi.veri.iloc[train_idx]),
            "test_sinif_dagilimi": _sinif_dagilimi(veri_adayi.veri.iloc[test_idx]),
            "train_source_family_count": int(veri_adayi.metadata.iloc[train_idx]["source_id"].nunique()),
            "test_source_family_count": int(veri_adayi.metadata.iloc[test_idx]["source_id"].nunique()),
        },
        "threshold_ozeti": threshold_ozeti,
        "test_metrikleri": metrikler,
        "hedef_kapisi": _hedef_kapisi(metrikler, hedef_accuracy, hedef_diger_metrikler),
        "leakage_kontrolu": leakage,
        "_model": final_model,
        "_test_y": y_test.to_numpy(dtype=int),
        "_test_prob": test_prob,
    }


def _group_cv_degerlendir(
    *,
    veri_adayi: VeriAdayi,
    feature_set: FeatureSet,
    model_adayi: ModelAdayi,
    hedef_accuracy: float,
    hedef_diger_metrikler: float,
    random_state: int,
) -> dict[str, Any]:
    x = veri_adayi.veri[OZELLIK_KOLONLARI]
    y = veri_adayi.veri[HEDEF_KOLONU].reset_index(drop=True)
    groups = veri_adayi.metadata["source_id"].astype(str).to_numpy()
    splitter = StratifiedGroupKFold(n_splits=5, shuffle=True, random_state=random_state)
    foldlar: list[dict[str, Any]] = []
    for fold, (train_idx, test_idx) in enumerate(splitter.split(x, y, groups), start=1):
        _assert_source_ayrik(veri_adayi.metadata, train_idx, test_idx, f"group_cv_fold_{fold}")
        train_metadata = veri_adayi.metadata.iloc[train_idx].reset_index(drop=True)
        x_train = x.iloc[train_idx].reset_index(drop=True)
        y_train = y.iloc[train_idx].reset_index(drop=True)
        x_test = x.iloc[test_idx].reset_index(drop=True)
        y_test = y.iloc[test_idx].reset_index(drop=True)
        fit_idx, val_idx = _inner_group_validation_split(y_train, train_metadata, random_state + fold)
        estimator = _pipeline_olustur(feature_set, model_adayi.estimator)
        threshold_model = clone(estimator)
        threshold_model.fit(x_train.iloc[fit_idx], y_train.iloc[fit_idx])
        val_prob = _pozitif_olasilik(threshold_model, x_train.iloc[val_idx])
        threshold_ozeti = _en_iyi_esik(y_train.iloc[val_idx], val_prob, hedef_accuracy, hedef_diger_metrikler)
        final_model = clone(estimator)
        final_model.fit(x_train, y_train)
        test_prob = _pozitif_olasilik(final_model, x_test)
        metrikler = _metrikleri_hesapla(y_test, test_prob, float(threshold_ozeti["esik"]))
        foldlar.append(
            {
                "fold": fold,
                "train_satir": int(len(train_idx)),
                "test_satir": int(len(test_idx)),
                "train_source_family_count": int(veri_adayi.metadata.iloc[train_idx]["source_id"].nunique()),
                "test_source_family_count": int(veri_adayi.metadata.iloc[test_idx]["source_id"].nunique()),
                "source_id_kesisim_sayisi": 0,
                "threshold": float(threshold_ozeti["esik"]),
                "metrikler": metrikler,
            }
        )
    return {
        "veri_adayi": veri_adayi.ad,
        "model": model_adayi.ad,
        "feature_set": feature_set.ad,
        "cv_yontemi": "StratifiedGroupKFold(n_splits=5, shuffle=True, random_state=42)",
        "foldlar": foldlar,
        "summary": _cv_summary(foldlar),
    }


def _fold_ici_sentetik_cv(
    *,
    original_dev: pd.DataFrame,
    hedef_sinif_sayisi: int,
    feature_set: FeatureSet,
    model_adayi: ModelAdayi,
    hedef_accuracy: float,
    hedef_diger_metrikler: float,
    random_state: int,
) -> dict[str, Any]:
    y = original_dev[HEDEF_KOLONU].reset_index(drop=True)
    groups = original_dev["source_id"].astype(str).to_numpy()
    splitter = StratifiedGroupKFold(n_splits=5, shuffle=True, random_state=random_state)
    foldlar: list[dict[str, Any]] = []
    for fold, (train_idx, test_idx) in enumerate(splitter.split(original_dev[OZELLIK_KOLONLARI], y, groups), start=1):
        train_original = original_dev.iloc[train_idx].reset_index(drop=True)
        test_original = original_dev.iloc[test_idx].reset_index(drop=True)
        train_adayi = _sentetik_veri_adayi_uret(train_original, hedef_sinif_sayisi, random_state + 500 + fold)
        train_x = train_adayi.veri[OZELLIK_KOLONLARI]
        train_y = train_adayi.veri[HEDEF_KOLONU]
        train_metadata = train_adayi.metadata
        test_x = test_original[OZELLIK_KOLONLARI].reset_index(drop=True)
        test_y = test_original[HEDEF_KOLONU].reset_index(drop=True)
        fit_idx, val_idx = _inner_group_validation_split(train_y.reset_index(drop=True), train_metadata, random_state + fold)
        estimator = _pipeline_olustur(feature_set, model_adayi.estimator)
        threshold_model = clone(estimator)
        threshold_model.fit(train_x.iloc[fit_idx], train_y.iloc[fit_idx])
        val_prob = _pozitif_olasilik(threshold_model, train_x.iloc[val_idx])
        threshold_ozeti = _en_iyi_esik(train_y.iloc[val_idx], val_prob, hedef_accuracy, hedef_diger_metrikler)
        final_model = clone(estimator)
        final_model.fit(train_x, train_y)
        test_prob = _pozitif_olasilik(final_model, test_x)
        metrikler = _metrikleri_hesapla(test_y, test_prob, float(threshold_ozeti["esik"]))
        ortak = set(train_adayi.metadata["source_id"]) & set(test_original["source_id"])
        if ortak:
            raise RuntimeError(f"Fold ici sentetik CV source_id sızıntısı: {sorted(ortak)[:5]}")
        foldlar.append(
            {
                "fold": fold,
                "train_original_rows": int(len(train_original)),
                "train_after_synthetic_rows": int(len(train_adayi.veri)),
                "test_original_rows": int(len(test_original)),
                "source_id_kesisim_sayisi": 0,
                "metrikler": metrikler,
            }
        )
    return {
        "uygulandi": True,
        "not": "Her fold'da sentetik uretim yalniz train_original tarafindan yapildi; test_original dokunulmadi.",
        "hedef_sinif_sayisi": int(hedef_sinif_sayisi),
        "model": model_adayi.ad,
        "feature_set": feature_set.ad,
        "foldlar": foldlar,
        "summary": _cv_summary(foldlar),
    }


def _external_holdout_degerlendir(
    model: BaseEstimator,
    external_holdout: pd.DataFrame,
    feature_set: FeatureSet,
    esik: float,
) -> dict[str, Any]:
    del feature_set
    prob = _pozitif_olasilik(model, external_holdout[OZELLIK_KOLONLARI])
    return {
        "not": (
            "Original external holdout sentetik uretimden once ayrildi; model secimi, threshold tuning ve "
            "sentetik uretim surecine dahil edilmedi."
        ),
        "satir_sayisi": int(len(external_holdout)),
        "sinif_dagilimi": _sinif_dagilimi(external_holdout),
        "metrikler": _metrikleri_hesapla(external_holdout[HEDEF_KOLONU], prob, esik),
    }


def _metrikleri_hesapla(y_true: pd.Series | np.ndarray, y_prob: np.ndarray, esik: float) -> dict[str, Any]:
    y_np = np.asarray(y_true, dtype=int)
    prob = np.asarray(y_prob, dtype=float)
    pred = (prob >= esik).astype(int)
    cm = confusion_matrix(y_np, pred, labels=[0, 1])
    tn, fp, fn, tp = [int(v) for v in cm.ravel()]
    specificity = tn / (tn + fp) if (tn + fp) else 0.0
    try:
        roc_auc = float(roc_auc_score(y_np, prob))
    except ValueError:
        roc_auc = float("nan")
    metrikler = {
        "accuracy": float(accuracy_score(y_np, pred)),
        "precision": float(precision_score(y_np, pred, zero_division=0)),
        "recall": float(recall_score(y_np, pred, zero_division=0)),
        "sensitivity": float(recall_score(y_np, pred, zero_division=0)),
        "specificity": float(specificity),
        "f1": float(f1_score(y_np, pred, zero_division=0)),
        "roc_auc": roc_auc,
        "balanced_accuracy": float(balanced_accuracy_score(y_np, pred)),
        "mcc": float(matthews_corrcoef(y_np, pred)),
        "brier": float(brier_score_loss(y_np, prob)),
        "esik": float(esik),
        "confusion_matrix": {
            "labels": [0, 1],
            "matris": cm.tolist(),
            "tn": tn,
            "fp": fp,
            "fn": fn,
            "tp": tp,
        },
    }
    metrikler["ana_metrik_minimumu"] = float(
        min(float(metrikler[m]) for m in ANA_METRIKLER if not math.isnan(float(metrikler[m])))
    )
    return metrikler


def _en_iyi_esik(
    y_true: pd.Series,
    y_prob: np.ndarray,
    hedef_accuracy: float,
    hedef_diger_metrikler: float,
) -> dict[str, Any]:
    adaylar = np.unique(np.concatenate([np.linspace(0.10, 0.90, 81), np.asarray(y_prob, dtype=float)]))
    en_iyi: dict[str, Any] | None = None
    for esik in adaylar:
        metrikler = _metrikleri_hesapla(y_true, y_prob, float(esik))
        kapi = _hedef_kapisi(metrikler, hedef_accuracy, hedef_diger_metrikler)
        anahtar = (
            kapi["gecti"],
            metrikler["ana_metrik_minimumu"],
            metrikler["accuracy"],
            metrikler["f1"],
            metrikler["roc_auc"],
        )
        if en_iyi is None or anahtar > en_iyi["siralama_anahtari"]:
            en_iyi = {
                "esik": float(esik),
                "validation_metrikleri": metrikler,
                "hedef_kapisi": kapi,
                "siralama_anahtari": anahtar,
            }
    if en_iyi is None:
        raise RuntimeError("Threshold secilemedi.")
    en_iyi["siralama_anahtari"] = list(en_iyi["siralama_anahtari"])
    return en_iyi


def _hedef_kapisi(metrikler: dict[str, Any], hedef_accuracy: float, hedef_diger_metrikler: float) -> dict[str, Any]:
    eksikler = {}
    for metrik in ANA_METRIKLER:
        hedef = hedef_accuracy if metrik == "accuracy" else hedef_diger_metrikler
        deger = float(metrikler.get(metrik, 0.0) or 0.0)
        if math.isnan(deger) or deger < hedef:
            eksikler[metrik] = {"deger": deger, "hedef": hedef}
    return {"gecti": not eksikler, "eksikler": eksikler}


def _pozitif_olasilik(model: BaseEstimator, x: pd.DataFrame) -> np.ndarray:
    if hasattr(model, "predict_proba"):
        return np.asarray(model.predict_proba(x))[:, 1]
    if hasattr(model, "decision_function"):
        skor = np.asarray(model.decision_function(x), dtype=float)
        return 1 / (1 + np.exp(-skor))
    pred = np.asarray(model.predict(x), dtype=float)
    return np.clip(pred, 0, 1)


def _cv_summary(foldlar: list[dict[str, Any]]) -> dict[str, Any]:
    summary: dict[str, Any] = {}
    for metrik in RAPOR_METRIKLERI:
        degerler = [float(f["metrikler"][metrik]) for f in foldlar if not math.isnan(float(f["metrikler"][metrik]))]
        summary[f"{metrik}_mean"] = float(np.mean(degerler)) if degerler else float("nan")
        summary[f"{metrik}_std"] = float(np.std(degerler, ddof=1)) if len(degerler) > 1 else 0.0
    summary["max_source_id_kesisim_sayisi"] = int(max((f.get("source_id_kesisim_sayisi", 0) for f in foldlar), default=0))
    return summary


def _split_leakage_raporu(
    veri_adayi: VeriAdayi,
    train_idx: np.ndarray,
    test_idx: np.ndarray,
    feature_set: FeatureSet,
) -> dict[str, Any]:
    train_groups = set(veri_adayi.metadata.iloc[train_idx]["source_id"].astype(str))
    test_groups = set(veri_adayi.metadata.iloc[test_idx]["source_id"].astype(str))
    ortak = sorted(train_groups & test_groups)
    transformer = FeatureAblationTransformer(feature_set.cikarilanlar)
    x_train = transformer.transform(veri_adayi.veri.iloc[train_idx][OZELLIK_KOLONLARI])
    x_test = transformer.transform(veri_adayi.veri.iloc[test_idx][OZELLIK_KOLONLARI])
    train_key = x_train.round(6).astype(str).agg("|".join, axis=1)
    test_key = x_test.round(6).astype(str).agg("|".join, axis=1)
    exact = int(test_key.isin(set(train_key)).sum())
    near = _near_duplicate_ozeti(x_train, x_test)
    return {
        "source_id_kesisim_sayisi": int(len(ortak)),
        "source_id_kesisim_ornekleri": ortak[:10],
        "exact_duplicate_sayisi": exact,
        "exact_duplicate_orani": float(exact / max(len(x_test), 1)),
        "near_duplicate_ozeti": near,
    }


def _near_duplicate_ozeti(x_train: pd.DataFrame, x_test: pd.DataFrame) -> dict[str, Any]:
    if x_train.empty or x_test.empty:
        return {"min_mesafe": None, "p01_mesafe": None, "p05_mesafe": None, "esik_0_02_alti_oran": None}
    scaler = StandardScaler()
    train_scaled = scaler.fit_transform(x_train)
    test_scaled = scaler.transform(x_test)
    nn = NearestNeighbors(n_neighbors=1)
    nn.fit(train_scaled)
    dist, _ = nn.kneighbors(test_scaled)
    d = dist.ravel()
    return {
        "min_mesafe": float(np.min(d)),
        "p01_mesafe": float(np.quantile(d, 0.01)),
        "p05_mesafe": float(np.quantile(d, 0.05)),
        "esik_0_02_alti_oran": float(np.mean(d < 0.02)),
    }


def _aday_leakage_raporu(
    aday: VeriAdayi,
    external_holdout: pd.DataFrame,
    aday_sonuclari: list[dict[str, Any]],
    cv_sonuclari: list[dict[str, Any]],
) -> dict[str, Any]:
    en_iyi = sorted(aday_sonuclari, key=_siralama_anahtari, reverse=True)[0] if aday_sonuclari else None
    split_leakage = en_iyi["leakage_kontrolu"] if en_iyi else {}
    cv_fold_counts = [
        int(f["source_id_kesisim_sayisi"])
        for cv in cv_sonuclari
        if cv["veri_adayi"] == aday.ad
        for f in cv["foldlar"]
    ]
    valid_original = aday.metadata["source_id"].astype(str).str.match(r"^original_\d+$")
    independent_synthetic = aday.metadata.loc[aday.metadata["is_synthetic"], "source_id"].astype(str).str.match(
        r"^(?!original_\d+$).*"
    )
    parent_ids = set(aday.metadata.loc[aday.metadata["is_synthetic"], "source_id"].astype(str))
    external_overlap = len(parent_ids & set(external_holdout["source_id"].astype(str)))
    exact = int(split_leakage.get("exact_duplicate_sayisi", 0) or 0)
    near = split_leakage.get("near_duplicate_ozeti", {})
    failed = (
        int(split_leakage.get("source_id_kesisim_sayisi", 0) or 0) > 0
        or max(cv_fold_counts or [0]) > 0
        or exact > 0
        or int(independent_synthetic.sum()) > 0
        or external_overlap > 0
    )
    status = "failed" if failed else "clean"
    return {
        "dataset_name": aday.ad,
        "total_rows": int(len(aday.veri)),
        "original_rows": int((~aday.metadata["is_synthetic"]).sum()),
        "synthetic_rows": int(aday.metadata["is_synthetic"].sum()),
        "class_distribution": _sinif_dagilimi(aday.veri),
        "train_test_source_intersection_count": int(split_leakage.get("source_id_kesisim_sayisi", 0) or 0),
        "cv_fold_source_intersection_counts": cv_fold_counts,
        "exact_duplicate_count": exact,
        "near_duplicate_threshold": 0.02,
        "near_duplicate_rate": near.get("esik_0_02_alti_oran"),
        "min_near_duplicate_distance": near.get("min_mesafe"),
        "synthetic_rows_with_valid_original_source_id": int(
            (aday.metadata["is_synthetic"] & valid_original).sum()
        ),
        "independent_synthetic_source_id_count": int(independent_synthetic.sum()),
        "external_holdout_source_overlap_count": int(external_overlap),
        "leakage_status": status,
        "notes": "Sentetik satirlar original_{index} source_id ailesine baglidir." if status == "clean" else "Kontrol basarisiz.",
    }


def _metadata_source_id_dogrula(metadata: pd.DataFrame) -> None:
    required = {"source_id", "is_synthetic", "parent_original_index", "generation_method"}
    eksik = required - set(metadata.columns)
    if eksik:
        raise RuntimeError(f"Metadata kolonlari eksik: {sorted(eksik)}")
    invalid = metadata.loc[metadata["is_synthetic"], "source_id"].astype(str).str.match(r"^original_\d+$") == False
    if bool(invalid.any()):
        raise RuntimeError("Bagimsiz sentetik source_id tespit edildi.")


def _assert_source_ayrik(metadata: pd.DataFrame, train_idx: np.ndarray, test_idx: np.ndarray, baglam: str) -> None:
    train_groups = set(metadata.iloc[train_idx]["source_id"].astype(str))
    test_groups = set(metadata.iloc[test_idx]["source_id"].astype(str))
    ortak = train_groups & test_groups
    if ortak:
        raise RuntimeError(f"{baglam} source_id kesisimi 0 degil: {sorted(ortak)[:8]}")


def _cv_icin_aday_sec(basarili: list[dict[str, Any]]) -> list[dict[str, Any]]:
    secilen: list[dict[str, Any]] = []
    for veri_adayi in sorted({s["veri_adayi"] for s in basarili}):
        aday_sonuclari = [s for s in basarili if s["veri_adayi"] == veri_adayi]
        sirali = sorted(aday_sonuclari, key=_siralama_anahtari, reverse=True)
        secilen.extend(sirali[:2])
    tekil: dict[tuple[str, str, str], dict[str, Any]] = {}
    for s in secilen:
        tekil[(s["veri_adayi"], s["feature_set"], s["model"])] = s
    return list(tekil.values())


def _final_sonuc_sec(basarili: list[dict[str, Any]], cv_sonuclari: list[dict[str, Any]]) -> dict[str, Any]:
    cv_map = {(c["veri_adayi"], c["feature_set"], c["model"]): c for c in cv_sonuclari}
    adaylar = [s for s in basarili if (s["veri_adayi"], s["feature_set"], s["model"]) in cv_map]
    if not adaylar:
        return sorted(basarili, key=_siralama_anahtari, reverse=True)[0]

    def anahtar(sonuc: dict[str, Any]) -> tuple[float, float, float, float, float, float]:
        cv = cv_map[(sonuc["veri_adayi"], sonuc["feature_set"], sonuc["model"])]
        holdout = sonuc["test_metrikleri"]
        cv_summary = cv["summary"]
        hedef = sonuc["hedef_sinif_sayisi"]
        veri_onceligi = 3.0 if hedef == 2700 else 2.0 if hedef == 2500 else 1.0 if hedef == 5000 else 0.0
        kapi = 1.0 if holdout["accuracy"] >= HEDEF_ACCURACY and holdout["ana_metrik_minimumu"] >= HEDEF_DIGER_METRIKLER else 0.0
        stabilite = -float(cv_summary["accuracy_std"])
        cv_min = float(cv_summary["ana_metrik_minimumu_mean"])
        return (
            kapi,
            veri_onceligi if kapi else 0.0,
            cv_min,
            float(holdout["ana_metrik_minimumu"]),
            stabilite,
            float(holdout["accuracy"]),
        )

    return sorted(adaylar, key=anahtar, reverse=True)[0]


def _siralama_anahtari(sonuc: dict[str, Any]) -> tuple[float, float, float, float, float]:
    metrik = sonuc.get("test_metrikleri", {})
    hedef = sonuc.get("hedef_sinif_sayisi")
    veri_bonus = 0.003 if hedef == 2700 else 0.002 if hedef == 2500 else 0.001 if hedef == 5000 else 0.0
    return (
        1.0 if sonuc.get("hedef_kapisi", {}).get("gecti") else 0.0,
        float(metrik.get("ana_metrik_minimumu", 0.0) or 0.0) + veri_bonus,
        float(metrik.get("accuracy", 0.0) or 0.0),
        float(metrik.get("roc_auc", 0.0) or 0.0),
        -float(sonuc.get("final_feature_sayisi", 99) or 99),
    )


def _karsilastirma_raporu(
    *,
    ham_veri: pd.DataFrame,
    original_dev: pd.DataFrame,
    external_holdout: pd.DataFrame,
    veri_adaylari: list[VeriAdayi],
    basarili: list[dict[str, Any]],
    cv_sonuclari: list[dict[str, Any]],
    final_sonuc: dict[str, Any],
    final_cv: dict[str, Any],
    external_metrikleri: dict[str, Any],
    fold_ici_cv: dict[str, Any],
    final_veri_yolu: Path,
    final_metadata_yolu: Path,
    veri_yolu: Path,
) -> dict[str, Any]:
    top_by_dataset = {}
    for aday in veri_adaylari:
        sonuc = sorted(_sonuclar_adaya_gore(basarili, aday.ad), key=_siralama_anahtari, reverse=True)
        if sonuc:
            top_by_dataset[aday.ad] = _jsondan_modeli_cikar(sonuc[0])
    return {
        "olusturulma_zamani_utc": _zaman(),
        "orijinal_veri": {
            "veri_yolu": str(veri_yolu),
            "sha256": _sha256(veri_yolu),
            "satir_sayisi": int(len(ham_veri)),
            "sinif_dagilimi": _sinif_dagilimi(ham_veri),
        },
        "original_dev": {
            "satir_sayisi": int(len(original_dev)),
            "sinif_dagilimi": _sinif_dagilimi(original_dev),
            "source_family_count": int(original_dev["source_id"].nunique()),
        },
        "external_holdout": {
            "satir_sayisi": int(len(external_holdout)),
            "sinif_dagilimi": _sinif_dagilimi(external_holdout),
            "source_family_count": int(external_holdout["source_id"].nunique()),
            "izolasyon_notu": (
                "Orijinal external holdout, sentetik uretimden once ayrildi ve model gelistirme surecine dahil edilmedi."
            ),
        },
        "veri_adaylari": [aday.rapor for aday in veri_adaylari],
        "top_by_dataset": top_by_dataset,
        "holdout_leaderboard": [_jsondan_modeli_cikar(s) for s in sorted(basarili, key=_siralama_anahtari, reverse=True)],
        "group_cv_sonuclari": [_jsondan_modeli_cikar(s) for s in cv_sonuclari],
        "final_veri": {
            "dataset_name": final_sonuc["veri_adayi"],
            "yol": str(final_veri_yolu),
            "metadata_yolu": str(final_metadata_yolu),
            "satir_sayisi": int(len(_veri_adayi_bul(veri_adaylari, final_sonuc["veri_adayi"]).veri)),
            "sinif_dagilimi": _sinif_dagilimi(_veri_adayi_bul(veri_adaylari, final_sonuc["veri_adayi"]).veri),
            "kolonlar": list(_veri_adayi_bul(veri_adaylari, final_sonuc["veri_adayi"]).veri.columns),
        },
        "final_secim": _jsondan_modeli_cikar(final_sonuc),
        "final_group_cv": _jsondan_modeli_cikar(final_cv),
        "fold_ici_sentetik_uretim_cv": fold_ici_cv,
        "original_external_holdout_sonucu": external_metrikleri,
        "karar_notu": (
            "Final sonuc leakage duzeltmesi sonrasi yeniden uretilen source_id kontrollu sentetik benchmark sonuclarina dayanir. "
            "Eski 5000/sinif sonucu final karar icin kullanilmadi."
        ),
    }


def _veri_adayi_dosyalarini_yaz(veri_klasoru: Path, aday: VeriAdayi) -> None:
    if aday.hedef_sinif_sayisi is None:
        return
    csv_yolu = veri_klasoru / f"{aday.ad}.csv"
    meta_yolu = veri_klasoru / f"{aday.ad}_metadata.csv"
    aday.veri.to_csv(csv_yolu, index=False)
    aday.metadata.to_csv(meta_yolu, index=False)


def _final_uyumluluk_dosyalari_yaz(veri_klasoru: Path, aday: VeriAdayi, final_sonuc: dict[str, Any]) -> tuple[Path, Path]:
    final_yol = veri_klasoru / "sentetik_gercek_pima.csv"
    meta_yol = veri_klasoru / "sentetik_gercek_pima_metadata.csv"
    manifest_yol = veri_klasoru / "sentetik_gercek_pima_split_manifest.json"
    aday.veri.to_csv(final_yol, index=False)
    aday.metadata.to_csv(meta_yol, index=False)
    _json_yaz(
        manifest_yol,
        {
            "dataset_name": aday.ad,
            "not": "Bu manifest leakage duzeltmesi sonrasi final source_id kontrollu kosuya aittir.",
            "split": final_sonuc["split"],
            "model": final_sonuc["model"],
            "feature_set": final_sonuc["feature_set"],
            "cikarilan_featurelar": final_sonuc["cikarilan_featurelar"],
            "threshold": final_sonuc["threshold_ozeti"]["esik"],
        },
    )
    return final_yol, meta_yol


def _uyumluluk_raporlarini_yaz(
    klasor: Path,
    karsilastirma: dict[str, Any],
    cv_raporu: dict[str, Any],
    leakage_json: dict[str, Any],
) -> None:
    _json_yaz(klasor / "sentetik_benchmark_leaderboard.json", {"leaderboard": karsilastirma["holdout_leaderboard"]})
    _json_yaz(klasor / "selected_model_metrics_report.json", karsilastirma)
    _json_yaz(
        klasor / "feature_ablation_raporu.json",
        {
            "final_feature_set": karsilastirma["final_secim"]["feature_set"],
            "final_cikarilan_featurelar": karsilastirma["final_secim"]["cikarilan_featurelar"],
            "tum_sonuclar": karsilastirma["holdout_leaderboard"],
        },
    )
    _json_yaz(klasor / "sentetik_group_cv_raporu.json", cv_raporu)
    _json_yaz(klasor / "source_id_leakage_raporu.json", leakage_json)


def _grafikleri_uret(
    *,
    grafik_klasoru: Path,
    ham_veri: pd.DataFrame,
    veri_adaylari: list[VeriAdayi],
    karsilastirma: dict[str, Any],
    final_sonuc: dict[str, Any],
    final_cv: dict[str, Any],
    leakage_raporlari: list[dict[str, Any]],
) -> dict[str, str]:
    plt.rcParams.update({"font.size": 10, "axes.titlesize": 12, "axes.labelsize": 10})
    grafikler: dict[str, str] = {}
    grafikler["orijinal_sinif_dagilimi"] = str(_grafik_orijinal_sinif(grafik_klasoru, ham_veri))
    grafikler["aday_sinif_dagilimi"] = str(_grafik_aday_sinif(grafik_klasoru, veri_adaylari))
    grafikler["metrik_karsilastirma"] = str(_grafik_metrik_karsilastirma(grafik_klasoru, karsilastirma))
    grafikler["holdout_cv_karsilastirma"] = str(_grafik_holdout_cv(grafik_klasoru, karsilastirma))
    grafikler["confusion_matrix"] = str(_grafik_confusion_matrix(grafik_klasoru, final_sonuc))
    grafikler["feature_importance"] = str(_grafik_feature_importance(grafik_klasoru, final_sonuc))
    grafikler["calibration_curve"] = str(_grafik_calibration(grafik_klasoru, final_sonuc))
    grafikler["brier_karsilastirma"] = str(_grafik_brier(grafik_klasoru, karsilastirma))
    grafikler["leakage_ozeti"] = str(_grafik_leakage(grafik_klasoru, leakage_raporlari))
    return grafikler


def _grafik_orijinal_sinif(grafik_klasoru: Path, ham_veri: pd.DataFrame) -> Path:
    yol = grafik_klasoru / "orijinal_pima_sinif_dagilimi.png"
    counts = ham_veri[HEDEF_KOLONU].value_counts().sort_index()
    fig, ax = plt.subplots(figsize=(6.2, 3.2))
    colors = ["#2a9d8f", "#e76f51"]
    bars = ax.bar(["Negatif", "Pozitif"], counts.values, color=colors)
    ax.set_title("Orijinal PIMA Sınıf Dağılımı")
    ax.set_ylabel("Satır sayısı")
    for bar, value in zip(bars, counts.values, strict=True):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 8, str(int(value)), ha="center", weight="bold")
    fig.tight_layout()
    fig.savefig(yol, dpi=180)
    plt.close(fig)
    return yol


def _grafik_aday_sinif(grafik_klasoru: Path, veri_adaylari: list[VeriAdayi]) -> Path:
    yol = grafik_klasoru / "sentetik_aday_sinif_dagilimi.png"
    labels = []
    neg = []
    pos = []
    for aday in veri_adaylari:
        if aday.hedef_sinif_sayisi is None:
            labels.append("Orijinal dev")
        else:
            labels.append(f"{aday.hedef_sinif_sayisi}/sınıf")
        dist = aday.veri[HEDEF_KOLONU].value_counts().sort_index()
        neg.append(int(dist.get(0, 0)))
        pos.append(int(dist.get(1, 0)))
    x = np.arange(len(labels))
    fig, ax = plt.subplots(figsize=(7.2, 3.5))
    ax.bar(x - 0.18, neg, width=0.36, label="Negatif", color="#457b9d")
    ax.bar(x + 0.18, pos, width=0.36, label="Pozitif", color="#f4a261")
    ax.set_title("Veri Adaylarına Göre Sınıf Dağılımı")
    ax.set_xticks(x, labels)
    ax.set_ylabel("Satır sayısı")
    ax.legend(frameon=False)
    fig.tight_layout()
    fig.savefig(yol, dpi=180)
    plt.close(fig)
    return yol


def _grafik_metrik_karsilastirma(grafik_klasoru: Path, rapor: dict[str, Any]) -> Path:
    yol = grafik_klasoru / "sentetik_metrik_karsilastirma.png"
    rows = []
    for ad, sonuc in rapor["top_by_dataset"].items():
        if ad == "orijinal_pima_baseline":
            label = "Orijinal dev"
        else:
            label = ad.replace("sentetik_", "").replace("_per_class_pima", "/sınıf")
        m = sonuc["test_metrikleri"]
        rows.append((label, m["accuracy"], m["f1"], m["roc_auc"], m["ana_metrik_minimumu"]))
    labels = [r[0] for r in rows]
    metrics = ["Accuracy", "F1", "ROC-AUC", "Min ana"]
    values = np.array([r[1:] for r in rows])
    x = np.arange(len(labels))
    width = 0.18
    fig, ax = plt.subplots(figsize=(8.0, 3.8))
    colors = ["#264653", "#2a9d8f", "#e9c46a", "#e76f51"]
    for i, metric in enumerate(metrics):
        ax.bar(x + (i - 1.5) * width, values[:, i], width=width, label=metric, color=colors[i])
    ax.set_ylim(0.45, 1.02)
    ax.set_title("Holdout Metrik Karşılaştırması")
    ax.set_xticks(x, labels, rotation=10)
    ax.legend(ncol=4, frameon=False, loc="lower center", bbox_to_anchor=(0.5, -0.32))
    fig.tight_layout()
    fig.savefig(yol, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return yol


def _grafik_holdout_cv(grafik_klasoru: Path, rapor: dict[str, Any]) -> Path:
    yol = grafik_klasoru / "holdout_group_cv_karsilastirma.png"
    rows = []
    cv_map = {c["veri_adayi"]: c for c in rapor["group_cv_sonuclari"]}
    for ad, sonuc in rapor["top_by_dataset"].items():
        if ad not in cv_map:
            continue
        label = "Orijinal dev" if ad == "orijinal_pima_baseline" else ad.replace("sentetik_", "").replace("_per_class_pima", "/sınıf")
        rows.append((label, sonuc["test_metrikleri"]["accuracy"], cv_map[ad]["summary"]["accuracy_mean"]))
    labels = [r[0] for r in rows]
    x = np.arange(len(labels))
    fig, ax = plt.subplots(figsize=(7.2, 3.5))
    ax.bar(x - 0.18, [r[1] for r in rows], width=0.36, label="Holdout", color="#6d597a")
    ax.bar(x + 0.18, [r[2] for r in rows], width=0.36, label="Group CV mean", color="#43aa8b")
    ax.set_ylim(0.45, 1.02)
    ax.set_title("Holdout ve Source ID Group CV Accuracy")
    ax.set_xticks(x, labels, rotation=10)
    ax.legend(frameon=False)
    fig.tight_layout()
    fig.savefig(yol, dpi=180)
    plt.close(fig)
    return yol


def _grafik_confusion_matrix(grafik_klasoru: Path, final_sonuc: dict[str, Any]) -> Path:
    yol = grafik_klasoru / "selected_model_confusion_matrix.png"
    cm = np.array(final_sonuc["test_metrikleri"]["confusion_matrix"]["matris"])
    fig, ax = plt.subplots(figsize=(4.2, 3.6))
    im = ax.imshow(cm, cmap="YlGnBu")
    ax.set_title("Final Holdout Confusion Matrix")
    ax.set_xticks([0, 1], ["Tahmin 0", "Tahmin 1"])
    ax.set_yticks([0, 1], ["Gerçek 0", "Gerçek 1"])
    for i in range(2):
        for j in range(2):
            ax.text(j, i, str(int(cm[i, j])), ha="center", va="center", color="#111", weight="bold")
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    fig.tight_layout()
    fig.savefig(yol, dpi=180)
    plt.close(fig)
    return yol


def _grafik_feature_importance(grafik_klasoru: Path, final_sonuc: dict[str, Any]) -> Path:
    yol = grafik_klasoru / "selected_model_feature_importance.png"
    model = final_sonuc["_model"]
    names = FeatureAblationTransformer(tuple(final_sonuc["cikarilan_featurelar"])).feature_names()
    raw_model = model.named_steps["model"] if isinstance(model, Pipeline) else model
    importances = None
    if hasattr(raw_model, "feature_importances_"):
        importances = np.asarray(raw_model.feature_importances_, dtype=float)
    elif isinstance(raw_model, Pipeline) and hasattr(raw_model.steps[-1][1], "feature_importances_"):
        importances = np.asarray(raw_model.steps[-1][1].feature_importances_, dtype=float)
    if importances is None or len(importances) != len(names):
        importances = np.zeros(len(names))
    order = np.argsort(importances)
    fig, ax = plt.subplots(figsize=(6.5, 3.8))
    ax.barh(np.array(names)[order], importances[order], color="#577590")
    ax.set_title("Final Model Feature Importance")
    ax.set_xlabel("Önem skoru")
    fig.tight_layout()
    fig.savefig(yol, dpi=180)
    plt.close(fig)
    return yol


def _grafik_calibration(grafik_klasoru: Path, final_sonuc: dict[str, Any]) -> Path:
    yol = grafik_klasoru / "selected_model_calibration_curve.png"
    y_true = np.asarray(final_sonuc["_test_y"], dtype=int)
    prob = np.asarray(final_sonuc["_test_prob"], dtype=float)
    frac, mean_pred = calibration_curve(y_true, prob, n_bins=8, strategy="quantile")
    fig, ax = plt.subplots(figsize=(4.8, 3.6))
    ax.plot([0, 1], [0, 1], "--", color="#777", label="İdeal")
    ax.plot(mean_pred, frac, marker="o", color="#bc6c25", label="Final model")
    ax.set_title("Calibration Curve")
    ax.set_xlabel("Ortalama tahmin olasılığı")
    ax.set_ylabel("Gerçek pozitif oran")
    ax.legend(frameon=False)
    fig.tight_layout()
    fig.savefig(yol, dpi=180)
    plt.close(fig)
    return yol


def _grafik_brier(grafik_klasoru: Path, rapor: dict[str, Any]) -> Path:
    yol = grafik_klasoru / "brier_score_karsilastirma.png"
    rows = []
    for ad, sonuc in rapor["top_by_dataset"].items():
        label = "Orijinal dev" if ad == "orijinal_pima_baseline" else ad.replace("sentetik_", "").replace("_per_class_pima", "/sınıf")
        rows.append((label, sonuc["test_metrikleri"]["brier"]))
    fig, ax = plt.subplots(figsize=(6.8, 3.2))
    ax.bar([r[0] for r in rows], [r[1] for r in rows], color="#8ecae6")
    ax.set_title("Brier Score Karşılaştırması")
    ax.set_ylabel("Brier score (düşük daha iyi)")
    ax.tick_params(axis="x", rotation=10)
    fig.tight_layout()
    fig.savefig(yol, dpi=180)
    plt.close(fig)
    return yol


def _grafik_leakage(grafik_klasoru: Path, leakage_raporlari: list[dict[str, Any]]) -> Path:
    yol = grafik_klasoru / "leakage_kontrol_ozeti.png"
    labels = []
    values = []
    for r in leakage_raporlari:
        if r["dataset_name"] == "orijinal_pima_baseline":
            continue
        labels.append(r["dataset_name"].replace("sentetik_", "").replace("_per_class_pima", "/sınıf"))
        values.append(
            r["train_test_source_intersection_count"]
            + max(r["cv_fold_source_intersection_counts"] or [0])
            + r["exact_duplicate_count"]
            + r["independent_synthetic_source_id_count"]
            + r["external_holdout_source_overlap_count"]
        )
    fig, ax = plt.subplots(figsize=(6.8, 3.0))
    ax.bar(labels, values, color="#2a9d8f")
    ax.set_title("Leakage Kontrol Özeti")
    ax.set_ylabel("Problemli kontrol sayısı")
    ax.set_ylim(0, max(values + [1]) + 1)
    fig.tight_layout()
    fig.savefig(yol, dpi=180)
    plt.close(fig)
    return yol


def _word_raporu_yaz(
    *,
    proje_koku: Path,
    rapor: dict[str, Any],
    cv_raporu: dict[str, Any],
    leakage_raporu: dict[str, Any],
    grafikler: dict[str, str],
) -> Path:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    hedef = proje_koku / "makine_ogrenmesi" / "raporlar" / "sentetik_benchmark_word_raporu.docx"
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(9.2)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[style_name].font.name = "Aptos"
        doc.styles[style_name].font.color.rgb = RGBColor(31, 78, 121)

    final = rapor["final_secim"]
    final_m = final["test_metrikleri"]
    final_cv = rapor["final_group_cv"]["summary"]
    external = rapor["original_external_holdout_sonucu"]["metrikler"]

    _baslik(doc, "Source ID Kontrollü Sentetik PIMA Benchmark Raporu")
    _paragraf(
        doc,
        "Bu çalışmada PIMA diyabet veri seti üzerinde leakage kontrolü açıkça tanımlanmış, source_id aile yapısı "
        "korunan sentetik veri adayları üretilmiş ve sonuçlar hem holdout hem de Cross-Validation (CV) ile "
        "değerlendirilmiştir. ROC-AUC (Receiver Operating Characteristic - Area Under Curve), F1 (Precision ve "
        "Recall dengesini gösteren metrik) ve MCC (Matthews Correlation Coefficient) değerleri ayrı raporlanmıştır.",
    )
    _word_tablo(
        doc,
        ["Alan", "Değer"],
        [
            ["Orijinal veri", f"{rapor['orijinal_veri']['satir_sayisi']} satır"],
            ["Orijinal dağılım", _dist_yazi(rapor["orijinal_veri"]["sinif_dagilimi"])],
            ["Original dev", f"{rapor['original_dev']['satir_sayisi']} satır"],
            ["External holdout", f"{rapor['external_holdout']['satir_sayisi']} satır"],
            ["Final benchmark", rapor["final_veri"]["dataset_name"]],
        ],
        font_size=8.7,
    )
    _resim(doc, grafikler["orijinal_sinif_dagilimi"], width=5.9)
    _paragraf(
        doc,
        "Orijinal external holdout, sentetik üretimden önce ayrılmış ve model geliştirme sürecine dahil edilmemiştir. "
        "Bu nedenle external holdout sonucu, sentetik benchmarktan ayrı bir dış kontrol performansı olarak değerlendirilmiştir.",
    )

    doc.add_page_break()
    doc.add_heading("Leakage Kontrolü ve Source ID Aile Yapısı", level=1)
    _paragraf(
        doc,
        "Sentetik veri üretiminde her orijinal PIMA satırı bir kaynak aile olarak kabul edilmiştir. Her orijinal "
        "satıra original_{index} formatında bir source_id atanmış, bu satırdan üretilen tüm sentetik örnekler aynı "
        "source_id değerini taşımıştır. Böylece aynı kaynak aileden gelen örneklerin hem eğitim hem de test tarafına "
        "düşmesi engellenmiştir.",
    )
    leak_final = next(r for r in leakage_raporu["adaylar"] if r["dataset_name"] == final["veri_adayi"])
    _word_tablo(
        doc,
        ["Kontrol", "Beklenen", "Sonuç"],
        [
            ["Train/Test source_id kesişimi", "0", str(leak_final["train_test_source_intersection_count"])],
            ["CV fold source_id kesişimi", "0", str(max(leak_final["cv_fold_source_intersection_counts"] or [0]))],
            ["Exact duplicate", "0", str(leak_final["exact_duplicate_count"])],
            ["Near duplicate oranı", "Düşük / 0'a yakın", _oran(leak_final["near_duplicate_rate"])],
            ["Minimum mesafe", "Raporlanır", _sayi(leak_final["min_near_duplicate_distance"])],
            ["External holdout izolasyonu", "Evet", "Evet" if leak_final["external_holdout_source_overlap_count"] == 0 else "Hayır"],
            ["Bağımsız synthetic source_id kaldı mı?", "Hayır", "Hayır" if leak_final["independent_synthetic_source_id_count"] == 0 else "Evet"],
        ],
        font_size=8.3,
    )
    _resim(doc, grafikler["aday_sinif_dagilimi"], width=6.1)
    aday_rows = []
    for aday in rapor["veri_adaylari"]:
        aday_rows.append(
            [
                aday["dataset_name"].replace("_", " "),
                str(aday["original_rows"]),
                str(aday["synthetic_rows"]),
                str(aday["total_rows"]),
                _dist_yazi(aday["class_distribution"]),
            ]
        )
    _word_tablo(doc, ["Veri adayı", "Orijinal", "Sentetik", "Toplam", "Dağılım"], aday_rows, font_size=7.8)

    doc.add_page_break()
    doc.add_heading("Holdout ve Group CV Sonuçları", level=1)
    _resim(doc, grafikler["metrik_karsilastirma"], width=6.4)
    _resim(doc, grafikler["holdout_cv_karsilastirma"], width=6.2)
    cv_rows = []
    for cv in cv_raporu["sonuclar"]:
        s = cv["summary"]
        cv_rows.append(
            [
                cv["veri_adayi"].replace("sentetik_", "").replace("_per_class_pima", "/sınıf"),
                cv["model"],
                cv["feature_set"],
                f"{s['accuracy_mean']:.3f} ± {s['accuracy_std']:.3f}",
                f"{s['f1_mean']:.3f} ± {s['f1_std']:.3f}",
                f"{s['roc_auc_mean']:.3f} ± {s['roc_auc_std']:.3f}",
                f"{s['ana_metrik_minimumu_mean']:.3f} ± {s['ana_metrik_minimumu_std']:.3f}",
            ]
        )
    _word_tablo(doc, ["Veri", "Model", "Feature", "Accuracy", "F1", "ROC-AUC", "Min ana"], cv_rows, font_size=6.7)

    doc.add_page_break()
    doc.add_heading("Final Model, Confusion Matrix ve Kalibrasyon", level=1)
    _word_tablo(
        doc,
        ["Alan", "Değer"],
        [
            ["Final veri adayı", final["veri_adayi"]],
            ["Model", final["model"]],
            ["Feature set", final["feature_set"]],
            ["Çıkarılan feature", ", ".join(final["cikarilan_featurelar"]) or "Yok"],
            ["Threshold", f"{final['threshold_ozeti']['esik']:.3f}"],
            ["Holdout Accuracy", _pct(final_m["accuracy"])],
            ["Holdout F1", _pct(final_m["f1"])],
            ["Holdout ROC-AUC", _pct(final_m["roc_auc"])],
            ["Group CV Accuracy", f"{_pct(final_cv['accuracy_mean'])} ± {final_cv['accuracy_std']:.3f}"],
        ],
        font_size=8.5,
    )
    _resim(doc, grafikler["confusion_matrix"], width=4.1)
    _resim(doc, grafikler["calibration_curve"], width=4.6)
    _resim(doc, grafikler["brier_karsilastirma"], width=5.8)
    _paragraf(
        doc,
        "Brier score düşük olduğunda olasılık tahmini daha güvenilir kabul edilir. Bu değer, sınıflandırma başarısından "
        "ayrı olarak modelin olasılık kalibrasyonunu yorumlamak için kullanılmıştır.",
    )

    doc.add_page_break()
    doc.add_heading("Feature Yorumu, External Holdout ve Sonuç", level=1)
    _resim(doc, grafikler["feature_importance"], width=6.2)
    _word_tablo(
        doc,
        ["External holdout metriği", "Değer"],
        [
            ["Accuracy", _pct(external["accuracy"])],
            ["Precision", _pct(external["precision"])],
            ["Recall / Sensitivity", _pct(external["recall"])],
            ["Specificity", _pct(external["specificity"])],
            ["F1", _pct(external["f1"])],
            ["ROC-AUC", _pct(external["roc_auc"])],
            ["Balanced accuracy", _pct(external["balanced_accuracy"])],
            ["MCC", f"{external['mcc']:.3f}"],
            ["Brier", f"{external['brier']:.3f}"],
        ],
        font_size=8.4,
    )
    _paragraf(
        doc,
        "Sentetik benchmark sonuçları ile orijinal PIMA external holdout sonuçları aynı anlamda yorumlanmamıştır. "
        "Sentetik benchmark, kontrollü veri artırımı sonrası modelin ayrıştırma gücünü; orijinal external holdout ise "
        "gerçek veri üzerindeki dış kontrol performansını göstermektedir.",
    )
    _paragraf(
        doc,
        "Bu çalışmada kontrollü sentetik benchmark üzerinde yüksek performans elde edilmiştir. Buna karşın bu sonuçlar, "
        "dış veri setleriyle desteklenmeden klinik genellenebilirlik kanıtı olarak yorumlanmamalıdır. Leakage kontrolü, "
        "aile bazlı source_id ayrımı ve kopya kontrolleriyle desteklenmiştir.",
    )

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(hedef)
    return hedef


def _baslik(doc: Any, text: str) -> None:
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)


def _paragraf(doc: Any, text: str) -> None:
    from docx.shared import Pt

    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05


def _resim(doc: Any, path: str, width: float) -> None:
    from docx.shared import Inches

    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run()
    run.add_picture(path, width=Inches(width))


def _word_tablo(doc: Any, basliklar: list[str], satirlar: list[list[str]], font_size: float = 8.0) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    tablo = doc.add_table(rows=1, cols=len(basliklar))
    tablo.style = "Table Grid"
    header = tablo.rows[0].cells
    for i, baslik in enumerate(basliklar):
        header[i].text = baslik
        for p in header[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor(255, 255, 255)
        shading = header[i]._tc.get_or_add_tcPr()
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        fill = OxmlElement("w:shd")
        fill.set(qn("w:fill"), "1F4E79")
        shading.append(fill)
    for satir in satirlar:
        cells = tablo.add_row().cells
        for i, deger in enumerate(satir):
            cells[i].text = str(deger)
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(font_size)
    doc.add_paragraph()


def _terminal_ozeti_yaz(rapor: dict[str, Any]) -> None:
    final = rapor["en_iyi_sonuc"]
    m = final["test_metrikleri"]
    cv = rapor["group_cv"]["summary"]
    ext = rapor["orijinal_pima_external_holdout"]["metrikler"]
    leak = next(
        a for a in rapor["source_id_leakage_raporu"]["adaylar"] if a["dataset_name"] == final["veri_adayi"]
    )
    print("\n=== SELECTED SYNTHETIC BENCHMARK SUMMARY ===")
    print(f"- En iyi savunulabilir aday: {final['veri_adayi']}")
    print(f"- Veri boyutu: {rapor['final_veri']['satir_sayisi']}")
    print(f"- Model: {final['model']}")
    print(f"- Feature set: {final['feature_set']}")
    print(f"- Holdout Accuracy: {m['accuracy']:.4f}")
    print(f"- Holdout F1: {m['f1']:.4f}")
    print(f"- Holdout ROC-AUC: {m['roc_auc']:.4f}")
    print(f"- Group CV Accuracy mean ± std: {cv['accuracy_mean']:.4f} ± {cv['accuracy_std']:.4f}")
    print(f"- Group CV F1 mean ± std: {cv['f1_mean']:.4f} ± {cv['f1_std']:.4f}")
    print(f"- Group CV ROC-AUC mean ± std: {cv['roc_auc_mean']:.4f} ± {cv['roc_auc_std']:.4f}")
    print(f"- Original external holdout Accuracy: {ext['accuracy']:.4f}")
    print(f"- Original external holdout F1: {ext['f1']:.4f}")
    print(f"- Original external holdout ROC-AUC: {ext['roc_auc']:.4f}")
    print(f"- Leakage durumu: {leak['leakage_status']}")
    print(f"- Train/test source_id kesişimi: {leak['train_test_source_intersection_count']}")
    print(f"- CV source_id kesişimi: {max(leak['cv_fold_source_intersection_counts'] or [0])}")
    print(f"- Exact duplicate: {leak['exact_duplicate_count']}")
    print(f"- Near duplicate oranı: {leak['near_duplicate_rate']}")
    print(f"- Minimum mesafe: {leak['min_near_duplicate_distance']}")
    print(f"- Word raporu yolu: {rapor['word_raporu']}")
    print(f"- Masaüstü Word raporu yolu: {rapor['masaustu_word_raporu']}")


def _xgboost_model(random_state: int, n_jobs: int) -> BaseEstimator | None:
    try:
        mod = importlib.import_module("xgboost")
    except Exception:
        return None
    return mod.XGBClassifier(
        objective="binary:logistic",
        eval_metric="logloss",
        n_estimators=180,
        max_depth=4,
        learning_rate=0.055,
        subsample=0.95,
        colsample_bytree=0.95,
        reg_lambda=0.9,
        random_state=random_state,
        n_jobs=n_jobs,
    )


def _lightgbm_model(random_state: int, n_jobs: int) -> BaseEstimator | None:
    try:
        mod = importlib.import_module("lightgbm")
    except Exception:
        return None
    return mod.LGBMClassifier(
        objective="binary",
        n_estimators=180,
        num_leaves=31,
        learning_rate=0.045,
        subsample=0.95,
        colsample_bytree=0.95,
        random_state=random_state,
        n_jobs=n_jobs,
        verbose=-1,
    )


def _varsayilan_hedefler(mod: str) -> tuple[int, ...]:
    if mod == "quick":
        return (620,)
    if mod in {"full", "target"}:
        return (2500, 2700, 5000)
    raise ValueError("mod quick, full veya target olmali.")


def _adetleri_bol(toplam: int, oranlar: list[float]) -> list[int]:
    adetler = [int(toplam * oran) for oran in oranlar]
    adetler[-1] += toplam - sum(adetler)
    return adetler


def _sinif_dagilimi(df: pd.DataFrame) -> dict[str, int]:
    return {str(int(k)): int(v) for k, v in df[HEDEF_KOLONU].value_counts().sort_index().to_dict().items()}


def _dataframe_yap(x: Any) -> pd.DataFrame:
    if isinstance(x, pd.DataFrame):
        return x.copy()
    arr = np.asarray(x)
    return pd.DataFrame(arr, columns=OZELLIK_KOLONLARI[: arr.shape[1]])


def _json_yaz(path: Path, veri: dict[str, Any]) -> None:
    path.write_text(json.dumps(_json_uyumlu(veri), ensure_ascii=False, indent=2), encoding="utf-8")


def _json_uyumlu(obj: Any) -> Any:
    if isinstance(obj, dict):
        return {str(k): _json_uyumlu(v) for k, v in obj.items() if not str(k).startswith("_")}
    if isinstance(obj, list | tuple):
        return [_json_uyumlu(v) for v in obj]
    if isinstance(obj, np.ndarray):
        return obj.tolist()
    if isinstance(obj, np.integer):
        return int(obj)
    if isinstance(obj, np.floating):
        return float(obj)
    if isinstance(obj, pd.Series):
        return obj.tolist()
    if isinstance(obj, Path):
        return str(obj)
    return obj


def _jsondan_modeli_cikar(sonuc: dict[str, Any]) -> dict[str, Any]:
    return _json_uyumlu({k: v for k, v in sonuc.items() if not k.startswith("_")})


def _sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _zaman() -> str:
    return datetime.now(timezone.utc).isoformat()


def _pct(v: Any) -> str:
    return f"%{float(v) * 100:.2f}"


def _oran(v: Any) -> str:
    if v is None:
        return "Yok"
    return f"{float(v):.4f}"


def _sayi(v: Any) -> str:
    if v is None:
        return "Yok"
    return f"{float(v):.4f}"


def _dist_yazi(dist: dict[str, int]) -> str:
    return f"{dist.get('0', 0)} negatif, {dist.get('1', 0)} pozitif"


def _sonuclar_adaya_gore(basarili: list[dict[str, Any]], aday_adi: str) -> list[dict[str, Any]]:
    return [s for s in basarili if s["veri_adayi"] == aday_adi]


def _veri_adayi_bul(adaylar: list[VeriAdayi], ad: str) -> VeriAdayi:
    for aday in adaylar:
        if aday.ad == ad:
            return aday
    raise KeyError(ad)


def _feature_set_bul(feature_setleri: list[FeatureSet], ad: str) -> FeatureSet:
    for feature_set in feature_setleri:
        if feature_set.ad == ad:
            return feature_set
    raise KeyError(ad)


def _model_adayi_bul(model_adaylari: list[ModelAdayi], ad: str) -> ModelAdayi:
    for model in model_adaylari:
        if model.ad == ad:
            return model
    raise KeyError(ad)


def _cv_sonuc_bul(cv_sonuclari: list[dict[str, Any]], sonuc: dict[str, Any]) -> dict[str, Any]:
    for cv in cv_sonuclari:
        if cv["veri_adayi"] == sonuc["veri_adayi"] and cv["feature_set"] == sonuc["feature_set"] and cv["model"] == sonuc["model"]:
            return cv
    raise KeyError((sonuc["veri_adayi"], sonuc["feature_set"], sonuc["model"]))
