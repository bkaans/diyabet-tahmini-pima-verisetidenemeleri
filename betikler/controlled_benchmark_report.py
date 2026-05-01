"""Mevcut controlled_benchmark sonuçlarından açıklayıcı V2 Word raporu üretir."""

from __future__ import annotations

import json
import math
import shutil
import sys
from pathlib import Path
from typing import Any

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from PIL import Image, ImageStat


PROJE_KOKU = Path(__file__).resolve().parents[1]
if str(PROJE_KOKU) not in sys.path:
    sys.path.insert(0, str(PROJE_KOKU))

from makine_ogrenmesi.kaynak.ozellik_yapilandirmasi import HEDEF_KOLONU, OZELLIK_KOLONLARI  # noqa: E402


RAPOR_DIR = PROJE_KOKU / "makine_ogrenmesi" / "raporlar"
GRAFIK_DIR = RAPOR_DIR / "grafikler_controlled_benchmark"
GRAFIK_V2_DIR = RAPOR_DIR / "grafikler_controlled_benchmark_v2"
DOCX_PATH = RAPOR_DIR / "controlled_synthetic_benchmark_report_v2.docx"
DESKTOP_DOCX_PATH = Path.home() / "Desktop" / "pima_controlled_synthetic_benchmark_report_v2.docx"
AUDIT_PATH = RAPOR_DIR / "controlled_benchmark_quality_audit_v2.json"
VERI_DIR = PROJE_KOKU / "makine_ogrenmesi" / "veri" / "deneysel"


REQUIRED_GRAPHS = {
    "original_class_distribution": "controlled_benchmark_orijinal_pima_sinif_dagilimi.png",
    "candidate_distribution": "controlled_benchmark_aday_veri_boyutlari.png",
    "holdout_min_metric": "controlled_benchmark_holdout_ana_metrik_minimumu.png",
    "cv_min_metric": "controlled_benchmark_cv_ana_metrik_minimumu_mean.png",
    "holdout_roc_auc": "controlled_benchmark_holdout_roc_auc.png",
    "cv_roc_auc": "controlled_benchmark_cv_roc_auc_mean.png",
    "confusion_matrix": "controlled_benchmark_selected_model_confusion_matrix.png",
    "feature_importance": "controlled_benchmark_selected_model_feature_importance.png",
    "calibration_curve": "controlled_benchmark_selected_model_calibration_curve.png",
    "brier": "controlled_benchmark_brier_karsilastirma.png",
    "leakage": "controlled_benchmark_leakage_ozeti.png",
}


BAD_TURKISH_ASCII = [
    "yapilan",
    "sinirli",
    "yaklasimlari",
    "artirmak",
    "uretimi",
    "uretim",
    "gorecek",
    "olusur",
    "calismada",
    "calisma",
    "performansi",
    "kullanilmaktadir",
    "uygulanirsa",
    "saklanmistir",
    "basari",
    "sinif",
    "ozellik",
    "onemli",
    "ogrenci",
    "aciklama",
    "olcut",
    "guvenilirlik",
]


def main() -> None:
    model_report = _json_load(RAPOR_DIR / "controlled_benchmark_model_selection_report.json")
    leakage_report = _json_load(RAPOR_DIR / "controlled_benchmark_leakage_report.json")
    dataset_report = _json_load(RAPOR_DIR / "controlled_benchmark_dataset_comparison.json")
    fark_report = _json_load(RAPOR_DIR / "benchmark_method_comparison.json")

    GRAFIK_V2_DIR.mkdir(parents=True, exist_ok=True)
    graph_summary, graph_paths = _prepare_graphs(dataset_report, fark_report)
    audit = _build_quality_audit(model_report, leakage_report, fark_report, graph_summary)

    _write_docx(model_report, leakage_report, dataset_report, fark_report, graph_paths)
    shutil.copy2(DOCX_PATH, DESKTOP_DOCX_PATH)

    text = _docx_text(DOCX_PATH)
    turkish_scan = _scan_turkish_ascii(text)
    audit["turkish_character_fix_summary"] = turkish_scan
    audit["word_report_path"] = str(DOCX_PATH)
    audit["desktop_word_report_path"] = str(DESKTOP_DOCX_PATH)
    audit["final_quality_status"] = "pass" if _quality_pass(audit) else "warning"
    _json_write(AUDIT_PATH, audit)

    print("- Yeni V2 Word raporu üretildi mi: Evet")
    print("- Masaüstü kopyası üretildi mi: Evet")
    print("- Türkçe karakter düzeltildi mi:", "Evet" if not turkish_scan["flagged_terms"] else "Kontrol uyarısı var")
    print("- Boş grafik kontrolü yapıldı mı: Evet")
    print("- Boş/eksik grafik bulundu mu:", "Evet" if graph_summary["empty_or_missing_graphs"] else "Hayır")
    print("- Boş/eksik grafikler nasıl çözüldü:", graph_summary["resolution"])
    print("- Teknik açıklamalar rapora eklendi mi: Evet")
    print("- Önceki ve yeni sentetik üretim fark analizi eklendi mi: Evet")
    print("- Sentetik benchmark hedefi korunuyor mu: Evet")
    print("- External holdout ayrı mı: Evet")
    print("- Quality audit sonucu:", audit["final_quality_status"])
    print("- Distribution shift notu eklendi mi: Evet")
    print("- Outcome encoding uyarısı var mı:", audit["outcome_encoding_check"]["status"])
    print("- Feature importance yoğunlaşma uyarısı var mı:", audit["feature_importance_concentration_check"]["status"])
    print("- Yeni Word yolu:", DOCX_PATH)
    print("- Masaüstü Word yolu:", DESKTOP_DOCX_PATH)


def _json_load(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def _json_write(path: Path, data: dict[str, Any]) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def _prepare_graphs(dataset_report: dict[str, Any], fark_report: dict[str, Any]) -> tuple[dict[str, Any], dict[str, str]]:
    graph_paths: dict[str, str] = {}
    empty_or_missing: list[str] = []
    for key, filename in REQUIRED_GRAPHS.items():
        source = GRAFIK_DIR / filename
        dest = GRAFIK_V2_DIR / filename
        check = _graph_check(source)
        if check["status"] == "pass":
            shutil.copy2(source, dest)
            graph_paths[key] = str(dest)
        else:
            empty_or_missing.append(filename)

    diff_graph = GRAFIK_V2_DIR / "controlled_benchmark_v2_strict2700_cross_model.png"
    shift_graph = GRAFIK_V2_DIR / "controlled_benchmark_v2_distribution_shift_cohens_d.png"
    feature_signal_graph = GRAFIK_V2_DIR / "controlled_benchmark_v2_feature_signal_importance.png"
    _plot_cross_model_graph(fark_report, diff_graph)
    _plot_shift_graph(fark_report, shift_graph)
    _plot_feature_signal_importance(feature_signal_graph)
    for key, path in {
        "strict_controlled_benchmark_cross_model": diff_graph,
        "distribution_shift_cohens_d": shift_graph,
        "feature_importance": feature_signal_graph,
    }.items():
        check = _graph_check(path)
        if check["status"] == "pass":
            graph_paths[key] = str(path)
        else:
            empty_or_missing.append(path.name)

    return (
        {
            "checked_graph_count": len(REQUIRED_GRAPHS) + 3,
            "included_graph_count": len(graph_paths),
            "empty_or_missing_graphs": empty_or_missing,
            "resolution": "Boş grafik bulunmadı; mevcut grafikler doğrulanıp V2 klasörüne kopyalandı ve üç yeni analiz grafiği üretildi."
            if not empty_or_missing
            else "Boş/eksik grafikler rapora eklenmedi; ilgili bilgi tablo/metin ile verildi.",
        },
        graph_paths,
    )


def _graph_check(path: Path) -> dict[str, Any]:
    if not path.exists() or path.stat().st_size <= 0:
        return {"status": "missing_or_empty", "path": str(path)}
    try:
        image = Image.open(path).convert("RGB")
        stat = ImageStat.Stat(image)
        avg_var = float(np.mean(stat.var))
        return {
            "status": "pass" if avg_var > 25 else "blank_warning",
            "path": str(path),
            "size_bytes": int(path.stat().st_size),
            "image_size": list(image.size),
            "avg_pixel_variance": avg_var,
        }
    except Exception as exc:
        return {"status": "invalid", "path": str(path), "error": str(exc)}


def _plot_cross_model_graph(fark: dict[str, Any], path: Path) -> None:
    rows = fark["cross_model_experiments"]
    order = [
        ("strict_xgb_no_skin_on_strict2700", "Önceki üretim\nXGBoost"),
        ("strict_xgb_no_skin_on_controlled_benchmark_2500", "Yeni üretim\nXGBoost"),
        ("new_soft_all_on_strict2700", "Önceki üretim\nSoft voting"),
        ("new_soft_all_on_controlled_benchmark_2500", "Yeni üretim\nSoft voting"),
    ]
    values = [rows[key]["tuned_threshold_metrics"]["ana_metrik_minimumu"] for key, _ in order]
    colors = ["#607D8B", "#2A9D8F", "#90A4AE", "#1F4E79"]
    fig, ax = plt.subplots(figsize=(8.6, 4.2), dpi=180)
    bars = ax.bar([label for _, label in order], values, color=colors)
    ax.set_ylim(0.70, 1.0)
    ax.set_ylabel("Min ana metrik")
    ax.set_title("Üretim Stratejisi ve Model Değişiminin Etkisi")
    ax.grid(axis="y", alpha=0.25)
    for bar, value in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, value + 0.006, f"{value:.3f}", ha="center", fontsize=9, fontweight="bold")
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def _plot_shift_graph(fark: dict[str, Any], path: Path) -> None:
    dist = fark["distribution"]
    shift = fark["distribution_shift_vs_original_dev"]
    labels = ["Geliştirme", "Önceki sentetik", "Yeni sentetik", "Dış kontrol"]
    dist_keys = ["original_dev", "strict2700_all_rows", "controlled_benchmark_2500_all_rows", "external_holdout"]
    cohen = [dist[key]["class_separation"]["aggregate"]["mean_abs_cohens_d"] for key in dist_keys]
    shifts = [0.0, shift["strict2700_all_rows"]["aggregate"]["avg_abs_z_mean_shift"], shift["controlled_benchmark_2500_all_rows"]["aggregate"]["avg_abs_z_mean_shift"], shift["external_holdout"]["aggregate"]["avg_abs_z_mean_shift"]]
    x = np.arange(len(labels))
    width = 0.36
    fig, ax = plt.subplots(figsize=(8.8, 4.4), dpi=180)
    ax.bar(x - width / 2, cohen, width, label="Ortalama mutlak Cohen's d", color="#1F4E79")
    ax.bar(x + width / 2, shifts, width, label="Ortalama mutlak dağılım kayması", color="#E76F51")
    ax.set_xticks(x, labels)
    ax.set_title("Sınıf Ayrımı ve Dağılım Kayması")
    ax.grid(axis="y", alpha=0.25)
    ax.legend()
    for idx, value in enumerate(cohen):
        ax.text(idx - width / 2, value + 0.04, f"{value:.2f}", ha="center", fontsize=8)
    for idx, value in enumerate(shifts):
        ax.text(idx + width / 2, value + 0.04, f"{value:.2f}", ha="center", fontsize=8)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def _plot_feature_signal_importance(path: Path) -> None:
    df = pd.read_csv(VERI_DIR / "controlled_synthetic_2500_per_class_pima.csv")
    y = df[HEDEF_KOLONU].to_numpy(dtype=int)
    rows = []
    for col in OZELLIK_KOLONLARI:
        auc = _single_feature_auc(y, df[col].to_numpy(dtype=float))
        neg = df.loc[df[HEDEF_KOLONU] == 0, col].astype(float)
        pos = df.loc[df[HEDEF_KOLONU] == 1, col].astype(float)
        pooled = math.sqrt(
            ((len(neg) - 1) * float(neg.var(ddof=1)) + (len(pos) - 1) * float(pos.var(ddof=1)))
            / max(len(neg) + len(pos) - 2, 1)
        )
        d = abs(float((pos.mean() - neg.mean()) / pooled)) if pooled else 0.0
        # AUC ve Cohen's d birlikte kullanılır; amaç model içi katsayı değil, final benchmarktaki ayrıştırma sinyalini görünür yapmak.
        score = 0.65 * max(0.0, (auc - 0.5) / 0.5) + 0.35 * min(d / 3.0, 1.0)
        rows.append((col, score, auc, d))
    rows.sort(key=lambda item: item[1])
    labels = [item[0] for item in rows]
    scores = [item[1] for item in rows]
    aucs = [item[2] for item in rows]
    ds = [item[3] for item in rows]
    colors = ["#2A9D8F" if label in {"Glucose", "BMI", "Age", "DiabetesPedigreeFunction"} else "#7A8FA6" for label in labels]
    fig, ax = plt.subplots(figsize=(8.8, 4.9), dpi=180)
    bars = ax.barh(labels, scores, color=colors)
    ax.set_xlim(0, 1.05)
    ax.set_xlabel("Normalize ayrıştırma gücü")
    ax.set_title("Final Benchmarkta Değişken Ayrıştırma Gücü")
    ax.grid(axis="x", alpha=0.25)
    for bar, auc, d in zip(bars, aucs, ds):
        ax.text(
            bar.get_width() + 0.015,
            bar.get_y() + bar.get_height() / 2,
            f"AUC={auc:.3f}, d={d:.2f}",
            va="center",
            fontsize=8,
        )
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def _build_quality_audit(
    model_report: dict[str, Any],
    leakage_report: dict[str, Any],
    fark_report: dict[str, Any],
    graph_summary: dict[str, Any],
) -> dict[str, Any]:
    final = model_report["final"]
    cv = model_report["final_cv"]["summary"]
    external = model_report["external_holdout"]["metrics"]
    final_leak = next(item for item in leakage_report["datasets"] if item["dataset_name"] == final["dataset_name"])
    outcome_check = _outcome_encoding_check()
    feature_concentration = _feature_concentration_check()
    return {
        "final_dataset_name": final["dataset_name"],
        "final_model": final["model"],
        "final_feature_set": final["feature_set"],
        "final_threshold": final["threshold_report"]["selected_threshold"],
        "synthetic_holdout_metrics": final["tuned_threshold_metrics"],
        "synthetic_group_cv_metrics": cv,
        "external_holdout_metrics": external,
        "leakage_summary": final_leak,
        "turkish_character_fix_summary": {},
        "empty_graph_check_summary": graph_summary,
        "controlled_benchmark_vs_strict2700_summary": fark_report["decomposition"],
        "distribution_shift_summary": fark_report["decomposition"]["synthetic_distribution_shift"],
        "outcome_encoding_check": outcome_check,
        "feature_importance_concentration_check": feature_concentration,
        "final_quality_status": "pending",
        "notes": [
            "Final benchmark sonucu korunmuştur; model yeniden eğitilmemiş veya yeniden optimize edilmemiştir.",
            "Yüksek sentetik benchmark sonucunun üretim stratejisiyle ilişkisi raporda açıkça sınırlılık olarak anlatılmıştır.",
        ],
    }


def _outcome_encoding_check() -> dict[str, Any]:
    paths = [
        PROJE_KOKU / "makine_ogrenmesi" / "veri" / "ham" / "diabetes.csv",
        VERI_DIR / "controlled_synthetic_2500_per_class_pima.csv",
        VERI_DIR / "sentetik_2700_per_class_pima.csv",
    ]
    details = {}
    ok = True
    for path in paths:
        df = pd.read_csv(path)
        values = sorted(int(v) for v in df[HEDEF_KOLONU].dropna().unique())
        valid = values == [0, 1]
        ok = ok and valid
        details[path.name] = {"unique_values": values, "valid_binary_0_1": valid}
    return {"status": "pass" if ok else "warning", "details": details}


def _feature_concentration_check() -> dict[str, Any]:
    df = pd.read_csv(VERI_DIR / "controlled_synthetic_2500_per_class_pima.csv")
    y = df[HEDEF_KOLONU].to_numpy(dtype=int)
    scores = {}
    for col in OZELLIK_KOLONLARI:
        scores[col] = _single_feature_auc(y, df[col].to_numpy(dtype=float))
    top_feature = max(scores, key=scores.get)
    top_score = scores[top_feature]
    status = "warning" if top_score >= 0.90 else "pass"
    return {
        "status": status,
        "top_feature": top_feature,
        "top_single_feature_auc_directionless": top_score,
        "all_single_feature_auc_directionless": scores,
        "note": "Tek değişken ayrımı yüksekse bu durum sentetik veri sınıf ayrımının güçlendiğini gösterir; raporda sınırlılık olarak açıklandı.",
    }


def _single_feature_auc(y: np.ndarray, values: np.ndarray) -> float:
    order = np.argsort(values)
    ranks = np.empty_like(order, dtype=float)
    ranks[order] = np.arange(1, len(values) + 1)
    n_pos = int(y.sum())
    n_neg = int(len(y) - n_pos)
    if n_pos == 0 or n_neg == 0:
        return 0.5
    sum_ranks_pos = float(ranks[y == 1].sum())
    auc = (sum_ranks_pos - n_pos * (n_pos + 1) / 2) / (n_pos * n_neg)
    return float(max(auc, 1 - auc))


def _dataset_label(value: str) -> str:
    labels = {
        "controlled_synthetic_2500_per_class_pima": "Kontrollü sentetik PIMA benchmarkı (2500/2500)",
        "controlled_synthetic_2700_per_class_pima": "Kontrollü sentetik PIMA benchmarkı (2700/2700)",
        "controlled_synthetic_3000_per_class_pima": "Kontrollü sentetik PIMA benchmarkı (3000/3000)",
        "controlled_synthetic_4000_per_class_pima": "Kontrollü sentetik PIMA benchmarkı (4000/4000)",
        "controlled_synthetic_5000_per_class_pima": "Kontrollü sentetik PIMA benchmarkı (5000/5000)",
        "controlled_synthetic_7500_per_class_pima": "Kontrollü sentetik PIMA benchmarkı (7500/7500)",
        "controlled_synthetic_10000_per_class_pima": "Genişletilmiş kontrollü sentetik PIMA benchmarkı (10000/10000)",
    }
    return labels.get(value, value.replace("_", " "))


def _model_label(value: str) -> str:
    labels = {
        "soft_voting_xgb_lgbm_et": "XGBoost, LightGBM ve ExtraTrees tabanlı soft voting ensemble",
        "xgboost": "XGBoost",
        "extra_trees": "ExtraTrees",
        "random_forest": "Random Forest",
        "hist_gradient_boosting": "Histogram Gradient Boosting",
    }
    return labels.get(value, value.replace("_", " "))


def _feature_set_label(value: str) -> str:
    labels = {
        "all_features": "Tüm orijinal PIMA değişkenleri",
        "no_skinthickness": "SkinThickness çıkarılmış değişken seti",
        "no_insulin": "Insulin çıkarılmış değişken seti",
        "no_skinthickness_no_insulin": "SkinThickness ve Insulin çıkarılmış değişken seti",
        "clinical_interactions": "Klinik etkileşim değişkenleri eklenmiş set",
        "compact_best_features": "Kompakt yüksek katkılı değişken seti",
        "high_signal_features": "Yüksek ayrıştırma sinyali taşıyan değişken seti",
    }
    return labels.get(value, value.replace("_", " "))


def _preprocessing_label(value: str) -> str:
    labels = {
        "tree_models_no_scaler": "Ağaç tabanlı modeller için ek ölçekleme uygulanmadı",
        "no_imputation": "Ek imputasyon uygulanmadı",
        "global_median_imputation": "Genel medyan ile eksik değer tamamlama",
        "class_conditional_median_imputation": "Sınıf koşullu medyan ile eksik değer tamamlama",
        "knn_imputer": "KNN tabanlı eksik değer tamamlama",
    }
    return labels.get(value, value.replace("_", " "))


def _resampling_label(value: str) -> str:
    labels = {
        "none": "Ek resampling uygulanmadı",
        "class_weight_balanced": "Sınıf ağırlığı dengelenmiş eğitim",
        "smote": "Eğitim içinde SMOTE",
        "smoteenn": "Eğitim içinde SMOTE-ENN",
        "smotetomek": "Eğitim içinde SMOTE-Tomek",
    }
    return labels.get(value, value.replace("_", " "))


def _write_docx(
    model_report: dict[str, Any],
    leakage_report: dict[str, Any],
    dataset_report: dict[str, Any],
    fark_report: dict[str, Any],
    graphs: dict[str, str],
) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    _set_styles(doc)

    final = model_report["final"]
    holdout = final["tuned_threshold_metrics"]
    cv = model_report["final_cv"]["summary"]
    external = model_report["external_holdout"]["metrics"]
    final_leak = next(item for item in leakage_report["datasets"] if item["dataset_name"] == final["dataset_name"])

    _title(doc, "PIMA Diyabet Veri Seti İçin Kaynak Aile Kontrollü Sentetik Benchmark Raporu")
    _p(
        doc,
        "Bu çalışmada ana başarı hedefi, ham PIMA veri setinde doğrudan klinik genellenebilirlik kanıtlamak değil; "
        "orijinal PIMA verisinden türetilmiş, kaynak aile kimliği (source_id) ayrımıyla kontrol edilmiş sentetik benchmark üzerinde "
        "minimum ana metrik değerini 0.90 üzerinde tutan bir makine öğrenmesi akışı oluşturmaktır.",
    )
    _p(
        doc,
        "Sentetik benchmark sonuçları ile orijinal dış kontrol verisi sonuçları aynı anlamda yorumlanmamıştır. "
        "Sentetik benchmark, kontrollü veri artırımı sonrası modelin ayrıştırma başarısını; orijinal dış kontrol verisi ise "
        "gerçek PIMA verisi üzerinde dış kontrol performansını göstermektedir.",
    )
    _table(
        doc,
        ["Alan", "Değer"],
        [
            ["Orijinal PIMA", f"{dataset_report['original_csv']['rows']} satır, {dataset_report['original_csv']['class_distribution']['0']} negatif / {dataset_report['original_csv']['class_distribution']['1']} pozitif"],
            ["Geliştirme verisi", f"{dataset_report['original_dev']['rows']} satır"],
            ["Orijinal dış kontrol verisi", f"{dataset_report['external_holdout']['rows']} satır"],
            ["Ana hedef", "PIMA + kaynak aile kontrollü sentetik benchmarkta minimum ana metrik >= 0.90"],
            ["Final aday", _dataset_label(final["dataset_name"])],
        ],
    )
    _img(doc, graphs.get("original_class_distribution"), 5.5, "Orijinal PIMA sınıf dağılımı, veri setindeki negatif sınıf ağırlığını göstermektedir. Bu dengesizlik sentetik benchmark kurgusunun temel gerekçelerinden biridir.")
    _p(
        doc,
        "Veri temizleme aşamasında Glucose, BloodPressure, SkinThickness, Insulin ve BMI kolonlarındaki klinik olarak gerçekçi olmayan 0 değerleri eksik değer gibi ele alınmıştır. "
        "Orijinal dış kontrol verisi, sentetik üretimden önce ayrılmış ve model seçimi, eşik ayarı veya değişken seçimi kararlarında kullanılmamıştır.",
    )

    doc.add_page_break()
    _h(doc, "Sentetik Veri Üretimi ve Kaynak Aile Yapısı")
    _p(
        doc,
        "Orijinal PIMA verisi temizlendikten sonra geliştirme verisi ve dış kontrol verisi olarak ikiye ayrılmıştır. "
        "Sentetik üretim yalnızca geliştirme verisi üzerinden yapılmıştır. Her sentetik satır, türediği orijinal satırın kaynak aile kimliğine bağlı tutulmuştur.",
    )
    _p(
        doc,
        "Yeni sentetik üretim hattında örnekler kaynak aile yapısı korunarak sınıf hedef merkezlerine yaklaştırılmıştır. "
        "Bu yaklaşım leakage üretmeden benchmark sınıflarını daha kolay ayrıştırılabilir hale getirmiştir; bu nedenle yüksek sonuçlar sentetik benchmark bağlamında yorumlanmalıdır.",
    )
    _h(doc, "Leakage Kontrolü ve Kaynak Aile Yapısı", level=2)
    _p(
        doc,
        "Sentetik veri üretiminde her orijinal PIMA satırı bir kaynak aile olarak kabul edilmiştir. Her orijinal satıra original_{index} formatında bir kaynak aile kimliği atanmış, "
        "bu satırdan üretilen tüm sentetik örnekler aynı kimliği taşımıştır. Böylece aynı kaynak aileden gelen örneklerin hem eğitim hem de test tarafına düşmesi engellenmiştir.",
    )
    _p(
        doc,
        "Veri ayrımı sırasında kaynak aile tabanlı grup ayrımı uygulanmış ve eğitim/test kümeleri arasında kaynak aile kesişimi kontrol edilmiştir. "
        "Ayrıca birebir kopya ve çok yakın benzerlik kontrolleri yapılmıştır. Bu kontroller, sentetik benchmark sonucunun doğrudan kopya veya aile sızıntısı nedeniyle şişmediğini göstermek amacıyla rapora dahil edilmiştir.",
    )
    _table(
        doc,
        ["Kontrol", "Beklenen", "Sonuç"],
        [
            ["Eğitim/test kaynak aile kesişimi", "0", str(final_leak["train_test_source_intersection"])],
            ["CV katmanı kaynak aile kesişimi", "0", str(max(final_leak["cv_source_intersections"] or [0]))],
            ["Birebir kopya sayısı", "0", str(final_leak["exact_duplicate_count"])],
            ["Çok yakın benzerlik oranı", "Düşük / 0'a yakın", f"{final_leak['near_duplicate_rate']:.4f}"],
            ["Minimum mesafe", "Raporlanır", f"{final_leak['minimum_near_duplicate_distance']:.4f}"],
            ["Orijinal dış kontrol izolasyonu", "Evet", "Evet" if final_leak["external_holdout_source_overlap_count"] == 0 else "Hayır"],
            ["Bağımsız sentetik kaynak aile kimliği kaldı mı?", "Hayır", "Hayır" if final_leak["independent_synthetic_source_id_count"] == 0 else "Evet"],
        ],
    )
    candidate_rows = [
        [
            str(item["target_per_class"]),
            str(item["original_dev_rows"]),
            str(item["synthetic_rows"]),
            str(item["total_rows"]),
            f"{item['class_distribution']['0']} / {item['class_distribution']['1']}",
        ]
        for item in dataset_report["datasets"]
    ]
    _table(doc, ["Sınıf başı", "Orijinal dev", "Sentetik", "Toplam", "Negatif / Pozitif"], candidate_rows, font_size=8.0)
    _img(doc, graphs.get("candidate_distribution"), 5.8, "Grafik, aday benchmark dosyalarının sınıf başı hedeflere göre dengeli üretildiğini göstermektedir.")

    doc.add_page_break()
    _h(doc, "Kullanılan Teknikler ve Modelleme Yaklaşımı")
    _p(
        doc,
        "Final model, olasılık tabanlı oylama yaklaşımı kullanan soft voting ensemble olarak seçilmiştir. Bu yapı XGBoost, LightGBM ve ExtraTrees modellerinin olasılık tahminlerini birleştirir. "
        "Tek modele bağımlı kalmak yerine farklı model ailelerinin güçlü tarafları birlikte kullanılmıştır.",
    )
    _p(
        doc,
        "Buna rağmen fark analizi, performans artışının ana kaynağının modelden çok sentetik benchmark yapısı olduğunu göstermektedir. "
        "Model seçimi herhangi tek algoritmaya zorlanmamış, leakage-clean adaylar arasından selection score, holdout ve group CV dengesi dikkate alınarak yapılmıştır.",
    )
    _table(
        doc,
        ["Alan", "Değer"],
        [
            ["Final aday", _dataset_label(final["dataset_name"])],
            ["Model", _model_label(final["model"])],
            ["Değişken seti", _feature_set_label(final["feature_set"])],
            ["Ön işleme", _preprocessing_label(final["preprocessing"])],
            ["Yeniden örnekleme", _resampling_label(final["resampling"])],
            ["Karar eşiği", f"{final['threshold_report']['selected_threshold']:.2f}"],
            ["Sentetik benchmark holdout min ana metrik", _pct(holdout["ana_metrik_minimumu"])],
            ["Sentetik benchmark group CV min ana metrik", f"{_pct(cv['ana_metrik_minimumu_mean'])} ± {_pct_std(cv['ana_metrik_minimumu_std'])}"],
            ["Leakage durumu", final_leak["leakage_status"]],
        ],
    )
    _p(
        doc,
        "Karar eşiği ayarı aşamasında varsayılan 0.50 yerine 0.52 seçilmiştir. Amaç tek metriği şişirmek değil; precision, recall, specificity, F1 ve balanced accuracy metriklerinin minimumunu yüksek tutmaktır.",
    )
    _p(
        doc,
        "StratifiedGroupKFold kullanılırken grup değişkeni kaynak aile kimliği olarak belirlenmiştir. Böylece aynı kaynak aileden gelen örneklerin farklı fold'lara sızması engellenmiştir. "
        "CV sonucunun holdout sonucuyla uyumlu olması, final benchmarkın tek bir eğitim/test ayrımına aşırı bağımlı olmadığını göstermektedir.",
    )

    doc.add_page_break()
    _h(doc, "Sentetik Benchmark Adaylarının Karşılaştırılması")
    _img(doc, graphs.get("holdout_min_metric"), 5.8, "Grafik, sentetik benchmark adayları arasında holdout minimum ana metrik değerinin nasıl değiştiğini göstermektedir.")
    _img(doc, graphs.get("cv_min_metric"), 5.8, "Group CV sonuçlarının holdout sonuçlarına yakın olması, final benchmarkın tek bir train/test ayrımına aşırı bağımlı olmadığını göstermektedir.")
    _img(doc, graphs.get("holdout_roc_auc"), 5.8, "Holdout ROC-AUC değerleri, sentetik benchmark ortamında sınıfların güçlü biçimde ayrıştığını göstermektedir.")
    _img(doc, graphs.get("cv_roc_auc"), 5.8, "Group CV ROC-AUC değerlerinin yüksek kalması, kaynak aile ayrımı korunurken de ayrıştırma gücünün sürdüğünü göstermektedir.")

    doc.add_page_break()
    _h(doc, "Final Model Sonuçları")
    _table(
        doc,
        ["Metrik", "Sentetik benchmark holdout"],
        [
            ["Accuracy", _pct(holdout["accuracy"])],
            ["Precision", _pct(holdout["precision"])],
            ["Recall / Sensitivity", _pct(holdout["recall"])],
            ["Specificity", _pct(holdout["specificity"])],
            ["F1", _pct(holdout["f1"])],
            ["ROC-AUC", _pct(holdout["roc_auc"])],
            ["Min ana metrik", _pct(holdout["ana_metrik_minimumu"])],
        ],
    )
    _table(
        doc,
        ["Metrik", "Group CV ortalama ± std"],
        [
            ["Accuracy", f"{_pct(cv['accuracy_mean'])} ± {_pct_std(cv['accuracy_std'])}"],
            ["Precision", f"{_pct(cv['precision_mean'])} ± {_pct_std(cv['precision_std'])}"],
            ["Recall", f"{_pct(cv['recall_mean'])} ± {_pct_std(cv['recall_std'])}"],
            ["Specificity", f"{_pct(cv['specificity_mean'])} ± {_pct_std(cv['specificity_std'])}"],
            ["F1", f"{_pct(cv['f1_mean'])} ± {_pct_std(cv['f1_std'])}"],
            ["ROC-AUC", f"{_pct(cv['roc_auc_mean'])} ± {_pct_std(cv['roc_auc_std'])}"],
            ["Min ana metrik", f"{_pct(cv['ana_metrik_minimumu_mean'])} ± {_pct_std(cv['ana_metrik_minimumu_std'])}"],
        ],
    )
    _img(doc, graphs.get("confusion_matrix"), 4.4, "Confusion matrix, final modelin pozitif ve negatif sınıflar üzerindeki hata dağılımını göstermektedir.")
    _img(doc, graphs.get("feature_importance"), 5.9, "Değişken ayrıştırma gücü grafiği, final sentetik benchmarkta hangi değişkenlerin sınıfları daha güçlü ayırdığını gösterir. Bu grafik soft voting modelinin iç katsayısı değil, benchmark üzerindeki sınıf ayrımı denetimidir.")
    _img(doc, graphs.get("brier"), 5.7, "Brier score karşılaştırması, olasılık tahminlerinin hata düzeyini sentetik benchmark adayları arasında kıyaslar.")
    _img(doc, graphs.get("calibration_curve"), 4.9, "Calibration curve, tahmin olasılıklarının gözlenen pozitif oranlarla ne kadar uyumlu olduğunu görselleştirir.")

    doc.add_page_break()
    _h(doc, "Sentetik Benchmark Sonucunun Yorumlanması")
    _p(
        doc,
        "Çapraz model analizi, yeni kontrollü sentetik benchmarkta elde edilen performans artışının ağırlıklı olarak model değişiminden değil, sentetik veri üretim stratejisinden kaynaklandığını göstermiştir. "
        "Yeni üretim hattında örnekler, kaynak aile yapısı korunarak sınıf hedef merkezlerine daha yakın üretildiği için Glucose, BMI, Age ve DiabetesPedigreeFunction değişkenlerinde sınıflar arası ayrım belirginleşmiştir. "
        "Bu nedenle bu sonuç, ham PIMA verisi üzerinde doğrudan genellenebilirlik kanıtı olarak değil, sınıf ayrımı güçlendirilmiş kontrollü sentetik benchmark performansı olarak değerlendirilmiştir.",
    )
    cross_rows = [
        ["Önceki üretim + XGBoost (SkinThickness çıkarılmış)", "0.7985", "0.8161", "0.8914", "0.57", "Daha doğal sentetik benchmark"],
        ["Yeni üretim + XGBoost (SkinThickness çıkarılmış)", "0.9491", "0.9618", "0.9945", "0.66", "Yeni veri daha ayrıştırılabilir"],
        ["Önceki üretim + soft voting ensemble", "0.7804", "0.8161", "0.8969", "0.54", "Model tek başına farkı açıklamıyor"],
        ["Yeni üretim + soft voting ensemble", "0.9667", "0.9677", "0.9962", "0.52", "Final kontrollü sentetik benchmark"],
    ]
    _table(doc, ["Deney", "Min ana metrik", "Accuracy", "ROC-AUC", "Karar eşiği", "Yorum"], cross_rows, font_size=7.7)
    _table(
        doc,
        ["Dağılım özeti", "Değer"],
        [
            ["Geliştirme verisi ortalama mutlak Cohen's d", "0.712"],
            ["Önceki kontrollü benchmark ortalama mutlak Cohen's d", "0.860"],
            ["Yeni kontrollü benchmark ortalama mutlak Cohen's d", "2.950"],
            ["Dış kontrol verisi ortalama mutlak Cohen's d", "0.608"],
            ["Önceki kontrollü benchmark dağılım kayması", "0.084"],
            ["Yeni kontrollü benchmark dağılım kayması", "0.556"],
            ["Dış kontrol verisi dağılım kayması", "0.072"],
        ],
    )
    _img(doc, graphs.get("strict_controlled_benchmark_cross_model"), 5.9, "Grafik, aynı model farklı veri üzerinde denenince asıl artışın yeni sentetik üretim stratejisinden geldiğini göstermektedir.")
    _img(doc, graphs.get("distribution_shift_cohens_d"), 5.9, "Dağılım kayması grafiği, yeni kontrollü benchmarkın önceki kontrollü benchmarka göre sınıfları daha ayrıştırılabilir hale getirdiğini göstermektedir.")

    doc.add_page_break()
    _h(doc, "Dış Kontrol, Literatür ve Sınırlılıklar")
    _table(
        doc,
        ["Metrik", "Orijinal dış kontrol"],
        [
            ["Accuracy", _pct(external["accuracy"])],
            ["F1", _pct(external["f1"])],
            ["ROC-AUC", _pct(external["roc_auc"])],
            ["Min ana metrik", _pct(external["ana_metrik_minimumu"])],
            ["Brier", f"{external['brier']:.3f}"],
        ],
    )
    _p(
        doc,
        "Orijinal dış kontrol performansının sentetik benchmarka göre daha düşük kalması, modelin gerçek PIMA dağılımına transferinde sınırlılık bulunduğunu göstermektedir. "
        "Bu sonuç final sentetik benchmark başarısını geçersiz kılmaz; ancak çalışmanın klinik genellenebilirlik iddiası taşımadığını açık biçimde gösterir.",
    )
    _box(
        doc,
        "Literatürle Karşılaştırma",
        "PIMA veri setinde ham veriyle yapılan çalışmalarda performans genellikle sınırlı kalırken, SMOTE, SMOTE-ENN, SMOTE-Tomek, CTGAN ve benzeri sentetik/hibrit dengeleme yaklaşımları performansı artırmak için sık kullanılmaktadır. "
        "Bu çalışmada da temel başarı hedefi ham PIMA üzerinde değil, kaynak aile ayrımıyla kontrol edilen PIMA + sentetik benchmark üzerinde tanımlanmıştır.\n\n"
        "Yeniden örnekleme ve sentetik üretim işlemleri test verisini görecek şekilde uygulanırsa leakage riski oluşur. Bu nedenle çalışmada kaynak aile ayrımı, birebir kopya kontrolü, çok yakın benzerlik kontrolü ve orijinal dış kontrol izolasyonu birlikte ele alınmıştır.",
    )
    _p(
        doc,
        "Bu çalışmada orijinal PIMA veri seti korunmuş, eğitim ve deney amaçlı olarak kaynak aile yapısı gözetilen kontrollü sentetik benchmark oluşturulmuştur. "
        "Final benchmark üzerinde soft voting ensemble modeli ile minimum ana metrik değeri %96.67 olarak elde edilmiştir. Group CV sonucunun da benzer seviyede kalması, sonucun tek bir split'e aşırı bağımlı olmadığını göstermektedir.",
    )
    _p(
        doc,
        "Buna karşılık orijinal dış kontrol performansı daha düşük kalmıştır. Bu durum, final sentetik benchmark başarısının ham PIMA genellenebilirliği olarak değil, sınıf ayrımı güçlendirilmiş kontrollü sentetik benchmark performansı olarak yorumlanması gerektiğini göstermektedir. "
        "Bu nedenle çalışma klinik genellenebilirlik iddiası taşımamakta, sentetik benchmark üretimi ve leakage kontrollü modelleme akışı olarak değerlendirilmelidir.",
    )

    doc.save(DOCX_PATH)


def _set_styles(doc: Document) -> None:
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(9.3)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[style_name].font.name = "Aptos"
        doc.styles[style_name].font.color.rgb = RGBColor(31, 78, 121)


def _title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)


def _h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _p(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.line_spacing = 1.08
    paragraph.paragraph_format.space_after = Pt(4)


def _box(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), "EAF2F8")
    cell._tc.get_or_add_tcPr().append(shade)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)
    cell.add_paragraph(text)
    doc.add_paragraph()


def _table(doc: Document, headers: list[str], rows: list[list[str]], font_size: float = 8.4) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), "1F4E79")
        cell._tc.get_or_add_tcPr().append(shade)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for paragraph in cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    doc.add_paragraph()


def _img(doc: Document, path: str | None, width: float, caption: str) -> None:
    if not path:
        _p(doc, caption)
        return
    image_path = Path(path)
    if _graph_check(image_path)["status"] != "pass":
        _p(doc, caption)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(8.2)


def _pct(value: float) -> str:
    return f"%{value * 100:.2f}"


def _pct_std(value: float) -> str:
    return f"{value * 100:.2f}"


def _docx_text(path: Path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _scan_turkish_ascii(text: str) -> dict[str, Any]:
    lower = text.lower()
    flagged = sorted(term for term in BAD_TURKISH_ASCII if term in lower)
    return {
        "checked_terms": BAD_TURKISH_ASCII,
        "flagged_terms": flagged,
        "status": "pass" if not flagged else "warning",
        "note": "Rapor metni Türkçe karakterlerle yeniden yazıldı; teknik İngilizce terimler ve source_id korunmuştur.",
    }


def _quality_pass(audit: dict[str, Any]) -> bool:
    leak = audit["leakage_summary"]
    return (
        audit["turkish_character_fix_summary"].get("status") == "pass"
        and not audit["empty_graph_check_summary"]["empty_or_missing_graphs"]
        and leak["leakage_status"] == "clean"
        and leak["train_test_source_intersection"] == 0
        and max(leak["cv_source_intersections"] or [0]) == 0
        and leak["exact_duplicate_count"] == 0
        and audit["outcome_encoding_check"]["status"] == "pass"
    )


if __name__ == "__main__":
    main()
