from __future__ import annotations

import json
import shutil
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RAW_CSV = ROOT / "makine_ogrenmesi" / "veri" / "ham" / "diabetes.csv"
SMALL_DIR = ROOT / "makine_ogrenmesi" / "veri" / "deneysel" / "benchmark_scale_sweep"
REPORT_DIR = ROOT / "makine_ogrenmesi" / "raporlar"
GRAPH_DIR = REPORT_DIR / "grafikler_pima_korelasyon_deney_dagilim"
OUT_DOCX = REPORT_DIR / "pima_korelasyon_deney_dagilim_raporu.docx"
DESKTOP_DOCX = Path("/Users/beratkaanseven/Desktop/PIMA_Korelasyon_Deney_Dagilim_Raporu.docx")
OUT_JSON = REPORT_DIR / "pima_korelasyon_deney_dagilim_ozeti.json"
SWEEP_RESULTS_JSON = REPORT_DIR / "benchmark_scale_sweep_results.json"

FEATURES = [
    "Pregnancies",
    "Glucose",
    "BloodPressure",
    "SkinThickness",
    "Insulin",
    "BMI",
    "DiabetesPedigreeFunction",
    "Age",
]

FEATURE_TR = {
    "Pregnancies": "Gebelik",
    "Glucose": "Glikoz",
    "BloodPressure": "Kan basıncı",
    "SkinThickness": "Deri kalınlığı",
    "Insulin": "İnsülin",
    "BMI": "Vücut kitle indeksi",
    "DiabetesPedigreeFunction": "Diyabet soy geçmişi fonksiyonu",
    "Age": "Yaş",
    "Outcome": "Hedef sınıf",
}

FEATURE_SHORT_TR = {
    "Pregnancies": "Gebelik",
    "Glucose": "Glikoz",
    "BloodPressure": "Kan basıncı",
    "SkinThickness": "Deri kalınlığı",
    "Insulin": "İnsülin",
    "BMI": "VKİ",
    "DiabetesPedigreeFunction": "Soy geçmişi",
    "Age": "Yaş",
    "Outcome": "Sınıf",
}

MODEL_TR = {
    "soft_voting_xgb_lgbm_et": "XGBoost + LightGBM + ExtraTrees soft voting",
    "extra_trees": "ExtraTrees",
    "lightgbm": "LightGBM",
    "xgboost": "XGBoost",
    "random_forest": "Random Forest",
    "hist_gradient_boosting": "HistGradientBoosting",
    "catboost": "CatBoost",
}

FEATURE_SET_TR = {
    "all_features": "Tüm orijinal değişkenler",
    "no_skinthickness": "SkinThickness çıkarılmış",
    "high_signal_features": "Yüksek sinyal değişken seti",
    "clinical_interactions": "Klinik etkileşim değişkenleri",
    "compact_best_features": "Kompakt seçilmiş değişkenler",
}

EXPERIMENTS = [
    {
        "label": "500/500",
        "dataset_name": "scale_sweep_500_per_class_strength_0p40_low",
        "target_per_class": 500,
        "strength": 0.40,
        "noise_profile": "low",
        "previous_best_model": "XGBoost + LightGBM + ExtraTrees soft voting",
        "previous_feature_set": "SkinThickness çıkarılmış",
        "previous_status": "Geçmedi",
    },
    {
        "label": "650/650",
        "dataset_name": "scale_sweep_650_per_class_strength_0p50_low",
        "target_per_class": 650,
        "strength": 0.50,
        "noise_profile": "low",
        "previous_best_model": "XGBoost + LightGBM + ExtraTrees soft voting",
        "previous_feature_set": "SkinThickness çıkarılmış",
        "previous_status": "Geçmedi",
    },
    {
        "label": "800/800",
        "dataset_name": "scale_sweep_800_per_class_strength_0p60_low",
        "target_per_class": 800,
        "strength": 0.60,
        "noise_profile": "low",
        "previous_best_model": "XGBoost + LightGBM + ExtraTrees soft voting",
        "previous_feature_set": "Tüm orijinal PIMA değişkenleri",
        "previous_status": "Geçmedi",
    },
    {
        "label": "1000/1000",
        "dataset_name": "scale_sweep_1000_per_class_strength_0p60_adaptive",
        "target_per_class": 1000,
        "strength": 0.60,
        "noise_profile": "adaptive",
        "previous_best_model": "ExtraTrees",
        "previous_feature_set": "Yüksek sinyal değişken seti",
        "previous_status": "Seçildi",
    },
]


def ensure_dirs() -> None:
    GRAPH_DIR.mkdir(parents=True, exist_ok=True)
    REPORT_DIR.mkdir(parents=True, exist_ok=True)


def load_data() -> tuple[pd.DataFrame, list[dict]]:
    pima = pd.read_csv(RAW_CSV)
    loaded = []
    for exp in EXPERIMENTS:
        data_path = SMALL_DIR / f"{exp['dataset_name']}.csv"
        meta_path = SMALL_DIR / f"{exp['dataset_name']}_metadata.csv"
        if not data_path.exists():
            raise FileNotFoundError(data_path)
        if not meta_path.exists():
            raise FileNotFoundError(meta_path)
        df = pd.read_csv(data_path)
        meta = pd.read_csv(meta_path)
        item = dict(exp)
        item["data_path"] = data_path
        item["metadata_path"] = meta_path
        item["data"] = df
        item["metadata"] = meta
        loaded.append(item)
    return pima, loaded


def class_counts(df: pd.DataFrame) -> dict[str, int]:
    return {str(k): int(v) for k, v in df["Outcome"].value_counts().sort_index().items()}


def pct(x: float) -> str:
    return f"{100 * x:.1f}%"


def pct2(x: float) -> str:
    if pd.isna(x):
        return "-"
    return f"%{100 * x:.2f}"


def fmt(x: float, ndigits: int = 3) -> str:
    if pd.isna(x):
        return "-"
    return f"{x:.{ndigits}f}"


def human_model(name: str | None) -> str:
    if not name:
        return "-"
    return MODEL_TR.get(name, name)


def human_feature_set(name: str | None) -> str:
    if not name:
        return "-"
    return FEATURE_SET_TR.get(name, name)


def load_sweep_results() -> dict:
    if not SWEEP_RESULTS_JSON.exists():
        return {}
    return json.loads(SWEEP_RESULTS_JSON.read_text(encoding="utf-8"))


def best_model_effects(sweep: dict, dataset_names: list[str]) -> dict[str, dict]:
    if not sweep:
        return {}

    result_rows: list[dict] = []
    for key in ["phase1_results", "phase2_results", "optuna_results"]:
        rows = sweep.get(key, [])
        if isinstance(rows, list):
            result_rows.extend([row for row in rows if row.get("status") == "completed"])

    cv_rows = sweep.get("cv_results", [])
    effects: dict[str, dict] = {}
    for dataset_name in dataset_names:
        rows = [row for row in result_rows if row.get("dataset_name") == dataset_name]
        rows.sort(
            key=lambda row: (
                row.get("tuned_threshold_metrics", {}).get("ana_metrik_minimumu", -1),
                row.get("tuned_threshold_metrics", {}).get("accuracy", -1),
                row.get("tuned_threshold_metrics", {}).get("roc_auc", -1),
            ),
            reverse=True,
        )
        best = rows[0] if rows else {}

        matching_cv = [
            row
            for row in cv_rows
            if row.get("dataset_name") == dataset_name
            and row.get("model") == best.get("model")
            and row.get("feature_set") == best.get("feature_set")
            and row.get("resampling") == best.get("resampling")
        ]
        matching_cv.sort(
            key=lambda row: row.get("summary", {}).get("ana_metrik_minimumu_mean", -1),
            reverse=True,
        )
        all_cv = [row for row in cv_rows if row.get("dataset_name") == dataset_name]
        all_cv.sort(
            key=lambda row: row.get("summary", {}).get("ana_metrik_minimumu_mean", -1),
            reverse=True,
        )
        best_cv_same_model = matching_cv[0] if matching_cv else {}
        best_cv_overall = all_cv[0] if all_cv else {}
        metrics = best.get("tuned_threshold_metrics", {})
        cm = metrics.get("confusion_matrix", {})
        effects[dataset_name] = {
            "holdout_model_raw": best.get("model"),
            "holdout_model": human_model(best.get("model")),
            "holdout_feature_set_raw": best.get("feature_set"),
            "holdout_feature_set": human_feature_set(best.get("feature_set")),
            "holdout_resampling": best.get("resampling", "-"),
            "holdout_threshold": metrics.get("threshold"),
            "holdout_accuracy": metrics.get("accuracy"),
            "holdout_precision": metrics.get("precision"),
            "holdout_recall": metrics.get("recall"),
            "holdout_specificity": metrics.get("specificity"),
            "holdout_f1": metrics.get("f1"),
            "holdout_roc_auc": metrics.get("roc_auc"),
            "holdout_min_main": metrics.get("ana_metrik_minimumu"),
            "holdout_confusion_matrix": cm,
            "cv_same_model": {
                "available": bool(best_cv_same_model),
                "model": human_model(best_cv_same_model.get("model")),
                "feature_set": human_feature_set(best_cv_same_model.get("feature_set")),
                **best_cv_same_model.get("summary", {}),
            },
            "cv_best_overall": {
                "available": bool(best_cv_overall),
                "model": human_model(best_cv_overall.get("model")),
                "feature_set": human_feature_set(best_cv_overall.get("feature_set")),
                **best_cv_overall.get("summary", {}),
            },
        }
    return effects


def plot_corr(df: pd.DataFrame, title: str, out: Path) -> None:
    corr = df[FEATURES + ["Outcome"]].corr(numeric_only=True)
    labels = [FEATURE_SHORT_TR.get(col, col) for col in corr.columns]
    fig, ax = plt.subplots(figsize=(9.8, 7.6))
    im = ax.imshow(corr.values, vmin=-1, vmax=1, cmap="RdBu_r")
    ax.set_title(title, fontsize=14, fontweight="bold", pad=14)
    ax.set_xticks(range(len(corr.columns)))
    ax.set_xticklabels(labels, rotation=35, ha="right", fontsize=9)
    ax.set_yticks(range(len(corr.index)))
    ax.set_yticklabels(labels, fontsize=9)
    for i in range(corr.shape[0]):
        for j in range(corr.shape[1]):
            val = corr.iloc[i, j]
            ax.text(j, i, f"{val:.2f}", ha="center", va="center", fontsize=7, color="black")
    ax.set_xticks(np.arange(-0.5, len(corr.columns), 1), minor=True)
    ax.set_yticks(np.arange(-0.5, len(corr.index), 1), minor=True)
    ax.grid(which="minor", color="white", linestyle="-", linewidth=0.8)
    fig.colorbar(im, ax=ax, shrink=0.75)
    plt.tight_layout()
    plt.savefig(out, dpi=220)
    plt.close()


def plot_pima_distributions(pima: pd.DataFrame, out: Path) -> None:
    fig, axes = plt.subplots(4, 2, figsize=(10.5, 11.6))
    axes = axes.ravel()
    palette = {0: "#2D6A4F", 1: "#C0392B"}
    for ax, feature in zip(axes, FEATURES):
        for outcome, color in palette.items():
            vals = pima.loc[pima["Outcome"] == outcome, feature].dropna()
            ax.hist(vals, bins=24, alpha=0.50, color=color, density=True, label=f"Outcome {outcome}")
        ax.set_title(FEATURE_TR[feature], fontsize=10, fontweight="bold")
        ax.set_xlabel("")
        ax.set_ylabel("Yoğunluk")
        ax.grid(alpha=0.18)
    handles, labels = axes[0].get_legend_handles_labels()
    fig.legend(handles, labels, loc="upper center", ncol=2, frameon=False)
    fig.suptitle("Orijinal PIMA: Tüm Özelliklerin Sınıf Bazlı Dağılımı", fontsize=15, fontweight="bold", y=0.995)
    fig.tight_layout(rect=(0, 0, 1, 0.965))
    fig.savefig(out, dpi=220)
    plt.close(fig)


def plot_class_counts(pima: pd.DataFrame, experiments: list[dict], out: Path) -> None:
    rows = [{"Aday": "Orijinal PIMA", "Negatif": int((pima["Outcome"] == 0).sum()), "Pozitif": int((pima["Outcome"] == 1).sum())}]
    for exp in experiments:
        df = exp["data"]
        rows.append(
            {
                "Aday": exp["label"],
                "Negatif": int((df["Outcome"] == 0).sum()),
                "Pozitif": int((df["Outcome"] == 1).sum()),
            }
        )
    chart_df = pd.DataFrame(rows).set_index("Aday")
    ax = chart_df.plot(kind="bar", stacked=False, figsize=(9, 4.6), color=["#2D6A4F", "#C0392B"], width=0.72)
    ax.set_title("Orijinal PIMA ve Sentetik Adaylarda Sınıf Dağılımı", fontsize=13, fontweight="bold")
    ax.set_ylabel("Satır sayısı")
    ax.grid(axis="y", alpha=0.20)
    ax.legend(frameon=False)
    plt.xticks(rotation=0)
    plt.tight_layout()
    plt.savefig(out, dpi=220)
    plt.close()


def plot_synthetic_additions(summary_rows: list[dict], out: Path) -> None:
    chart_df = pd.DataFrame(
        [
            {
                "Aday": row["label"],
                "Eklenen negatif": row["added_negative"],
                "Eklenen pozitif": row["added_positive"],
            }
            for row in summary_rows
        ]
    ).set_index("Aday")[["Eklenen negatif", "Eklenen pozitif"]]
    ax = chart_df.plot(kind="bar", stacked=True, figsize=(8.5, 4.5), color=["#74C69D", "#F39C12"], width=0.68)
    ax.set_title("Adaylara Göre Eklenen Sentetik Satır Kırılımı", fontsize=13, fontweight="bold")
    ax.set_ylabel("Eklenen sentetik satır")
    ax.grid(axis="y", alpha=0.20)
    ax.legend(frameon=False)
    plt.xticks(rotation=0)
    plt.tight_layout()
    plt.savefig(out, dpi=220)
    plt.close()


def plot_model_effects(summaries: list[dict], out: Path) -> None:
    rows = []
    for row in summaries:
        effect = row.get("model_effect", {})
        cv = effect.get("cv_best_overall", {})
        rows.append(
            {
                "Aday": row["label"],
                "Accuracy": effect.get("holdout_accuracy", np.nan),
                "Holdout min ana": effect.get("holdout_min_main", np.nan),
                "Group CV min ana": cv.get("ana_metrik_minimumu_mean", np.nan),
                "ROC-AUC": effect.get("holdout_roc_auc", np.nan),
            }
        )
    chart_df = pd.DataFrame(rows).set_index("Aday") * 100
    ax = chart_df.plot(
        kind="bar",
        figsize=(10.2, 5.2),
        color=["#1B4332", "#2D6A4F", "#74C69D", "#F39C12"],
        width=0.76,
    )
    ax.axhline(92, color="#C0392B", linestyle="--", linewidth=1.2, label="Min ana hedefi %92")
    ax.axhline(93, color="#6C3483", linestyle=":", linewidth=1.2, label="Accuracy hedefi %93")
    ax.set_title("Adaylara Göre Model Etkisi", fontsize=13, fontweight="bold")
    ax.set_ylabel("Metrik (%)")
    ax.set_ylim(80, 100)
    ax.grid(axis="y", alpha=0.22)
    ax.legend(frameon=False, fontsize=8, ncol=3)
    plt.xticks(rotation=0)
    plt.tight_layout()
    plt.savefig(out, dpi=220)
    plt.close()


def plot_feature_change_summary(summaries: list[dict], out: Path) -> None:
    rows = []
    for row in summaries:
        rows.append(
            {
                "Aday": row["label"],
                "PIMA'ya net kayma": row["feature_change_summary"]["mean_abs_standardized_delta_vs_pima"],
                "Parent'a ort. oynama": row["feature_change_summary"]["mean_abs_parent_change_in_std"],
            }
        )
    chart_df = pd.DataFrame(rows).set_index("Aday")
    ax = chart_df.plot(kind="bar", figsize=(8.8, 4.8), color=["#2D6A4F", "#F39C12"], width=0.70)
    ax.set_title("Adaylara Göre Özellik Oynama Özeti", fontsize=13, fontweight="bold")
    ax.set_ylabel("Standart sapma birimi")
    ax.grid(axis="y", alpha=0.22)
    ax.legend(frameon=False, fontsize=8)
    plt.xticks(rotation=0)
    plt.tight_layout()
    plt.savefig(out, dpi=220)
    plt.close()


def standardized_mean_diff(candidate: pd.DataFrame, reference: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for feature in FEATURES:
        ref_std_all = reference[feature].std(ddof=0) or 1.0
        rows.append(
            {
                "feature": feature,
                "Genel": (candidate[feature].mean() - reference[feature].mean()) / ref_std_all,
                "Negatif": (
                    candidate.loc[candidate["Outcome"] == 0, feature].mean()
                    - reference.loc[reference["Outcome"] == 0, feature].mean()
                )
                / (reference.loc[reference["Outcome"] == 0, feature].std(ddof=0) or 1.0),
                "Pozitif": (
                    candidate.loc[candidate["Outcome"] == 1, feature].mean()
                    - reference.loc[reference["Outcome"] == 1, feature].mean()
                )
                / (reference.loc[reference["Outcome"] == 1, feature].std(ddof=0) or 1.0),
            }
        )
    return pd.DataFrame(rows).set_index("feature")


def feature_change_table(df: pd.DataFrame, pima: pd.DataFrame) -> tuple[list[dict], dict]:
    synthetic = df[df["is_synthetic"].astype(bool)].copy()
    original_dev = df[~df["is_synthetic"].astype(bool)].copy()
    parent_lookup = original_dev.drop_duplicates("source_id").set_index("source_id")
    aligned_parent = parent_lookup.reindex(synthetic["source_id"])

    rows: list[dict] = []
    parent_scaled_changes: list[float] = []
    standardized_deltas: list[float] = []
    for feature in FEATURES:
        pima_mean = float(pima[feature].mean())
        candidate_mean = float(df[feature].mean())
        pima_std = float(pima[feature].std(ddof=0)) or 1.0
        original_dev_std = float(original_dev[feature].std(ddof=0)) or 1.0
        net_delta = candidate_mean - pima_mean
        standardized_delta = net_delta / pima_std
        parent_delta = synthetic[feature].to_numpy(dtype=float) - aligned_parent[feature].to_numpy(dtype=float)
        synthetic_parent_abs_change = float(np.abs(parent_delta).mean())
        synthetic_parent_abs_change_std = synthetic_parent_abs_change / original_dev_std
        parent_scaled_changes.append(abs(synthetic_parent_abs_change_std))
        standardized_deltas.append(abs(standardized_delta))
        rows.append(
            {
                "feature": feature,
                "feature_tr": FEATURE_TR[feature],
                "pima_mean": pima_mean,
                "candidate_mean": candidate_mean,
                "net_delta": net_delta,
                "net_delta_percent": (net_delta / pima_mean) if abs(pima_mean) > 1e-9 else np.nan,
                "standardized_delta_vs_pima": standardized_delta,
                "synthetic_parent_abs_change": synthetic_parent_abs_change,
                "synthetic_parent_abs_change_std": synthetic_parent_abs_change_std,
                "negative_mean": float(df.loc[df["Outcome"] == 0, feature].mean()),
                "positive_mean": float(df.loc[df["Outcome"] == 1, feature].mean()),
            }
        )
    summary = {
        "mean_abs_standardized_delta_vs_pima": float(np.mean(standardized_deltas)),
        "mean_abs_parent_change_in_std": float(np.mean(parent_scaled_changes)),
        "max_changed_feature_by_parent": max(rows, key=lambda row: abs(row["synthetic_parent_abs_change_std"]))["feature"],
        "max_shift_feature_vs_pima": max(rows, key=lambda row: abs(row["standardized_delta_vs_pima"]))["feature"],
    }
    return rows, summary


def plot_diff_heatmap(diff_df: pd.DataFrame, title: str, out: Path) -> None:
    values = diff_df[["Genel", "Negatif", "Pozitif"]]
    lim = max(0.5, float(np.nanmax(np.abs(values.values))))
    labels = [FEATURE_SHORT_TR.get(feature, feature) for feature in values.index]
    fig, ax = plt.subplots(figsize=(7.4, 5.8))
    im = ax.imshow(values.values, cmap="BrBG", vmin=-lim, vmax=lim)
    ax.set_title(title, fontsize=13, fontweight="bold", pad=12)
    ax.set_xticks(range(len(values.columns)))
    ax.set_xticklabels(values.columns, fontsize=9)
    ax.set_yticks(range(len(values.index)))
    ax.set_yticklabels(labels, fontsize=9)
    for i in range(values.shape[0]):
        for j in range(values.shape[1]):
            ax.text(j, i, f"{values.iloc[i, j]:.2f}", ha="center", va="center", fontsize=8, color="black")
    ax.set_xticks(np.arange(-0.5, len(values.columns), 1), minor=True)
    ax.set_yticks(np.arange(-0.5, len(values.index), 1), minor=True)
    ax.grid(which="minor", color="white", linestyle="-", linewidth=0.8)
    cbar = fig.colorbar(im, ax=ax, shrink=0.75)
    cbar.set_label("Standartlaştırılmış ortalama fark")
    plt.tight_layout()
    plt.savefig(out, dpi=220)
    plt.close()


def top_differences(diff_df: pd.DataFrame, n: int = 4) -> list[str]:
    vals = diff_df["Genel"].abs().sort_values(ascending=False).head(n)
    return [f"{feature}: {diff_df.loc[feature, 'Genel']:+.2f}" for feature in vals.index]


def summarize_experiment(exp: dict, pima: pd.DataFrame, model_effect: dict | None = None) -> dict:
    df = exp["data"]
    meta = exp["metadata"]
    synthetic_mask = meta["is_synthetic"].astype(bool)
    synthetic_rows = df.loc[synthetic_mask.to_numpy()]
    original_rows = df.loc[~synthetic_mask.to_numpy()]

    added_by_class = synthetic_rows["Outcome"].value_counts().sort_index().to_dict()
    original_by_class = original_rows["Outcome"].value_counts().sort_index().to_dict()
    target = exp["target_per_class"]
    diff_df = standardized_mean_diff(df, pima)
    feature_rows, feature_summary = feature_change_table(df, pima)
    generation_counts = meta["generation_method"].value_counts().to_dict()
    parent_unique_total = int(meta.loc[synthetic_mask, "source_id"].nunique())
    parent_unique_neg = int(df.loc[synthetic_mask.to_numpy() & (df["Outcome"] == 0), "source_id"].nunique())
    parent_unique_pos = int(df.loc[synthetic_mask.to_numpy() & (df["Outcome"] == 1), "source_id"].nunique())
    independent_source_id_count = int(
        meta.loc[synthetic_mask, "source_id"].astype(str).str.startswith(("synthetic_", "gaussian_")).sum()
    )

    return {
        "label": exp["label"],
        "dataset_name": exp["dataset_name"],
        "target_per_class": target,
        "total_rows": int(len(df)),
        "original_dev_rows": int((~synthetic_mask).sum()),
        "synthetic_rows": int(synthetic_mask.sum()),
        "class_distribution": class_counts(df),
        "original_dev_distribution": {str(k): int(v) for k, v in original_by_class.items()},
        "added_negative": int(added_by_class.get(0, 0)),
        "added_positive": int(added_by_class.get(1, 0)),
        "generation_method_counts": {str(k): int(v) for k, v in generation_counts.items()},
        "synthetic_strength": exp["strength"],
        "noise_profile": exp["noise_profile"],
        "synthetic_parent_family_count": parent_unique_total,
        "synthetic_parent_family_negative": parent_unique_neg,
        "synthetic_parent_family_positive": parent_unique_pos,
        "independent_synthetic_source_id_count": independent_source_id_count,
        "previous_best_model": exp["previous_best_model"],
        "previous_feature_set": exp["previous_feature_set"],
        "previous_status": exp["previous_status"],
        "top_standardized_mean_differences": top_differences(diff_df),
        "standardized_mean_difference": {
            feature: {col: float(diff_df.loc[feature, col]) for col in diff_df.columns}
            for feature in diff_df.index
        },
        "feature_change_rows": feature_rows,
        "feature_change_summary": feature_summary,
        "model_effect": model_effect or {},
    }


def apply_doc_style(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.color.rgb = RGBColor(27, 67, 50)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(11)


def set_margins(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(27, 67, 50)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(45, 106, 79)


def add_note(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = cell.paragraphs[0].add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(45, 106, 79)
    shading = cell._tc.get_or_add_tcPr()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "D8F3DC")
    shading.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(8.5)
        shade_cell(hdr[i], "1B4332")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
            if len(table.rows) % 2 == 0:
                shade_cell(cells[i], "F4F6F6")
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Cm(width)


def shade_cell(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_image(doc: Document, path: Path, caption: str, width: float = 6.6) -> None:
    if not path.exists() or path.stat().st_size == 0:
        doc.add_paragraph(f"Görsel üretilemedi: {path.name}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(80, 80, 80)


def add_page_break(doc: Document) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    section = doc.sections[-1]
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)


def build_report(pima: pd.DataFrame, experiments: list[dict], summaries: list[dict], graph_paths: dict[str, Path]) -> None:
    doc = Document()
    set_margins(doc)
    apply_doc_style(doc)

    add_title(
        doc,
        "PIMA Veri Dağılımı ve Sentetik Aday Korelasyon Raporu",
        "Orijinal PIMA korelasyonu, özellik dağılımları ve 500/650/800/1000 sentetik benchmark adaylarının veri farkı",
    )
    add_note(
        doc,
        "Bu rapor model skoru şişirmek için değil, veri tarafını açık göstermek için üretildi. "
        "Sentetik adaylarda ham PIMA kolonlarına yeni klinik özellik eklenmedi; CSV düzeyinde orijinal PIMA değişkenleri korundu. "
        "Modelleme tarafında bazı denemelerde özellik seti filtreleri kullanılmış olsa da, veri dosyalarındaki temel kolon yapısı aynıdır.",
    )

    doc.add_heading("1. Orijinal PIMA Veri Seti", level=1)
    doc.add_paragraph(
        "Orijinal PIMA veri seti 768 satırdan oluşmaktadır. Hedef sınıf dağılımı 500 negatif ve 268 pozitif örnektir. "
        "Aşağıdaki grafiklerde önce ham veri dağılımı, ardından özelliklerin hedef sınıfla ve kendi aralarındaki ilişkisi gösterilmiştir."
    )
    add_table(
        doc,
        ["Alan", "Değer"],
        [
            ["Toplam satır", str(len(pima))],
            ["Negatif / Pozitif", f"{(pima['Outcome'] == 0).sum()} / {(pima['Outcome'] == 1).sum()}"],
            ["Kullanılan özellik sayısı", str(len(FEATURES))],
            ["Temel özellikler", ", ".join(FEATURE_TR[feature] for feature in FEATURES)],
        ],
        widths=[4, 12],
    )
    add_image(
        doc,
        graph_paths["pima_distributions"],
        "Şekil 1. Orijinal PIMA veri setindeki tüm özelliklerin negatif ve pozitif sınıfa göre dağılımı.",
        width=6.8,
    )
    add_image(
        doc,
        graph_paths["pima_corr"],
        "Şekil 2. Orijinal PIMA korelasyon matrisi. Glucose değişkeninin Outcome ile ilişkisi diğer özelliklere göre daha belirgindir.",
        width=6.2,
    )

    add_page_break(doc)
    doc.add_heading("2. Sentetik Adayların Kurulumu", level=1)
    doc.add_paragraph(
        "500/500, 650/650, 800/800 ve 1000/1000 adayları original_dev kümesinden türetilmiştir. "
        "Her sentetik satır, kaynak aldığı orijinal satırın source_id ailesine bağlıdır. "
        "External holdout bu üretime dahil edilmemiştir. Bu nedenle rapordaki eklenen satır sayıları, original_dev üzerine eklenen kontrollü sentetik satırları ifade eder."
    )
    add_table(
        doc,
        [
            "Aday",
            "Toplam",
            "Orijinal dev",
            "Eklenen sentetik",
            "Eklenen negatif",
            "Eklenen pozitif",
            "Strength",
            "Gürültü",
            "Önceki rapor modeli",
        ],
        [
            [
                s["label"],
                str(s["total_rows"]),
                str(s["original_dev_rows"]),
                str(s["synthetic_rows"]),
                str(s["added_negative"]),
                str(s["added_positive"]),
                fmt(s["synthetic_strength"], 2),
                s["noise_profile"],
                s["previous_best_model"],
            ]
            for s in summaries
        ],
    )
    add_image(
        doc,
        graph_paths["class_counts"],
        "Şekil 3. Orijinal PIMA ve sentetik adaylarda sınıf dengesi. Adaylarda negatif ve pozitif sınıflar eşitlenmiştir.",
        width=6.4,
    )
    add_image(
        doc,
        graph_paths["synthetic_additions"],
        "Şekil 4. Her adayda original_dev üzerine eklenen sentetik negatif ve pozitif satır sayısı.",
        width=6.2,
    )

    doc.add_heading("3. Adayların Modele Etkisi", level=1)
    doc.add_paragraph(
        "Bu bölümde her adayın en iyi holdout sonucu ve aynı aday için kaydedilen Group CV doğrulaması birlikte gösterilmiştir. "
        "Buradaki amaç, veri boyutu arttığında modelin sadece satır sayısı nedeniyle mi, yoksa değişken dağılımı daha ayrıştırıcı hale geldiği için mi güçlendiğini görünür kılmaktır."
    )
    add_table(
        doc,
        [
            "Aday",
            "En iyi holdout modeli",
            "Özellik seti",
            "Eşik",
            "Accuracy",
            "Min ana",
            "ROC-AUC",
            "Aynı model CV min",
            "En iyi CV min",
        ],
        [
            [
                s["label"],
                s.get("model_effect", {}).get("holdout_model", "-"),
                s.get("model_effect", {}).get("holdout_feature_set", "-"),
                fmt(s.get("model_effect", {}).get("holdout_threshold"), 2),
                pct2(s.get("model_effect", {}).get("holdout_accuracy")),
                pct2(s.get("model_effect", {}).get("holdout_min_main")),
                pct2(s.get("model_effect", {}).get("holdout_roc_auc")),
                (
                    pct2(s.get("model_effect", {}).get("cv_same_model", {}).get("ana_metrik_minimumu_mean"))
                    + " ± "
                    + pct2(s.get("model_effect", {}).get("cv_same_model", {}).get("ana_metrik_minimumu_std"))
                )
                if s.get("model_effect", {}).get("cv_same_model", {}).get("available")
                else "-",
                (
                    pct2(s.get("model_effect", {}).get("cv_best_overall", {}).get("ana_metrik_minimumu_mean"))
                    + " ± "
                    + pct2(s.get("model_effect", {}).get("cv_best_overall", {}).get("ana_metrik_minimumu_std"))
                )
                if s.get("model_effect", {}).get("cv_best_overall", {}).get("available")
                else "-",
            ]
            for s in summaries
        ],
    )
    add_image(
        doc,
        graph_paths["model_effects"],
        "Şekil 5. Adaylara göre model etkisi. 1000/1000 adayında Accuracy ve min ana metrik hedef bandına çıkmaktadır.",
        width=6.8,
    )

    doc.add_heading("4. Özellik Düzeyinde Ne Kadar Oynama Yapıldı?", level=1)
    doc.add_paragraph(
        "CSV düzeyinde yeni klinik kolon eklenmedi; oynama mevcut PIMA değişkenlerinin sentetik satırlarda aldığı değerler üzerinden yapıldı. "
        "Aşağıdaki özet iki farklı şeyi ayırır: PIMA'ya net kayma, aday veri setinin PIMA ortalamasından ne kadar uzaklaştığını; "
        "parent'a ortalama oynama ise sentetik satırların kendi kaynak ailesinden ortalama ne kadar uzak üretildiğini gösterir."
    )
    add_table(
        doc,
        ["Aday", "PIMA'ya net kayma", "Parent'a ort. oynama", "En çok oynayan parent özelliği", "PIMA'ya en çok kayan özellik"],
        [
            [
                s["label"],
                fmt(s["feature_change_summary"]["mean_abs_standardized_delta_vs_pima"], 3),
                fmt(s["feature_change_summary"]["mean_abs_parent_change_in_std"], 3),
                FEATURE_TR[s["feature_change_summary"]["max_changed_feature_by_parent"]],
                FEATURE_TR[s["feature_change_summary"]["max_shift_feature_vs_pima"]],
            ]
            for s in summaries
        ],
    )
    add_image(
        doc,
        graph_paths["feature_change_summary"],
        "Şekil 6. Adaylara göre özellik oynama özeti. Daha büyük adaylarda sınıf ayrımı güçlenirken veri müdahalesi de artar.",
        width=6.4,
    )

    doc.add_heading("5. PIMA'ya Göre Farkın Okunması", level=1)
    doc.add_paragraph(
        "Aşağıdaki fark grafikleri, her adayın tam PIMA veri setine göre standartlaştırılmış ortalama farkını gösterir. "
        "Pozitif değer, aday veri setinde ilgili özelliğin PIMA ortalamasından daha yüksek olduğunu; negatif değer ise daha düşük olduğunu gösterir. "
        "Bu tablo, sadece skor değil veri yapısı da savunulacaksa kritik önemdedir."
    )

    for section_no, (exp, summary) in enumerate(zip(experiments, summaries), start=6):
        add_page_break(doc)
        doc.add_heading(f"{section_no}. {summary['label']} Adayı: Veri Ekleme ve Korelasyon", level=1)
        doc.add_paragraph(
            f"{summary['label']} adayı {summary['total_rows']} satırdan oluşmaktadır. "
            f"Bu yapıda original_dev tarafındaki {summary['original_dev_rows']} satır korunmuş, "
            f"{summary['added_negative']} negatif ve {summary['added_positive']} pozitif olmak üzere "
            f"toplam {summary['synthetic_rows']} sentetik satır eklenmiştir. "
            f"Üretim kaynağı original_dev satırlarıdır; bağımsız synthetic_id kullanılmamıştır."
        )
        effect = summary.get("model_effect", {})
        if effect:
            doc.add_paragraph(
                f"Model etkisi açısından bu adayda en iyi holdout sonucu {effect.get('holdout_model', '-')} modeliyle alındı. "
                f"Accuracy {pct2(effect.get('holdout_accuracy'))}, min ana metrik {pct2(effect.get('holdout_min_main'))}, "
                f"ROC-AUC {pct2(effect.get('holdout_roc_auc'))} ve karar eşiği {fmt(effect.get('holdout_threshold'), 2)} olarak raporlandı."
            )
        add_table(
            doc,
            ["Alan", "Değer"],
            [
                ["Aday kurulumu", f"{summary['label']} - strength {fmt(summary['synthetic_strength'], 2)}, {summary['noise_profile']} gürültü"],
                ["Sınıf dağılımı", f"{summary['class_distribution'].get('0', 0)} negatif / {summary['class_distribution'].get('1', 0)} pozitif"],
                ["Sentetik kaynak aile sayısı", str(summary["synthetic_parent_family_count"])],
                ["Negatif sentetik kaynak aile", str(summary["synthetic_parent_family_negative"])],
                ["Pozitif sentetik kaynak aile", str(summary["synthetic_parent_family_positive"])],
                ["Bağımsız sentetik source_id", str(summary["independent_synthetic_source_id_count"])],
                ["Önceki rapordaki model", summary["previous_best_model"]],
                ["Önceki rapordaki özellik seti", summary["previous_feature_set"]],
                ["En belirgin PIMA farkları", "; ".join(summary["top_standardized_mean_differences"])],
            ],
            widths=[5, 11],
        )
        add_table(
            doc,
            ["Özellik", "PIMA ort.", "Aday ort.", "Net fark", "Net %", "Parent oynama", "Neg. ort.", "Poz. ort."],
            [
                [
                    row["feature_tr"],
                    fmt(row["pima_mean"], 2),
                    fmt(row["candidate_mean"], 2),
                    fmt(row["net_delta"], 2),
                    pct2(row["net_delta_percent"]),
                    fmt(row["synthetic_parent_abs_change"], 2),
                    fmt(row["negative_mean"], 2),
                    fmt(row["positive_mean"], 2),
                ]
                for row in summary["feature_change_rows"]
            ],
        )
        add_image(
            doc,
            graph_paths[f"{summary['dataset_name']}_diff"],
            f"Şekil. {summary['label']} adayının PIMA'ya göre genel, negatif ve pozitif sınıf bazlı standartlaştırılmış ortalama farkı.",
            width=5.8,
        )
        add_image(
            doc,
            graph_paths[f"{summary['dataset_name']}_corr"],
            f"Şekil. {summary['label']} adayının korelasyon matrisi. Bu matris, sentetik üretim sonrası değişken ilişkilerinin ne ölçüde değiştiğini gösterir.",
            width=6.2,
        )

    add_page_break(doc)
    doc.add_heading("10. Kısa Teknik Yorum", level=1)
    doc.add_paragraph(
        "Bu adaylarda veri artırımı sınıf dengesini sağlamak için yapılmıştır. 500/500 adayında eklenen veri miktarı daha sınırlı olduğu için PIMA'ya yakınlık daha savunulabilir kalır; "
        "1000/1000 adayında ise eklenen sentetik satır sayısı arttığı için sınıf ayrımı daha belirginleşir. "
        "Bu durum korelasyon ve standartlaştırılmış fark grafiklerinde doğrudan görülebilir."
    )
    doc.add_paragraph(
        "Raporun ana teknik sonucu şudur: Skor yükseldikçe veri üretiminin PIMA dağılımına müdahale etkisi de artmaktadır. "
        "Bu nedenle 500/500, veri sadakati açısından daha temkinli; 1000/1000 ise performans açısından daha güçlü fakat daha müdahaleci adaydır."
    )

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    shutil.copy2(OUT_DOCX, DESKTOP_DOCX)


def main() -> None:
    ensure_dirs()
    plt.rcParams["font.family"] = "Arial"
    plt.rcParams["axes.grid"] = False
    pima, experiments = load_data()
    sweep = load_sweep_results()
    model_effect_map = best_model_effects(sweep, [exp["dataset_name"] for exp in experiments])

    pima_corr = GRAPH_DIR / "pima_korelasyon_matrisi.png"
    pima_dist = GRAPH_DIR / "pima_tum_ozellik_dagilimlari.png"
    class_counts_png = GRAPH_DIR / "aday_sinif_dagilimlari.png"
    synthetic_add_png = GRAPH_DIR / "aday_sentetik_eklenen_satirlar.png"
    model_effect_png = GRAPH_DIR / "aday_model_etkisi.png"
    feature_change_png = GRAPH_DIR / "aday_ozellik_oynama_ozeti.png"

    plot_corr(pima, "Orijinal PIMA Korelasyon Matrisi", pima_corr)
    plot_pima_distributions(pima, pima_dist)
    plot_class_counts(pima, experiments, class_counts_png)

    summaries = []
    graph_paths: dict[str, Path] = {
        "pima_corr": pima_corr,
        "pima_distributions": pima_dist,
        "class_counts": class_counts_png,
        "synthetic_additions": synthetic_add_png,
        "model_effects": model_effect_png,
        "feature_change_summary": feature_change_png,
    }

    for exp in experiments:
        summary = summarize_experiment(exp, pima, model_effect_map.get(exp["dataset_name"], {}))
        summaries.append(summary)
        diff_df = standardized_mean_diff(exp["data"], pima)
        corr_path = GRAPH_DIR / f"{exp['dataset_name']}_korelasyon_matrisi.png"
        diff_path = GRAPH_DIR / f"{exp['dataset_name']}_pima_fark_heatmap.png"
        plot_corr(exp["data"], f"{exp['label']} Adayı Korelasyon Matrisi", corr_path)
        plot_diff_heatmap(diff_df, f"{exp['label']} Adayı - PIMA'ya Göre Standartlaştırılmış Fark", diff_path)
        graph_paths[f"{exp['dataset_name']}_corr"] = corr_path
        graph_paths[f"{exp['dataset_name']}_diff"] = diff_path

    plot_synthetic_additions(summaries, synthetic_add_png)
    plot_model_effects(summaries, model_effect_png)
    plot_feature_change_summary(summaries, feature_change_png)

    output = {
        "raw_pima": {
            "path": str(RAW_CSV),
            "rows": int(len(pima)),
            "class_distribution": class_counts(pima),
            "features": FEATURES,
            "correlation_matrix": pima[FEATURES + ["Outcome"]].corr(numeric_only=True).round(6).to_dict(),
        },
        "experiments": summaries,
        "graphs": {k: str(v) for k, v in graph_paths.items()},
        "word_report": str(OUT_DOCX),
        "desktop_copy": str(DESKTOP_DOCX),
    }
    OUT_JSON.write_text(json.dumps(output, ensure_ascii=False, indent=2), encoding="utf-8")
    build_report(pima, experiments, summaries, graph_paths)

    print("PIMA korelasyon ve deney dağılım raporu üretildi.")
    print(f"Word: {OUT_DOCX}")
    print(f"Masaüstü: {DESKTOP_DOCX}")
    print(f"JSON: {OUT_JSON}")
    print(f"Grafik klasörü: {GRAPH_DIR}")


if __name__ == "__main__":
    main()
