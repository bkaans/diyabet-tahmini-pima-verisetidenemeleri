"""Sentetik benchmark icin saglamlik ve risk audit DOCX raporu uretir."""

from __future__ import annotations

import json
import shutil
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

PROJE_KOKU = Path(__file__).resolve().parents[1]
if str(PROJE_KOKU) not in sys.path:
    sys.path.insert(0, str(PROJE_KOKU))

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


RAPOR_DIR = PROJE_KOKU / "makine_ogrenmesi" / "raporlar"
GRAFIK_DIR = RAPOR_DIR / "grafikler_min90_risk_audit"
SWEEP_JSON = RAPOR_DIR / "min90_small_scale_sweep_results.json"
SELECTION_JSON = RAPOR_DIR / "min90_small_scale_selection_report.json"
AUDIT_JSON = RAPOR_DIR / "sentetik_benchmark_risk_audit.json"
DOCX_PATH = RAPOR_DIR / "sentetik_benchmark_risk_audit_raporu.docx"
DESKTOP_DOCX = Path("/Users/beratkaanseven/Desktop/Sentetik_Benchmark_Risk_Audit_Raporu.docx")


def main() -> None:
    RAPOR_DIR.mkdir(parents=True, exist_ok=True)
    GRAFIK_DIR.mkdir(parents=True, exist_ok=True)
    sweep = _read_json(SWEEP_JSON)
    selection = _read_json(SELECTION_JSON)
    candidate_rows = _candidate_rows(sweep, selection)
    focus_rows = [row for row in candidate_rows if row["target"] <= int(selection["final"]["target_per_class"])]
    risk_summary = _risk_summary(selection, candidate_rows)
    graphics = _make_graphics(focus_rows, selection, risk_summary)
    audit = {
        "created_at_utc": datetime.now(timezone.utc).isoformat(),
        "source_files": {"sweep": str(SWEEP_JSON), "selection": str(SELECTION_JSON)},
        "candidate_rows": candidate_rows,
        "focus_rows_500_to_1000": focus_rows,
        "risk_summary": risk_summary,
        "graphics": {k: str(v) for k, v in graphics.items()},
        "quality_metric_definitions": {
            "cohens_d": "Outcome=0 ve Outcome=1 siniflari arasindaki ortalama mutlak standartlastirilmis fark. Deger yukseldikce siniflar daha kolay ayrisir.",
            "distribution_shift": "Aday benchmark ortalamasinin original_dev ortalamasindan sapmasi. Her ozellik icin original_dev standart sapmasina bolunur ve ortalamasi alinir.",
            "note": "Cohen's d ile shift ayni sey degildir; siniflar arasi ayrim artarken genel ortalama kaymasi dusuk kalabilir.",
        },
        "notes": [
            "Bu rapor yeni model egitmez; mevcut small-scale sweep sonuclarini risk denetimi icin ozetler.",
            "Original dış kontrol sentetik benchmark basari hedefi olarak yorumlanmamistir.",
            "Rapor ana karar hattini 500/500 adayindan 1000/1000 adayina kadar odaklar.",
            "500/500 daha muhafazakar gorunse de hedef metrikleri tam karsilamadigi icin final aday degildir.",
        ],
    }
    AUDIT_JSON.write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    _write_docx(sweep, selection, focus_rows, risk_summary, graphics)
    shutil.copy2(DOCX_PATH, DESKTOP_DOCX)
    print(json.dumps({"docx": str(DOCX_PATH), "desktop_docx": str(DESKTOP_DOCX), "audit_json": str(AUDIT_JSON)}, ensure_ascii=False, indent=2))


def _read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def _candidate_rows(sweep: dict[str, Any], selection: dict[str, Any]) -> list[dict[str, Any]]:
    best_by_target: dict[int, dict[str, Any]] = {}
    for row in [*sweep["phase1_results"], *sweep["phase2_results"], *sweep["optuna_results"]]:
        if row.get("status") != "completed":
            continue
        target = int(row["target_per_class"])
        current = best_by_target.get(target)
        if current is None or _rank_key(row) > _rank_key(current):
            best_by_target[target] = row

    cv_by_target: dict[int, dict[str, Any]] = {}
    for row in sweep["cv_results"]:
        target = int(row.get("target_per_class") or _target_from_dataset(row["dataset_name"]))
        current = cv_by_target.get(target)
        if current is None or row["summary"]["ana_metrik_minimumu_mean"] > current["summary"]["ana_metrik_minimumu_mean"]:
            cv_by_target[target] = row

    final_target = int(selection["final"]["target_per_class"])
    quality_by_dataset = sweep["quality_by_dataset"]
    rows: list[dict[str, Any]] = []
    final_target = int(selection["final"]["target_per_class"])
    final_row = selection["final"]
    final_cv_summary = selection["final_cv"]["summary"]
    for target in sorted(best_by_target):
        row = final_row if target == final_target else best_by_target[target]
        metrics = row["tuned_threshold_metrics"]
        cv = final_cv_summary if target == final_target else cv_by_target.get(target, {}).get("summary", {})
        quality = selection["quality_summary"] if target == final_target else quality_by_dataset[row["dataset_name"]]
        passes_holdout = metrics["ana_metrik_minimumu"] >= 0.92 and metrics["accuracy"] >= 0.93
        status = "Geçmedi"
        if target == final_target:
            status = "Seçildi"
        elif passes_holdout and target >= 2500:
            status = "Referans / daha büyük"
        elif passes_holdout:
            status = "Geçti ama daha büyük"
        rows.append(
            {
                "target": target,
                "candidate": f"{target}/{target}",
                "dataset_name": row["dataset_name"],
                "best_model": _model_label(row["model"]),
                "accuracy": metrics["accuracy"],
                "min_main": metrics["ana_metrik_minimumu"],
                "roc_auc": metrics["roc_auc"],
                "confusion_matrix": metrics["confusion_matrix"],
                "cv_min_main": cv.get("ana_metrik_minimumu_mean"),
                "cv_min_main_std": cv.get("ana_metrik_minimumu_std"),
                "cv_accuracy": cv.get("accuracy_mean"),
                "cohens_d": quality["class_separation"]["aggregate"]["mean_abs_cohens_d"],
                "shift": quality["distribution_shift_vs_original_dev"]["aggregate"]["avg_abs_z_mean_shift"],
                "status": status,
            }
        )
    return rows


def _rank_key(row: dict[str, Any]) -> tuple[float, float, float]:
    m = row["tuned_threshold_metrics"]
    return (float(m["ana_metrik_minimumu"]), float(m["accuracy"]), float(m["roc_auc"]))


def _target_from_dataset(name: str) -> int:
    # small_t1000_s0p6_adaptive -> 1000
    return int(name.split("_")[1][1:])


def _risk_summary(selection: dict[str, Any], rows: list[dict[str, Any]]) -> dict[str, Any]:
    final = selection["final"]
    final_cv = selection["final_cv"]["summary"]
    external = selection["external_holdout"]["metrics"]
    quality = selection["quality_summary"]
    leak = selection["leakage_summary"]
    final_metrics = final["tuned_threshold_metrics"]
    row_500 = next((r for r in rows if r["target"] == 500), None)
    selected = next(r for r in rows if r["target"] == int(final["target_per_class"]))
    return {
        "selected_candidate": selected,
        "external_holdout": external,
        "synthetic_holdout_accuracy_gap": final_metrics["accuracy"] - external["accuracy"],
        "synthetic_cv_accuracy_gap": final_cv["accuracy_mean"] - external["accuracy"],
        "synthetic_min_main_gap": final_metrics["ana_metrik_minimumu"] - external["ana_metrik_minimumu"],
        "external_error_count": int(sum(final_metrics["confusion_matrix"]["matrix"][0]) + sum(final_metrics["confusion_matrix"]["matrix"][1])) if False else None,
        "selected_external_false_negative": external["confusion_matrix"]["fn"],
        "selected_synthetic_false_negative": final_metrics["confusion_matrix"]["fn"],
        "selected_shift": quality["distribution_shift_vs_original_dev"]["aggregate"]["avg_abs_z_mean_shift"],
        "selected_cohens_d": quality["class_separation"]["aggregate"]["mean_abs_cohens_d"],
        "candidate_500": row_500,
        "leakage_status": leak["leakage_status"],
        "source_intersection": leak["train_test_source_intersection"],
        "exact_duplicate_count": leak["exact_duplicate_count"],
        "near_duplicate_rate": leak["near_duplicate_rate"],
        "risk_levels": {
            "synthetic_overfitting": "yüksek",
            "imputation_leakage": "düşük-orta",
            "man_made_separation": "orta",
            "small_data_variance": "orta",
            "ensemble_overfit": "orta",
        },
    }


def _make_graphics(rows: list[dict[str, Any]], selection: dict[str, Any], risk: dict[str, Any]) -> dict[str, Path]:
    plt.rcParams.update({"font.size": 9, "axes.titlesize": 11, "axes.labelsize": 9})
    paths = {
        "gap": _plot_gap(selection),
        "candidate_risk": _plot_candidate_risk(rows),
        "risk_level": _plot_risk_level(risk),
        "candidate_confusion": _plot_candidate_confusion(rows),
    }
    return paths


def _plot_gap(selection: dict[str, Any]) -> Path:
    path = GRAFIK_DIR / "risk_audit_gap_analizi.png"
    final = selection["final"]["tuned_threshold_metrics"]
    cv = selection["final_cv"]["summary"]
    external = selection["external_holdout"]["metrics"]
    labels = ["Sentetik holdout", "Grup çapraz doğrulama", "Dış kontrol"]
    acc = [final["accuracy"], cv["accuracy_mean"], external["accuracy"]]
    min_main = [final["ana_metrik_minimumu"], cv["ana_metrik_minimumu_mean"], external["ana_metrik_minimumu"]]
    x = range(len(labels))
    fig, ax = plt.subplots(figsize=(7.4, 3.6))
    ax.bar([i - 0.18 for i in x], acc, width=0.36, label="Doğruluk", color="#1f4e79")
    ax.bar([i + 0.18 for i in x], min_main, width=0.36, label="Min ana metrik", color="#f4a261")
    ax.set_xticks(list(x), labels)
    ax.set_ylim(0.50, 1.02)
    ax.set_title("Sentetik Benchmark ile Dış Kontrol Arasındaki Skor Farkı")
    ax.legend()
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _plot_candidate_risk(rows: list[dict[str, Any]]) -> Path:
    path = GRAFIK_DIR / "risk_audit_aday_karsilastirma.png"
    labels = [str(r["target"]) for r in rows]
    min_main = [r["min_main"] for r in rows]
    shift = [r["shift"] for r in rows]
    fig, ax1 = plt.subplots(figsize=(7.6, 3.6))
    ax1.bar(labels, min_main, color="#1f4e79", label="Holdout min ana")
    ax1.axhline(0.92, color="#d62828", linestyle="--", linewidth=1)
    ax1.set_ylim(0.80, 1.0)
    ax1.set_xlabel("Sınıf başı hedef")
    ax1.set_ylabel("Min ana metrik")
    ax2 = ax1.twinx()
    ax2.plot(labels, shift, color="#e76f51", marker="o", label="Dağılım kayması")
    ax2.set_ylabel("Shift")
    ax1.set_title("Skor Artışı ve Dağılım Kayması Birlikte Değerlendirme")
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _plot_risk_level(risk: dict[str, Any]) -> Path:
    path = GRAFIK_DIR / "risk_audit_risk_duzeyleri.png"
    labels = ["Sentetik aşırı uyum", "İmputasyon", "Yapay ayrışma", "Varyans", "Ensemble"]
    mapping = {"düşük": 1, "düşük-orta": 2, "orta": 3, "yüksek": 4}
    values = [
        mapping[risk["risk_levels"]["synthetic_overfitting"]],
        mapping[risk["risk_levels"]["imputation_leakage"]],
        mapping[risk["risk_levels"]["man_made_separation"]],
        mapping[risk["risk_levels"]["small_data_variance"]],
        mapping[risk["risk_levels"]["ensemble_overfit"]],
    ]
    fig, ax = plt.subplots(figsize=(7.2, 3.2))
    colors = ["#d62828" if v == 4 else "#f4a261" if v == 3 else "#2a9d8f" for v in values]
    ax.bar(labels, values, color=colors)
    ax.set_ylim(0, 4.5)
    ax.set_yticks([1, 2, 3, 4], ["Düşük", "Düşük-orta", "Orta", "Yüksek"])
    ax.tick_params(axis="x", rotation=15)
    ax.set_title("Risk Denetimi Özet Düzeyleri")
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _plot_candidate_confusion(rows: list[dict[str, Any]]) -> Path:
    path = GRAFIK_DIR / "risk_audit_aday_hata_matrisi_heatmap.png"
    fig, axes = plt.subplots(2, 2, figsize=(7.4, 5.8))
    axes_flat = axes.ravel()
    vmax = max(
        max(row["confusion_matrix"]["tn"], row["confusion_matrix"]["fp"], row["confusion_matrix"]["fn"], row["confusion_matrix"]["tp"])
        for row in rows
    )
    for ax, row in zip(axes_flat, rows):
        cm = row["confusion_matrix"]
        matrix = [[cm["tn"], cm["fp"]], [cm["fn"], cm["tp"]]]
        ax.imshow(matrix, cmap="YlGnBu", vmin=0, vmax=vmax)
        ax.set_title(f"{row['candidate']} - {row['best_model']}", fontsize=9)
        ax.set_xticks([0, 1], ["Tahmin 0", "Tahmin 1"])
        ax.set_yticks([0, 1], ["Gerçek 0", "Gerçek 1"])
        for i in range(2):
            for j in range(2):
                value = matrix[i][j]
                ax.text(j, i, str(value), ha="center", va="center", color="black", fontsize=10, fontweight="bold")
    for ax in axes_flat[len(rows):]:
        ax.axis("off")
    fig.suptitle("Adaylara Göre Sentetik Holdout Hata Matrisleri", fontsize=12, fontweight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.94])
    fig.savefig(path, dpi=180)
    plt.close(fig)
    return path


def _write_docx(
    sweep: dict[str, Any],
    selection: dict[str, Any],
    rows: list[dict[str, Any]],
    risk: dict[str, Any],
    graphics: dict[str, Path],
) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    doc.styles["Normal"].font.name = "Aptos"
    doc.styles["Normal"].font.size = Pt(8.8)
    for style in ["Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[style].font.name = "Aptos"
        doc.styles[style].font.color.rgb = RGBColor(31, 78, 121)

    final = selection["final"]
    final_metrics = final["tuned_threshold_metrics"]
    final_cv = selection["final_cv"]["summary"]
    external = selection["external_holdout"]["metrics"]

    _title(doc, "PIMA Sentetik Benchmark Sağlamlık ve Risk Denetimi Raporu")
    _p(
        doc,
        "Bu rapor, küçük ölçekli sentetik benchmark sonuçlarını yalnız skor odaklı değil; sentetik veriye aşırı uyum, veri sızıntısı, yapay ayrışma, küçük veri varyansı ve model karmaşıklığı açısından denetlemek için hazırlanmıştır.",
    )
    _p(
        doc,
        "Ana sonuç değişmemiştir: 1000/1000 sınıf dengesi hedef metrikleri geçen en küçük adaydır. 500/500 adayı daha muhafazakar görünmesine rağmen hedefleri tam karşılamadığı ve grup çapraz doğrulama tarafında zayıf kaldığı için final aday yapılmamıştır.",
    )
    _table(
        doc,
        ["Alan", "Değer"],
        [
            ["Seçilen aday", "1000/1000 sınıf dengesi"],
            ["Model", "ExtraTrees"],
            ["Sentetik holdout doğruluk", _pct(final_metrics["accuracy"])],
            ["Sentetik holdout min ana", _pct(final_metrics["ana_metrik_minimumu"])],
            ["Grup çapraz doğrulama min ana", f"{_pct(final_cv['ana_metrik_minimumu_mean'])} ± {final_cv['ana_metrik_minimumu_std']:.3f}"],
            ["Dış kontrol doğruluk", _pct(external["accuracy"])],
            ["Dış kontrol ROC-AUC", _pct(external["roc_auc"])],
            ["Sızıntı durumu", _leakage_label(selection["leakage_summary"]["leakage_status"])],
        ],
    )
    _image(doc, graphics["gap"], 6.3, "Grafik, sentetik benchmark skorları ile orijinal dış kontrol sonucu arasındaki farkı göstermektedir. Bu fark, sonucun ham PIMA genellenebilirliği olarak değil sentetik benchmark başarısı olarak yorumlanması gerektiğini gösterir.")

    doc.add_page_break()
    doc.add_heading("500/500'den 1000/1000'e Karar Hattı", level=1)
    _p(
        doc,
        "Bu bölüm yalnız 500/500 ile 1000/1000 arasındaki karar hattını gösterir. Amaç, en küçük adaydan başlayarak hangi noktada hem sentetik holdout hedeflerinin hem de grup çapraz doğrulama kararlılığının yeterli seviyeye geldiğini açıkça göstermektir.",
    )
    _p(
        doc,
        "Cohen's d ve dağılım kayması aynı ölçü değildir. Cohen's d, sınıfların birbirinden ne kadar ayrıldığını gösterir; dağılım kayması ise aday verinin original_dev ortalamasından ne kadar uzaklaştığını ölçer. Bu yüzden 1000/1000 adayında sınıf ayrımı artarken genel dağılım kayması 500/500'e çok yakın, hatta biraz daha düşük görünebilir.",
    )
    _table(
        doc,
        ["Aday", "En iyi model", "Doğruluk", "Min ana", "Grup çapraz doğrulama min ana", "Cohen's d", "Shift", "Durum"],
        [
            [
                r["candidate"],
                r["best_model"],
                _pct(r["accuracy"]),
                _pct(r["min_main"]),
                _cv_text(r),
                f"{r['cohens_d']:.3f}",
                f"{r['shift']:.3f}",
                r["status"],
            ]
            for r in rows
        ],
        font_size=6.0,
    )
    _image(doc, graphics["candidate_risk"], 6.4, "Grafik, 500/500'den 1000/1000'e geçerken minimum ana metriğin hedefe yaklaştığını gösterir. 1000/1000 hedefi geçen ilk adaydır; daha büyük adaylar bu risk denetiminde ana karar eksenine alınmamıştır.")

    doc.add_heading("Aday Bazlı Hata Matrisleri", level=2)
    _p(
        doc,
        "Aşağıdaki tablo ve sıcaklık haritası, her adayın sentetik holdout üzerindeki hata dağılımını gösterir. Özellikle yanlış negatif (FN) değeri önemlidir; çünkü pozitif sınıfı kaçırma eğilimini doğrudan gösterir.",
    )
    _table(
        doc,
        ["Aday", "Model", "TN", "FP", "FN", "TP", "Kısa yorum"],
        [
            [
                r["candidate"],
                r["best_model"],
                str(r["confusion_matrix"]["tn"]),
                str(r["confusion_matrix"]["fp"]),
                str(r["confusion_matrix"]["fn"]),
                str(r["confusion_matrix"]["tp"]),
                _confusion_comment(r),
            ]
            for r in rows
        ],
        font_size=6.6,
    )
    _image(doc, graphics["candidate_confusion"], 6.4, "Sıcaklık haritaları, 500/500 ile 1000/1000 arasındaki adaylarda hataların nasıl dağıldığını görsel olarak gösterir. 1000/1000 adayında hem yanlış pozitif hem de yanlış negatif sayısı hedeflere göre daha dengeli kalmıştır.")

    doc.add_page_break()
    doc.add_heading("Temel Riskler ve Teknik Teşhis", level=1)
    _table(
        doc,
        ["Risk", "Gözlenen kanıt", "Düzey", "Rapor dili / önlem"],
        [
            [
                "Sentetik veriye aşırı uyum",
                f"Sentetik holdout doğruluk {_pct(final_metrics['accuracy'])}, dış kontrol doğruluk {_pct(external['accuracy'])}. Fark {_pct(risk['synthetic_holdout_accuracy_gap'])}.",
                "Yüksek",
                "Sonuç sentetik benchmark bağlamında yorumlanmalı; dış kontrol ayrı dış kontrol olarak kalmalı.",
            ],
            [
                "İmputasyon sızıntısı",
                "Dış kontrol en başta ayrılmıştır; ancak daha sıkı çapraz doğrulama için fold-içi imputasyon ayrıca önerilir.",
                "Düşük-orta",
                "Final raporda dış kontrol izolasyonu ve fold-içi imputasyon sınırlılığı açık yazılmalı.",
            ],
            [
                "Yapay ayrışma",
                f"500/500 Cohen's d {risk['candidate_500']['cohens_d']:.3f}, shift {risk['candidate_500']['shift']:.3f}; 1000/1000 Cohen's d {risk['selected_cohens_d']:.3f}, shift {risk['selected_shift']:.3f}.",
                "Orta",
                "500/500 daha muhafazakar ancak hedef altı; 1000/1000 hedefi geçen ilk eşik olarak sunulmalı.",
            ],
            [
                "Küçük veri varyansı",
                f"Grup çapraz doğrulama min ana {_pct(final_cv['ana_metrik_minimumu_mean'])} ± {final_cv['ana_metrik_minimumu_std']:.3f}.",
                "Orta",
                "Repeated grup çapraz doğrulama veya seed audit ek kontrol olarak önerilmeli.",
            ],
            [
                "Ensemble overfit",
                "Final küçük adayda kompleks soft voting yerine ExtraTrees seçildi.",
                "Orta",
                "Basit modelin hedefi geçmesi, gereksiz ensemble karmaşıklığını azaltır.",
            ],
        ],
        font_size=6.6,
    )
    _image(doc, graphics["risk_level"], 6.2, "Risk düzeyleri, raporda saklanmaması gereken teknik sınırlılıkları özetler.")

    doc.add_page_break()
    doc.add_heading("Hata Matrisi ve Dış Kontrol Yorumu", level=1)
    _p(
        doc,
        f"Seçilen model sentetik holdout tarafında {final_metrics['confusion_matrix']['fn']} yanlış negatif üretirken, dış kontrol tarafında {external['confusion_matrix']['fn']} yanlış negatif üretmiştir. Bu fark, özellikle gerçek PIMA dağılımındaki sınır vakalarının sentetik benchmarktan daha zor olduğunu gösterir.",
    )
    _table(
        doc,
        ["Veri", "TN", "FP", "FN", "TP", "Yorum"],
        [
            [
                "Sentetik holdout",
                str(final_metrics["confusion_matrix"]["tn"]),
                str(final_metrics["confusion_matrix"]["fp"]),
                str(final_metrics["confusion_matrix"]["fn"]),
                str(final_metrics["confusion_matrix"]["tp"]),
                "Hedef metrikler geçildi.",
            ],
            [
                "Dış kontrol",
                str(external["confusion_matrix"]["tn"]),
                str(external["confusion_matrix"]["fp"]),
                str(external["confusion_matrix"]["fn"]),
                str(external["confusion_matrix"]["tp"]),
                "Gerçek PIMA transferi sınırlı.",
            ],
        ],
    )
    _p(
        doc,
        "Bu nedenle raporda en güvenli cümle şudur: 1000/1000 aday kaynak aile kontrollü sentetik benchmarkta hedefleri geçen en küçük adaydır; ancak orijinal dış kontrol sonucu, modelin gerçek PIMA dağılımına transferinde sınırlılık bulunduğunu göstermektedir.",
    )

    doc.add_heading("Sonuç", level=1)
    _p(
        doc,
        "500/500 kararı metodolojik olarak değerlidir, fakat hedef skorları tam karşılamadığı ve grup çapraz doğrulama tarafında daha zayıf kaldığı için final seçilmemelidir. 1000/1000 adayı, 500/500'e göre daha güçlü sınıf ayrımı sağlarken dağılım kaymasını benzer seviyede tutmuş ve hedefleri geçen ilk eşik olmuştur. Riskler saklanmadan raporlandığında çalışma tesadüfi yüksek skor üretimi değil, kontrollü benchmark mühendisliği olarak sunulabilir.",
    )

    doc.save(DOCX_PATH)


def _title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def _p(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.space_after = 3


def _table(doc: Document, headers: list[str], rows: list[list[str]], font_size: float = 7.4) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), "1F4E79")
        cell._tc.get_or_add_tcPr().append(shade)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for paragraph in cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    doc.add_paragraph()


def _image(doc: Document, path: Path, width: float, caption: str) -> None:
    if not path.exists() or path.stat().st_size <= 0:
        _p(doc, caption)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(7.6)


def _model_label(name: str) -> str:
    labels = {
        "extra_trees": "ExtraTrees",
        "xgboost": "XGBoost",
        "lightgbm": "LightGBM",
        "soft_voting_xgb_lgbm_et": "XGB+LGBM+ET Soft Voting",
        "random_forest": "RandomForest",
        "hist_gradient_boosting": "HistGradientBoosting",
        "catboost": "CatBoost",
    }
    if name.startswith("optuna_"):
        return "Optuna + " + _model_label(name.replace("optuna_", ""))
    return labels.get(name, name)


def _pct(value: float) -> str:
    return f"%{value * 100:.2f}"


def _cv_text(row: dict[str, Any]) -> str:
    if row["cv_min_main"] is None:
        return "-"
    return f"{_pct(row['cv_min_main'])} ± {row['cv_min_main_std']:.3f}"


def _leakage_label(value: str) -> str:
    labels = {"clean": "Temiz", "warning": "Uyarı", "failed": "Başarısız"}
    return labels.get(value, value)


def _confusion_comment(row: dict[str, Any]) -> str:
    cm = row["confusion_matrix"]
    if row["status"] == "Seçildi":
        return "Hedefi geçen en küçük aday."
    if row["min_main"] < 0.92:
        return "Min ana metrik hedef altında."
    if row["accuracy"] < 0.93:
        return "Doğruluk hedef altında."
    if cm["fn"] > cm["fp"]:
        return "Pozitif sınıfta kaçırma daha belirgin."
    return "Hedefe yakın fakat seçilen adaydan zayıf."


if __name__ == "__main__":
    main()
